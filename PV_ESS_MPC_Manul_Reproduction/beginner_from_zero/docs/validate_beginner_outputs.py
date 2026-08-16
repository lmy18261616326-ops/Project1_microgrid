from pathlib import Path
from docx import Document

root = Path(r"D:\PV_MPPT\PV_ESS_MPC_Paper_Reproduction\beginner_from_zero")
docx = root / "docs" / "PV_ESS_MPC_From_Zero_Beginner_Manual.docx"
required = [
    root / "scripts" / "pvess_beginner_init.m",
    root / "controllers" / "pvess_mpvc_block.m",
    root / "controllers" / "pvess_mppc_block.m",
]
for path in [docx, *required]:
    assert path.exists() and path.stat().st_size > 0, path

doc = Document(docx)
text = "\n".join(p.text for p in doc.paragraphs)
assert "�" not in text
for stage in range(16):
    assert f"阶段 {stage}" in text, stage
for token in ["附录 A", "附录 B", "附录 C", "附录 D", "完成判据"]:
    assert token in text, token

headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
images = len(doc.inline_shapes)
assert len(doc.tables) >= 80
assert images >= 6
assert len(headings) >= 70
print({"paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "headings": len(headings), "images": images})

