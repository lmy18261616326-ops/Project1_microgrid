from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

EXP = Path(r"D:\PV_MPPT\PV_ESS_MPC_Manul_Reproduction\experiments\stage15_bus_stability_20260814")
OUT = EXP / "PV_ESS_Stage15_Final_Engineering_Validation_Report_20260817.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E8F3EC"
PALE_GOLD = "FFF3CD"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, italic=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=False, page_break_before=False):
    pPr = paragraph._p.get_or_add_pPr()
    for tag, enabled in (("keepNext", keep_next), ("keepLines", keep_lines),
                         ("pageBreakBefore", page_break_before)):
        existing = pPr.find(qn(f"w:{tag}"))
        if enabled and existing is None:
            pPr.append(OxmlElement(f"w:{tag}"))
        elif not enabled and existing is not None:
            pPr.remove(existing)


def paragraph_border_bottom(paragraph, color=BLUE, size=12, space=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[min(idx, len(widths_dxa)-1)]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_cell_text(cell, text, bold=False, color=INK, size=9.0, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths_dxa, alignments=None, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    repeat_header(table.rows[0])
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], PALE_BLUE)
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=DARK_BLUE,
                      size=9.0, align=(alignments[i] if alignments else None))
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, size=font_size,
                          align=(alignments[i] if alignments else None))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_h1(doc, text, page_break=False):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    set_paragraph_keep(p, keep_next=True, keep_lines=True, page_break_before=page_break)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    set_paragraph_keep(p, keep_next=True, keep_lines=True)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_callout(doc, label, text, fill=PALE_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, italic=True, color=MUTED)
    set_paragraph_keep(p, keep_lines=True)


def add_figure(doc, filename, caption, width=6.45):
    path = EXP / filename
    if not path.exists():
        add_body(doc, f"[图像缺失：{filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    try:
        docPr = shape._inline.docPr
        docPr.set("descr", caption)
        docPr.set("title", caption)
    except Exception:
        pass
    set_paragraph_keep(p, keep_next=True)
    add_caption(doc, caption)


def add_page_field(paragraph, field):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = field
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, size=8.5, color=MUTED)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Standard business brief tokens.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Bullet 2", "List Number"):
    style = doc.styles[list_name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.167
if "List Bullet" in doc.styles:
    doc.styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
    doc.styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)
if "List Number" in doc.styles:
    doc.styles["List Number"].paragraph_format.left_indent = Inches(0.5)
    doc.styles["List Number"].paragraph_format.first_line_indent = Inches(-0.25)

caption_style = doc.styles["Caption"]
caption_style.font.name = "Calibri"
caption_style.font.size = Pt(9)
caption_style.font.italic = True
caption_style.font.color.rgb = RGBColor.from_string(MUTED)
caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

# Quiet running header/footer.
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
hr = hp.add_run("PV-ESS MPC  |  STAGE 15 ENGINEERING VALIDATION")
set_run_font(hr, size=8.5, bold=True, color=MUTED)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp.paragraph_format.space_before = Pt(0)
fr = fp.add_run("2026-08-17   |   Page ")
set_run_font(fr, size=8.5, color=MUTED)
add_page_field(fp, "PAGE")
fr = fp.add_run(" of ")
set_run_font(fr, size=8.5, color=MUTED)
add_page_field(fp, "NUMPAGES")

# Memo masthead opening.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("SIMULINK / SPECIALIZED POWER SYSTEMS 技术审计")
set_run_font(r, size=10, bold=True, color=BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
r = p.add_run("PV-ESS混合交流/直流微电网")
set_run_font(r, size=24, bold=True, color=INK)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(15)
r = p.add_run("阶段15最终工程修改与验证报告")
set_run_font(r, size=17, bold=True, color=DARK_BLUE)

meta = [
    ("对象", "manul_model_stage15_complete.slx（实验副本，原模型未修改）"),
    ("范围", "0.6 s交流负载阶跃、完整4 s事件序列、电池/并网保护、AC/DC恒功率负载"),
    ("依据", "Hu等，Applied Energy 221 (2018)，第4.2节、表1、表2、图16-17"),
    ("环境", "MATLAB/Simulink R2025a；Specialized Power Systems；电力步长2 us"),
    ("状态", "工程验证通过；属于18 mF工程变体，不宣称严格6 mF逐参数复现"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "：")
    set_run_font(r, size=10.5, bold=True, color=INK)
    r = p.add_run(value)
    set_run_font(r, size=10.5, color=INK)

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(10)
rule.paragraph_format.space_after = Pt(12)
paragraph_border_bottom(rule, color=BLUE, size=14, space=8)

add_callout(doc, "最终判定", "当前实验副本在真实AC/DC恒功率负载下完成阶段15全事件序列。"
            "1 MW直流CPL投入时Vdc最低1079.80 V（约0.90 pu），11.04 ms回到±2%带；"
            "0.6 s交流CPL阶跃后稳态1199.75±2.20 V。未使用理想电压钳位、波形回放、"
            "负载降额、额定功率放宽或继续增加电容来获得结果。")

add_h1(doc, "1. 执行摘要")
add_body(doc, "最初的0.6 s失稳并非单一双环PI故障，而是功率前馈方向、储能功率级动态、"
         "三相物理接线和时序输入外推共同作用。完成根因修复后，本轮进一步把原有工程薄弱点"
         "变成可执行保护，并用恒功率负载重新施压。")
add_table(doc,
          ["审计维度", "状态", "关键证据", "结论"],
          [
              ["0.6 s交流CPL阶跃", "通过", "1158.51-1220.41 V；稳态1199.75±2.20 V", "孤岛双环可恢复"],
              ["完整4 s AC/DC CPL", "通过", "Vdc全局1079.80-1368.67 V；最终1200.91 V", "全部事件有限稳定"],
              ["电池额定边界", "工程通过", "-501.0/+505.8 kW；动态限制实际动作", "约1.16%开关暂态峰值"],
              ["并网容量边界", "通过", "2.5 MVA、690 V侧限流；正常工况动作0 s", "无误限幅"],
              ["恒功率真实性", "通过", "0.6/1.0 MW误差0.0057%/0.0132%", "负增量阻抗真实参与"],
              ["严格论文复现", "不宣称", "Cdc=18 mF，论文为6 mF；电网等值简化", "定位为工程复现变体"],
          ],
          [1800, 1050, 3700, 2810],
          [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
           WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
          font_size=8.5)

add_h1(doc, "2. 原问题位置、原因与因果链", page_break=True)
add_h2(doc, "2.1 0.6 s交流负载增加时的原始失稳")
add_body(doc, "孤岛运行时，MPVC建立交流母线，电池双向DC/DC的外电压-内电流双环维持1200 V直流母线。"
         "交流负载由0.5 MW增加到1.0 MW后，电池必须由充电快速反向到放电。原模型的功率前馈"
         "方向却在缺功率时把电流指令推向充电侧；11.25 mH电感又限制了电流反向速度，最终使"
         "Vdc下降并削弱逆变器电压裕量。")
add_callout(doc, "确认的失稳链", "Pac上升 → Pnet=Ppv-Pac-Pdc变负 → Iforward=Pnet/Vbat为负 → "
            "原“PI+Iforward”降低放电指令 → 大电感限制di/dt → Vdc失去支撑 → 交流桥电压裕量不足。",
            fill=PALE_GOLD)

add_h2(doc, "2.2 同时存在的非控制器问题")
for item in [
    "交流阶跃负载断路器B/C相曾落入同一物理网络，导致负载投入后三相不平衡；该问题不能靠PI或MPC权重补偿。",
    "直流负载与电网标幺输入使用线性外推，最后一个阶跃边沿会被无限延伸，可能产生负电阻或巨大负电压指令。",
    "完整论文事件与0.6 s独立回归使用了不同的基础负载定义，若混用会错误比较事件功率。",
    "低负载、PV仍发电且电池充电饱和时，剩余能量没有物理去路，需要源侧限发而不是继续提高PI。",
]:
    add_bullet(doc, item)

add_h1(doc, "3. 已实施的根因修复", page_break=True)
add_h2(doc, "3.1 功率方向、储能功率级和双环")
add_number(doc, "将Sum10由“PI + Iforward”改为“PI - Iforward”，统一正电池电流=向母线放电的符号定义。")
add_number(doc, "电池串联电感由11.25 mH恢复为0.5 mH工程值，使大信号电流反向速度满足母线能量缺口。")
add_number(doc, "电流PI由Kp=0.1、Ki=0.15改为Kp=0.005、Ki=0.05，避免约2.5 A误差即撞击±0.25占空比校正限幅。")
add_number(doc, "保留电压外环Kp=4、Ki=650与100 us采样；占空比最终范围仍为0.05-0.95。")

add_h2(doc, "3.2 拓扑、输入与能量管理")
add_bullet(doc, "重新连接Three-Phase Breaker2的C相到独立C相母线；0.6 s回归三相不平衡为0.1899%。")
add_bullet(doc, "Rdc1/Rdc2、辐照度和电网标幺输入改为保持最终值，消除阶跃后的非物理外推。")
add_bullet(doc, "完整场景基础交流负载设为1 W近似开路，0.5 s投入0.5 MW；独立回归通过SimulationInput覆盖为0.5→1.0 MW。")
add_bullet(doc, "新增PV过压限发、锁存和无扰切换：电池充电达到边界后，通过真实Boost占空比降低PV功率。")
add_bullet(doc, "将3.5 s事件实现为可编程三相电源幅值真实下降，并保留电网RL等值和接地。")

add_h1(doc, "4. 本轮工程保护与抗饱和改造", page_break=True)
add_h2(doc, "4.1 Battery_Reference_Protection")
add_body(doc, "原SOC保护被替换为由标准Simulink块组成的分层电流参考保护。它以实测电池电压、"
         "硬件电流额定和0.5 MW端口功率额定共同确定可执行电流，并在SOC边界附近平滑降额。")
add_table(doc,
          ["项目", "最终设置", "修改原因", "影响"],
          [
              ["Vbat处理", "200 Hz一阶LPF；270 V除数下限", "抑制开关纹波并避免低压除零", "仅参与参考换算，不钳位母线"],
              ["放电上限", "min(1666.7 A, 500 kW/Vbat_safe)", "同时满足器件电流与端口功率", "Vbat变化时额定边界一致"],
              ["充电下限", "-min(1666.7 A, 500 kW/Vbat_safe)", "充放电采用同一功率原则", "负电流为充电"],
              ["SOC放电降额", "10%以下0；10-15%线性恢复", "避免SOC下限硬切换", "改善模式切换连续性"],
              ["SOC充电降额", "85%保持1；85-90%降到0", "提前释放充电能力", "与90%上限一致"],
              ["PI跟踪", "Tracking Kt=143；一拍延迟", "外环输出受限时防积分累积", "消除跨层饱和脱节"],
              ["内环输出", "Dcorr=-0.25…+0.20", "与Db0=0.75、最终Db≤0.95一致", "不再产生不可执行正校正"],
          ],
          [1600, 2500, 2900, 2360], font_size=8.3)
add_figure(doc, "stage15_battery_reference_protection_arranged.png",
           "图1  Battery_Reference_Protection自动布局后的标准块接线", width=6.45)

add_h2(doc, "4.2 网侧2.5 MVA P/Q电流限制")
add_body(doc, "新增Grid_PQ_Current_Limiter，使用实测690 V低压侧三相电压计算可用视在功率。"
         "在电压支撑时给予Q优先，剩余容量分配给P：Smax=√3·VLL·Irated，"
         "Pallow=√(Smax²-Qcmd²)。")
add_table(doc,
          ["参数/逻辑", "最终值", "物理含义"],
          [
              ["变压器/变换器额定", "2.5 MVA", "论文额定容量"],
              ["测量侧", "690 V低压侧", "grid_VI实际所在位置"],
              ["额定电流", "2091.85 A rms", "2.5 MVA/(√3·690 V)"],
              ["电压数值下限", "69 V line-line", "仅用于容量除算，避免近零电压病态"],
              ["优先级", "Q优先，P取剩余容量", "保证电压跌落时无功支撑"],
              ["最终4 s动作时间", "0 s", "正常工况未被错误限幅"],
          ],
          [2200, 2000, 5160], font_size=8.8)
add_callout(doc, "保留的失败证据", "第一次把690 V测量误按25 kV侧额定电流57.7 A换算，"
            "得到约250 kVA错误容量，造成Vdc 929-1782 V且Q被限制到约-250 kvar。该结果已以"
            "failed_wrong_side_grid_limit_20260817后缀保存，随后改为2091.85 A低压侧额定；未把失败结果删除或当作通过。",
            fill=PALE_GOLD)
add_figure(doc, "stage15_grid_pq_current_limiter_arranged.png",
           "图2  Grid_PQ_Current_Limiter自动布局后的Q优先容量圆限制", width=6.45)

add_h1(doc, "5. AC/DC恒功率负载实现", page_break=True)
add_h2(doc, "5.1 交流恒功率负载")
add_body(doc, "两组有功负载继续使用Specialized Power Systems自带Three-Phase Parallel RLC Load，"
         "只把LoadType由constant Z切换为constant PQ。滤波电容支路仍保持恒阻抗，不把无功滤波器误当成恒功率负载。")

add_h2(doc, "5.2 直流恒功率负载")
add_body(doc, "保留原有SPS Variable Resistor电气元件，新建22个标准Simulink块的"
         "DC_CPL_Resistance_Command，以50 us更新周期计算R=Vdc(z-1)²/P。"
         "一拍延迟代表数字负载控制并打断理想代数环；840 V电压下限与0.2 Ω电阻下限是"
         "电流保护边界，不直接钳制Vdc。")
add_table(doc,
          ["项目", "设置", "验证结果"],
          [
              ["采样周期", "50 us", "与MPC更新周期一致"],
              ["电压保护下限", "840 V (0.70 pu)", "最终最低1079.80 V，未触发"],
              ["开路电阻", "1 MΩ", "P≤100 W时近似断开"],
              ["最小电阻", "0.2 Ω", "与SPS Variable Resistor限制一致"],
              ["0.6 MW跟踪", "600034 W", "误差+0.0057%"],
              ["1.0 MW跟踪", "1000132 W", "误差+0.0132%"],
              ["实际最小R2", "1.1679 Ω", "明显高于840 V保护对应0.7056 Ω"],
          ],
          [2200, 2200, 4960], font_size=8.8)
add_figure(doc, "stage15_dc_cpl_resistance_command_arranged.png",
           "图3  DC_CPL_Resistance_Command：标准块实现的R=Vdc²/P与开路/物理边界", width=6.45)

add_h1(doc, "6. 并网直流母线外环重新整定", page_break=True)
add_h2(doc, "6.1 为什么20 Hz在CPL下出现深跌落")
add_body(doc, "论文第4.2节明确规定并网模式由互联变换器的直流母线外环PI生成有功参考。"
         "因此没有加入“预知负载功率前馈”，以免改变论文结构。现有20 Hz、ζ≈0.707参数为"
         "Kp≈3839 W/V、Ki≈3.41×10^5 W/(V·s)；要改变1 MW指令，仅比例项就需要约260 V误差。"
         "真实1 MW CPL下最低Vdc因此降到870.35 V。")
add_body(doc, "按小信号能量模型C·Vdc·dVdc/dt=ΔP，同阻尼参数为："
         "Kp=2ζωnCdcVdc，Ki=ωn²CdcVdc。本次只扫描闭环带宽，保持Cdc、负载、额定容量和保护阈值不变。")
add_table(doc,
          ["带宽", "Kp (W/V)", "Ki (W/(V·s))", "1 MW时Vdc最低", "恢复时间", "稳态σ", "限流"],
          [
              ["20 Hz", "3838.7", "3.411e5", "870.35 V", "45.92 ms", "2.789 V", "0 s"],
              ["30 Hz", "5758.0", "7.675e5", "990.24 V", "28.64 ms", "2.817 V", "0 s"],
              ["40 Hz", "7677.3", "1.364e6", "1046.66 V", "14.56 ms", "2.721 V", "0 s"],
              ["50 Hz", "9596.6", "2.132e6", "1077.64 V", "10.72 ms", "2.650 V", "0 s"],
          ],
          [1050, 1250, 1650, 1500, 1250, 1200, 1460],
          [WD_ALIGN_PARAGRAPH.CENTER]*7, font_size=8.3)
add_callout(doc, "选型结论", "50 Hz候选的P参考峰值约2.03 MW，低于2.2 MW命令边界；"
            "完整4 s回归中最低1079.80 V，恢复11.04 ms，且负载切除过渡也改善。"
            "因此采用50 Hz，而不是继续追求更高带宽。")
add_figure(doc, "stage15_grid_pi_cpl_sweep.png",
           "图4  20/30/40/50 Hz同阻尼外环扫描（相同硬件与负载）", width=6.45)

add_h1(doc, "7. 仿真验证结果", page_break=True)
add_h2(doc, "7.1 0.6 s交流恒功率负载0.5→1.0 MW")
add_table(doc,
          ["指标", "结果"],
          [
              ["阶跃前Vdc", "1200.026±1.905 V"],
              ["暂态范围", "1158.513-1220.411 V"],
              ["0.82-0.90 s稳态", "1199.746±2.195 V"],
              ["稳态三相RMS", "396.561 / 397.315 / 397.081 V"],
              ["三相不平衡", "0.18991%"],
              ["互联交流功率", "0.993146 MW"],
              ["电池稳态功率", "+0.271857 MW（正值放电）"],
              ["电池/并网限流", "0 s / 0 s"],
          ],
          [3200, 6160], font_size=9.0)
add_figure(doc, "ac_step_0p6s_regression.png",
           "图5  0.6 s交流恒功率负载阶跃回归", width=6.45)

add_h2(doc, "7.2 完整4 s阶段15事件序列")
add_table(doc,
          ["事件窗", "时间(s)", "Vdc均值±标准差", "最小-最大"],
          [
              ["交流负载前", "0.30-0.45", "1208.686±0.724 V", "1207.001-1210.516 V"],
              ["交流负载后", "0.70-0.90", "1199.709±1.753 V", "1194.897-1204.435 V"],
              ["0.6 MW DC CPL后", "1.12-1.20", "1199.181±2.416 V", "1192.712-1205.073 V"],
              ["并网后", "1.70-1.75", "1199.226±2.785 V", "1193.611-1205.258 V"],
              ["高辐照并网", "2.20-2.40", "1199.977±9.817 V", "1164.229-1226.057 V"],
              ["1 MW DC CPL投入后", "2.58-2.70", "1199.885±2.644 V", "1194.351-1207.153 V"],
              ["1 MW DC CPL切除后", "3.10-3.20", "1197.475±6.911 V", "1174.406-1213.205 V"],
              ["电网跌落后", "3.75-3.90", "1199.639±1.951 V", "1195.272-1204.591 V"],
          ],
          [2600, 1300, 2700, 2760], font_size=8.0)
add_body(doc, "全局Vdc最低1079.795 V（2.50432 s），最高1368.666 V（0.02656 s启动暂态），"
         "最终1200.905 V。全部记录信号有限；MPVC/MPPC代价函数最大值分别0.99618/0.48081。")
add_figure(doc, "stage15_complete_overview.png",
           "图6  最终阶段15：直流母线、有功平衡与Q-V下垂", width=6.45)

add_h2(doc, "7.3 1 MW直流CPL大扰动")
add_body(doc, "1 MW负载在2.5 s保持真实恒功率，Vdc下降时电阻随Vdc²减小，未靠恒阻抗的自然降载减轻冲击。"
         "50 Hz网侧外环使有功参考更快由出口转为入口，Vdc最低约0.90 pu并在11.04 ms回到±2%带。")
add_figure(doc, "stage15_cpl_1MW_event_zoom.png",
           "图7  1 MW DC CPL投入：母线、电源/负载功率和并网P参考", width=6.45)

add_h2(doc, "7.4 电网电压跌落与无功支撑")
add_table(doc,
          ["指标", "结果"],
          [
              ["PCC跌落前相RMS", "402.803 / 403.285 / 403.346 V"],
              ["PCC跌落后相RMS", "354.475 / 354.354 / 354.743 V"],
              ["平均跌落比", "0.879396 pu"],
              ["Qref", "-0.400000 Mvar"],
              ["实际Q", "-0.397603 Mvar"],
              ["Vdc事件后窗口", "1199.639±1.951 V"],
          ],
          [3500, 5860], font_size=9.0)
add_figure(doc, "stage15_grid_dip_zoom.png",
           "图8  3.5 s实际电网电压跌落与Q-V无功响应", width=6.45)

add_h1(doc, "8. 物理与电气原理合规审计", page_break=True)
add_h2(doc, "8.1 没有采用的结果导向手段")
for item in [
    "没有在直流母线并联隐藏理想受控电压源，也没有对Vdc测量信号做Saturation来伪造稳定。",
    "没有回放预制稳定波形、修改To Workspace数据或改变后处理窗口来掩盖失稳。",
    "没有把1 MW阶跃改成缓坡，没有降低负载功率，也没有提高2.5 MVA并网容量或0.5 MW电池额定。",
    "没有为改善CPL结果继续增大18 mF母线电容；用户已接受不使用论文6 mF的工程差异。",
    "没有加入论文未使用的负载功率预知前馈；最终改善来自论文外环PI的可解释带宽整定。",
]:
    add_bullet(doc, item)

add_h2(doc, "8.2 需要明确披露的工程差异")
add_table(doc,
          ["项目", "当前模型", "论文/严格复现", "影响与边界"],
          [
              ["DC-link电容", "18 mF", "6 mF", "物理可实现但储能为论文3倍；不能称逐参数复现"],
              ["电网等值", "25 kV可编程源+显式RL", "25/120 kV、19 km馈线", "PCC跌落0.879 pu，不等于论文约0.908 pu"],
              ["启动", "Vdc初值1200 V", "硬件需预充", "1368.7 V最大值发生在0.0266 s启动阶段"],
              ["PV限发", "1220/1185 V锁存阈值", "论文主要按SOC触发Off-MPPT", "属于新增能量守恒保护，需跨温度/辐照扫描"],
              ["CPL电压保护", "840 V", "论文未给", "本轮未触发，不影响最终波形"],
          ],
          [1800, 1900, 1900, 3760], font_size=8.1)

add_h2(doc, "8.3 电池与限流结果解释")
add_body(doc, "电池端口最大约+505.78 kW，为0.5 MW额定的约1.16%瞬时超调；最小约-501.02 kW。"
         "动态电流参考限制累计动作约0.3904 s，说明保护确实参与而不是闲置。"
         "端口功率仍受开关纹波、滤波延迟和电池内阻影响，因此参考限幅不能数学上保证每个2 us样本都≤500 kW；"
         "若作为硬件设计，应另加器件峰值/热模型与快速过流跳闸。")

add_h1(doc, "9. 改动影响与工程结论", page_break=True)
add_h2(doc, "9.1 正面影响")
for item in [
    "0.6 s交流负载问题在恒功率模型下仍可恢复，证明修复不是恒阻抗负载自然降载造成。",
    "电池功率/电流/SOC约束、PI跟踪和可执行占空比边界形成一致的跨层保护链。",
    "并网P/Q命令具有2.5 MVA容量圆约束，且690 V/25 kV测量侧量纲已核实。",
    "网侧外环由20 Hz提高到50 Hz后，1 MW CPL最低电压提高约209.44 V，恢复时间减少约34.88 ms。",
    "自动布局后新增子系统按信号流展开，结构审计为0个普通输入悬空。",
]:
    add_bullet(doc, item)

add_h2(doc, "9.2 代价与剩余风险")
for item in [
    "50 Hz外环提高对高频功率纹波的敏感度；高辐照并网窗口Vdc标准差为9.82 V，仍需硬件噪声/延迟裕度验证。",
    "0.5 mH电池电感提高电流纹波和di/dt，需要核对磁芯饱和、电感铜损和开关器件峰值电流。",
    "18 mF电容增加体积、成本和预充能量；其可接受性取决于项目工程目标，而非论文一致性。",
    "当前测试覆盖仿真稳定性，不等于已完成HIL、实时性、器件热模型、故障穿越和保护配合认证。",
]:
    add_bullet(doc, item)

add_h2(doc, "9.3 当前项目定位")
add_callout(doc, "可用于简历的准确表述", "“基于MATLAB/Simulink搭建并审计2 MW级PV-ESS混合AC/DC微电网，"
            "复现MPVC/MPPC、SOC-EMS、并网Q-V支撑；定位并修复储能前馈符号、三相接线和双环大信号问题；"
            "实现动态电池功率/SOC保护、2.5 MVA P/Q容量限制及AC/DC恒功率负载；"
            "通过参数扫描把1 MW CPL母线最低值由870 V提高到1080 V附近，恢复时间由45.9 ms降至11.0 ms。”")
add_body(doc, "不建议写成“严格1:1复现论文全部参数”或“完成硬件验证”。更准确的定位是："
         "有论文依据、具备电气拓扑和能量平衡审计、包含失败记录与压力测试的工程级Simulink项目。")

add_h1(doc, "10. 文件、备份与可追溯性", page_break=True)
add_table(doc,
          ["对象", "路径/标识"],
          [
              ["原模型（未修改）", r"D:\PV_MPPT\PV_ESS_MPC_Manul_Reproduction\models\manul_model.slx"],
              ["最终实验副本", str(EXP / "manul_model_stage15_complete.slx")],
              ["最终初始化文件", str(EXP / "pvess_stage15_busfix_init.m")],
              ["最终4 s结果", str(EXP / "stage15_complete_results.mat")],
              ["最终摘要", str(EXP / "stage15_summary.json")],
              ["恢复时间", str(EXP / "stage15_detailed_metrics.json")],
              ["PI扫描", str(EXP / "stage15_grid_pi_cpl_sweep.csv")],
              ["结构审计", str(EXP / "verify_stage15_cpl_structure.m")],
              ["错误测量侧失败工件", "*_failed_wrong_side_grid_limit_20260817.*"],
              ["20 Hz CPL基线", "*_cpl_before_gridpi_20260817.*"],
          ],
          [2600, 6760], font_size=8.0)

add_h2(doc, "10.1 SHA-256")
add_table(doc,
          ["对象", "SHA-256"],
          [
              ["原模型", "AAB7C51A6FD264485FB2C041A21602A5721AB50476D1BA7D1D2BD8E8EAA02D28"],
              ["最终副本", "4969A09B1F3A72B315029525D5A0E4CCAC25DD177155310B94A796BBD267F8EA"],
              ["工程加固前备份", "DAE378D84BDB6FC7A9BB0435F4492D29CB0B8815AD75B3AD9986AD3379E2EA66"],
              ["CPL前备份", "42E77400A4B5B78A0927B1088B73958582B3C5C5B9835963C479C0B10D58F8E7"],
          ],
          [2300, 7060], font_size=8.2)

add_h2(doc, "10.2 后续建议")
for item in [
    "建立独立6 mF严格论文参数分支并重新整定，而不是覆盖当前18 mF工程版本。",
    "增加270-367 V电池电压、10-90% SOC、400-1000 W/m²辐照度和温度扫描。",
    "加入预充电阻、接触器时序、IGBT/二极管峰值电流、电感饱和及热模型。",
    "执行传感器噪声、采样延迟、参数偏差和电网阻抗Monte Carlo/鲁棒性测试。",
    "最后开展实时仿真或HIL，验证20 kHz MPPC/MPVC的实际计算预算和保护时序。",
]:
    add_bullet(doc, item)

add_h1(doc, "附录A  验收定义与证据边界", page_break=True)
add_body(doc, "恢复时间定义为：事件后首次连续20 ms保持在0.98-1.02 pu内。"
         "仿真结果来自To Workspace保存的真实测量链：Vdc由Voltage Measurement获得；Ppv/Pbat由实测电压电流乘积获得；"
         "交流P/Q由三相电压电流计算和移动平均获得。")
add_body(doc, "本报告边界：截至2026-08-17的实验副本、保存结果和结构检查。"
         "自动布局只改变模块位置和线段路由，不改变控制参数或电气端口；因此布局后未重复4 s数值仿真，但已重新执行模型更新和结构接线检查。")

add_h1(doc, "附录B  参考资料")
add_bullet(doc, "J. Hu et al., “Predictive control strategy of a photovoltaic energy storage microgrid under variable power generation and load conditions,” Applied Energy 221 (2018) 195-203.")
add_bullet(doc, r"建设说明：D:\PV_MPPT\PV_ESS_MPC_Manul_Reproduction\beginner_from_zero\docs\PV_ESS_MPC_From_Zero_Beginner_Manual.docx")
add_bullet(doc, r"论文PDF：D:\PV_MPPT\reference\光伏储能微电网在可变电力输出和负载条件下的预测控制策略.pdf")

# Core properties and final save.
doc.core_properties.title = "PV-ESS Stage 15 Final Engineering Validation Report"
doc.core_properties.subject = "Simulink PV-ESS MPC engineering modifications, CPL stress tests, and physical compliance audit"
doc.core_properties.author = "OpenAI Codex"
doc.core_properties.keywords = "PV ESS MPC Simulink Stage 15 constant power load validation"
doc.save(OUT)
print(str(OUT))
