from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\PV_MPPT\PV_ESS_MPC_Paper_Reproduction")
BEGINNER = ROOT / "beginner_from_zero"
DOC_DIR = BEGINNER / "docs"
ASSETS = DOC_DIR / "assets"
OUT = DOC_DIR / "PV_ESS_MPC_From_Zero_Beginner_Manual.docx"
INIT_SCRIPT = BEGINNER / "scripts" / "pvess_beginner_init.m"
MPVC_SCRIPT = BEGINNER / "controllers" / "pvess_mpvc_block.m"
MPPC_SCRIPT = BEGINNER / "controllers" / "pvess_mppc_block.m"
TEMPLATE_IMAGE = ROOT / "docs" / "model_root.png"
PAPER = Path(r"D:\PV_MPPT\reference\光伏储能微电网在可变电力输出和负载条件下的预测控制策略.pdf")

BLUE = "2E74B5"
DARK_BLUE = "17365D"
MID_BLUE = "5B9BD5"
PALE_BLUE = "E8EEF5"
LIGHT_BLUE = "DDEBF7"
PALE_GREEN = "E2F0D9"
PALE_YELLOW = "FFF2CC"
PALE_RED = "FCE4D6"
LIGHT_GRAY = "F2F4F7"
GRAY = "666666"
BLACK = "222222"


def rgb(value):
    return RGBColor.from_string(value)


def set_run_font(run, east_asia="Microsoft YaHei", ascii_font="Calibri", size=None):
    run.font.name = ascii_font
    if size is not None:
        run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        item = margins.find(qn("w:" + name))
        if item is None:
            item = OxmlElement("w:" + name)
            margins.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    node = tcpr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tcpr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_border(cell, color="B7C9DC", size="4"):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(Inches(width).twips for width in widths)
    props = table._tbl.tblPr
    width_node = props.first_child_found_in("w:tblW")
    if width_node is None:
        width_node = OxmlElement("w:tblW")
        props.append(width_node)
    width_node.set(qn("w:w"), str(total))
    width_node.set(qn("w:type"), "dxa")
    indent = props.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        props.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = props.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        props.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = list(table._tbl.tblGrid)
    for index, width in enumerate(widths):
        twips = Inches(width).twips
        if index < len(grid):
            grid[index].set(qn("w:w"), str(twips))
        for row in table.rows:
            cell_props = row.cells[index]._tc.get_or_add_tcPr()
            cell_width = cell_props.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_props.append(cell_width)
            cell_width.set(qn("w:w"), str(twips))
            cell_width.set(qn("w:type"), "dxa")


def repeat_header(row):
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    props.append(node)


def prevent_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True

    for name in ["List Bullet", "List Number", "List Bullet 2", "List Number 2"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375 if "2" not in name else 0.65)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in doc.styles:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.10)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(5)
    code_style.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("PV–ESS MPC 从零搭建手册")
    set_run_font(run, size=8.5)
    run.bold = True
    run.font.color.rgb = rgb(DARK_BLUE)
    run = header.add_run("    功率电路 · 分阶段送电 · 标准模块控制")
    set_run_font(run, size=8.2)
    run.font.color.rgb = rgb(GRAY)
    ppr = header._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), MID_BLUE)
    border.append(bottom)
    ppr.append(border)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_run_font(run, size=8)
    run.font.color.rgb = rgb(GRAY)
    add_field(footer, "PAGE")
    run = footer.add_run(" 页")
    set_run_font(run, size=8)
    run.font.color.rgb = rgb(GRAY)


def paragraph(doc, text="", style=None, align=None):
    item = doc.add_paragraph(style=style)
    run = item.add_run(text)
    set_run_font(run)
    if align is not None:
        item.alignment = align
    return item


def bullet(doc, text, level=0):
    return paragraph(doc, text, "List Bullet" if level == 0 else "List Bullet 2")


def number(doc, text, level=0):
    return paragraph(doc, text, "List Number" if level == 0 else "List Number 2")


def heading(doc, text, level=1):
    item = doc.add_paragraph(text, style="Heading " + str(level))
    for run in item.runs:
        set_run_font(run)
    return item


def code(doc, text):
    item = doc.add_paragraph(style="Code Block")
    run = item.add_run(text)
    set_run_font(run, ascii_font="Consolas", size=8.5)
    ppr = item._p.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), LIGHT_GRAY)
    ppr.append(shade)
    return item


def callout(doc, title, body, fill=PALE_YELLOW, border=MID_BLUE):
    box = doc.add_table(rows=1, cols=1)
    set_table_geometry(box, [6.5])
    cell = box.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 100, 140, 100, 140)
    set_cell_border(cell, border, "8")
    item = cell.paragraphs[0]
    run = item.add_run(title)
    set_run_font(run)
    run.bold = True
    run.font.color.rgb = rgb(DARK_BLUE)
    run = item.add_run("\n" + body)
    set_run_font(run)
    item.paragraph_format.space_after = Pt(0)
    prevent_split(box.rows[0])
    paragraph(doc, "")
    return box


def table(doc, headers, rows, widths, font_size=8.1, header_fill=PALE_BLUE):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    set_table_geometry(result, widths)
    for index, text in enumerate(headers):
        cell = result.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        set_cell_border(cell)
        item = cell.paragraphs[0]
        item.paragraph_format.space_after = Pt(0)
        item.paragraph_format.keep_with_next = True
        run = item.add_run(str(text))
        set_run_font(run, size=font_size)
        run.bold = True
        run.font.color.rgb = rgb(DARK_BLUE)
    repeat_header(result.rows[0])
    prevent_split(result.rows[0])
    for row_index, row in enumerate(rows):
        cells = result.add_row().cells
        for index, text in enumerate(row):
            cell = cells[index]
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2:
                set_cell_shading(cell, "F8FAFC")
            item = cell.paragraphs[0]
            item.paragraph_format.space_after = Pt(0)
            run = item.add_run(str(text))
            set_run_font(run, size=font_size)
        prevent_split(result.rows[-1])
    paragraph(doc, "")
    return result


def caption(doc, text):
    item = paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER)
    item.paragraph_format.space_before = Pt(2)
    item.paragraph_format.space_after = Pt(8)
    for run in item.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = rgb(GRAY)


def add_picture(doc, path, width=6.3, label=None):
    path = Path(path)
    if not path.exists():
        callout(doc, "图像缺失", str(path), fill=PALE_RED)
        return
    item = doc.add_paragraph()
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = item.add_run().add_picture(str(path), width=Inches(width))
    alt = label if label else path.stem
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)
    if label:
        caption(doc, label)


def load_font(size=28, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def draw_center(draw, box, text, font, fill=BLACK):
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + (len(lines) - 1) * 6
    y = box[1] + (box[3] - box[1] - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        x = box[0] + (box[2] - box[0] - w) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += h + 6


def draw_box(draw, box, text, fill=PALE_BLUE, outline=BLUE, font=None):
    draw.rounded_rectangle(box, radius=14, fill="#" + fill, outline="#" + outline, width=3)
    draw_center(draw, box, text, font or load_font(24), "#" + BLACK)


def draw_arrow(draw, start, end, color=BLUE, width=5):
    color_hex = "#" + color
    draw.line([start, end], fill=color_hex, width=width)
    import math
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    a1 = angle + 2.55
    a2 = angle - 2.55
    p1 = (end[0] + length * math.cos(a1), end[1] + length * math.sin(a1))
    p2 = (end[0] + length * math.cos(a2), end[1] + length * math.sin(a2))
    draw.polygon([end, p1, p2], fill=color_hex)


def create_assets():
    ASSETS.mkdir(parents=True, exist_ok=True)
    title_font = load_font(30, True)
    box_font = load_font(23, True)
    small_font = load_font(20)

    # Overall physical topology.
    img = Image.new("RGB", (1800, 650), "white")
    d = ImageDraw.Draw(img)
    boxes = [
        (60, 120, 300, 260, "PV Array"),
        (380, 120, 650, 260, "Boost\n5 kHz"),
        (750, 80, 1040, 300, "1200 V DC Bus\nCdc = 6 mF"),
        (1140, 120, 1430, 260, "Three-phase\n2-level bridge"),
        (1510, 120, 1750, 260, "690 V AC bus"),
        (740, 390, 1050, 540, "Battery 300 V\nBidirectional DC/DC"),
        (1480, 390, 1770, 540, "Transformer 2.5 MVA\n25 kV utility"),
    ]
    for x1, y1, x2, y2, label in boxes:
        draw_box(d, (x1, y1, x2, y2), label, font=box_font)
    draw_arrow(d, (300, 190), (380, 190))
    draw_arrow(d, (650, 190), (750, 190))
    draw_arrow(d, (1040, 190), (1140, 190))
    draw_arrow(d, (1430, 190), (1510, 190))
    draw_arrow(d, (895, 390), (895, 300))
    draw_arrow(d, (1630, 390), (1630, 260))
    d.text((60, 35), "最终功率电路拓扑（所有支路均先放隔离器）", font=title_font, fill="#" + DARK_BLUE)
    img.save(ASSETS / "physical_topology.png")

    # Commissioning ladder.
    img = Image.new("RGB", (1800, 760), "white")
    d = ImageDraw.Draw(img)
    labels = ["空白模型\npowergui", "预充直流母线", "PV + 开环 Boost", "增量电导 MPPT",
              "电池电流环", "电池母线电压环", "逆变器 + LC", "MPVC 孤岛",
              "同步与并网", "MPPC + EMS"]
    positions = []
    for idx, label in enumerate(labels):
        col = idx % 5
        row = idx // 5
        x1 = 55 + col * 345
        y1 = 105 + row * 330
        box = (x1, y1, x1 + 285, y1 + 150)
        positions.append(box)
        fill = PALE_GREEN if idx < 4 else (PALE_BLUE if idx < 8 else PALE_YELLOW)
        draw_box(d, box, f"{idx}. {label}", fill=fill, font=box_font)
    for idx in range(4):
        draw_arrow(d, (positions[idx][2], 180), (positions[idx + 1][0], 180))
    draw_arrow(d, (positions[4][0] + 140, positions[4][1]),
               (positions[4][0] + 140, positions[0][3] + 70))
    for idx in range(5, 9):
        draw_arrow(d, (positions[idx - 1][2], 510), (positions[idx][0], 510))
    d.text((55, 30), "从零搭建的唯一正确节奏：上一阶段通过后才接下一支路", font=title_font, fill="#" + DARK_BLUE)
    img.save(ASSETS / "commissioning_ladder.png")

    # MPPT logic.
    img = Image.new("RGB", (1800, 780), "white")
    d = ImageDraw.Draw(img)
    items = [
        ((60, 120, 310, 250), "Vpv, Ipv\nZOH 1 ms"),
        ((390, 120, 680, 250), "Unit Delay\nV(k-1), I(k-1)"),
        ((760, 120, 1050, 250), "ΔV, ΔI\nInc = ΔI/ΔV + I/V"),
        ((1130, 90, 1450, 280), "判断方向\n左侧 MPP: −ΔD\n右侧 MPP: +ΔD"),
        ((1510, 120, 1750, 250), "D(k)\nSaturation"),
        ((760, 430, 1050, 580), "|ΔV| < ε\n改用 ΔI 符号"),
        ((1170, 430, 1450, 580), "D(k−1) + ΔD"),
    ]
    for box, label in items:
        draw_box(d, box, label, fill=PALE_BLUE, font=small_font)
    for a, b in [((310,185),(390,185)),((680,185),(760,185)),((1050,185),(1130,185)),((1450,185),(1510,185))]:
        draw_arrow(d, a, b)
    draw_arrow(d, (900,250), (900,430))
    draw_arrow(d, (1050,505), (1170,505))
    draw_arrow(d, (1450,505), (1630,250))
    d.text((60, 35), "增量电导 MPPT：用标准模块搭建，不需要 MPPT 脚本", font=title_font, fill="#" + DARK_BLUE)
    img.save(ASSETS / "mppt_logic.png")

    # Battery loops.
    img = Image.new("RGB", (1800, 690), "white")
    d = ImageDraw.Draw(img)
    labels = [
        ((70, 140, 330, 280), "Vdc_ref − Vdc"),
        ((410, 140, 680, 280), "外环 PI\n→ Ibat_ref"),
        ((760, 140, 1040, 280), "Ibat_ref − Ibat"),
        ((1120, 140, 1390, 280), "内环 PI + Db0\n→ Duty"),
        ((1470, 140, 1740, 280), "互补门极\n双向 DC/DC"),
        ((390, 430, 700, 570), "并网时：SOC/EMS\n替代 Vdc 外环"),
    ]
    for box, label in labels:
        draw_box(d, box, label, fill=PALE_GREEN if box[1] < 300 else PALE_YELLOW, font=small_font)
    for a, b in [((330,210),(410,210)),((680,210),(760,210)),((1040,210),(1120,210)),((1390,210),(1470,210))]:
        draw_arrow(d, a, b)
    draw_arrow(d, (545,430), (545,280))
    d.text((70, 35), "电池控制：先调电流内环，再闭合母线电压外环", font=title_font, fill="#" + DARK_BLUE)
    img.save(ASSETS / "battery_loops.png")

    # FCS cycle.
    img = Image.new("RGB", (1800, 650), "white")
    d = ImageDraw.Draw(img)
    labels = ["采样 V/I", "枚举 8 个\n开关状态", "预测 k+1", "计算代价 J", "选择最小 J", "输出 6 门极"]
    coords = []
    for idx, label in enumerate(labels):
        x1 = 45 + idx * 295
        box = (x1, 190, x1 + 240, 350)
        coords.append(box)
        draw_box(d, box, label, fill=PALE_BLUE if idx < 4 else PALE_GREEN, font=small_font)
    for idx in range(len(coords)-1):
        draw_arrow(d, (coords[idx][2],270), (coords[idx+1][0],270))
    d.text((45, 55), "每 50 μs 执行一次；FCS 门极直接送桥臂，不再经过 PWM", font=title_font, fill="#" + DARK_BLUE)
    img.save(ASSETS / "fcs_cycle_beginner.png")


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("从空白 Simulink 模型开始")
    set_run_font(r, size=16)
    r.bold = True
    r.font.color.rgb = rgb(MID_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("PV–电池混合 AC/DC 微电网\n功率电路与 MPC 复现手册")
    set_run_font(r, size=27)
    r.bold = True
    r.font.color.rgb = rgb(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("鼠标搭建为主 · 分阶段送电 · 参数逐项解释 · 脚本最小化")
    set_run_font(r, size=13)
    r.font.color.rgb = rgb(BLUE)

    table(doc, ["交付项", "位置/用途"], [
        ["本手册", str(OUT)],
        ["参数初始化", str(INIT_SCRIPT)],
        ["孤岛 MPVC 算法核", str(MPVC_SCRIPT)],
        ["并网 MPPC 算法核", str(MPPC_SCRIPT)],
        ["论文", str(PAPER)],
    ], [1.55, 4.95], font_size=8.4, header_fill=PALE_GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(f"适用：MATLAB/Simulink R2025a；生成日期：{date.today().isoformat()}")
    set_run_font(r, size=9.5)
    r.font.color.rgb = rgb(GRAY)
    doc.add_page_break()


def add_toc(doc):
    heading(doc, "目录", 1)
    paragraph(doc, "以下索引无需刷新即可阅读；在 Word 中按 Ctrl+A，再按 F9，可刷新可点击目录和页码。")
    table(doc, ["阶段", "从零搭建任务", "阶段", "从零搭建任务"], [
        ["0", "空白模型与 powergui", "8", "SOC 与 Vdc 外环"],
        ["1", "预充 1200 V 直流母线", "9", "三相桥、LC 与交流负载"],
        ["2", "PV Array 静态验证", "10", "Clarke 与 P/Q 计算"],
        ["3", "开环 Boost 固定占空比", "11", "孤岛 MPVC"],
        ["4", "标准模块增量电导 MPPT", "12", "变压器、电网与同步"],
        ["5", "两路直流负载", "13", "并网 MPPC 与 Vdc 功率环"],
        ["6", "电池双向 DC/DC 功率级", "14", "标准模块 EMS"],
        ["7", "电池电流内环", "15", "论文事件与分段运行"],
        ["附录 A/B", "最小代码、发散因果链", "附录 C/D", "调参顺序、最终布局"],
    ], [0.55, 2.70, 0.55, 2.70], font_size=8.0, header_fill=PALE_GREEN)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()


def stage(doc, number_text, title, meaning, blocks, actions, connections,
          first_run, tuning, faults, checks, formulas=None):
    heading(doc, f"阶段 {number_text}：{title}", 1)
    callout(doc, "本阶段的意义", meaning, fill=LIGHT_BLUE)
    heading(doc, "需要拖入的模块与精确设置", 2)
    table(doc, ["模块名", "Library Browser 搜索/路径", "关键设置", "作用"], blocks,
          [1.35, 1.75, 1.75, 1.65], font_size=7.25)
    heading(doc, "鼠标搭建步骤", 2)
    for item in actions:
        number(doc, item)
    if connections:
        heading(doc, "逐线连接清单", 2)
        table(doc, ["起点", "终点", "端口/方向", "为什么这样接"], connections,
              [1.35, 1.35, 1.45, 2.35], font_size=7.35)
    if formulas:
        heading(doc, "本阶段必须理解的公式", 2)
        for item in formulas:
            code(doc, item)
    heading(doc, "第一次运行：只看这些量", 2)
    callout(doc, "验收方法", first_run, fill=PALE_GREEN, border="70AD47")
    if tuning:
        heading(doc, "参数改变会发生什么", 2)
        table(doc, ["参数", "增大", "减小", "首先观察"], tuning,
              [1.05, 1.80, 1.80, 1.85], font_size=7.35, header_fill=PALE_YELLOW)
    if faults:
        heading(doc, "常见问题与修正", 2)
        table(doc, ["现象", "最可能原因", "检查与修正"], faults,
              [1.45, 2.00, 3.05], font_size=7.25, header_fill=PALE_RED)
    heading(doc, "通过后再进入下一阶段", 2)
    for item in checks:
        bullet(doc, "□ " + item)


def build_manual():
    create_assets()
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)

    heading(doc, "先读这一章：这版手册怎样使用", 1)
    callout(
        doc,
        "本版与旧版最重要的区别",
        "你不需要先运行任何“自动建模”脚本。第一个动作是点击 Simulink 的 Blank Model；之后每个功率器件、测量块、控制块和连线都由你亲手建立。"
        "唯一的初始化脚本只负责把参数放入工作区；两份控制器文件只保存论文中必须枚举 8 个开关状态的算法核。",
        fill=PALE_GREEN,
        border="70AD47",
    )
    add_picture(doc, ASSETS / "physical_topology.png", 6.35, "图 1  最终要得到的功率电路拓扑")
    add_picture(doc, ASSETS / "commissioning_ladder.png", 6.35, "图 2  从零搭建与送电顺序")

    heading(doc, "三条硬规则", 2)
    bullet(doc, "每个功率支路都先放隔离器。新支路接好后，先保持隔离器断开，让旧模型仍能运行。")
    bullet(doc, "每次只改变一件事：先结构编译，再短时运行，再闭环；不要一次接上 PV、电池、逆变器和电网。")
    bullet(doc, "电流方向先于控制器增益。若正方向定义错，任何正增益闭环都会变成正反馈并发散。")

    heading(doc, "论文值、推导值和工程补充值", 2)
    table(doc, ["类别", "含义", "本手册示例"], [
        ["论文原值", "文献明确给出，可直接作为复现目标。", "Vdc=1200 V，Cdc=6 mF，Lf=0.6 mH，Cf=1338 μF，Rf=0.019 Ω。"],
        ["可推导值", "由论文模块/额定值计算，但不是论文逐字给出的设置。", "SunPower 阵列取 Ns=10、Np=656，约 2 MW。"],
        ["工程补充值", "开关仿真必须有、论文却未公开；只能作为起始值。", "Lpv=0.2 mH、Lbat=0.5 mH、寄生电阻、PI 参数、网侧短路容量。"],
    ], [1.25, 2.45, 2.80], font_size=8.0)

    heading(doc, "建议保存的模型版本", 2)
    table(doc, ["文件名", "保存时机", "回退用途"], [
        ["PV_ESS_00_blank.slx", "只完成求解器和 powergui", "全局设置出错时回退"],
        ["PV_ESS_01_dc_bus.slx", "直流母线单独通过", "排除母线初始化问题"],
        ["PV_ESS_02_pv_boost.slx", "PV 开环 Boost 通过", "排除 MPPT"],
        ["PV_ESS_03_battery.slx", "电池内外环通过", "排除逆变器"],
        ["PV_ESS_04_island.slx", "MPVC 孤岛通过", "排除电网与 MPPC"],
        ["PV_ESS_05_grid.slx", "同步并网通过", "最终 EMS 联调基线"],
    ], [2.05, 2.15, 2.30], font_size=8.0)

    heading(doc, "最小代码如何使用", 2)
    paragraph(doc, "打开 MATLAB，将 Current Folder 切换到 D:\\PV_MPPT\\PV_ESS_MPC_Paper_Reproduction\\beginner_from_zero，然后在命令行执行：")
    code(doc, "addpath('D:/PV_MPPT/PV_ESS_MPC_Paper_Reproduction/beginner_from_zero/scripts')\n"
              "addpath('D:/PV_MPPT/PV_ESS_MPC_Paper_Reproduction/beginner_from_zero/controllers')\n"
              "run('D:/PV_MPPT/PV_ESS_MPC_Paper_Reproduction/beginner_from_zero/scripts/pvess_beginner_init.m')")
    paragraph(doc, "看到“parameters and profiles are ready”后才运行模型。可在 Model Properties > Callbacks > InitFcn 中粘贴最后一行，使模型每次启动自动读参数；这只是一行调用，不会自动改模型。")

    stage(
        doc, "0", "创建空白模型、配置离散电力网络",
        "建立一个可重复、可调试的数值环境。此时不放任何电源和控制器；目标只是让空模型编译通过。",
        [
            ["powergui", "搜索 powergui；Specialized Power Systems", "Simulation type=Discrete；Sample time=Ts_power；Solver=Tustin/Backward Euler (TBE)", "Specialized Power Systems 网络必须且只能有一个。"],
            ["Scope", "Simulink > Sinks > Scope", "默认；后续按需增加输入端口", "观察阶段信号。"],
            ["Display", "Simulink > Sinks > Display", "Format=shortG", "快速看标量是否有限。"],
        ],
        [
            "MATLAB 命令行运行 pvess_beginner_init.m。Workspace 中确认 Vdc_ref=1200、Ts_power=2e-6、Ts_mpc=5e-5。",
            "点击 Simulink > Blank Model，立即 Save As 为 PV_ESS_00_blank.slx。不要从现有模板复制。",
            "打开 Modeling > Model Settings > Solver：Type 选 Fixed-step；Solver 选 ode3 (Bogacki-Shampine)；Fixed-step size 填 Ts_power；Stop time 先填 0.01。",
            "在 Library Browser 搜索 powergui，拖入顶层。双击按表格设置。若同名 Simscape Electrical 与 Specialized Power Systems 块混用，优先选择图标和本手册截图一致的 Specialized Power Systems 版本。",
            "按 Ctrl+D 更新模型。此时应无红色错误。保存。",
        ],
        [],
        "按 Ctrl+D 后 Diagnostic Viewer 没有“缺少 powergui”“连续/离散网络不一致”错误即可。本阶段无需按 Run；如果运行，0.01 s 应立即结束。",
        [
            ["Ts_power", "更准确但运行更慢；对 5 kHz 开关应保留足够采样点。", "更快但开关边沿和 LC 动态失真，可能数值发散。", "一次 PWM 周期至少约 50–100 个电力步长。"],
            ["Solver", "不适用", "不适用", "本手册固定 ode3 + 离散 powergui，联调前不要改。"],
        ],
        [
            ["找不到 powergui", "未安装 Simscape Electrical Specialized Power Systems。", "在 Add-On/安装器确认 Simscape Electrical；搜索完整英文名。"],
            ["FixedStep 变量未定义", "未先运行初始化脚本。", "运行 pvess_beginner_init.m；或暂填 2e-6 验证。"],
            ["模型一运行就很慢", "Stop time 误填 4 s，或步长比 2e-6 更小。", "空模型阶段改回 0.01 s。"],
        ],
        ["模型从 Blank Model 建立。", "只有一个 powergui。", "Ctrl+D 无错误。", "已保存 00_blank 版本。"],
    )

    stage(
        doc, "1", "建立预充电的 1200 V 直流母线",
        "先把直流母线作为独立储能节点验证。后面 PV、电池和逆变器全部向这一对正/负母线接入；若母线极性或初值错，所有控制都会看似发散。",
        [
            ["Parallel RLC Branch", "搜索 Parallel RLC Branch（SPS）", "Branch type=C；C= Cdc；Set initial voltage=on；Initial voltage=Vdc0", "论文 6 mF 直流母线电容。"],
            ["Voltage Measurement", "SPS > Measurements", "默认", "把物理电压转换为 Simulink 信号。"],
            ["Ground", "SPS > Fundamental Blocks > Elements > Ground", "默认", "定义直流负母线参考点。"],
            ["Parallel RLC Branch", "同上，另放一个", "Branch type=R；R=1e6", "泄放/数值参考电阻；几乎不带载。"],
            ["Scope", "Simulink > Sinks", "输入端口=1", "观察 Vdc。"],
        ],
        [
            "从 Cdc 模块上端画一条水平线作为 DC+，下端画一条水平线作为 DC−；以后只在这些线上创建分支。",
            "Voltage Measurement 的 + 端接 DC+，− 端接 DC−。如果模块图标端子方向不明显，旋转模块而不是交叉接线。",
            "Ground 接 DC−。1 MΩ 电阻跨接 DC+ 与 DC−。",
            "Voltage Measurement 的 Simulink 输出接 Scope，并把信号命名 Vdc。",
            "保存为 PV_ESS_01_dc_bus.slx，Stop time=0.01，按 Run。",
        ],
        [
            ["Cdc 上端", "DC+", "物理电气线", "定义正母线。"],
            ["Cdc 下端", "DC− 与 Ground", "物理电气线", "定义负母线和参考。"],
            ["Voltage Measurement +/−", "DC+/DC−", "极性必须一致", "保证 Vdc 正值。"],
            ["Voltage Measurement 输出", "Scope", "Simulink 线", "只用于观测，不承载功率。"],
        ],
        "Vdc 在 t=0 附近应从 1200 V 初值开始，10 ms 内基本保持；因 1 MΩ 泄放电阻仅有约 1.2 mA，下降应极小。若显示 −1200 V，只交换测量端，暂时不要改控制符号。",
        [
            ["Cdc", "母线更“硬”、纹波小、动态慢、储能和启动冲击增大。", "母线变化快、纹波大，控制更难。", "Vdc 纹波、启动电流、仿真刚性。"],
            ["Vdc0", "初始能量 0.5CdcV² 增大。", "启动时更依赖源支路充电。", "t=0 的 Vdc 与支路浪涌。"],
        ],
        [
            ["Vdc=0", "没有勾选初始电压或参数未载入。", "检查 Set initial voltage 和 Initial voltage=Vdc0。"],
            ["Vdc=-1200", "Voltage Measurement 极性反。", "只交换 +/− 物理端口。"],
            ["拓扑奇异/浮地", "DC− 未接 SPS Ground。", "增加 Ground；不要用 Simscape Electrical Reference 代替。"],
        ],
        ["Vdc 初值为正且约 1200 V。", "10 ms 内没有 NaN/Inf。", "DC+/DC− 已清楚标注。", "已保存 01_dc_bus。"],
        ["E_dc = 0.5*Cdc*Vdc^2 = 0.5*0.006*1200^2 = 4320 J"],
    )

    stage(
        doc, "2", "单独建立并验证 PV Array",
        "先证明静态 I–V/P–V 特性和阵列规模正确，再接 Boost。若 PV 阵列本身错误，任何 MPPT 都只会追踪错误曲线。",
        [
            ["PV Array", "搜索 PV Array（SPS）", "Module=SunPower SPR-305E-WHT-D；Series modules=Ns_pv；Parallel strings=Np_pv；Irradiance/temperature inputs=on；Sample time=Ts_power", "产生真实非线性 PV I–V 特性。"],
            ["Constant ×2", "Simulink > Sources", "G=1000；T=25", "先用 STC 固定输入。"],
            ["Current Measurement", "SPS > Measurements", "箭头从 PV+ 指向负载", "测 Ipv，暂定发电为正。"],
            ["Voltage Measurement", "SPS > Measurements", "+ 接 PV+；− 接 PV−", "测 Vpv。"],
            ["Variable Resistor", "搜索 Variable Resistor（SPS）", "Minimum resistance=0.05；控制输入先 1e6，后 0.15", "临时扫描/加载 PV；阶段结束移除。"],
            ["Product", "Simulink > Math Operations", "Multiplication=Element-wise", "Ppv=Vpv×Ipv。"],
        ],
        [
            "在模型左侧放 PV Array。把 Irradiance 输入接 Constant=1000，把 Temperature 输入接 Constant=25；测量总线输出暂接 Terminator。",
            "PV+ 依次接 Current Measurement 和临时 Variable Resistor 上端；PV− 接电阻下端。Voltage Measurement 跨接 PV+/PV−。",
            "Variable Resistor 控制端接 Constant=1e6，先运行 2 ms 看开路电压；再改成 0.15 Ω，运行 10 ms 看接近 MPP 的输出。",
            "Vpv、Ipv 接 Product 得 Ppv；三者各接 Scope。不要在本阶段把 PV 接入 1200 V 母线。",
            "验证完成后删除临时 Variable Resistor，保留 PV、测量块和 Ppv 计算。",
        ],
        [
            ["G Constant", "PV Array irradiance 端", "Simulink 输入 1", "辐照度主要改变电流和功率。"],
            ["T Constant", "PV Array temperature 端", "Simulink 输入 2", "温度主要改变开路电压。"],
            ["PV+", "Current Measurement → 临时负载", "电流箭头朝外", "使发电电流为正。"],
            ["Vpv、Ipv", "Product", "两个标量输入", "Ppv 正值表示 PV 发电。"],
        ],
        "G=1000 W/m²、T=25°C、R≈0.15 Ω 时，Vpv 应在约 500–600 V 区域，Ppv 应接近 2 MW；开路时电流接近 0、Vpv 接近约 640 V。这里看数量级和曲线形状，不要求逐点等于论文图。",
        [
            ["Np_pv", "电流和功率近似等比例增大，电压基本不变。", "电流和功率下降。", "Ipv、Ppv。"],
            ["Ns_pv", "MPP/开路电压近似增大，电流基本不变。", "电压下降。", "Vpv。"],
            ["G", "短路电流与最大功率显著增加。", "电流和功率下降。", "Ipv、Ppv。"],
            ["温度", "Vmp/Voc 下降，功率通常下降。", "电压上升。", "Vpv、Voc。"],
        ],
        [
            ["功率只有几 kW", "串并联数仍是 1，或模块选择错。", "确认 Ns_pv=10、Np_pv=656、模块全名一致。"],
            ["Ppv 为负", "Current Measurement 箭头反。", "翻转测量块；本项目规定 PV 发电为正。"],
            ["输入端口不存在", "PV Array 未启用外部 G/T 输入。", "在块参数中启用 irradiance and temperature inputs。"],
            ["一加载就数值崩溃", "临时电阻过小，产生过大电流。", "先 1e6 Ω，再逐步降到 0.3、0.2、0.15 Ω。"],
        ],
        ["开路电压数量级正确。", "0.15 Ω 附近功率约 2 MW。", "Ppv 为正。", "临时负载已移除。"],
        ["P_STC ≈ 305 W × Ns_pv × Np_pv = 305×10×656 ≈ 2.0008 MW"],
    )

    stage(
        doc, "3", "接入开关级 Boost，并用固定占空比送电",
        "这一阶段只验证主电路和电感电压平衡，不使用 MPPT。固定占空比能稳定后，才允许闭合追踪算法。",
        [
            ["Parallel RLC Branch", "SPS；Branch type=C", "C=Cpv；Initial voltage=Vpv0", "PV 端缓冲电容。"],
            ["Series RLC Branch", "SPS；Branch type=RL", "R=RLpv；L=Lpv；Initial current=ILpv0", "Boost 输入电感。"],
            ["Breaker ×2", "搜索 Breaker（SPS）", "External control=on；Ron=1e-3；Rs=1e6；Cs=inf；Initial=open", "PV 输入和 DC 输出调试隔离。"],
            ["Boost Converter", "搜索 Boost Converter（SPS）", "Model=Switching devices；Ron=Ron_switch；diode Ron=Ron_switch；Vf=Vf_diode；snubber Rs=1e6,Cs=inf", "真实开关和二极管升压。"],
            ["PWM Generator (DC-DC)", "搜索 PWM Generator (DC-DC)", "Switching frequency=Fsw_pv；Sample time=Ts_power", "把占空比转换为 5 kHz 门极。"],
            ["Constant", "Simulink > Sources", "Value=Dpv0", "固定占空比冒烟测试。"],
            ["Step ×2", "Simulink > Sources", "PV input connect: 0→1 at 0.002 s；DC output connect: 0→1 at 0.004 s", "按顺序闭合隔离器。"],
        ],
        [
            "Cpv 跨接 PV+ 与 PV−。PV+ 经 Ipv Measurement → Lpv → PV_Input_Isolator → Boost 输入正端；PV− 接 Boost 输入负端。",
            "Boost 输出正端经 PV_DC_Isolator 接 DC+；输出负端接 DC−。确认 DC 母线仍由 Cdc 预充至 1200 V。",
            "Constant Dpv0 接 PWM Generator 输入，PWM 输出接 Boost gate；Boost blocking 端接 Constant=0（若块有该端口）。",
            "先保持两个隔离器开路按 Ctrl+D；再运行 20 ms，按 2 ms/4 ms 顺序闭合。",
            "观察 Vpv、Ipv、iLpv、Vdc、Ppv。第一次不要接 DC 负载。",
        ],
        [
            ["PV+", "Ipv → Lpv → 输入隔离器 → Boost +", "串联功率线", "电感位于开关节点之前。"],
            ["PV−", "Boost 输入−、DC−", "公共负母线", "建立非隔离 Boost 回路。"],
            ["Boost 输出+", "输出隔离器 → DC+", "正母线", "允许独立切断 PV。"],
            ["Dpv0", "PWM Generator", "0–1 标量", "5 kHz PWM 生成。"],
            ["PWM", "Boost gate", "Simulink 脉冲", "控制开关器件。"],
        ],
        "隔离器闭合后，Vpv 应向 (1−Dpv0)×Vdc≈547 V 附近移动；Vdc 不应出现大幅上冲，所有电流有限。若辐照度为 400 W/m²，功率约为 STC 的四成数量级。",
        [
            ["Lpv", "电流纹波下降、响应变慢、体积等效增大。", "纹波和峰值增大，数值更敏感。", "iLpv 峰峰值。"],
            ["Cpv", "Vpv 纹波减小、MPPT 动态变慢。", "Vpv 更快但纹波大。", "Vpv 纹波与 MPPT 稳定性。"],
            ["Dpv", "在 Vdc 刚性时 Vpv≈(1−D)Vdc 下降，PV 工作点向低电压移动。", "Vpv 上升。", "Vpv 与 Ppv。"],
            ["Fsw_pv", "纹波下降但计算量和开关损耗上升。", "纹波上升。", "电感纹波、仿真速度。"],
        ],
        [
            ["Vpv 立即变 0", "Boost 输入极性/公共负线错，或隔离器逻辑反。", "先断 PV_DC，检查 PV 单独输出；用 Display 看断路器控制 0/1。"],
            ["Vdc 突然数倍上冲", "母线无负载且 PV 持续注入，或占空比过大。", "保持 PV 支路短时运行；加限幅/负载；D 从 0.4 缓慢增。"],
            ["iLpv 巨大", "L 单位误填成 0.2e-6，或初值/极性错。", "确认 Lpv=0.2e-3 H；断开输出隔离器逐段查。"],
            ["PWM 有波形但器件不切换", "blocking=1 或接错 gate 端。", "把 blocking 明确设 0；检查 PWM 接 Simulink 三角端口。"],
        ],
        ["固定占空比下 Vpv 接近 547 V。", "Vdc、Ipv、iLpv 均有限。", "两个隔离器能独立断开。", "已保存 02_pv_boost。"],
        ["开关闭合：vL,on = Vpv − RLpv·iL", "开关断开：vL,off = Vpv − Vdc − RLpv·iL", "平均电压平衡：D·vL,on + (1−D)·vL,off = 0", "因此：L·diL/dt = Vpv − (1−D)Vdc − RLpv·iL；稳态时导数为 0，并不是把同一个 Vpv 重复相减。"],
    )

    stage(
        doc, "4", "用标准模块搭建增量电导 MPPT",
        "控制周期与 PWM 周期分开：功率开关仍以 5 kHz 工作，而 MPPT 每 1 ms 才根据新的 V/I 改一次占空比。这能避免把开关纹波误判成工作点移动。",
        [
            ["Zero-Order Hold ×2", "Simulink > Discrete", "Sample time=Ts_mppt", "在实际 MPPT 周期采样 Vpv/Ipv。"],
            ["Unit Delay ×3", "Simulink > Discrete", "Vprev IC=Vpv0；Iprev IC=0；Duty IC=Dpv0；Sample time=Ts_mppt", "保存上一采样值和上一占空比。"],
            ["Sum ×4", "Simulink > Math Operations", "ΔV/ΔI: Inputs=+-；Inc/Dnext: 按文字设置", "计算差分与占空比累加。"],
            ["Abs", "Simulink > Math Operations", "默认", "取得 |ΔV|。"],
            ["Compare To Constant ×5", "Simulink > Logic and Bit Operations", "|ΔV|<eps_dV；ΔI>eps_dI；ΔI<−eps_dI；Inc>eps_inc；Inc<−eps_inc", "产生分支布尔量。"],
            ["Switch ×5", "Simulink > Signal Routing", "Criteria=u2 ~= 0", "用布尔量选择 +step、−step 或 0。"],
            ["Divide ×2", "Simulink > Math Operations", "dI/safe_dV；I/Vsafe", "计算增量和瞬时电导。"],
            ["Saturation ×2", "Simulink > Discontinuities", "Vsafe: [1,inf]；Duty: [Dpv_min,Dpv_max]", "防除零与限制占空比。"],
            ["Constant ×4", "Simulink > Sources", "0；Dpv_step；−Dpv_step；1", "算法常量。"],
        ],
        [
            "将 Vpv、Ipv 分别接 Zero-Order Hold，输出命名 V(k)、I(k)。各自同时接 Unit Delay，Unit Delay 输出为 V(k−1)、I(k−1)。",
            "两个 Sum(+-) 计算 ΔV=V(k)−V(k−1)、ΔI=I(k)−I(k−1)。Abs(ΔV) 接 Compare To Constant，Operator 选 <，Constant=eps_dV，得到 DV_small。",
            "为了避免除以 0：放 Switch_SafeDV，u1=Constant 1，u2 接现有 DV_small，u3=ΔV；DV_small 为真时输出 1，否则输出 ΔV，得到 safe_dV。V(k) 经 Saturation[1,inf] 得 Vsafe。",
            "Divide1=ΔI/safe_dV，Divide2=I/Vsafe，用 Sum(++) 得 Inc=ΔI/ΔV+I/V。",
            "普通分支：Inc>eps_inc 时选择 −Dpv_step；Inc<−eps_inc 时选择 +Dpv_step；否则 0。原因是 Boost 增大 D 会降低 Vpv。",
            "ΔV≈0 分支：ΔI>eps_dI 时选择 −Dpv_step；ΔI<−eps_dI 时选择 +Dpv_step；否则 0。",
            "最终 Switch 的控制端接 DV_small：真时选 ΔV≈0 分支，假时选普通分支。将得到的 ΔD 与 Duty Unit Delay 输出 D(k−1) 相加，经 Saturation[Dpv_min,Dpv_max] 后反馈到 Duty Unit Delay 输入。",
            "Duty Saturation 输出替换阶段 3 的 Constant Dpv0，接 PWM Generator。开始运行前仍可用 Manual Switch 在固定 D 与 MPPT D 之间切换。",
        ],
        [
            ["Vpv/Ipv", "各自 ZOH → 差分与电导计算", "Ts_mppt 采样", "隔离 5 kHz 纹波。"],
            ["ΔV、ΔI", "DV_small 与 Inc 逻辑", "标量", "决定使用哪条判据。"],
            ["D(k−1)、ΔD", "Sum → Saturation → Duty Delay", "离散反馈", "每次只改变一个 Dpv_step。"],
            ["Duty", "PWM Generator (DC-DC)", "0.35–0.75", "驱动 Boost。"],
        ],
        "G=400 W/m²、T=25°C、无突变时，D 应从约 0.544 缓慢移动并在小范围往复；Ppv 先上升后在最大值附近抖动。重点是 D 每 1 ms 才更新一次，而 PWM 每 0.2 ms 重复。",
        [
            ["Ts_mppt", "更新变慢、抗纹波好但辐照度突变跟踪慢。", "更新快但容易追踪开关纹波并振荡。", "Duty 更新台阶、Ppv 稳态波动。"],
            ["Dpv_step", "跟踪快、MPP 附近振荡和功率损失增大。", "稳态细但跟踪慢。", "Ppv 上升时间与稳态摆幅。"],
            ["eps_dV/eps_dI", "更容易判作“无变化”，抗噪强但形成死区。", "更灵敏但可能被噪声触发。", "D 是否无故来回切换。"],
        ],
        [
            ["功率越追越低", "Boost 占空比方向写反。", "固定改变 D 验证：D 增大应使 Vpv 降低；然后交换 ±Dpv_step。"],
            ["D 每个电力步都变化", "缺少 ZOH/Unit Delay，继承了 Ts_power。", "所有 MPPT 状态块 Sample time 明确填 Ts_mppt。"],
            ["出现 Inf/NaN", "ΔV 或 Vpv 做除数时为 0。", "确认 SafeDV Switch 与 Vsafe Saturation；先查看两个 Divide 输入。"],
            ["D 卡在上下限", "步进方向、测量极性或 PV/DC 电压范围不匹配。", "切回固定 D；确认 Vpv、Ipv、Ppv 均为正且约 547 V 工作点可达。"],
        ],
        ["D 只按 1 ms 更新。", "Ppv 能从偏离点向峰值移动。", "D 不长期卡限幅。", "切回固定 D 后主电路仍稳定。"],
        ["MPP 条件：dP/dV = I + V·dI/dV = 0", "左侧 MPP：dP/dV>0，需要升高 Vpv，因此 Boost 占空比 D 减小。", "右侧 MPP：dP/dV<0，需要降低 Vpv，因此 D 增大。"],
    )

    add_picture(doc, ASSETS / "mppt_logic.png", 6.35, "图 3  标准模块实现增量电导 MPPT")

    stage(
        doc, "5", "加入两路可变直流负载并核对功率平衡",
        "负载让直流母线真正出现充放电动态。先用电阻负载复现论文功率台阶，避免恒功率负载的负阻抗特性在控制器尚未调好时引发发散。",
        [
            ["Variable Resistor ×2", "搜索 Variable Resistor（SPS）", "Minimum R=0.2；输入来自 Rdc profile", "0.6 MW 与 1 MW 直流支路。"],
            ["Current Measurement ×2", "SPS > Measurements", "箭头从 DC+ 指向负载", "测每路负载电流。"],
            ["From Workspace ×2", "Simulink > Sources", "Variable=Rdc1_profile / Rdc2_profile；Interpolate=on", "按论文时刻改变等效电阻。"],
            ["Sum", "Simulink > Math Operations", "Inputs=++", "Idc_load=I1+I2。"],
            ["Product", "Simulink > Math Operations", "默认", "Pdc=Vdc×Idc_load。"],
        ],
        [
            "每一路结构均为 DC+ → Current Measurement → Variable Resistor → DC−；不要把两路串联。",
            "Rdc1_Profile 接第一路控制端，Rdc2_Profile 接第二路控制端。初始化脚本把未投入状态设为 1 MΩ。",
            "两个电流相加后与 Vdc 相乘，得到 Pdc。将 Vdc、Ppv、Pdc 放在同一 Scope 的不同轴。",
            "先把 Stop time=0.02，并临时将 From Workspace 换为 Constant=1e6；确认不开负载。再把第一路固定 2.4 Ω，短时运行。",
            "只有 PV/电池能提供足够功率时才运行论文完整负载时序；否则 4320 J 母线储能会在毫秒级耗尽。",
        ],
        [
            ["DC+", "I1 → R1 → DC−", "并联支路 1", "0.6 MW 目标。"],
            ["DC+", "I2 → R2 → DC−", "并联支路 2", "1.0 MW 目标。"],
            ["I1、I2", "Sum → Product with Vdc", "Simulink 信号", "形成 Pdc 测量。"],
        ],
        "Vdc≈1200 V 时，R1=2.4 Ω 对应约 0.6 MW，R2=1.44 Ω 对应约 1 MW。投入前后电流应分别约 500 A、833 A。若没有源支撑，母线下降是物理正确结果，不是求解器故障。",
        [
            ["负载电阻 R", "功率 P=V²/R 下降。", "负载功率和冲击增大。", "Pdc、Vdc。"],
            ["Rabsmin", "限制较弱。", "允许更低阻但潜在巨大电流。", "启动峰值和数值稳定。"],
            ["负载投入斜率", "更慢、更容易稳定。", "更接近理想台阶、冲击更强。", "Vdc 下陷与电池电流。"],
        ],
        [
            ["Pdc 为负", "电流测量箭头反。", "翻转两路 Current Measurement；负载消耗定义为正。"],
            ["投入 1 MW 后立刻欠压", "PV 功率不足且电池未闭环。", "本阶段只做极短验证；先完成电池阶段。"],
            ["Variable Resistor 报非正输入", "时序过渡或变量中出现 0/负值。", "检查 profile；R 只能为正且不小于 0.2 Ω。"],
        ],
        ["两路电阻确为并联。", "Pdc 与 Vdc²/R 数量级一致。", "负载电流方向为正。", "大功率台阶暂不用于长时间运行。"],
        ["R_0.6MW = 1200²/600000 = 2.4 Ω", "R_1MW = 1200²/1000000 = 1.44 Ω"],
    )

    stage(
        doc, "6", "建立电池与双向 DC/DC 功率支路",
        "先只搭功率电路和互补门极，不闭合电压环。该支路既能把 300 V 电池升压放电到 1200 V 母线，也能从母线降压给电池充电；方向由电感平均电压和电流控制决定。",
        [
            ["Battery", "搜索 Battery（SPS）", "Type=Lithium-Ion；Nominal V=V_bat_nom；Rated capacity=Q_bat_Ah；Initial SOC=SOC0；R=R_bat_internal；Full V=330；Min V=270", "论文 300 V、1300 Ah 储能。"],
            ["Two-Quadrant DC/DC Converter", "搜索 Two-Quadrant DC/DC Converter（SPS）", "Model=Switching devices；Ron=Ron_switch；snubber Rs=1e6,Cs=inf；diode Vf=0", "同步半桥，可双向传能。"],
            ["Series RLC Branch", "SPS；Branch type=RL", "R=RLbat；L=Lbat；Initial current=Ibat0", "电池侧平波电感。"],
            ["Breaker ×2", "SPS Breaker", "External=on；Initial=open；Ron=1e-3；Rs=1e6；Cs=inf", "电池侧和母线侧隔离。"],
            ["Current Measurement", "SPS > Measurements", "方向先定义为电池→变换器为正", "Ibat>0 表示放电。"],
            ["PWM Generator (DC-DC)", "SPS pulse generators", "Fsw=Fsw_bat；Ts=Ts_power", "产生一个基本 PWM。"],
            ["Logical Operator NOT", "Simulink > Logic and Bit Operations", "Operator=NOT", "产生互补逻辑。"],
            ["Unit Delay ×2", "Simulink > Discrete", "Ts=Ts_power；IC=0", "实现一个电力步长死区。"],
            ["Logical Operator AND ×2", "Simulink > Logic and Bit Operations", "Operator=AND；2 inputs", "要求当前与延迟后信号都为 1，避免直通。"],
            ["Mux + Data Type Conversion", "Signal Routing / Signal Attributes", "Mux inputs=2；output data type=double", "形成双门极向量。"],
        ],
        [
            "把双向变换器直流端上端经 Battery_DC_Isolator 接 DC+，下端接 DC−；中点端经 Battery_Isolator → Lbat → Current Measurement → Battery+。Battery− 接 DC−。",
            "观察 Current Measurement 图标箭头。把“电池流向变换器”的方向定义为正；若首次小放电测试得到负值，旋转测量块，不要先改 PI 符号。",
            "Constant Db0 接 PWM Generator。PWM 输出记为 p。构造 lower=p AND z⁻¹(p)；upper=NOT(p) AND z⁻¹(NOT(p))。",
            "Mux 顺序先取 [upper,lower]，转成 double 后接 Two-Quadrant Converter gate。Blocking 端接 Constant=1，两个隔离器保持开路，先 Ctrl+D。",
            "运行 2 ms 仅验证断开态。然后让电池侧隔离器 2 ms 闭合、母线侧 4 ms 闭合，最后在 6 ms 把 blocking 从 1 变 0。先只运行到 10 ms。",
            "若电流方向和电压都合理，再保存 PV_ESS_03_battery_power.slx；此时仍没有电池闭环。",
        ],
        [
            ["DC+", "母线隔离器 → 变换器 DC 上端", "物理线", "接入 1200 V 正母线。"],
            ["变换器 DC 下端", "DC−", "物理线", "公共负母线。"],
            ["变换器中点", "隔离器 → Lbat → Ibat → Battery+", "电池支路", "电感限制充放电纹波。"],
            ["Battery−", "DC−", "物理线", "非隔离双向拓扑回路。"],
            ["upper/lower", "Mux → gate", "[上管,下管] 起始假设", "必须通过小电流测试确认顺序。"],
        ],
        "在 Db0≈0.75、母线和电池初值接近理论变比时，Ibat 应保持有限且接近 0；允许有短时充放电冲击。把 Db0 暂增 0.002 后，按本手册正方向 Ibat 应向放电方向变化；若相反，先核对门极顺序和电流箭头。",
        [
            ["Lbat", "电流纹波小、内环慢。", "电流响应快但纹波和峰值增大。", "Ibat 峰峰值。"],
            ["Db0", "通常提高升压作用并增加放电趋势。", "降低放电/增强充电趋势。", "Ibat 的符号。"],
            ["死区（1×Ts_power）", "直通风险下降但占空比有效值误差增大。", "波形更理想但有上下管同时导通风险。", "桥臂电流尖峰。"],
        ],
        [
            ["一解封就数值爆炸", "上下管同时导通或门极顺序错误。", "blocking=1 回退；Scope 看两门极是否同为 1；检查 NOT/Delay/AND。"],
            ["Ibat 初值数千安", "Db0 与实际 Vbat/Vdc 不符，或 Lbat 单位错。", "Db0=1−Vbat/Vdc≈0.75；Lbat=0.5e-3 H。"],
            ["Battery 报参数不一致", "预设模型仍启用，手填参数被锁定。", "Preset model=off，再输入容量、电压和放电曲线。"],
            ["Ibat 符号与预期反", "测量箭头或门极定义反。", "固定微小 Duty 扰动分离判断：先确认箭头，再确认 gate 顺序。"],
        ],
        ["两门极不会同时为 1。", "小幅 Duty 增量对应的 Ibat 方向已记录。", "Ibat 正方向确定为放电。", "支路可由两个隔离器和 blocking 安全切断。"],
        ["理想稳态变比：Vdc ≈ Vbat/(1−Db)，所以 Db0≈1−300/1200=0.75"],
    )

    stage(
        doc, "7", "先调电池电流内环",
        "电流内环是双闭环的地基。它必须在固定 Ibat_ref 下稳定、方向正确、无持续饱和；否则外层 Vdc 环只会把错误放大。",
        [
            ["Sum", "Simulink > Math Operations", "Inputs=+-", "eI=Ibat_ref−Ibat。"],
            ["Discrete PID Controller", "Simulink > Discrete", "Controller=PI；Form=Parallel；P=Kpi_bat；I=Kii_bat；D=0；Ts=Ts_bat_ctrl；Output limits=±0.25；Anti-windup=Clamping", "产生占空比修正 ΔD。"],
            ["Constant", "Simulink > Sources", "Value=Db0", "理论稳态占空比前馈。"],
            ["Sum", "Simulink > Math Operations", "Inputs=++", "D=Db0+ΔD。"],
            ["Saturation", "Simulink > Discontinuities", "Lower=Db_min；Upper=Db_max", "最终占空比保护。"],
            ["Step", "Simulink > Sources", "0→100 A at 0.01 s；sample=Ts_bat_ctrl", "小电流参考测试。"],
            ["Zero-Order Hold", "Simulink > Discrete", "Ts=Ts_bat_ctrl", "把 Ibat 送入控制采样域。"],
        ],
        [
            "把 Ibat 经 Zero-Order Hold 后接 Sum 负端，把 Step Ibat_ref 接正端。",
            "Sum 输出接 Discrete PID Controller。先关闭积分：I=0，只用很小 P 验证符号；确认负反馈后再填 Kpi_bat、Kii_bat。",
            "PI 输出与 Db0 相加，经 Saturation 接电池 PWM Generator，替代阶段 6 的 Constant Db0。",
            "让功率支路按阶段 6 顺序接通；Ibat_ref 在 10 ms 由 0 变 100 A，Stop time=0.05 s。",
            "100 A 跟踪通过后分别测试 +300 A、−100 A。每次改参考前先保存，负值测试表示充电。",
        ],
        [
            ["Ibat_ref", "eI Sum 正端", "正=放电", "定义目标方向。"],
            ["Ibat ZOH", "eI Sum 负端", "反馈", "构成负反馈。"],
            ["eI", "PI → +Db0 → Saturation", "控制链", "把电流误差变为 Duty。"],
            ["Duty", "Battery PWM", "0.05–0.95", "驱动双向功率级。"],
        ],
        "Ibat_ref=+100 A 后，Ibat 应沿正方向收敛，Duty 从 Db0 附近小幅变化；不应立即顶到 0.05/0.95。先接受较慢响应，确认方向后再提速。",
        [
            ["Kpi_bat", "上升更快、超调和噪声放大增加。", "更慢但平稳。", "Ibat 上升时间、Duty 跳变。"],
            ["Kii_bat", "消除稳态误差快，但易低频振荡/积分饱和。", "稳态误差消除慢。", "Ibat 尾部与 PI 积分状态。"],
            ["PI 输出限幅", "允许更强动作，也可能打到最终 Duty 限幅。", "保护更强但大阶跃跟不上。", "ΔD 与 Duty。"],
        ],
        [
            ["Iref 为正，Ibat 向负方向跑", "反馈符号、测量方向或门极顺序错误。", "把 PI 的 I 设 0、小 P 测试；先修物理符号，不能靠负增益掩盖。"],
            ["Duty 长期在 0.95", "电流目标不可达、方向错或母线未预充。", "降到 50 A；核对 Vdc、Vbat、Db0 和隔离器。"],
            ["Ibat 高频锯齿很大", "Lbat 太小或观察的是开关波形。", "先对 Ibat 做一个开关周期移动平均；不要用 MPPT 周期替代电力步长。"],
            ["低频来回振荡", "Kii 过大或饱和后积分未释放。", "先把 Kii 降 5–10 倍，启用 Clamping。"],
        ],
        ["+100 A 和 −100 A 均能按定义方向响应。", "Duty 不长期饱和。", "积分开启后稳态误差缩小。", "内环带宽明显高于后续外环。"],
        ["eI = Ibat_ref − Ibat", "D = sat[Db_min,Db_max]{Db0 + Kpi·eI + Kii·∫eI dt}"],
    )

    stage(
        doc, "8", "建立 SOC 与直流母线电压外环",
        "孤岛时电池承担功率缺口并维持 1200 V；并网时电网侧变换器接管 Vdc，电池改由 SOC/EMS 给出电流参考。用 Switch 明确分工，避免两个外环同时争夺母线。",
        [
            ["Gain", "Simulink > Math Operations", "Gain=−100/(3600*Q_bat_Ah)", "把 A 转换为 %SOC/s。"],
            ["Discrete-Time Integrator", "Simulink > Discrete", "Ts=Ts_bat_ctrl；IC=SOC0；limits 0–100", "库仑计量 SOC。"],
            ["Sum", "Simulink > Math Operations", "Inputs=+-；u+=Vdc_ref，u−=Vdc", "eV=Vdc_ref−Vdc。"],
            ["Discrete PID Controller", "Simulink > Discrete", "PI Parallel；P=Kpv_bat；I=Kiv_bat；Ts=Ts_bat_ctrl；output limits=±Ibat_max；anti-windup=Clamping", "孤岛 Vdc 外环输出 Ibat_ref。"],
            ["Saturation", "Simulink > Discontinuities", "±Ibat_max", "额定 0.5 MW 电池电流限制。"],
            ["Switch", "Simulink > Signal Routing", "Criteria=u2 >= Threshold；Threshold=0.5", "grid=0 选 Vdc 外环；grid=1 选 EMS/SOC 电流。"],
            ["Relational Operator ×2", "Simulink > Logic and Bit Operations", "SOC>SOC_min；SOC<SOC_max", "禁止过放/过充。"],
        ],
        [
            "Ibat 测量接 Gain，再接 Discrete-Time Integrator。因本项目 Ibat>0 表示放电，所以增益必须为负。",
            "Vdc_ref 与 Vdc 做差，接外环 PI；PI 输出再 Saturation 到 ±Ibat_max。先将外环 P、I 都降低为表中值的 1/10 做第一次闭环。",
            "Switch 的 u3（下端）接孤岛 Vdc 外环，u1（上端）接暂时 Constant=0，u2 接 GridConnected=0。Switch 输出替换阶段 7 的 Step Ibat_ref。",
            "只投入 0.6 MW 负载做测试。先让 PV 功率略低于负载，Vdc 下陷时 Ibat_ref 应变正、电池放电；PV 功率过剩时 Ibat_ref 应变负、充电。",
            "逐步恢复 Kpv_bat、Kiv_bat。外环应比电流内环慢至少 5–10 倍。",
            "在 Ibat_ref 后增加 SOC 保护逻辑：SOC≤SOC_min 时禁止正放电参考；SOC≥SOC_max 时禁止负充电参考。",
        ],
        [
            ["Ibat", "Gain → SOC Integrator", "正放电导致 SOC 降", "能量状态估计。"],
            ["Vdc_ref、Vdc", "eV → 外环 PI", "正误差要求放电", "孤岛稳压。"],
            ["外环/EMS Iref", "Mode Switch → 电流内环", "GridConnected 选择", "避免双控制权。"],
        ],
        "投入 0.6 MW 负载时 Vdc 允许短时下陷，但应回到 1200 V 附近；Ibat_ref 和 Ibat 应同向，SOC 以非常缓慢的速率下降。因为 1300 Ah 很大，几十毫秒内 SOC 几乎不变是正确的。",
        [
            ["Kpv_bat", "Vdc 恢复快，但 Iref 尖峰和耦合振荡增大。", "母线更软、恢复慢。", "Vdc 下陷、Iref 峰值。"],
            ["Kiv_bat", "稳态误差消除快，但易与 Cdc/Lbat 低频振荡。", "稳态偏差保持更久。", "Vdc 低频摆动。"],
            ["Ibat_max", "可支撑更大功率但超过电池额定。", "保护电池但更易母线欠压/切负载。", "Pbat≈Vbat·Ibat。"],
        ],
        [
            ["负载投入后 Vdc 越降 Ibat 越负", "外环误差或电流定义反。", "必须是 eV=Vref−Vdc，低压时 Iref>0；回到阶段 7 确认正电流为放电。"],
            ["Vdc 振荡但电流环单独稳定", "外环过快。", "Kpv、Kiv 同时降 5 倍；确认电流内环先稳定。"],
            ["SOC 快速掉到 0", "漏乘 3600 或单位把 A 当 Ah/s。", "Gain 必须为 −100/(3600*1300)。"],
            ["并网后两个 PI 积分继续累积", "模式切换只切输出，未冻结非活动积分。", "给外环 PID 增加 Enable/Reset，或并网时输入误差置 0；先用手动复位验证。"],
        ],
        ["低 Vdc 会产生正放电参考。", "高 Vdc 会产生负充电参考。", "SOC 方向正确且变化缓慢。", "孤岛/并网只激活一个母线外环。"],
        ["dSOC/dt = −100·Ibat/(3600·Q_Ah)", "Ibat_max = 0.5 MW / 300 V ≈ 1666.7 A"],
    )

    add_picture(doc, ASSETS / "battery_loops.png", 6.35, "图 4  电池内外环及模式分工")

    stage(
        doc, "9", "建立三相两电平桥、LC 滤波器与交流负载",
        "这一步只完成逆变器功率电路。先用临时 SPWM 和小负载确认三相桥、门极顺序、滤波器和电压测量正确，再用 FCS MPC 取代 PWM。",
        [
            ["Universal Bridge", "搜索 Universal Bridge（SPS）", "Arms=3；Device=IGBT/Diodes；Ron=Ron_switch；Forward voltages=[Vf_diode,Vf_diode]；snubber R=1e5,C=inf", "三相两电平互联变换器。"],
            ["Three-Phase Series RLC Branch", "搜索完整名称（SPS）", "Branch type=RL；R=Rf；L=Lf", "每相 0.6 mH/0.019 Ω 滤波支路。"],
            ["Three-Phase V-I Measurement ×2", "SPS > Measurements", "Voltage=phase-to-ground；Current=yes", "分别测滤波电流/母线电压与负载电流。"],
            ["Three-Phase Parallel RLC Load", "SPS；作为 Filter_Cf", "Y grounded；Vn=Vac_ll；fn=f_grid；P=0；QL=0；Qc=Qc_filter；constant Z", "以额定无功表示每相等效电容。"],
            ["Three-Phase Parallel RLC Load", "SPS；作为 AC_Load_Base", "Y grounded；Vn=Vac_ll；fn=f_grid；P 先 50e3，最终 P_ac_base；Q=0；constant Z", "孤岛交流负载。"],
            ["Breaker", "SPS Breaker", "External=on；Initial=open", "逆变器 DC 侧隔离。"],
            ["PWM Generator (2-Level)", "搜索 PWM Generator (2-Level)", "Carrier frequency≈3.3e3；sample time=Ts_power；3 arms", "临时 SPWM 冒烟测试，之后删除。"],
            ["Sine Wave ×3 + Mux", "Simulink > Sources / Signal Routing", "Amplitude=0.90；ω=2πf_grid；phases 0,−2π/3,+2π/3；sample=Ts_power", "临时归一化三相调制波。"],
        ],
        [
            "Universal Bridge DC+ 经 Inverter_DC_Isolator 接 DC+，DC− 接 DC−。三个 AC 端依次接 Three-Phase Series RLC Branch 的 A/B/C。不要交换相序。",
            "RL Branch 输出接 Filter_VI；Filter_VI 输出形成 690 V AC bus。Filter_Cf 与 AC_Load_Base 均从该三相母线并联接入。",
            "再串联一个 Load_VI 到负载支路，确保其电流只表示负载电流 IL，而 Filter_VI 电流表示电感电流 If。",
            "三路幅值 0.90 的正弦波经 Mux 接 PWM Generator (2-Level)，其 6 路门极直接接 Universal Bridge。",
            "第一次可暂用 1200 V DC Voltage Source 代替 PV/电池供电，或让已验证的 DC 支路工作；AC 负载先设 50 kW，Stop time=0.02 s。",
            "确认三相波形和门极顺序后，把负载恢复 P_ac_base=0.5 MW；删除临时 DC 源，保留 SPWM 作为故障排查备用。",
        ],
        [
            ["DC+ / DC−", "隔离器/桥 DC 端", "物理线", "给逆变器供能。"],
            ["Bridge A/B/C", "Filter RL A/B/C", "相序一致", "产生三相电感电流。"],
            ["Filter RL", "Filter_VI → AC bus", "串联", "测 Vc 与 If。"],
            ["AC bus", "Filter_Cf、Load_VI/Load", "并联", "形成 LC 滤波与负载。"],
            ["6 gates", "Universal Bridge gate", "[A上,A下,B上,B下,C上,C下]", "与后续 FCS 一致。"],
        ],
        "三相 Vcabc 应相差 120°，平均值接近 0；线电压 RMS 在数百伏范围。SPWM 幅值 0.90 时理想线电压约 0.612×0.90×1200≈661 V，滤波和负载会略改变。Ifabc 不应含直流偏置。",
        [
            ["Lf", "电流纹波减小但动态变慢、相位延迟增大。", "电流纹波和 MPC 敏感度增大。", "If THD、相位。"],
            ["Cf/Qc", "电压纹波小，但无功/谐振电流增大。", "滤波能力下降。", "Vc THD、Cf 电流。"],
            ["调制幅值", "交流基波电压升高，过 1 进入过调制。", "电压下降。", "Vll_rms。"],
            ["AC 负载 P", "电流与 DC 功率需求增大。", "调试更安全。", "If、Vdc。"],
        ],
        [
            ["桥一使能就短路", "6 门极顺序不对应互补器件对。", "用 Scope 验证 (1,2)、(3,4)、(5,6) 互补；先 blocking/隔离。"],
            ["三相不是 120°", "正弦相位或桥相序错误。", "相位用弧度；B=−2π/3、C=+2π/3；A/B/C 不交叉。"],
            ["输出约 0 V", "DC 隔离器未闭合、门极未使能或 Cf/负载接线短路。", "分层查看 Vdc、gates、桥端电压、滤波后电压。"],
            ["交流电压持续增大", "孤岛没有电压闭环且负载太轻。", "SPWM 仅短时测试；下一阶段必须接 MPVC。"],
        ],
        ["6 门极互补对顺序已验证。", "三相相序和幅值正确。", "Filter_VI=Vcabc/Ifabc，Load_VI=ILabc。", "50 kW 小负载下 20 ms 无 NaN/Inf。"],
        ["V_LL,rms ≈ 0.612·m·Vdc（线性 SPWM）", "Qc_filter = 2π·f·Cf·Vac_ll² ≈ 200 kvar"],
    )

    stage(
        doc, "10", "用标准模块建立 Clarke 变换与 P/Q 计算",
        "论文的 MPVC/MPPC 都在静止 αβ 坐标中预测。这个阶段只做信号处理，不改功率电路；先用已知三相正弦验证坐标变换和功率符号。",
        [
            ["Gain ×2", "Simulink > Math Operations", "Gain=Kclarke；Multiplication=Matrix(K*u)", "abc→αβ；分别处理电压和电流。"],
            ["Product ×4", "Simulink > Math Operations", "Element-wise scalar products", "VαIα、VβIβ、VβIα、VαIβ。"],
            ["Sum ×2", "Simulink > Math Operations", "P: ++；Q:+−", "形成 P/Q 括号项。"],
            ["Gain ×2", "Simulink > Math Operations", "Gain=1.5", "三相瞬时功率系数。"],
            ["Mux/Demux", "Simulink > Signal Routing", "按 3 相或 2 轴维度", "整理向量。"],
            ["Zero-Order Hold", "Simulink > Discrete", "Ts=Ts_mpc", "给预测控制提供同步采样。"],
        ],
        [
            "在初始化脚本后，工作区可增加或直接在 Gain 中填写 Kclarke=[2/3 -1/3 -1/3;0 1/sqrt(3) -1/sqrt(3)]。",
            "Filter_VI 的 Vabc 接一个 Matrix Gain 得 Vc_ab；Ifabc 接另一个得 If_ab；Load_VI 电流接第三个同参数 Gain 得 IL_ab；Grid_VI 之后也用同样方法得 Vg_ab、Ig_ab。",
            "以网侧/滤波侧约定电流计算 P=1.5(VαIα+VβIβ)。Q=1.5(VβIα−VαIβ)，严格保持论文符号。",
            "所有送入 MPC 的 αβ 量先经过 Zero-Order Hold(Ts_mpc)。P/Q 计算既可在 ZOH 前做并再采样，也可先采样后做，但全组必须同一时刻。",
            "使用临时 SPWM 测试：纯阻性负载下平均 Q 应接近 0，P 应与交流负载额定功率同数量级。",
        ],
        [
            ["Vabc", "Matrix Gain → V_ab", "3×1 到 2×1", "电压空间矢量。"],
            ["Iabc", "Matrix Gain → I_ab", "3×1 到 2×1", "电流空间矢量。"],
            ["Vα,Vβ,Iα,Iβ", "4 Products → 2 Sums → 1.5 Gains", "标量", "实时 P/Q。"],
        ],
        "平衡三相正弦的 αβ 分量应近似正交、幅值恒定并以 50 Hz 旋转；纯阻性负载时平均 Q≈0。若 P 的幅值正确但符号为负，说明电流测量方向与功率正方向相反。",
        [
            ["Ts_mpc", "控制更新慢、预测误差增大。", "更新快、计算量增加。", "P/Q 纹波与 FCS 切换。"],
            ["Clarke 归一化", "若误用 2/3 以外系数，P/Q 标度错误。", "同左。", "P 与三相直接点积比较。"],
        ],
        [
            ["αβ 波形失衡", "输入 abc 顺序错或中性点/相电压定义错。", "确认 V-I Measurement 输出为 [Va,Vb,Vc]，相序 A-B-C。"],
            ["P 比额定值差 1.5 倍", "Clarke 与功率系数使用了不同归一化。", "本手册固定 2/3 Clarke 与 3/2 功率公式。"],
            ["纯阻负载 Q 很大", "Q 符号项接错或 Filter Cf 无功占主导。", "先断 Cf 做信号验证；核对 VβIα−VαIβ。"],
        ],
        ["Kclarke 的矩阵维度正确。", "P 与直接 abc 点积数量级一致。", "Q 符号约定已记录。", "所有 MPC 输入同为 Ts_mpc 采样。"],
        ["[xα;xβ] = (2/3)[1 −1/2 −1/2; 0 √3/2 −√3/2][xa;xb;xc]", "P=3/2(VαIα+VβIβ)", "Q=3/2(VβIα−VαIβ)"],
    )

    stage(
        doc, "11", "加入孤岛模式 MPVC 有限控制集预测控制",
        "孤岛时电网不提供电压基准，三相逆变器必须像“可控交流电压源”一样建立 690 V/50 Hz 母线。MPVC 对 8 个桥开关状态逐一预测下一采样时刻的电容电压，选择误差最小者。",
        [
            ["MATLAB Function", "Simulink > User-Defined Functions", "把 pvess_mpvc_block.m 全文粘入；入口函数名改为 fcn", "唯一不可避免的 8 状态枚举算法核。"],
            ["Constant ×3", "Simulink > Sources", "Ad_mpvc；Bd_mpvc；lambda_sw", "精确离散矩阵和切换惩罚。"],
            ["Sine Wave ×3", "Simulink > Sources", "Amplitude=Vac_phase_peak；ω=2πf_grid；phases 0,−2π/3,+2π/3；Ts=Ts_mpc", "孤岛 690 V 线电压参考。"],
            ["Mux + Clarke Gain", "Signal Routing / Math Operations", "3→1 vector；Kclarke", "Vrefabc→Vref_ab。"],
            ["Selector", "Simulink > Signal Routing", "Index vector=[1 3 5]", "从 6 门极取上管状态 Sa,Sb,Sc。"],
            ["Unit Delay", "Simulink > Discrete", "Ts=Ts_mpc；IC=[0;0;0]", "保存上一开关状态用于 Jsw。"],
            ["Product", "Simulink > Math Operations", "gate vector × enable", "统一封锁 6 门极。"],
            ["Data Type Conversion", "Simulink > Signal Attributes", "Output=double", "匹配 Universal Bridge gate 输入。"],
        ],
        [
            "打开 pvess_mpvc_block.m，全选复制。双击 MATLAB Function 块，替换默认代码，并把第一行函数名 pvess_mpvc_block 改成 fcn；其余代码不改。",
            "在 Symbols/Ports and Data Manager 中确认输入尺寸：Vdc=[1 1]；Vc_ab、If_ab、IL_ab、Vref_ab=[2 1]；Ad、Bd=[2 2]；lambda_sw=[1 1]；previousState=[3 1]。输出 gates=[6 1]，stateIndex/Jmin 为标量。",
            "连接 Constant Ad_mpvc、Bd_mpvc、lambda_sw。所有测量向量必须来自阶段 10 的 Ts_mpc Zero-Order Hold。",
            "三路 Vref 正弦经 Mux、Clarke Gain 和 ZOH 得 Vref_ab。其相电压峰值 Vac_phase_peak=√2·690/√3≈563 V。",
            "gates 接 Selector [1 3 5]，再接 Unit Delay，反馈到 previousState。gates 同时经 enable Product 和 double conversion 接 Universal Bridge。",
            "先将逆变器 DC 隔离器打开、enable=0，按 Ctrl+D。编译通过后移除临时 SPWM，enable=1，按阶段 9 的安全次序接通，Stop time 先 0.02 s。",
            "MPVC 稳定后逐步把 AC_Load_Base 从 50 kW 恢复到 0.5 MW，再测试 0.5 MW 交流台阶。",
        ],
        [
            ["Vdc", "MPVC 输入 1", "标量", "决定 8 个候选电压矢量幅值。"],
            ["Vc_ab/If_ab/IL_ab", "MPVC 对应输入", "2×1；同一采样时刻", "状态与扰动。"],
            ["Vref_ab/Ad/Bd", "MPVC 对应输入", "2×1/2×2", "目标与精确离散模型。"],
            ["gates", "enable → double → Universal Bridge", "6×1", "FCS 直接开关，不经过 PWM。"],
            ["gates[1,3,5]", "Unit Delay → previousState", "3×1", "计算切换次数。"],
        ],
        "Vcabc 应建立近似 690 V 线电压、50 Hz 三相波形；stateIndex 在 1–8 间切换，Jmin 始终有限。第一目标是电压不发散，再看 THD 与论文波形。FCS 的门极每 50 μs 决策一次，但实际平均开关频率不等于 20 kHz。",
        [
            ["lambda_sw", "减少换相、降低开关频率，但电压跟踪变差。", "跟踪更积极、开关频率和损耗增加。", "stateIndex 变化率、Vc THD。"],
            ["Ts_mpc", "预测跨度大、模型误差大，等效开关慢。", "性能可提高但计算量更大。", "Jmin、THD、运行时间。"],
            ["Cf", "电压平滑但 LC 共振频率下降。", "电压纹波增大。", "Vc 和 If。"],
        ],
        [
            ["MATLAB Function 端口不出现", "粘贴代码后尺寸推断失败或函数头未改。", "检查入口名 fcn；在 Edit Data 中手动设置固定尺寸。"],
            ["输出门极全 0", "enable=0、Product 维度错误或 Jmin 为 Inf。", "先把 gates/stateIndex/Jmin 接 Display；检查 Vdc>0、Ad/Bd 有值。"],
            ["电压越控越大", "Vc/If/IL 测量位置混淆，或桥相序/门极映射错。", "Vc 取电容节点，If 取滤波电感电流，IL 取负载电流；回退 SPWM 验证桥。"],
            ["50 Hz 波形存在但幅值不对", "参考使用了线电压峰值 976 V 而不是相电压峰值 563 V。", "Sine Wave Amplitude 必须是 Vac_phase_peak。"],
            ["代价忽大忽小并发散", "输入未统一采样或 Ad/Bd 用了错误 Ts。", "所有 MPC 输入 ZOH=Ts_mpc；重新运行初始化脚本。"],
        ],
        ["MPVC 函数固定尺寸编译通过。", "门极顺序为 [A上,A下,B上,B下,C上,C下]。", "Jmin 有限、stateIndex 在 1–8。", "50 kW 后 0.5 MW 负载均可建立三相电压。"],
        ["x=[Vc;If]，u=[Vi;IL]", "A=[0 1/Cf;−1/Lf −Rf/Lf]，B=[0 −1/Cf;1/Lf 0]", "x(k+1)=Ad·x(k)+Bd·u(k)，Ad=exp(A·Ts)，Bd=A⁻¹(Ad−I)B", "Jv=(Vcα,ref−Vcα,pred)²+(Vcβ,ref−Vcβ,pred)²"],
    )

    add_picture(doc, ASSETS / "fcs_cycle_beginner.png", 6.35, "图 5  MPVC/MPPC 每个控制周期的共同计算流程")

    stage(
        doc, "12", "加入变压器、电网和同步并网顺序",
        "并网断路器不能直接闭合。先在断路器开路时测到网侧电压，让 MPVC 的参考从内部正弦平滑切换为电网三相电压；微网电压同相同幅后才闭合断路器。",
        [
            ["Three-Phase Breaker", "搜索 Three-Phase Breaker（SPS）", "External control=on；Initial=open；Ron=1e-3；Rs=1e6；Cs=inf", "并网点 PCC 断路器。"],
            ["Three-Phase V-I Measurement", "SPS > Measurements", "phase-to-ground；current=yes", "断路器网侧 Vgabc/Igabc。"],
            ["Three-Phase Transformer (Two Windings)", "搜索完整名称（SPS）", "Yg/Yg；pu；Nominal=[S_tr,f_grid]；W1=[V_tr_lv,0.002,0.06]；W2=[V_tr_hv,0.002,0.06]；Rm=Lm=500", "690 V/25 kV、2.5 MVA 变压器。"],
            ["Three-Phase Source", "SPS", "Yg；V=V_tr_hv；f=f_grid；phase=0；nonideal；Ssc=S_sc_grid；X/R=XR_grid", "25 kV 公用电网等值。"],
            ["Switch", "Simulink > Signal Routing", "Criteria=u2 ~=0", "在内部 Vrefabc 与实测 Vgabc 之间切换。"],
            ["Step ×2", "Simulink > Sources", "Sync: 0→1 at T_sync；Breaker: 0→1 at T_grid_connect；Ts=Ts_mpc", "1.4 s 同步，1.6 s 并网。"],
            ["Sum + RMS/Abs", "Math Operations / Measurements", "ΔVabc=Vcabc−Vgabc", "并网前检查幅值和相位误差。"],
        ],
        [
            "AC bus 经 Grid_Breaker 接 Grid_VI，再接变压器低压绕组；高压绕组接 25 kV Three-Phase Source。相序 A-B-C 全程一致。",
            "Grid_VI 必须放在断路器的电网侧，否则断路器开路时测不到电网参考。",
            "Synchronization_Reference_Select 的 u1 接 Grid_VI 电压，u3 接内部 Vrefabc，u2 接 Sync Step。输出再做 Clarke 作为 MPVC Vref_ab。",
            "Grid_Breaker 外部控制接 Breaker Step。先临时把 T_sync=0.05、T_grid_connect=0.15 做 0.2 s 短测试，确认顺序后再恢复论文 1.4/1.6 s。",
            "在断路器两侧测 Vcabc 与 Vgabc，Scope 叠加。并网前至少观察若干周期，确认同频、同相、相序一致。",
            "若条件不满足，禁止闭合；先让 MPVC 持续跟随 Vgabc。正式并网时可加入一个逻辑许可 syncOK 与 Breaker Step 做 AND。",
        ],
        [
            ["AC bus", "Grid_Breaker", "微网侧", "PCC 开断。"],
            ["Grid_Breaker", "Grid_VI → Transformer LV", "网侧", "断路器开路仍测 Vg。"],
            ["Transformer HV", "25 kV Source", "A/B/C", "升压并网。"],
            ["Vgabc/Internal Vref", "Sync Switch → MPVC reference", "先内部后电网", "闭环同步。"],
            ["Breaker Step AND syncOK", "Grid_Breaker control", "0=open,1=closed", "避免非同期合闸。"],
        ],
        "t<T_sync：MPVC 跟内部 50 Hz 参考；T_sync<t<T_grid_connect：断路器仍开路，但 Vcabc 应逐步与 Vgabc 重合；到并网时刻两侧差压应很小。闭合瞬间不应出现远大于额定的冲击电流。",
        [
            ["T_grid_connect−T_sync", "同步观察时间更长、更安全。", "可能还未锁相就合闸。", "断路器两侧差压。"],
            ["Source Ssc", "电网更强，电压刚性大、并网冲击更大。", "电网更弱、P/Q 耦合与电压波动增大。", "Ig 峰值、PCC 电压。"],
            ["Transformer leakage", "限制冲击但压降和动态延迟增大。", "并网电流更尖锐。", "Ig、P/Q 响应。"],
        ],
        [
            ["断路器开路时 Vg=0", "Grid_VI 放在了微网侧。", "移动到断路器和变压器之间。"],
            ["闭合瞬间巨大电流", "相位/相序/幅值未同步，或断路器先于参考切换。", "立即开断；叠加 Vc/Vg；确认 T_sync<T_grid_connect。"],
            ["A 相对上但 B/C 对不上", "相序在变压器/连线处交换。", "逐相追线，保持 A-A、B-B、C-C。"],
            ["变压器电压比相反", "低高压绕组接反或额定电压顺序错。", "W1=690 V 接微网，W2=25 kV 接电网。"],
        ],
        ["断路器开路时可测 Vgabc。", "同步参考先切换，断路器后闭合。", "并网前 Vcabc/Vgabc 同相序、同频率、近似同幅。", "合闸冲击在可接受范围。"],
    )

    stage(
        doc, "13", "加入并网 MPPC 与直流母线功率外环",
        "并网后逆变器改为受控功率接口：外层 Vdc PI 决定有功参考，Q–V 下垂决定无功参考，MPPC 从 8 个开关状态中选择下一步 P/Q 最接近参考的状态。功率可双向，取决于参考和符号定义。",
        [
            ["MATLAB Function", "Simulink > User-Defined Functions", "粘贴 pvess_mppc_block.m；入口改 fcn；固定尺寸", "并网 8 状态预测算法核。"],
            ["Sum", "Simulink > Math Operations", "Inputs=+-；Vdc_ref−Vdc", "按本手册 P>0=电网→变换器定义生成误差。"],
            ["Discrete PID Controller", "Simulink > Discrete", "PI；P=Kp_dc；I=Ki_dc；Ts=Ts_mpc；Output limits=[Pref_min,Pref_max]；anti-windup", "Vdc 外环输出 Pref。"],
            ["Bias + Gain", "Math Operations", "Bias=−1；Gain=m_var", "Qref=(Vg_pu−1)m_var。"],
            ["From Workspace", "Simulink > Sources", "Variable=GridVpu_profile", "论文 0.9 pu 电压跌落。"],
            ["Switch", "Simulink > Signal Routing", "GridConnected 选择 MPVC/MPPC gates", "孤岛/并网门极模式切换。"],
            ["Unit Delay/Selector", "Simulink > Discrete / Signal Routing", "与 MPVC 相同", "MPPC 上一状态。"],
            ["Saturation", "Simulink > Discontinuities", "Pref limits", "限制变流器功率命令。"],
        ],
        [
            "粘贴 pvess_mppc_block.m 到第二个 MATLAB Function，入口名改 fcn。输入尺寸：Vdc 标量；Vg_ab=[2 1]；P/Q/Pref/Qref/Ts/Lf/Rf/omega/Pbase/Qbase/lambda 均标量；previousState=[3 1]；输出同 MPVC。",
            "严格采用一个功率方向：本手册与该函数注释规定 P>0 为电网→变换器（从电网吸收），P<0 为变换器→电网（送电）。若 Grid_VI 箭头相反，在进入 P/Q 计算前对 Ig 乘 −1。",
            "Vdc 外环用 e=Vdc_ref−Vdc。于是 Vdc 偏高时 e<0、Pref<0，命令逆变器向电网送电；Vdc 偏低时 Pref>0，从电网吸收补充母线。",
            "Vg_pu 经 Bias(−1)、Gain(m_var) 得 Qref。0.9 pu 时 Qref=(0.9−1)×4e6=−0.4 Mvar；不能再把 100 当成数值倍率。",
            "GridConnected Switch：u3 接 MPVC gates，u1 接 MPPC gates，u2 接并网状态。输出再经过统一 enable 和 Data Type Conversion 接桥。",
            "第一次并网先将 Pref=0、Qref=0 固定，确认无冲击；再启用 Vdc PI；最后启用 Q–V 下垂。每增加一环都单独保存和测试。",
        ],
        [
            ["Vg_ab、P/Q", "MPPC Function", "同 Ts_mpc", "预测当前候选状态下的下一步功率。"],
            ["Vdc_ref−Vdc", "PI → Pref", "P>0 吸收", "维持直流母线能量。"],
            ["Vg_pu−1", "m_var Gain → Qref", "pu→var", "电网电压支持。"],
            ["MPVC/MPPC gates", "Mode Switch → gate enable", "6×1", "模式切换。"],
        ],
        "并网后 Vdc 应回到 1200 V 附近。若 PV/电池/负载使 DC 侧功率过剩，Pref 应为负并向电网送电；若 DC 侧缺功率，Pref 应为正、从电网吸收。电网 0.9 pu 时 Qref 应显示 −4e5 var。",
        [
            ["Kp_dc", "Vdc 响应更快但 Pref 尖峰和并网电流冲击增大。", "响应慢。", "Vdc、Pref、Ig。"],
            ["Ki_dc", "稳态误差消除快但可能低频功率摆动。", "母线保留静差更久。", "Pref 低频振荡。"],
            ["Pbase/Qbase", "对应误差权重减小。", "对应误差权重增大。", "MPPC 偏向 P 还是 Q。"],
            ["m_var", "电压跌落时无功动作更强。", "电压支持弱。", "Qref、Ig、PCC 电压。"],
        ],
        [
            ["Vdc 高时 Pref 正且 Vdc 更高", "P 正方向和外环误差不一致。", "按函数规定 P>0=电网→变换器；e=Vref−Vdc；必要时只翻转网侧电流测量。"],
            ["P 能跟踪但 Q 完全相反", "Q 公式或相序符号反。", "固定使用 Q=1.5(VβIα−VαIβ)；核对 A-B-C。"],
            ["模式切换时桥短暂全乱", "两算法状态未同步或同一采样时刻切换。", "可在 Detect Change 后封锁 1 个 Ts_mpc，并把上一状态复位为 000。"],
            ["P/Q 值巨大", "Vg 使用了高压侧 25 kV 而 Lf 是低压侧参数。", "MPPC 全部使用 690 V 低压侧测量和参数。"],
        ],
        ["P 的正方向与函数注释一致。", "Vdc PI 高压时命令向网送电。", "0.9 pu 得 Qref=−0.4 Mvar。", "MPVC/MPPC 可由 GridConnected 明确切换。"],
        ["P(k+1)=P+Ts[−(Rf/Lf)P−ωQ+3/(2Lf)(|Vg|²−Re{VgVi*})]", "Q(k+1)=Q+Ts[ωP−(Rf/Lf)Q−3/(2Lf)Im{VgVi*}]", "Jp=((Pref−Ppred)/Pbase)²+((Qref−Qpred)/Qbase)²"],
    )

    stage(
        doc, "14", "用 Relational Operator 与 Switch 搭建 EMS",
        "EMS 不直接产生高频门极，只决定能量流向：电池充/放/停、PV 是否退出 MPPT、低优先级负载是否切除。把论文流程图拆成布尔判断后，完全可以用标准模块实现。",
        [
            ["Sum", "Simulink > Math Operations", "Inputs=+--", "Pnet=Ppv−Pac−Pdc。"],
            ["Relational Operator ×5", "Logic and Bit Operations", "Pnet<0；SOC>SOC_min；SOC<SOC_max；SOC<SOC_mid；GridConnected>0.5", "流程图条件。"],
            ["Logical Operator AND/NOT", "Logic and Bit Operations", "按真值表组合", "产生各工作条件。"],
            ["Switch ×6", "Signal Routing", "Criteria=u2 ~=0", "在 +1/0/−1 和保护动作间选择。"],
            ["Constant", "Sources", "+1、0、−1、Dpv_off_mppt、Ibat_max", "放电/停/充电命令及限值。"],
            ["Saturation", "Discontinuities", "[0,1]", "SOC 归一化比例。"],
            ["Gain/Sum/Product", "Math Operations", "见步骤", "battCmd→连续 Ibat_ref。"],
            ["Breaker", "SPS Three-Phase Breaker", "External=on", "低优先级负载切除执行器。"],
        ],
        [
            "先计算 Pnet=Ppv−Pac−Pdc，所有三项都按“产生/消耗为正”的测量约定。Pnet<0 表示本地发电不足。",
            "孤岛真值表：Pnet<0 且 SOC>SOC_min → battCmd=+1；Pnet<0 且 SOC≤SOC_min → loadShed=1；Pnet≥0 且 SOC<SOC_max → battCmd=−1；Pnet≥0 且 SOC≥SOC_max → offMPPT=1。",
            "并网真值表：SOC<SOC_mid(=(SOC_min+SOC_max)/2) → battCmd=−1；否则若 Pnet<0 且 SOC>SOC_min → battCmd=+1；其余 battCmd=0。并网时 loadShed/offMPPT 通常为 0。",
            "用 +1、0、−1 Constant 和级联 Switch 实现 battCmd。每个 Switch 的 u2 接布尔条件，u1 接条件为真输出，u3 接条件为假输出。",
            "把 battCmd 转为 Iref：frac=sat((SOC−SOC_min)/(SOC_max−SOC_min),0,1)；放电 Iref=+Ibat_max·frac；充电 Iref=−Ibat_max·(1−frac)；停止为 0。输出接阶段 8 的并网模式 Switch u1。",
            "offMPPT 进入最终 Duty Switch：真时选择 Dpv_off_mppt，假时选择增量电导 Duty。固定 0.46 只是起始削减点，要根据 Ppv 实测调。",
            "loadShed 经 NOT 后与负载投入命令做 AND，控制低优先级 AC/DC 负载断路器。先用 Display 验证真值表，再连接物理断路器。",
        ],
        [
            ["Ppv/Pac/Pdc", "Pnet Sum", "+--", "本地功率裕量。"],
            ["Pnet/SOC/Grid", "比较与逻辑网络", "标量", "选择 EMS 分支。"],
            ["battCmd", "SOC 电流映射 → Battery mode Switch", "−1/0/+1", "并网电池目标。"],
            ["offMPPT", "PV Duty final Switch", "0/1", "满 SOC 且功率过剩时削减 PV。"],
            ["loadShed", "NOT/AND → load breaker", "0/1", "欠功率且低 SOC 时切负载。"],
        ],
        "先断开所有物理执行器，人工改变 Constant Pnet、SOC、GridConnected，逐格核对真值表。四种孤岛组合和三种并网组合全部正确后，才连接 Duty、Iref 和 Breaker。",
        [
            ["SOC_min", "更保守，较早停止放电/切负载。", "可用能量增加但深放电。", "loadShed 时刻。"],
            ["SOC_max", "允许充得更高。", "较早削减 PV。", "offMPPT 时刻。"],
            ["Dpv_off_mppt", "通常降低 Vpv、可能把工作点移到 MPP 左侧；削减量非单调需实测。", "通常提高 Vpv；也需实测。", "Ppv 而不是只看 Duty。"],
        ],
        [
            ["EMS 不停抖动", "Pnet 在 0 或 SOC 阈值附近来回跨越。", "用 Relay 带滞环替代简单比较，或加 10–50 ms 持续判定。"],
            ["SOC 低仍放电", "比较符号/逻辑优先级错误。", "把 SOC 保护放在最终 Iref 前，形成硬限幅而非仅 EMS 建议。"],
            ["offMPPT 后功率反而升高", "固定 Duty 位于错误方向。", "以小步扫描 Duty，找能可靠降低 Ppv 的值；PV 曲线两侧响应不同。"],
            ["loadShed=1 但负载仍接通", "Breaker 逻辑需要 1=closed。", "实际控制应为 LoadCommand AND NOT(loadShed)。"],
        ],
        ["全部真值表用 Constant 人工验证。", "SOC 保护在最终 Iref 处仍有效。", "负载断路器逻辑为 1=闭合。", "EMS 只在验证后连接物理执行器。"],
        ["Pnet = Ppv − Pac − Pdc", "frac = sat((SOC−SOC_min)/(SOC_max−SOC_min),0,1)", "Iref,dis = Ibat_max·frac；Iref,ch = −Ibat_max·(1−frac)"],
    )

    stage(
        doc, "15", "装入论文事件序列、记录信号并分段运行",
        "最后才复现论文表 2 的负载、辐照度、同步、并网和电压跌落事件。完整 4 s 开关仿真很重，必须用短段检查点逐步扩大时长。",
        [
            ["From Workspace", "Simulink > Sources", "G_profile、Rdc1_profile、Rdc2_profile、GridVpu_profile", "论文时序输入。"],
            ["Step", "Simulink > Sources", "AC load 0→1 at T_ac_step；Sync at T_sync；Grid at T_grid_connect", "离散事件。"],
            ["Three-Phase Breaker", "SPS", "AC step load external control", "0.5 MW 交流负载台阶。"],
            ["To Workspace ×若干", "Simulink > Sinks", "Save format=Timeseries；Decimation=10/25/50", "保存 Vdc/P/Q/SOC/门极等。"],
            ["Scope", "Simulink > Sinks", "分组：DC、PV、Battery、AC、MPC", "在线观察关键量。"],
        ],
        [
            "Irradiance Constant 替换为 From Workspace G_profile；两路 DC 负载控制接对应 Rdc profile；GridVpu 接 Qref 支路。",
            "给第二个 0.5 MW AC_Load_Step 串 Three-Phase Breaker，外部控制接 Step(T_ac_step)。初始开路。",
            "至少记录：Vdc、Vpv、Ipv、Ppv、Ibat、Ibat_ref、SOC、Vcabc、Ifabc、Vgabc、Igabc、P、Q、Pref、Qref、stateIndex、Jmin、battCmd/offMPPT/loadShed。",
            "先运行 0.02 s：只看启动；再 0.2 s：看孤岛稳态；再 0.6 s：越过 AC 负载台阶；再 1.2 s：越过 PV/DC 事件；再 1.7 s：完成并网；最后才 4 s。",
            "每次扩大时长前保存一个模型版本和一张 Scope 截图。若第一次出现 NaN，在刚新增的事件前后缩短 Stop time 二分定位。",
            "论文定量对比时，把工程假设参数和调参记录单独列出；不可把“结构复现通过”直接宣称为“曲线完全复现”。",
        ],
        [
            ["G_profile", "PV Array irradiance", "W/m²", "400→800→400 变化。"],
            ["Rdc profiles", "Variable Resistors", "Ω", "1 MΩ 表示断开，额定值表示投入。"],
            ["GridVpu", "Qref calculation", "pu", "3.5 s 后 0.9 pu。"],
            ["关键测量", "To Workspace", "timeseries", "离线绘图与对比。"],
        ],
        "每个事件之后信号都应重新达到有限稳态：Vdc 围绕 1200 V；SOC 连续无跳变；并网前后模式按时切换；0.9 pu 时 Qref=−0.4 Mvar。允许暂态尖峰，但不得出现 NaN/Inf 或持续增幅振荡。",
        [
            ["日志 Decimation", "文件小、细节少。", "数据大、绘图慢。", "是否仍能看到开关/控制动态。"],
            ["Stop time", "覆盖更多事件但运行显著变长。", "便于定位启动问题。", "先通过最近事件。"],
            ["负载阶跃边沿", "采用斜坡时数值更稳但偏离理想台阶。", "理想台阶冲击大。", "Vdc/If 峰值。"],
        ],
        [
            ["0.2 s 正常、0.6 s 发散", "0.5 s AC 负载台阶触发。", "只检查 AC 负载功率、MPVC 余量、电池 Vdc 环，不改网侧参数。"],
            ["1.7 s 才发散", "同步/并网模式或功率符号。", "固定 Pref=Qref=0 并网；检查 Vc/Vg 差压和 P 正方向。"],
            ["3.5 s 发散", "Qref 单位/符号错误。", "确认 −0.4e6 var，而不是 −40e6 或乘 100。"],
            ["仿真内存很大", "所有 2 μs 信号无抽取记录。", "只有门极/开关细节用小 decimation；慢变量用 25–100。"],
        ],
        ["短时启动、孤岛负载台阶、并网、低电压事件分别通过。", "日志均有时间向量且无 NaN/Inf。", "事件时刻与论文表 2 一致。", "论文值和工程补充值在结果中分开报告。"],
    )

    heading(doc, "阶段验收总表", 1)
    table(doc, ["阶段", "必须通过的可测条件", "未通过时禁止做什么"], [
        ["0–1", "Ctrl+D 无错；Vdc0≈1200 V；极性为正。", "禁止接任一功率源。"],
        ["2", "PV 静态功率/电压数量级正确，Ppv 为正。", "禁止接 Boost 闭环。"],
        ["3–4", "固定 D 稳定；MPPT 以 1 ms 更新并向峰值移动。", "禁止加入大负载。"],
        ["5", "负载 P≈V²/R，方向为正。", "禁止长时靠 Cdc 单独供电。"],
        ["6–8", "门极无直通；Ibat 正方向正确；内环后外环稳定。", "禁止并网时让电池继续控制 Vdc。"],
        ["9–11", "三相相序正确；MPVC 的 J 有限；孤岛 690 V/50 Hz。", "禁止合并网断路器。"],
        ["12", "并网前两侧同相序、同频、近似同幅。", "禁止非同期合闸。"],
        ["13–15", "P/Q 符号一致；Vdc PI 方向正确；事件分段通过。", "禁止直接跑完整 4 s 后盲目调所有参数。"],
    ], [0.8, 3.5, 2.2], font_size=8.0)

    heading(doc, "附录 A：三份必要代码的使用与修改边界", 1)
    callout(doc, "代码占比说明", "初始化脚本不建模；MPVC/MPPC 两个文件各只完成 8 状态枚举。MPPT、PI、SOC、P/Q、同步、EMS 均由可见标准模块搭建。", fill=PALE_GREEN)
    table(doc, ["文件", "运行/复制方法", "允许改什么", "不要随意改什么"], [
        [str(INIT_SCRIPT), "仿真前 run 一次；也可写入 InitFcn。", "工程补充参数、事件时刻、PI 初值。", "论文原值应保留记录；不要在模型和脚本各维护一套同名参数。"],
        [str(MPVC_SCRIPT), "全文粘入 MATLAB Function；入口名改 fcn。", "lambda_sw、代价归一化；确认后可加约束。", "8 状态表、桥电压矢量、Ad/Bd 输入顺序。"],
        [str(MPPC_SCRIPT), "全文粘入第二个 MATLAB Function；入口名改 fcn。", "P/Q 权重和开关惩罚。", "P/Q 正方向未统一前不要改预测方程符号。"],
    ], [2.35, 1.35, 1.30, 1.50], font_size=7.2)

    heading(doc, "MATLAB Function 粘贴操作", 2)
    number(doc, "在 MATLAB 编辑器打开控制器 .m 文件，Ctrl+A、Ctrl+C。")
    number(doc, "双击对应 MATLAB Function 块，Ctrl+A、Ctrl+V。")
    number(doc, "只把第一行的函数名改成 fcn，例如 pvess_mpvc_block 改为 fcn；输入参数名和顺序保持不变。")
    number(doc, "点击 Edit Data/Ports and Data Manager，按阶段 11/13 设置固定尺寸。")
    number(doc, "先把 gates、stateIndex、Jmin 接 Display/Scope，隔离器仍开路时 Ctrl+D。")

    heading(doc, "参数脚本修改后的唯一正确流程", 2)
    code(doc, "% 1) 保存 pvess_beginner_init.m\n"
              "run('D:/PV_MPPT/PV_ESS_MPC_Paper_Reproduction/beginner_from_zero/scripts/pvess_beginner_init.m')\n"
              "% 2) 在 Workspace 确认变量新值\n"
              "% 3) Ctrl+D 更新模型\n"
              "% 4) 只运行包含该参数的最短阶段测试")
    paragraph(doc, "不要同时在 Constant 块里写死数值、又在脚本里维护同一参数。例如 Lf 块应填 Lf，而不是 0.0006；这样改脚本后模型才会同步。")

    heading(doc, "附录 B：最常见的发散因果链", 1)
    table(doc, ["最先出现的异常", "真正根因候选", "最短隔离试验"], [
        ["Vdc 单调上升", "输入功率大于负载；Vdc 外环 P 方向反；PV 无削减。", "断开 PV/电池/逆变器逐支路；固定 Pref=0。"],
        ["Vdc 单调下降", "负载功率超过源；电池放电符号反；并网吸收/送电方向反。", "移除负载；给 Ibat_ref=+100 A；看 Vdc 是否上升。"],
        ["Ibat 交替增长", "电流反馈正反馈；上下管直通；PI 积分饱和。", "I=0、只留小 P；blocking 检查门极。"],
        ["Vcabc 幅值每周变大", "MPVC 的 Vc/If/IL 接错；参考幅值错；桥相序错。", "恢复临时 SPWM；只带 50 kW。"],
        ["并网瞬间 Ig 巨大", "非同期合闸；变压器电压比/相序错。", "断路器开路叠加两侧 abc 电压。"],
        ["0.9 pu 事件才发散", "Qref 单位多乘 100；Q 符号反；无功电流超额定。", "显示 Qref，应为 −4e5 var；先限制 ±0.5 Mvar。"],
        ["MPPT 越追越偏", "Boost D→Vpv 方向反；更新周期继承了 Ts_power。", "固定 D±0.01 扫描；查看 Duty 台阶间隔。"],
    ], [1.45, 2.85, 2.20], font_size=7.7, header_fill=PALE_RED)

    heading(doc, "附录 C：调参顺序与停止条件", 1)
    table(doc, ["顺序", "只调哪些量", "固定哪些量", "停止条件"], [
        ["1 PV 功率级", "Lpv、Cpv、固定 D", "MPPT 关闭", "电压电流有限、变比正确。"],
        ["2 MPPT", "Ts_mppt、Dstep、dead bands", "PV/负载工况固定", "功率趋近峰值且 D 不饱和。"],
        ["3 电池内环", "Kpi、Kii", "Vdc 外环断开，Iref 小阶跃", "正负参考稳定跟踪。"],
        ["4 电池外环", "Kpv、Kiv", "网侧断开、内环固定", "Vdc 恢复且无持续振荡。"],
        ["5 MPVC", "lambda_sw、必要时 Lf/Cf 假设", "小 AC 负载、网断开", "690 V/50 Hz，J 有限。"],
        ["6 同步", "参考切换时间/许可", "Pref=Qref=0", "合闸前差压小。"],
        ["7 MPPC", "Kp_dc、Ki_dc、P/Q 权重", "EMS 先固定", "Vdc 与 P/Q 稳定。"],
        ["8 EMS", "阈值、滞环、动作延时", "所有底层环固定", "真值表和物理动作一致。"],
    ], [0.55, 1.65, 2.35, 1.95], font_size=7.7, header_fill=PALE_YELLOW)

    heading(doc, "附录 D：最终顶层布局参考（只看结构，不要复制）", 1)
    paragraph(doc, "下图来自已有开关级模板，用于你完成全部阶段后的顶层布局自检。它不是本手册的起点，也不要求模块摆放位置一致；只比较功率路径、测量位置和控制分层。")
    if TEMPLATE_IMAGE.exists():
        add_picture(doc, TEMPLATE_IMAGE, 6.45, "图 6  完整模型顶层布局参考")
    else:
        callout(doc, "布局参考图缺失", str(TEMPLATE_IMAGE), fill=PALE_RED)

    heading(doc, "完成判据", 1)
    callout(
        doc,
        "你可以开始论文曲线对比的条件",
        "每个阶段都保留了可回退版本；主电路使用开关器件；PV、电池、逆变器和电网均可独立隔离；所有电流/功率方向有书面定义；"
        "孤岛 MPVC、同步合闸、并网 MPPC 和 EMS 分别通过短时测试；完整事件序列无 NaN/Inf。满足这些条件后，差异才主要来自论文未公开参数和调参，而不是搭建错误。",
        fill=PALE_GREEN,
        border="70AD47",
    )

    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_manual())
