from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(r"D:\PV_MPPT\PV_ESS_MPC_Paper_Reproduction")
DOC_DIR = PROJECT / "docs"
ASSET_DIR = DOC_DIR / "manual_assets"
MODEL_IMAGE = DOC_DIR / "model_root.png"
OUT = DOC_DIR / "PV_ESS_MPC_Paper_Reproduction_Manual_Detailed.docx"
MODEL = PROJECT / "models" / "PV_ESS_MPC_Switching_Template.slx"
SCRIPTS = PROJECT / "scripts"
PAPER = Path(r"D:\PV_MPPT\reference\光伏储能微电网在可变电力输出和负载条件下的预测控制策略.pdf")

BLUE = "2E74B5"
DARK_BLUE = "17365D"
MID_BLUE = "5B9BD5"
PALE_BLUE = "E8EEF5"
LIGHT_BLUE = "DDEBF7"
PALE_GREEN = "E2F0D9"
PALE_YELLOW = "FFF2CC"
PALE_RED = "FCE4D6"
LIGHT_GRAY = "F2F4F7"
GRAY = "666666"
BLACK = "222222"


def rgb(value):
    return RGBColor.from_string(value)


def set_run_font(run, east_asia="Microsoft YaHei", ascii_font="Calibri", size=None):
    run.font.name = ascii_font
    if size is not None:
        run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)


def keep(paragraph, next_item=False):
    ppr = paragraph._p.get_or_add_pPr()
    lines = OxmlElement("w:keepLines")
    lines.set(qn("w:val"), "1")
    ppr.append(lines)
    if next_item:
        node = OxmlElement("w:keepNext")
        node.set(qn("w:val"), "1")
        ppr.append(node)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        item = margins.find(qn("w:" + name))
        if item is None:
            item = OxmlElement("w:" + name)
            margins.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    node = tcpr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tcpr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_border(cell, color="B7C9DC", size="4"):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(Inches(width).twips for width in widths)
    props = table._tbl.tblPr
    width_node = props.first_child_found_in("w:tblW")
    if width_node is None:
        width_node = OxmlElement("w:tblW")
        props.append(width_node)
    width_node.set(qn("w:w"), str(total))
    width_node.set(qn("w:type"), "dxa")
    indent = props.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        props.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = props.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        props.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = list(table._tbl.tblGrid)
    for index, width in enumerate(widths):
        twips = Inches(width).twips
        if index < len(grid):
            grid[index].set(qn("w:w"), str(twips))
        for row in table.rows:
            cell_props = row.cells[index]._tc.get_or_add_tcPr()
            cell_width = cell_props.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_props.append(cell_width)
            cell_width.set(qn("w:w"), str(twips))
            cell_width.set(qn("w:type"), "dxa")


def repeat_header(row):
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    props.append(node)


def prevent_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def add_field(paragraph, code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True

    for name in ["List Bullet", "List Number", "List Bullet 2", "List Number 2"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375 if "2" not in name else 0.65)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in doc.styles:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.10)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(5)
    code_style.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("PV–Battery 混合 AC/DC 微电网复现")
    set_run_font(run, size=8.5)
    run.bold = True
    run.font.color.rgb = rgb(DARK_BLUE)
    run = header.add_run("    功率电路 · MPVC/MPPC · EMS · 分阶段验收")
    set_run_font(run, size=8.2)
    run.font.color.rgb = rgb(GRAY)
    ppr = header._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), MID_BLUE)
    border.append(bottom)
    ppr.append(border)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_run_font(run, size=8)
    run.font.color.rgb = rgb(GRAY)
    add_field(footer, "PAGE")
    run = footer.add_run(" 页")
    set_run_font(run, size=8)
    run.font.color.rgb = rgb(GRAY)


def paragraph(doc, text="", style=None, align=None):
    item = doc.add_paragraph(style=style)
    run = item.add_run(text)
    set_run_font(run)
    if align is not None:
        item.alignment = align
    return item


def bullet(doc, text, level=0):
    return paragraph(doc, text, "List Bullet" if level == 0 else "List Bullet 2")


def number(doc, text, level=0):
    return paragraph(doc, text, "List Number" if level == 0 else "List Number 2")


def heading(doc, text, level=1):
    item = doc.add_paragraph(text, style="Heading " + str(level))
    for run in item.runs:
        set_run_font(run)
    return item


def code(doc, text):
    item = doc.add_paragraph(style="Code Block")
    run = item.add_run(text)
    set_run_font(run, ascii_font="Consolas", size=8.5)
    ppr = item._p.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), LIGHT_GRAY)
    ppr.append(shade)
    return item


def callout(doc, title, body, fill=PALE_YELLOW, border=MID_BLUE):
    box = doc.add_table(rows=1, cols=1)
    set_table_geometry(box, [6.5])
    cell = box.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 100, 140, 100, 140)
    set_cell_border(cell, border, "8")
    item = cell.paragraphs[0]
    run = item.add_run(title)
    set_run_font(run)
    run.bold = True
    run.font.color.rgb = rgb(DARK_BLUE)
    run = item.add_run("\n" + body)
    set_run_font(run)
    item.paragraph_format.space_after = Pt(0)
    prevent_split(box.rows[0])
    paragraph(doc, "")
    return box


def table(doc, headers, rows, widths, font_size=8.2, header_fill=PALE_BLUE):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    set_table_geometry(result, widths)
    for index, text in enumerate(headers):
        cell = result.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        set_cell_border(cell)
        item = cell.paragraphs[0]
        item.paragraph_format.space_after = Pt(0)
        item.paragraph_format.keep_with_next = True
        run = item.add_run(str(text))
        set_run_font(run, size=font_size)
        run.bold = True
        run.font.color.rgb = rgb(DARK_BLUE)
    repeat_header(result.rows[0])
    prevent_split(result.rows[0])
    for row_index, row in enumerate(rows):
        cells = result.add_row().cells
        for index, text in enumerate(row):
            cell = cells[index]
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2:
                set_cell_shading(cell, "F8FAFC")
            item = cell.paragraphs[0]
            item.paragraph_format.space_after = Pt(0)
            run = item.add_run(str(text))
            set_run_font(run, size=font_size)
        prevent_split(result.rows[-1])
    paragraph(doc, "")
    return result


def caption(doc, text):
    item = paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER)
    item.paragraph_format.space_before = Pt(2)
    item.paragraph_format.space_after = Pt(8)
    for run in item.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = rgb(GRAY)


def add_picture(doc, path, width=6.3, label=None):
    path = Path(path)
    if not path.exists():
        callout(doc, "图像缺失", str(path), fill=PALE_RED)
        return
    item = doc.add_paragraph()
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = item.add_run().add_picture(str(path), width=Inches(width))
    alt = label if label else path.stem
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)
    if label:
        caption(doc, label)


def font(size=28, bold=False):
    choices = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for choice in choices:
        if choice.exists():
            return ImageFont.truetype(str(choice), size)
    return ImageFont.load_default()


def center_text(draw, rect, text, face, color=BLACK):
    x1, y1, x2, y2 = rect
    bounds = draw.multiline_textbbox((0, 0), text, font=face, spacing=5, align="center")
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    draw.multiline_text(((x1 + x2 - width) / 2, (y1 + y2 - height) / 2),
                        text, font=face, fill="#" + color, spacing=5, align="center")


def draw_box(draw, rect, text, fill=PALE_BLUE, border=BLUE):
    draw.rounded_rectangle(rect, radius=16, fill="#" + fill, outline="#" + border, width=4)
    center_text(draw, rect, text, font(24, True))


def draw_arrow(draw, start, end, color=BLUE, width=5):
    draw.line([start, end], fill="#" + color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 14
    points = [
        end,
        (x2 - size * ux + 0.55 * size * px, y2 - size * uy + 0.55 * size * py),
        (x2 - size * ux - 0.55 * size * px, y2 - size * uy - 0.55 * size * py),
    ]
    draw.polygon(points, fill="#" + color)


def create_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    if MODEL_IMAGE.exists():
        image = Image.open(MODEL_IMAGE).convert("RGB")
        width, height = image.size
        crops = {
            "model_pv_dc.png": (0, 0, int(width * 0.61), int(height * 0.48)),
            "model_battery.png": (int(width * 0.48), 0, width, int(height * 0.46)),
            "model_inverter_control.png": (0, int(height * 0.33), int(width * 0.75), int(height * 0.82)),
            "model_grid.png": (int(width * 0.56), int(height * 0.49), width, height),
        }
        for name, box in crops.items():
            image.crop(box).save(ASSET_DIR / name)

    canvas = Image.new("RGB", (1800, 980), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((60, 30), "论文复现的物理与控制分层", font=font(40, True), fill="#" + DARK_BLUE)
    boxes = [
        ((70, 160, 350, 330), "PV Array\nSunPower 305 W", PALE_GREEN),
        ((430, 160, 710, 330), "Boost + Lpv/Cpv\n5 kHz PWM", PALE_GREEN),
        ((790, 160, 1070, 330), "1200 V DC 母线\nCdc = 6 mF", PALE_YELLOW),
        ((1150, 80, 1450, 250), "双向半桥 + Battery\n300 V / 0.5 MW", LIGHT_BLUE),
        ((1150, 290, 1450, 460), "两电平三相桥\n20 kHz FCS", LIGHT_BLUE),
        ((1510, 290, 1750, 460), "LC + AC 负载\n电网/变压器", PALE_GREEN),
    ]
    for rect, text, fill in boxes:
        draw_box(draw, rect, text, fill)
    for start, end in [
        ((350, 245), (430, 245)), ((710, 245), (790, 245)),
        ((1070, 210), (1150, 165)), ((1070, 280), (1150, 375)),
        ((1450, 375), (1510, 375)),
    ]:
        draw_arrow(draw, start, end)
    draw.rounded_rectangle((90, 600, 1710, 900), radius=22, fill="#" + PALE_BLUE,
                           outline="#" + MID_BLUE, width=4)
    draw.text((130, 625), "控制与能量管理", font=font(30, True), fill="#" + DARK_BLUE)
    controls = [
        (140, "增量电导 MPPT"),
        (450, "SOC 定向 + 双环 PI"),
        (800, "离网 MPVC"),
        (1100, "并网 MPPC"),
        (1400, "EMS / 模式逻辑"),
    ]
    for x, label in controls:
        draw_box(draw, (x, 700, x + 250, 830), label, "FFFFFF")
    canvas.save(ASSET_DIR / "architecture.png")

    canvas = Image.new("RGB", (1800, 850), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((60, 30), "有限控制集 MPVC / MPPC 每个采样周期做什么", font=font(39, True),
              fill="#" + DARK_BLUE)
    labels = [
        ("采样\nVdc, Vc, If, IL, Vg", PALE_GREEN),
        ("枚举 8 个状态\n000…111", PALE_BLUE),
        ("计算 Viαβ\n预测 k+1", PALE_BLUE),
        ("计算代价\nJV 或 JP", PALE_YELLOW),
        ("选最小值\n直接输出 6 门极", PALE_GREEN),
    ]
    x = 70
    for index, (label, fill) in enumerate(labels):
        draw_box(draw, (x, 220, x + 280, 430), label, fill)
        if index < len(labels) - 1:
            draw_arrow(draw, (x + 280, 325), (x + 340, 325))
        x += 350
    draw.text((95, 535), "MPVC：精确离散 LC 模型预测 Vc(k+1)，最小化电压 αβ 误差。",
              font=font(28), fill="#" + BLACK)
    draw.text((95, 610), "MPPC：预测 If(k+1)→P(k+1), Q(k+1)，最小化有功/无功误差。",
              font=font(28), fill="#" + BLACK)
    draw.text((95, 685), "Ts = 50 μs；模板增加可关闭的换相惩罚 λsw，论文原始代价对应 λsw=0。",
              font=font(28, True), fill="#" + DARK_BLUE)
    canvas.save(ASSET_DIR / "fcs_cycle.png")

    canvas = Image.new("RGB", (1800, 830), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((60, 30), "论文表 2 的 4 s 事件序列", font=font(40, True), fill="#" + DARK_BLUE)
    x0, x1, y = 120, 1700, 250
    draw.line((x0, y, x1, y), fill="#" + DARK_BLUE, width=7)
    events = [
        (0.5, "AC +0.5 MW"),
        (1.0, "PV 上升\nDC +0.6 MW"),
        (1.4, "开始同步"),
        (1.6, "并网"),
        (2.5, "DC +1 MW"),
        (3.0, "PV 下降\nDC -1 MW"),
        (3.5, "电网 -10%\nQ 支撑"),
    ]
    for index, (time_value, label) in enumerate(events):
        x = x0 + time_value / 4.0 * (x1 - x0)
        draw.line((x, y - 30, x, y + 30), fill="#" + BLUE, width=5)
        top = 340 if index % 2 == 0 else 510
        draw.line((x, y + 30, x, top - 10), fill="#" + MID_BLUE, width=3)
        draw_box(draw, (x - 115, top, x + 115, top + 135), label, PALE_BLUE)
        draw.text((x - 30, y - 85), str(time_value) + " s", font=font(23, True),
                  fill="#" + DARK_BLUE)
    draw.text((125, 720), "先分段运行 0.2 s → 1.0 s → 2.0 s → 4.0 s；不要首次就运行完整开关级 4 s。",
              font=font(28, True), fill="#" + DARK_BLUE)
    canvas.save(ASSET_DIR / "timeline.png")

    canvas = Image.new("RGB", (1800, 820), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((60, 30), "模板中的调试隔离开关：默认断开，论文场景脚本自动闭合",
              font=font(37, True), fill="#" + DARK_BLUE)
    draw_box(draw, (80, 190, 410, 360), "PV 源与 Boost", PALE_GREEN)
    draw_box(draw, (735, 190, 1065, 360), "1200 V 母线", PALE_YELLOW)
    draw_box(draw, (1390, 90, 1720, 260), "Battery 半桥", LIGHT_BLUE)
    draw_box(draw, (1390, 330, 1720, 500), "三相逆变器", LIGHT_BLUE)
    for rect, label, yline in [
        ((470, 215, 670, 335), "PV_DC_Isolator", 275),
        ((1120, 100, 1330, 220), "Battery_DC_Isolator", 160),
        ((1120, 370, 1330, 490), "Inverter_DC_Isolator", 430),
    ]:
        draw_box(draw, rect, label, "FFFFFF", MID_BLUE)
        draw.line((rect[0] + 65, yline - rect[1], rect[2] - 65, yline - rect[1]),
                  fill="#" + BLACK, width=4)
    draw_arrow(draw, (410, 275), (470, 275))
    draw_arrow(draw, (670, 275), (735, 275))
    draw_arrow(draw, (1065, 225), (1120, 160))
    draw_arrow(draw, (1330, 160), (1390, 175))
    draw_arrow(draw, (1065, 315), (1120, 430))
    draw_arrow(draw, (1330, 430), (1390, 415))
    draw.text((105, 600), "意义：门极封锁并不能切断反并联二极管；分阶段调试必须能够真正断开功率支路。",
              font=font(29, True), fill="#" + DARK_BLUE)
    draw.text((105, 680), "警告：隔离器断开时不要给 Lpv/Lbat 设置非零初始电流，否则电感无续流路径会产生非物理过压。",
              font=font(27, True), fill="#9B1C1C")
    canvas.save(ASSET_DIR / "isolators.png")


def add_cover(doc):
    paragraph(doc, "")
    item = doc.add_paragraph()
    item.paragraph_format.space_before = Pt(75)
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = item.add_run("PV–Battery 混合 AC/DC 微电网")
    set_run_font(run, size=28)
    run.bold = True
    run.font.color.rgb = rgb(DARK_BLUE)
    item = doc.add_paragraph()
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = item.add_run("开关级功率电路模板与论文复现手册")
    set_run_font(run, size=22)
    run.bold = True
    run.font.color.rgb = rgb(BLUE)
    item = doc.add_paragraph()
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = item.add_run("模块、端口、参数、接线、控制原理、脚本修改、验收与排错")
    set_run_font(run, size=13)
    run.font.color.rgb = rgb(GRAY)
    item.paragraph_format.space_after = Pt(32)
    callout(
        doc,
        "复现对象",
        "J. Hu 等发表于 Applied Energy 221 (2018) 195–203 的论文："
        "A model predictive control strategy of PV-Battery microgrid under variable power generations and load conditions。"
        "本手册以论文图 1、2、5、6、7、8、9，表 1 和表 2 为建模依据。",
        fill=PALE_GREEN,
    )
    table(
        doc,
        ["交付项", "位置", "状态"],
        [
            ["开关级模板模型", str(MODEL), "已生成；保存态为隔离调试模式"],
            ["参数与场景脚本", str(SCRIPTS), "已生成；论文值与工程假设分区"],
            ["论文源文件", str(PAPER), "9 页，已逐页核对"],
            ["手册版本", date.today().isoformat(), "详细操作版"],
        ],
        [1.35, 4.25, 0.90],
        font_size=8.0,
    )
    paragraph(doc, "适用环境：MATLAB/Simulink R2025a、Simscape Electrical Specialized Power Systems。"
                   "建议同时安装 Model Predictive Control Toolbox，但本模板的有限控制集控制器由 Level-2 MATLAB S-Function 实现。",
              align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_toc(doc):
    item = doc.add_paragraph()
    run = item.add_run("目录")
    set_run_font(run, size=22)
    run.bold = True
    run.font.color.rgb = rgb(DARK_BLUE)
    note = doc.add_paragraph()
    note_run = note.add_run("静态目录（不依赖 Word 域更新；详细导航请使用 Word 的标题窗格）。")
    set_run_font(note_run, size=9)
    note_run.italic = True
    note_run.font.color.rgb = rgb(GRAY)
    entries = [
        "先读这一页：模板是什么，不是什么",
        "1 论文整体意义、原理与可复现边界",
        "2 数学原理与符号约定",
        "3 参数总表：论文值、推导值、工程补充值",
        "阶段 0：建立独立项目并确认工具箱",
        "阶段 1：参数脚本与安全保存态",
        "阶段 2：PV Array 与三维 I–V/P–V 验证",
        "阶段 3：开关级 PV Boost 与增量电导 MPPT",
        "阶段 4：1200 V 母线、6 mF 电容与直流负载",
        "阶段 5：电池与双向 Buck-Boost 功率电路",
        "阶段 6：电池离网双环与并网 SOC 定向控制",
        "阶段 7：三相互联变流器、LC 滤波器和交流负载",
        "阶段 8：离网 MPVC",
        "阶段 9：同步、断路器和 690 V/25 kV 电网接口",
        "阶段 10：并网 MPPC 与无功支撑",
        "阶段 11：EMS、Off-MPPT 与切负载逻辑",
        "阶段 12：论文表 2 场景与运行脚本",
        "13 分阶段接入顺序与实操命令",
        "14 如何修改脚本：改哪里、为什么、改后验证什么",
        "15 验证计划：从子系统到论文图 10–19",
        "16 综合排错树",
        "附录 A：完整模块清单与关键参数",
        "附录 B：关键接线总表",
        "附录 C：门极状态表与符号检查",
        "附录 D：已完成验证与已知限制",
        "附录 E：参考文献与可追溯来源",
    ]
    for entry in entries:
        bullet(doc, entry)
    doc.add_page_break()


def stage(doc, number_text, title, meaning, blocks, steps, connections, tuning, faults, checks):
    heading(doc, "阶段 " + number_text + "：" + title, 1)
    callout(doc, "本阶段的意义", meaning, fill=LIGHT_BLUE)
    heading(doc, "使用模块与精确设置", 2)
    table(doc, ["模块名", "Library Browser 路径", "关键参数", "作用"], blocks,
          [1.40, 2.05, 1.60, 1.45], font_size=7.5)
    heading(doc, "搭建步骤", 2)
    for item in steps:
        number(doc, item)
    if connections:
        heading(doc, "接线清单", 2)
        table(doc, ["起点", "终点", "方向/端口", "为什么这样接"], connections,
              [1.35, 1.35, 1.45, 2.35], font_size=7.6)
    if tuning:
        heading(doc, "参数改变会发生什么", 2)
        table(doc, ["参数", "增大时", "减小时", "首要观察量"], tuning,
              [1.05, 1.80, 1.80, 1.85], font_size=7.6, header_fill=PALE_YELLOW)
    if faults:
        heading(doc, "常见问题与修正", 2)
        table(doc, ["现象", "原因", "检查与修正"], faults,
              [1.45, 2.05, 3.00], font_size=7.5, header_fill=PALE_RED)
    heading(doc, "阶段检查点", 2)
    for item in checks:
        bullet(doc, "□ " + item)


def build_manual():
    create_assets()
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)

    heading(doc, "先读这一页：模板是什么，不是什么", 1)
    callout(
        doc,
        "交付结论",
        "本项目交付的是一套可打开、可编译、含真实开关器件和完整功率路径的论文复现模板，"
        "不是已经调到与论文所有曲线逐点重合的最终模型。论文未公开多个决定性参数，"
        "因此手册把“论文原值、可推导值、工程补充值”严格分开，并提供逐阶段校准办法。",
        fill=PALE_GREEN,
    )
    add_picture(doc, ASSET_DIR / "architecture.png", 6.35, "图 1  论文拓扑在模板中的物理与控制分层")
    heading(doc, "两个验收层级", 2)
    table(doc, ["层级", "判断标准", "本次状态", "你下一步要做什么"], [
        ["A 结构模板", "模型可编译；开关级电路存在；端口、参数和脚本可追踪；隔离态数值稳定",
         "已通过。保存态 10 ms：Vdc=1199.93–1200.00 V", "按阶段 0–7 独立验收支路"],
        ["B 论文定量复现", "图 10–19 的时序、幅值、潮流方向、纹波和动态误差满足定义的容差",
         "尚需参数辨识、预充和控制器整定", "按阶段 8–14 分段运行并对齐论文曲线"],
    ], [1.10, 2.35, 1.65, 1.40], font_size=7.8)
    heading(doc, "为什么必须使用调试隔离器", 2)
    add_picture(doc, ASSET_DIR / "isolators.png", 6.35, "图 2  模板保存态的三个主支路隔离器")
    paragraph(doc, "论文原理图通常不画调试隔离器，但开关器件带有反并联二极管，"
                   "仅把门极设为 0 并不能保证电源与母线电气隔离。模板增加 PV_DC_Isolator、"
                   "Battery_DC_Isolator 和 Inverter_DC_Isolator；论文场景脚本将它们闭合，"
                   "因此不改变最终运行拓扑。")

    heading(doc, "1 论文整体意义、原理与可复现边界", 1)
    heading(doc, "1.1 论文解决的核心问题", 2)
    paragraph(doc, "系统同时存在三种不确定性：PV 功率随辐照度变化、交流和直流负载阶跃、"
                   "微电网在离网与并网模式之间切换。作者没有用一个控制器包办全部任务，"
                   "而是把职责分成变换器局部控制和系统级 EMS：PV Boost 负责 MPPT/限发，"
                   "电池双向变换器负责母线功率缺口与 SOC，三相互联变流器负责交流电压或 P/Q，"
                   "EMS 决定充放电、弃光和切负载。")
    table(doc, ["层级", "离网模式", "并网模式", "工程意义"], [
        ["PV Boost", "MPPT；当 SOC 上限且功率过剩时 Off-MPPT", "同左", "尽可能利用光能，同时允许限发"],
        ["电池变换器", "Vdc 外环 + Ibat 内环维持 1200 V", "SOC 定向电流参考", "离网时电池是母线平衡执行器"],
        ["互联变流器", "MPVC 建立 690 V 交流母线", "MPPC 调节有功/无功和 Vdc", "同一桥在模式切换后更换控制目标"],
        ["EMS", "缺电→放电/切负载；过剩→充电/限发", "允许与电网交换功率", "避免局部控制器互相争抢"],
    ], [1.25, 1.75, 1.75, 1.75], font_size=7.8)
    heading(doc, "1.2 论文给出了什么", 2)
    table(doc, ["证据", "论文内容", "在模板中的落点"], [
        ["图 1", "混合 AC/DC 微电网总结构", "完整 DC/AC 功率路径、负载、并网接口"],
        ["图 2–4", "PV Array、Boost、P–V 特性与增量电导 MPPT", "PV_Array、Lpv、Cpv、PV_Boost、IncCond_MPPT"],
        ["图 5、7", "300 V 电池、双向 Buck-Boost、离网双环和并网 SOC 定向", "Battery_Leg、Lbat、Battery_Controller"],
        ["图 6、8", "三相两电平桥、LC、MPVC/MPPC", "Interlinking_Converter、Filter_Lf_Rf、Filter_Cf、MPVC_MPPC_FCS"],
        ["图 9", "EMS 逻辑", "EMS_Fig9 Level-2 MATLAB S-Function"],
        ["表 1", "Vdc、Vac、PV 型号、电池、Cdc、LC", "pvess_init.m 的 Paper values 区"],
        ["表 2", "0.5–3.5 s 事件序列", "run_paper_scenario.m 和 timeseries"],
    ], [0.90, 2.65, 2.95], font_size=7.9)
    heading(doc, "1.3 论文没有给出的量", 2)
    callout(doc, "不能假装知道",
            "论文未公开 Lpv、Cpv、Lbat、器件 Ron/Vf、PI 增益、电池完整放电曲线、"
            "电网低压侧等值阻抗和 FCS 细节（门极映射、延时补偿、归一化）。"
            "这些量集中放在 pvess_init.m 的 Engineering assumptions 区，"
            "在没有作者模型或实验数据前只能称为工程补充值。", fill=PALE_RED)

    heading(doc, "2 数学原理与符号约定", 1)
    heading(doc, "2.1 PV Boost", 2)
    paragraph(doc, "开关周期平均关系为 Vpv≈(1−Dpv)Vdc；因此 Dpv 增大通常使 PV 端电压降低。"
                   "增量电导 MPPT 使用最大功率点条件 dP/dV=I+V·dI/dV=0，"
                   "即 dI/dV=−I/V。位于 MPP 左侧时 dP/dV>0，应升高 Vpv，模板通过减小 Dpv 实现；"
                   "位于右侧则增大 Dpv。")
    heading(doc, "2.2 电池支路", 2)
    paragraph(doc, "规定 Ibat>0 为电池向母线放电。半桥中 PWM 原始脉冲驱动下管，"
                   "上管使用互补并延迟上升沿；中点平均电压约为 (1−Db)Vdc。"
                   "电感方程可写为 Lbat·dIbat/dt=Vbat−(1−Db)Vdc−R·Ibat。"
                   "所以在当前符号下 Db 增大使中点电压下降，推动放电电流增加。")
    paragraph(doc, "SOC 采用库仑计量：dSOC/dt=−100·Ibat/(3600·QAh)。"
                   "论文约束 10%≤SOC≤90%，|Ibat|≤Pbat,N/Vbat。")
    heading(doc, "2.3 离网 MPVC", 2)
    code(doc, "dVc/dt = (If − IL)/Cf\n"
              "dIf/dt = (Vi − Rf·If − Vc)/Lf\n"
              "x(k+1) = Ad·x(k) + Bd·[Vi(k); IL(k)]\n"
              "JV = (Vcα,ref − Vcα,pred)^2 + (Vcβ,ref − Vcβ,pred)^2")
    paragraph(doc, "若直接用前向欧拉预测 Vc(k+1)，Vi 在一步内不会进入 Vc 方程，"
                   "8 个候选状态可能得到相同代价。模板采用论文式 (11) 的精确离散矩阵："
                   "Ad=exp(A·Ts)，Bd=A⁻¹(Ad−I)B，使 Vi 通过 LC 耦合影响 Vc(k+1)。")
    heading(doc, "2.4 并网 MPPC", 2)
    code(doc, "If,pred = If + Ts/Lf · (Vi − Vg − Rf·If)\n"
              "Ppred = 1.5·(Vgα·Ifα,pred + Vgβ·Ifβ,pred)\n"
              "Qpred = 1.5·(Vgβ·Ifα,pred − Vgα·Ifβ,pred)\n"
              "JP = ((Pref−Ppred)/Pbase)^2 + ((Qref−Qpred)/Qbase)^2 + λsw·Nswitch")
    paragraph(doc, "论文原始代价是未归一化的平方误差。模板使用 Pbase/Qbase 归一化，"
                   "避免 P、Q 数值量级直接改变权重；将 λsw=0 可回到论文未含换相惩罚的形式。")
    heading(doc, "2.5 电压跌落无功参考", 2)
    paragraph(doc, "模板采用 Qref=(Vg,pu−1)·mvar，mvar=4×10⁶。"
                   "当电网电压跌落 10% 时 Qref=−0.4 Mvar，与论文事件 9 一致。"
                   "正负号取决于电流方向；模板约定负 Q 为向电网注入无功。")
    add_picture(doc, ASSET_DIR / "fcs_cycle.png", 6.35, "图 3  每个 50 μs 采样周期的有限控制集计算")

    heading(doc, "3 参数总表：论文值、推导值、工程补充值", 1)
    table(doc, ["参数", "模板值", "类别/来源", "修改位置", "主要影响"], [
        ["Vdc_ref", "1200 V", "论文表 1", "pvess_init.m", "所有 DC 负载电阻、占空比与电压应力"],
        ["Vac_ll", "690 V", "论文表 1 低压侧", "pvess_init.m", "交流参考幅值和电容无功"],
        ["Cdc", "6 mF", "论文表 1", "Cdc", "母线纹波与功率失衡暂态"],
        ["Lf / Cf / Rf", "0.6 mH / 1338 μF / 0.019 Ω", "论文表 1", "Filter_Lf_Rf / Filter_Cf", "MPVC 模型、谐振和阻尼"],
        ["PV module", "SunPower SPR-305E-WHT-D", "论文表 1", "PV_Array", "单组件 I–V/P–V"],
        ["Ns/Np", "10 / 656", "由图 3 约 2 MW、Vmp≈550 V 推导", "pvess_init.m", "阵列电压和功率规模"],
        ["Battery", "300 V, 0.5 MW, 1300 Ah", "论文表 1", "Battery", "SOC 斜率与电流上限"],
        ["Fsw_pv/bat", "5 kHz", "论文第 6 节", "PV_PWM/Battery_PWM", "DC/DC 纹波与计算量"],
        ["Ts_mpc", "50 μs", "论文 20 kHz 采样", "MPVC_MPPC_FCS", "预测矩阵和开关更新"],
        ["Lpv/Cpv", "0.2 mH / 5 mF", "工程补充", "pvess_init.m", "PV 电流纹波和 MPPT 暂态"],
        ["Lbat", "0.5 mH", "工程补充", "pvess_init.m", "电池电流纹波和内环带宽"],
        ["PI gains", "见脚本", "工程补充，未整定到论文", "pvess_init.m", "母线跌落、超调和电流峰值"],
        ["Ts_power", "2 μs", "数值选择", "powergui / Solver", "开关解析度与运行时间"],
    ], [1.15, 1.45, 1.65, 1.35, 0.90], font_size=7.2)

    stage(
        doc, "0", "建立独立项目并确认工具箱",
        "避免覆盖旧论文模型；先确保路径、库和文件权限正确，再开始任何电气接线。",
        [
            ["MATLAB Project/普通文件夹", "D:\\PV_MPPT\\PV_ESS_MPC_Paper_Reproduction", "保持 models/scripts/docs/results 四级", "独立复现空间"],
            ["powergui", "Simscape > Electrical > Specialized Power Systems", "Discrete；Ts=Ts_power", "SPS 网络求解与初始化"],
            ["Model Settings", "Simulation > Model Configuration Parameters", "Fixed-step；ode3；2e-6", "兼容 Battery 连续内部状态"],
        ],
        [
            "复制或保留旧项目，不删除旧模型。新项目文件夹已经建立。",
            "运行 ver，确认 Simulink 和 Simscape Electrical 可用。",
            "在 MATLAB Current Folder 中进入项目根目录，运行 START_HERE.m。",
            "第一次仅执行到 run_short_validation，不要直接运行 4 s 场景。",
        ],
        [],
        [
            ["Ts_power", "运行变慢、开关波形更准确", "运行更快、峰值和纹波误差增大", "门极与电感电流"],
            ["Solver", "更高阶不一定更好，SPS 离散网络仍由 powergui 控制", "FixedStepDiscrete 会因 Battery 连续状态报错", "编译诊断"],
        ],
        [
            ["找不到 sps...Lib", "Simscape Electrical 未安装或路径缓存异常", "确认 toolbox 安装；不要手工复制库文件；重启 MATLAB 后 rehash toolboxcache"],
            ["模型名 shadowing 警告", "同名模型已打开或路径上有缓存文件", "关闭同名模型，确保 models 文件夹只保留一个 .slx；清理无关 .slxc 不影响源文件"],
        ],
        [
            "START_HERE.m 能定位 scripts 和 models。",
            "模板模型打开后 powergui 显示 Discrete 2e-6 s。",
            "未修改或覆盖旧项目文件。",
        ],
    )

    stage(
        doc, "1", "参数脚本与安全保存态",
        "把论文值与工程假设集中管理；让模型直接按 Run 时处于隔离调试态，而论文场景脚本再显式闭合支路。",
        [
            ["pvess_init.m", "项目 scripts", "Paper values / Engineering assumptions / Event schedule", "唯一参数入口"],
            ["Model Workspace", "Model Explorer > Model Workspace", "保存所有参数和 timeseries", "模型脱离 base workspace 也能打开"],
            ["Constant", "Simulink > Sources > Constant", "PV_connect=0 等保存态值", "控制调试隔离器"],
        ],
        [
            "打开 pvess_init.m，先阅读 Paper values 区；论文原值不要在其他脚本重复定义。",
            "阅读 Engineering assumptions 区。任何替换都同时记录来源、日期和理由。",
            "执行 build_pvess_template。注意：它会重建并覆盖同名模板 .slx；手工修改模型后不要再次运行，除非已同步修改 builder。",
            "确认模型工作区中的 PV_connect、Battery_connect、Inverter_connect 为 0，三个 blocking 为 1，Inverter_enable 为 0。",
        ],
        [],
        [
            ["ILpv0", "非零可减少闭合支路启动时间，但隔离器断开时会产生无续流过压", "零启动稳健但达到稳态更慢", "Vpv、Ipv、Vdc"],
            ["Vdc0", "高于参考会使电池吸能或并网侧导出功率", "低于参考会触发电池放电/并网输入", "启动电流、Pref"],
            ["SOC0", "并网 SOC 逻辑更倾向放电", "更倾向充电", "IbatRef、EMS_cmd"],
        ],
        [
            ["改了 init 但模型没变化", "只修改了函数，没有重建或 SimulationInput 覆盖", "运行 build_pvess_template，或用 setVariable(...,'Workspace',model) 覆盖"],
            ["重建后手工接线消失", "builder 是模型结构的生成源", "把结构变化写回 build_pvess_template.m，再重建；不要双源维护"],
        ],
        [
            "模型保存态三个主隔离器断开。",
            "直接运行 10 ms，Vdc 保持约 1200 V 且无 Inf/NaN。",
            "论文值和工程补充值在脚本中分区清晰。",
        ],
    )

    stage(
        doc, "2", "PV Array 与三维 I–V/P–V 验证",
        "先证明光伏源在不同辐照度下具有正确的电气特性，再连接 Boost；静态曲线不对，后续 MPPT 没有意义。",
        [
            ["PV Array", "Simscape > Electrical > Specialized Power Systems > Renewable Energy", "Module=SunPower SPR-305E-WHT-D；Ns=10；Np=656", "论文 PV 源"],
            ["From Workspace", "Simulink > Sources", "G_profile", "辐照度输入"],
            ["Constant", "Simulink > Sources", "25 °C", "温度输入"],
            ["Voltage/Current Measurement", "SPS > Sensors and Measurements", "方向按图标 +/i", "输出 Vpv、Ipv"],
            ["Parallel RLC Branch", "SPS > Passive Components", "C；5 mF；IC=547 V", "PV 输入端去耦"],
        ],
        [
            "保持 PV_Isolator 和 PV_DC_Isolator 断开；把 Lpv 初始电流设为 0。",
            "单独运行 PV 支路，G 依次设 250、500、750、1000 W/m²，温度 25 °C。",
            "若要生成静态 P–V 曲线，暂时不用 Boost，给 PV 端接可扫描电阻并记录 V·I。",
            "检查 1000 W/m² 时阵列功率峰值约 2 MW、Vmp 约 550 V、Voc 约 640 V，趋势应与论文图 3 一致。",
        ],
        [
            ["Irradiance_Profile/1", "PV_Array/Ir", "Simulink", "改变光生电流"],
            ["PV_Temperature/1", "PV_Array/T", "Simulink", "改变 Voc 和 Pmax"],
            ["PV_Array +", "Cpv + / Vpv +", "SPS 电气", "建立 PV 正端节点"],
            ["PV_Array −", "Cpv − / DC_Ground", "SPS 电气", "公共负端"],
        ],
        [
            ["Np_pv", "电流和功率近似按比例增加", "电流和功率降低", "Isc、Pmax"],
            ["Ns_pv", "Vmp/Voc 增加，Boost 占空比降低", "阵列电压降低，电流压力增大", "Vmp、Dpv"],
            ["Cpv", "Vpv 更平滑、MPPT 变慢", "纹波增加、数值更敏感", "Vpv 纹波"],
            ["温度", "Voc 和 Pmax 通常下降", "Voc 上升", "P–V 峰位置"],
        ],
        [
            ["Ipv 达到 MA 级且 Vpv 反向", "隔离器断开但 Lpv 仍有非零初始电流，电感无续流路径", "将 ILpv0 设为 0；或先闭合完整续流路径再应用稳态初值"],
            ["脚本无参数报错", "把生成查表脚本当无参脚本调用", "本项目不依赖旧 pvbatt_generate_pv_lookup；如沿用，必须显式传入电压、温度、辐照度网格和参数结构"],
            ["Pmax 明显不是 2 MW", "Ns/Np 或辐照度单位错误", "PV Array 输入使用 W/m²；核对 Ns=10、Np=656 和组件预设"],
        ],
        [
            "I–V 曲线单调合理，P–V 每条只有一个主峰。",
            "不同辐照度的 Pmax 近似按比例变化。",
            "静态 PV 测试不连接 Boost 输出母线。",
        ],
    )

    stage(
        doc, "3", "开关级 PV Boost 与增量电导 MPPT",
        "把论文图 2 从平均公式变成真实 IGBT/二极管、Lpv、Cpv 和 5 kHz PWM；验证占空比到 Vpv 的实际方向。",
        [
            ["Series RLC Branch", "SPS > Passive Components", "RL；Lpv=0.2 mH；R=2 mΩ", "Boost 输入电感"],
            ["Boost Converter", "SPS > Power Electronics", "Switching devices；Ron=1 mΩ；Vf=1 V", "升压主电路"],
            ["PWM Generator (DC-DC)", "SPS > Control and Measurements > PWM", "5 kHz；Ts=2 μs", "占空比→门极"],
            ["Level-2 MATLAB S-Function", "Simulink > User-Defined Functions", "pvess_mppt_sfun；Ts=1 ms", "增量电导算法"],
            ["Breaker", "SPS > Switches and Breakers", "PV_Isolator / PV_DC_Isolator", "分阶段接入"],
        ],
        [
            "先闭合 PV_Isolator，保持 PV_DC_Isolator 断开，验证输入侧不会出现反向电压。",
            "用固定 Dpv=Dpv0≈0.544 代替 MPPT，闭合 PV_DC_Isolator，母线接预充电容或受控负载。",
            "固定占空比方向确认后，再恢复 IncCond_MPPT 输出。",
            "把 MPPT 采样周期保持为 1 ms；不要让算法在每个 2 μs 电力步更新。",
            "观察 Vpv、Ipv、Ppv 和 Dpv，而不是只看 Ppv。",
        ],
        [
            ["PV +", "Ipv_Measurement→Lpv→PV_Isolator→PV_Boost 输入", "正向电流", "对应论文 id"],
            ["PV_Boost +", "PV_DC_Isolator→Cdc +", "升压输出", "向 1200 V 母线送能"],
            ["PV_Boost −", "DC_Ground", "公共负端", "闭合功率回路"],
            ["IncCond_MPPT", "PV_PWM→PV_Boost/g", "Simulink", "5 kHz 门极"],
        ],
        [
            ["Dpv_step", "跟踪更快、稳态振荡更大", "跟踪更慢、纹波更小", "Ppv、Dpv"],
            ["Ts_mppt", "更新变慢、抗纹波更强", "更快但可能追逐开关纹波", "Vpv/Ipv 差分"],
            ["Lpv", "电流纹波降低、动态变慢", "纹波和峰值增加", "iLpv"],
            ["Dpv_off_mppt", "Vpv 被拉低，可能不是限发方向", "Vpv 升高靠近 Voc，功率下降", "Vpv、Ppv"],
        ],
        [
            ["Dpv 增大反而 Vpv 升高", "测量极性或 Boost 接线方向错误", "先固定占空比做两点试验；确认 Dpv 增大使 Vpv 降低，再启用 MPPT"],
            ["MPPT 发散", "使用原始开关纹波做差分、采样过快或步长过大", "在 Ts_mppt 采样；减小 Dpv_step；对 V/I 加低通或周期平均"],
            ["母线过压", "PV 已送能但电池/逆变器/负载未接入", "调试时使用受控耗能支路；不要在无负载母线上长时间闭合 PV"],
        ],
        [
            "固定占空比时 Vpv≈(1−Dpv)Vdc 的趋势成立。",
            "启用 MPPT 后 Ppv 朝静态 P–V 峰值移动。",
            "占空比始终位于 0.35–0.75。",
        ],
    )

    stage(
        doc, "4", "1200 V 母线、6 mF 电容与直流负载",
        "建立功率平衡的共同节点。母线电压不是独立电源，而是所有支路净电流对 Cdc 积分的结果。",
        [
            ["Parallel RLC Branch", "SPS > Passive Components", "C；6 mF；IC=1200 V", "论文 Cdc"],
            ["Variable Resistor", "SPS > Passive Components", "Rabsmin=0.2 Ω", "可变 DC 负载"],
            ["From Workspace", "Simulink > Sources", "Rdc1_profile / Rdc2_profile", "0.6/1 MW 事件"],
            ["Current Measurement", "SPS > Sensors", "Idc1、Idc2", "计算 Pdc"],
            ["Product/Sum", "Simulink > Math Operations", "Pdc=Vdc·(Idc1+Idc2)", "EMS 测量"],
        ],
        [
            "仅保留 Cdc、Vdc 测量和 1 MΩ 等效断开负载，运行 10 ms；Vdc 应近似保持 1200 V。",
            "把 0.6 MW 负载电阻计算为 R=1200²/0.6e6=2.4 Ω。",
            "把 1 MW 负载电阻计算为 R=1200²/1e6=1.44 Ω。",
            "用 timeseries 在目标时刻从 1 MΩ 变到工作电阻；理想阶跃会引入大冲击，必要时用 50–200 μs 斜坡。",
        ],
        [
            ["Cdc +", "Vdc + / Idc1 + / Idc2 + /三个 DC 隔离器", "公共正母线", "所有支路在此交换功率"],
            ["Cdc −", "Vdc − / 两负载 − / DC_Ground", "公共负母线", "统一参考"],
            ["Rdc profile", "Variable Resistor/R", "Simulink", "以电阻实现恒阻负载"],
        ],
        [
            ["Cdc", "电压变化更慢、纹波小、故障能量大", "动态更快、纹波和控制难度增大", "Vdc、iCdc"],
            ["负载电阻", "功率降低", "功率按 V²/R 增大", "Pdc、Vdc"],
            ["负载阶跃斜率", "冲击减小但不再是理想阶跃", "更接近论文事件但数值更尖锐", "Idc 峰值"],
        ],
        [
            ["Vdc 自发上升", "某电源支路仍通过二极管连接或初始化电感向母线放能", "打开 DC 隔离器；把断开支路电感初值设为 0；逐支路检查"],
            ["Variable Resistor 报极小电阻", "输入短暂经过 0 或负值", "timeseries 不要插值穿零；Rabsmin 设置正值"],
            ["功率不是标称值", "负载是恒阻而非恒功率", "论文给出 MW 事件但未明确负载模型；若需恒功率，换 Controlled Current Source 并加限幅"],
        ],
        [
            "隔离态 10 ms 内 Vdc 变化小于 0.1 V。",
            "2.4 Ω 和 1.44 Ω 在 1200 V 附近分别约为 0.6 MW、1 MW。",
            "所有功率符号与 EMS 约定一致。",
        ],
    )

    stage(
        doc, "5", "电池与双向 Buck-Boost 功率电路",
        "复现论文图 5 的 300 V 电池、Lbat 和上下管半桥，并明确 Ibat 的正方向与门极互补关系。",
        [
            ["Battery", "SPS > Sources", "Li-ion；300 V；1300 Ah；SOC0", "论文储能"],
            ["Two-Quadrant DC/DC Converter", "SPS > Power Electronics", "Switching devices；Ron=1 mΩ", "双向半桥"],
            ["Series RLC Branch", "SPS > Passive Components", "RL；0.5 mH；2 mΩ；IC=0 A", "电池侧电感"],
            ["Current Measurement", "SPS > Sensors", "正方向 Battery→Lbat→半桥", "Ibat>0 放电"],
            ["PWM Generator (DC-DC)", "SPS > PWM", "5 kHz", "下管原始脉冲"],
            ["NOT + Unit Delay + AND", "Simulink Logic/Discrete", "1 个 Ts_power 死区", "互补与防直通"],
        ],
        [
            "保持 Battery_DC_Isolator 和 Battery_Isolator 断开，确认 Battery 初始电压和 SOC。",
            "闭合电池侧隔离器但保持 DC 侧断开，用固定 Db 验证中点平均电压。",
            "闭合完整支路后，用 Db0=1−300/1200=0.75 起步。",
            "下管接 PWM 原始脉冲，上管接 NOT 后的脉冲；两路上升沿都经过 Unit Delay+AND。",
            "Gate Mux 顺序为 [upper; lower]，再转换为 double 接 Battery_Leg/g。",
        ],
        [
            ["Cdc +", "Battery_DC_Isolator→Battery_Leg +", "高压侧", "半桥正端"],
            ["Battery_Leg Out", "Battery_Isolator→Lbat→Ibat_Measurement→Battery +", "中点到电池", "双向电流路径"],
            ["Battery −", "DC_Ground", "公共负端", "完成回路"],
            ["Battery_PWM P", "NOT→Upper；直接→Lower", "互补", "避免两个独立 PWM 相位不一致"],
        ],
        [
            ["Lbat", "纹波降低、内环变慢", "纹波与峰值增加", "Ibat"],
            ["dead time", "直通裕量增加、平均电压误差增加", "效率/线性变好但直通风险上升", "上下管门极"],
            ["R_bat_internal", "端电压跌落和损耗增加", "更理想但可能高估功率能力", "Vbat、Pbat"],
            ["Db", "在当前符号下放电电流趋势增加", "放电减小/充电增强", "Ibat"],
        ],
        [
            ["上下管同时为 1", "用了两个独立 PWM 或 Mux 顺序错误", "只用一个 PWM 产生互补；记录 6–10 个周期检查重叠"],
            ["充电时 Ibat 为正", "传感器方向与约定相反", "旋转 Current Measurement；不要在控制器内随意乘 −1 掩盖拓扑"],
            ["Battery 与电容并联报错", "理想电压源直接并理想电容", "不要在电池端直接并理想 C；若必须，使用带 ESR 的 RC 支路"],
        ],
        [
            "固定 Db 时 Ibat 有界且纹波频率为 5 kHz。",
            "Ibat>0 对应 SOC 下降，Ibat<0 对应 SOC 上升。",
            "上下管不存在重叠导通。",
        ],
    )

    stage(
        doc, "6", "电池离网双环与并网 SOC 定向控制",
        "离网时电池维持 Vdc；并网时互联变流器维持 Vdc，电池改为按 SOC 决定充放电速率，避免两个外环争抢。",
        [
            ["Battery_Controller", "Level-2 MATLAB S-Function", "pvess_battery_sfun", "两种模式统一实现"],
            ["SOC_Integrator", "Simulink > Discrete", "Ts=100 μs；IC=SOC0；0–100%", "库仑计量"],
            ["Mux/Demux", "Simulink > Signal Routing", "输入 5；输出 duty/Iref", "控制接口"],
            ["Saturation", "控制器内部", "Iref±1666.7 A；Db 0.05–0.95", "约束"],
        ],
        [
            "离网模式输入 gridTied=0。外环误差使用 Vdc_ref−Vdc，使 Vdc 低时 Iref 为正放电。",
            "内环误差 Iref−Ibat 经 PI 叠加到 Db0。",
            "并网模式 gridTied=1，关闭 Vdc 外环积分更新，Iref 由 EMS battCmd 和 SOC 线性曲线生成。",
            "充电参考：Iref=−Imax·(SOCmax−SOC)/(SOCmax−SOCmin)；放电参考：Iref=Imax·(SOC−SOCmin)/(SOCmax−SOCmin)。",
            "先整定电流内环，再整定 Vdc 外环；外环带宽应至少比内环低 5–10 倍。",
        ],
        [
            ["Grid_Connected/Vdc/Ibat/SOC/EMS_cmd", "Battery_Control_Mux", "5 输入", "模式和测量"],
            ["Battery_Controller duty", "Battery_PWM", "0.05–0.95", "门极调制"],
            ["Ibat", "SOC_Coulomb_Gain→SOC_Integrator", "−100/(3600Q)", "SOC 更新"],
        ],
        [
            ["Kpi_bat/Kii_bat", "电流跟踪快，过冲/噪声放大", "跟踪慢、母线跌落大", "Ibat、Db"],
            ["Kpv_bat/Kiv_bat", "Vdc 恢复快但 Iref 峰值大", "Vdc 恢复慢", "Vdc、Iref"],
            ["Ibat_max", "功率支撑增强、热/电流应力上升", "保护更强但可能无法平衡功率", "Pbat、Vdc"],
        ],
        [
            ["Vdc 环发散", "内环方向错误、Db 到 Ibat 符号未验证或外环比内环快", "退回固定 Db；确认方向；只闭合电流内环；最后接 Vdc 外环"],
            ["切到并网时积分突跳", "离网积分器残留", "切换时冻结或回算积分状态；模板停止更新外环积分，但仍需做无扰切换整定"],
            ["SOC 很快跑满/跑空", "QAh 单位忘乘 3600 或 Ibat 符号错误", "使用 dSOC/dt=−100I/(3600QAh)"],
        ],
        [
            "离网负载增加时 Iref 和 Ibat 变正，Vdc 恢复。",
            "并网低 SOC 充电电流绝对值更大，高 SOC 时趋零。",
            "SOC 始终限制在 0–100%，EMS 使用 10–90% 工作窗。",
        ],
    )

    stage(
        doc, "7", "三相互联变流器、LC 滤波器和交流负载",
        "建立论文图 6 的真实三相桥和 LC，被控量 Vc 是滤波电容/交流母线电压，不是桥臂开关电压。",
        [
            ["Universal Bridge", "SPS > Power Electronics", "3 arms；IGBT/Diodes；Ron=1 mΩ", "三相两电平桥"],
            ["Three-Phase Series RLC Branch", "SPS > Passive Components", "RL；0.019 Ω；0.6 mH", "论文 Lf/Rf"],
            ["Three-Phase Parallel RLC Load", "SPS > Passive Components", "Y grounded；Qc=ωCfVLL²", "等效 1338 μF 三相电容"],
            ["Three-Phase V-I Measurement", "SPS > Sensors", "phase-to-ground；current=yes", "Vcabc、Ifabc"],
            ["Three-Phase Parallel RLC Load", "SPS > Loads", "0.5 MW base + 0.5 MW switched", "论文 AC 负载"],
        ],
        [
            "保持 Inverter_DC_Isolator 断开，核对桥、Lf/Rf、Filter_VI、Cf 和负载的三相顺序 A/B/C。",
            "闭合 DC 隔离器前，把 Inverter_enable=0，确认所有 6 门极为 0。",
            "先用标准 PWM Generator (2-Level) 和小调制度建立 50 Hz 正弦，再替换为 FCS 门极，便于排除功率电路问题。",
            "Universal Bridge 的门极配对是 (1,2)、(3,4)、(5,6)；模板输出 [Sa,¬Sa,Sb,¬Sb,Sc,¬Sc]。",
            "交流电容用 Qc=2πfCfVLL² 配置，约 0.2 Mvar；不要把 1338 μF 误填为每相总线电容的三倍。",
        ],
        [
            ["Cdc ±", "Inverter_DC_Isolator→Universal Bridge ±", "DC 端", "桥直流电源"],
            ["Bridge A/B/C", "Filter_Lf_Rf A/B/C", "三相顺序一致", "形成 Vi→Lf"],
            ["Filter_Lf_Rf a/b/c", "Filter_VI A/B/C", "三相", "测量电感电流"],
            ["Filter_VI a/b/c", "Filter_Cf、Load_VI、Grid_Breaker", "交流母线分支", "Vc 节点"],
        ],
        [
            ["Lf", "电流纹波小、动态慢、压降大", "电流变化快、纹波和 FCS 灵敏度高", "If、THD"],
            ["Cf", "电压平滑、谐振频率降低、无功增大", "Vc 纹波增加", "Vc、谐振"],
            ["Rf", "阻尼增加、损耗增加", "效率高但 LC 振荡增强", "Vc 振铃"],
            ["Ts_power", "开关解析度变差", "计算量增加", "桥电流峰值"],
        ],
        [
            ["Vabc 相序错误", "A/B/C 接线或 Vref 相位错误", "检查 Vref_A=0、B=−120°、C=+120°；逐相追线"],
            ["桥直流侧电流爆炸", "门极配对错误造成同桥臂直通", "用标准 PWM 发生器确认配对；记录每对门极互补性"],
            ["LC 强烈振荡", "Rf 太小、初值为零却直接施加全幅电压", "逐步提高参考幅值；核对 Rf=0.019 Ω；必要时增加启动阻尼"],
        ],
        [
            "标准 PWM 下能建立三相对称 50 Hz 电压。",
            "每相桥臂上下管互补，无直通。",
            "Vc 测量位置在 Cf/负载节点，不在桥输出端。",
        ],
    )

    stage(
        doc, "8", "离网 MPVC",
        "离网时互联变流器必须像电压源一样建立稳定交流母线，电池同时在 DC 侧维持 Vdc。",
        [
            ["MPVC_MPPC_FCS", "Level-2 MATLAB S-Function", "pvess_fcs_sfun；mode=0", "8 状态枚举"],
        ["Sine Wave ×3", "Simulink > Sources", "相电压峰值 √2·690/√3≈563 V；50 Hz", "离网 Vrefabc"],
            ["Mux/Switch", "Simulink > Signal Routing", "同步前内部参考；同步时 Vg", "参考切换"],
            ["Zero-order behavior", "S-Function sample time", "Ts=50 μs", "门极保持一个预测周期"],
        ],
        [
            "先在空载或小负载下闭合 Inverter_DC_Isolator，mode=0，参考从 0 斜坡到额定幅值。",
            "确认 Clarke 变换使用幅值不变形式 2/3·[1 −1/2 −1/2; 0 √3/2 −√3/2]。",
            "对每个候选状态计算 Viαβ，再用 Ad/Bd 预测 α、β 两轴 Vc(k+1)。",
            "代价最小的状态保持 50 μs；记录 FCSState_log 和 FCSCost_log。",
            "接入 0.5 MW 基础负载后再做 0.5 MW 阶跃，比较论文图 10。",
        ],
        [
            ["mode,Vdc,Vcabc,Ifabc,ILabc,Vgabc,Vrefabc,Pref,Qref", "FCS_Input_Mux", "总维度 19", "统一控制器输入"],
            ["FCS gates", "Inverter_Gate_Enable→Universal Bridge/g", "6 维 double", "直接开关状态"],
            ["Filter_VI V/I", "FCS Vc/If", "abc", "状态反馈"],
            ["Load_VI I", "FCS IL", "abc", "负载扰动输入"],
        ],
        [
            ["Ts_mpc", "预测间隔大、纹波与误差增加", "性能提高但计算量增加", "Vc 误差、开关频率"],
            ["lambda_sw", "换相少、纹波可能增大", "更接近论文平方误差、开关频率上升", "FCSState、THD"],
            ["Vref 幅值斜坡", "启动慢、冲击小", "启动快、LC/电流冲击大", "If 峰值"],
        ],
        [
            ["8 个候选代价相同", "使用前向欧拉只预测 Vc 一步，Vi 未进入 Vc 方程", "使用精确离散 Ad/Bd，或做两步欧拉预测"],
            ["输出频率不是 50 Hz", "参考相位/频率或同步开关错误", "先断开电网，固定内部 Vref，检查三个 Sine Wave"],
            ["控制器运算太慢", "在每个 2 μs 电力步执行 FCS", "S-Function SampleTimes 必须为 50 μs"],
        ],
        [
            "空载和 0.5 MW 负载下 Vc 为对称 50 Hz 正弦。",
            "0.5 MW 负载阶跃后电压幅值恢复，If 随负载增加。",
            "FCSState 在 1–8 内，门极仅由候选状态产生。",
        ],
    )

    stage(
        doc, "9", "同步、断路器和 690 V/25 kV 电网接口",
        "在闭合并网断路器前，让离网 MPVC 的参考跟随实际电网电压，减小相角和幅值差造成的冲击电流。",
        [
            ["Three-Phase Breaker", "SPS > Switches and Breakers", "open；T=1.6 s", "并网开关"],
            ["Three-Phase V-I Measurement", "SPS > Sensors", "电网侧 Vg/Ig", "同步和潮流测量"],
            ["Three-Phase Transformer (Two Windings)", "SPS > Transformers", "2.5 MVA；690 V/25 kV；50 Hz", "论文接口"],
            ["Three-Phase Source", "SPS > Sources", "25 kV；47 MVA 短路容量；X/R=10", "工程等值电网"],
            ["Switch", "Simulink > Signal Routing", "Tsync=1.4 s", "Vref 内部/电网切换"],
        ],
        [
            "断路器保持开路，Grid_VI 放在电源侧，以便断开时仍能测到 Vg。",
            "1.4 s 时把 MPVC 的 Vref 从内部正弦切换到 Vgabc。",
            "至少比较 |Vac−Vg|、相角差和频率差；论文在约 5 ms 内完成跟踪。",
            "满足同步条件后 1.6 s 闭合 Grid_Breaker，再把 mode 切到 MPPC。",
            "模板变压器使用 Yg/Yg 简化；若要重现实验相移，必须按作者实际接线替换并补偿相角。",
        ],
        [
            ["AC bus", "Grid_Breaker", "微电网侧", "受控交流母线"],
            ["Grid_Breaker", "Grid_VI→Transformer 690 V", "电网侧", "断开时仍测 Vg"],
            ["Transformer 25 kV", "Utility_25kV", "高压侧", "电网等值"],
            ["Grid_VI Vabc", "Synchronization_Reference_Select", "参考输入", "同步"],
        ],
        [
            ["电网短路容量", "电网更强，电压受微网影响更小、冲击电流更大", "电网更弱，电压支撑效果更明显", "Ig、Vg"],
            ["变压器漏抗", "并网电流变化慢、压降大", "冲击和 FCS 灵敏度高", "Ig、P/Q"],
            ["同步容差", "更容易并网但冲击大", "更平滑但等待时间长", "断路器瞬间 Ig"],
        ],
        [
            ["断路器闭合时大电流", "相角/幅值/相序未对齐", "记录闭合前 5 ms 的 Vac 与 Vg；不要只比较 RMS"],
            ["断开时 Vg=0", "V-I Measurement 放在微网侧", "移到电源侧，顺序为 Transformer→Grid_VI→Breaker→AC bus"],
            ["变压器报饱和/初始磁通问题", "直接在非零相角闭合且无初始化", "使用 powergui 初始化、预磁化或在电压过零附近闭合"],
        ],
        [
            "断开状态可测到 25 kV 等值源折算后的 690 V 三相电压。",
            "闭合前 Vac 与 Vg 的相序、频率、幅值和相角满足设定阈值。",
            "闭合瞬间 Ig 无不可接受峰值。",
        ],
    )

    stage(
        doc, "10", "并网 MPPC 与无功支撑",
        "并网后电网固定交流电压，互联变流器改为调节 P/Q；Vdc 外环给 Pref，电压偏差给 Qref。",
        [
            ["Vdc_Error + PI", "Simulink Math/Discrete", "Kp_dc=2500；Ki_dc=2e5；±2.2 MW", "生成 Pref"],
            ["Grid_Voltage_Pu_Profile", "From Workspace", "1 pu→0.9 pu at 3.5 s", "电压跌落命令"],
            ["Bias + Gain", "Simulink > Math Operations", "−1；m_var=4e6", "生成 Qref"],
            ["MPVC_MPPC_FCS", "Level-2 MATLAB S-Function", "mode=1", "P/Q 预测与状态选择"],
        ],
        [
            "并网闭合后 mode=1，冻结电池 Vdc 外环，启用互联变流器 Pref PI。",
            "确认 If 的正方向为互联变流器→交流母线；用该方向计算 P/Q。",
            "先设 Qref=0，只整定 Vdc/Pref；Vdc 稳定后再加入无功支撑。",
            "3.5 s 时把等值电网电压源替换为 Three-Phase Programmable Voltage Source 或在高压侧施加 0.9 pu；模板的 GridVpu_profile 目前只生成 Qref，不自动改变物理源幅值。",
            "比较论文图 17 的 Vc 从 0.9 pu 提升到约 0.908 pu 的趋势，而不是期望弱小微网完全恢复电网电压。",
        ],
        [
            ["Vdc−Vdc_ref", "Pref PI", "正误差→正 Pref", "母线过压时向电网送有功"],
            ["GridVpu−1", "m_var Gain", "负偏差→负 Qref", "按论文符号注入无功"],
            ["Pref/Qref", "FCS inputs 18/19", "W/var", "MPPC 目标"],
        ],
        [
            ["Kp_dc/Ki_dc", "Vdc 恢复快、P 峰值与振荡增加", "恢复慢、稳态偏差可能大", "Vdc、Pref"],
            ["m_var", "无功响应增强、桥电流和容量占用增加", "支撑较弱", "Q、Ig、Vc"],
            ["Pbase/Qbase", "对应误差权重降低", "对应误差权重提高", "P/Q 跟踪"],
        ],
        [
            ["Vdc 越高 Pref 越负", "误差符号与功率方向不一致", "固定候选状态或小 Pref 试验确定正功率方向，再决定 Vdc−Vref 或相反"],
            ["Qref 有变化但 Vg 不跌落", "模板 profile 只驱动控制参考", "用可编程三相电源施加真实 0.9 pu，再评估电压支撑"],
            ["P/Q 控制互相干扰", "归一化、符号或电网坐标错误", "先 Qref=0；核对 Clarke 与 Q 定义；再逐步增加 m_var"],
        ],
        [
            "并网后 Vdc 由 MPPC 调节，电池不再同时运行 Vdc 外环。",
            "Pref 正负与微网向电网送/受有功方向一致。",
            "0.9 pu 物理电压跌落下 Qref≈−0.4 Mvar。",
        ],
    )

    stage(
        doc, "11", "EMS、Off-MPPT 与切负载逻辑",
        "把局部控制器协调起来，避免 SOC 越界，并在发电不足或过剩时决定放电、充电、限发和切负载。",
        [
            ["EMS_Fig9", "Level-2 MATLAB S-Function", "pvess_ems_sfun；1 ms", "论文图 9"],
            ["PV_Power", "Product", "Vpv·Ipv", "Ppv 估计"],
            ["AC_Load_Power", "Dot Product", "Vabc·Iabc", "瞬时 Pac"],
            ["DC_Load_Power", "Product/Sum", "Vdc·(Idc1+Idc2)", "Pdc"],
            ["Demux", "Simulink > Signal Routing", "battCmd/offMPPT/loadShed", "执行命令"],
        ],
        [
            "定义 Pnet=Ppv−Pac−Pdc；模板暂时忽略 Ploss，论文式 (17) 包含 Ploss。",
            "离网且 Pnet<0：若 SOC>SOCmin 则 battCmd=+1 放电，否则 loadShed=1。",
            "离网且 Pnet>0：若 SOC<SOCmax 则 battCmd=−1 充电，否则 offMPPT=1。",
            "并网：模板按 SOC 中点和功率缺口选择充放电；如要复现电价调度，应替换这一简化策略。",
            "把 loadShed 接到非关键负载断路器前，先做逻辑验证；当前模板记录命令但未自动开断所有负载。",
        ],
        [
            ["Ppv/Pac/Pdc/SOC/gridTied", "EMS_Input_Mux", "5 输入", "状态估计"],
            ["battCmd", "Battery_Controller", "−1/0/+1", "充/停/放"],
            ["offMPPT", "IncCond_MPPT", "0/1", "限发"],
            ["loadShed", "非关键负载 breaker", "建议后续连接", "切负载执行"],
        ],
        [
            ["SOC_min", "更早禁止放电、可靠性高", "可用能量多但深放风险高", "loadShed、SOC"],
            ["SOC_max", "允许更多充电容量", "更早限发", "offMPPT"],
            ["功率滤波", "EMS 反应慢但不抖动", "快速但可能因 2ω/开关纹波频繁切换", "battCmd"],
        ],
        [
            ["EMS 命令频繁跳变", "使用瞬时 Pac，未做低通/平均", "对 Ppv/Pac/Pdc 做一周波或低通平均并增加滞环"],
            ["满 SOC 仍充电", "SOC 单位 0–1 与 0–100 混用", "模板使用百分数 0–100，限值 10/90"],
            ["loadShed 有信号但负载没断", "命令未接到物理断路器", "添加 External 三相/单相 Breaker，并验证 0/1 开合语义"],
        ],
        [
            "四个主要分支逻辑与论文图 9 一致。",
            "SOC 到达上下限时不会继续同方向充放电。",
            "EMS 功率输入经过适当平均，不追逐开关纹波。",
        ],
    )

    stage(
        doc, "12", "论文表 2 场景与运行脚本",
        "按论文事件顺序验证模式切换、功率平衡和无功支撑；开关级 4 s 计算量很大，必须分段。",
        [
            ["run_paper_scenario.m", "项目 scripts", "SimulationInput；StopTime 参数", "场景入口"],
            ["G_profile", "timeseries", "400→800→400 W/m²", "PV 变化"],
            ["Rdc profiles", "timeseries", "0.6 MW / 1 MW 事件", "DC 负载"],
            ["Grid_Connected / Synchronization", "Step", "1.6 s / 1.4 s", "模式切换"],
            ["GridVpu_profile", "timeseries", "3.5 s: 0.9 pu", "Qref 生成"],
        ],
        [
            "先运行 out=run_paper_scenario(0.2)，只验证启动和基础负载。",
            "再运行 1.0 s，加入 AC 负载阶跃；确认离网 MPVC 和电池双环。",
            "运行 2.0 s，覆盖 PV 上升、0.6 MW DC 负载、同步与并网。",
            "运行 3.2 s，加入 1 MW DC 负载、PV 下降和负载切除。",
            "最后运行 4.0 s，并使用可编程物理电源施加 3.5 s 的 10% 跌落。",
            "每段都保存 MAT 结果，不要依赖 Scope 肉眼截图作为唯一证据。",
        ],
        [],
        [
            ["StopTime", "覆盖更多事件、运行和文件显著增加", "便于定位，但看不到后续模式", "结果文件大小"],
            ["Ts_power", "运行更快但不再可靠解析 20 kHz 状态", "运行慢、波形更可信", "总步数"],
            ["日志数量", "结果文件大、诊断全面", "运行更快但排错证据少", "MAT 文件"],
        ],
        [
            ["4 s 运行太慢", "2 μs 步长意味着约 200 万电力步且含非线性 PV/Battery", "先分段；完成开关级校准后，可建立平均模型跑系统级长时场景"],
            ["1.4 s 同步时参考突变", "内部参考与 Vg 初始相角不同", "让内部参考锁相后再切，或加入无扰参考渐变"],
            ["3.5 s 论文波形复现不了", "物理电源没有实际电压跌落", "替换为可编程三相源；GridVpu_profile 不能代替物理扰动"],
        ],
        [
            "每个事件时刻与论文表 2 一致。",
            "模式切换时无 NaN、直通或不可接受冲击。",
            "结果文件包含 Vdc、SOC、Ibat、Vc/If、Vg/Ig、P/Q 参考和 FCS 状态。",
        ],
    )
    add_picture(doc, ASSET_DIR / "timeline.png", 6.35, "图 4  论文表 2 事件序列与推荐分段运行")

    heading(doc, "13 分阶段接入顺序与实操命令", 1)
    callout(doc, "不要跳步",
            "完整连接态的 5 ms 试跑能够完成且信号有限，但在没有预充和稳态初始化时，"
            "Vdc 会从 1200 V 明显下跌、Ibat 会出现约 0.9 kA 启动响应。"
            "这不是论文稳态波形，只说明拓扑可编译；定量复现必须按下列顺序接入。",
            fill=PALE_RED)
    table(doc, ["步骤", "PV_connect", "Battery_connect", "Inverter_connect/enable", "目标"], [
        ["A 母线", "0", "0", "0/0", "Cdc 1200 V 保持；检查测量与负载"],
        ["B PV 输入", "输入侧 1，DC 侧 0", "0", "0/0", "验证 PV、Lpv、固定占空比方向"],
        ["C PV→母线", "1", "0", "0/0", "母线接耗能支路；验证 Boost 功率"],
        ["D 电池", "按需", "1", "0/0", "固定 Db→电流内环→Vdc 外环"],
        ["E 逆变器空载", "按需", "1", "1/1", "参考幅值斜坡；建立 50 Hz"],
        ["F AC 负载", "按需", "1", "1/1", "0.5 MW→阶跃"],
        ["G 同步并网", "1", "SOC 逻辑", "1/1", "先同步再闭合 breaker，再切 MPPC"],
    ], [0.85, 1.05, 1.20, 1.60, 1.80], font_size=7.4)
    code(doc, "projectRoot = 'D:/PV_MPPT/PV_ESS_MPC_Paper_Reproduction';\n"
              "addpath(fullfile(projectRoot,'scripts'));\n"
              "validation = run_short_validation();\n"
              "% 分段论文场景\n"
              "out = run_paper_scenario(0.2);\n"
              "% out = run_paper_scenario(1.0);\n"
              "% out = run_paper_scenario(2.0);\n"
              "% out = run_paper_scenario(4.0);")
    add_picture(doc, MODEL_IMAGE, 6.45, "图 5  生成的完整开关级模板顶层视图")
    add_picture(doc, ASSET_DIR / "model_pv_dc.png", 6.35, "图 6  模板中的 PV、Boost 和直流母线区域")
    add_picture(doc, ASSET_DIR / "model_battery.png", 6.35, "图 7  模板中的电池半桥与门极逻辑区域")
    add_picture(doc, ASSET_DIR / "model_inverter_control.png", 6.35, "图 8  模板中的 FCS、三相桥、LC 和 EMS 区域")
    add_picture(doc, ASSET_DIR / "model_grid.png", 6.35, "图 9  模板中的交流负载、并网断路器和变压器区域")

    heading(doc, "14 如何修改脚本：改哪里、为什么、改后验证什么", 1)
    table(doc, ["脚本", "修改场景", "应改字段/函数", "修改后必须重做"], [
        ["pvess_init.m", "换电压等级、组件串并联、LC、电池、采样时间", "对应参数字段；保持 Paper/Assumption 分区", "静态电气检查、占空比、预测矩阵"],
        ["build_pvess_template.m", "增加/删除模块、改变接线和保存态", "add_block/add_line/ec 调用", "重建会覆盖 .slx；重新编译与截图"],
        ["pvess_mppt_sfun.m", "改变 MPPT 判据、步长、限发策略", "update()", "P–V 扫描、辐照度阶跃"],
        ["pvess_battery_sfun.m", "改变 PI、SOC 曲线、抗饱和", "outputs()/update()", "固定 Db、内环、外环、模式切换"],
        ["pvess_fcs_sfun.m", "改变预测、代价、门极映射、延时补偿", "outputs()", "8 状态单元测试、桥臂互补、P/Q 符号"],
        ["pvess_ems_sfun.m", "增加损耗、电价、滞环、切负载", "outputs()", "图 9 全分支真值表"],
        ["run_paper_scenario.m", "改变事件时刻/工况", "SimulationInput 和 cfg", "事件表、结果命名、StopTime"],
    ], [1.25, 1.65, 2.05, 1.55], font_size=7.4)
    heading(doc, "14.1 沿用旧脚本时的规则", 2)
    bullet(doc, "旧 PV 查表脚本只适合静态模型或快速预测器，不能替换本模板中的 PV Array 功率电路。")
    bullet(doc, "如果继续使用 pvbatt_generate_pv_lookup.m，必须把它改成有明确输入的函数，例如 "
                "lookup=pvbatt_generate_pv_lookup(arrayVoltage,temperature,irradiance,p)，并由调用脚本传参。")
    bullet(doc, "旧六状态平均模型可以作为控制器内部预测器或长时平均模型；不要把它当作本次功率级交付。")
    bullet(doc, "任何脚本改动都要在手册的参数来源表中记录：论文、数据手册、测量、辨识或工程假设。")

    heading(doc, "15 验证计划：从子系统到论文图 10–19", 1)
    table(doc, ["测试 ID", "对象/工况", "主要输出", "通过标准", "论文对应"], [
        ["T0", "隔离态母线 10 ms", "Vdc", "1199.9–1200.1 V；无 Inf/NaN", "结构测试"],
        ["T1", "PV 静态 G=250…1000", "I–V/P–V", "峰值/电压趋势与图 3 一致", "图 3"],
        ["T2", "PV 辐照度下降再上升", "Vpv/Ppv/Dpv", "Ppv 跟随辐照度，MPPT 不发散", "图 4"],
        ["T3", "电池固定 Db 和电流内环", "Ibat/Db", "方向正确、纹波有界、无直通", "图 5/7"],
        ["T4", "离网 0.5→1 MW AC", "Vc/If/Vdc/Ibat/SOC", "Vac 稳定；电池由充转放", "图 10/11"],
        ["T5", "PV 400→700 W/m² + 0.3 MW DC", "Vdc/Ibat/SOC", "母线恢复；充电速率变化", "图 12"],
        ["T6", "同步与并网", "Vac/Vg/If", "约 5 ms 内对齐；闭合无大冲击", "图 13"],
        ["T7", "并网 AC/DC 负载阶跃", "Vc/Ig/If/Vdc/Ibat", "潮流方向与论文一致", "图 14/15"],
        ["T8", "表 2 完整事件", "Ppv/Pbat/Pg/Pdc/Ig/Vdc/Vc", "事件响应与模式正确", "图 16"],
        ["T9", "物理电网 0.9 pu", "Q/Vc", "Q≈−0.4 Mvar；Vc 有提升趋势", "图 17"],
        ["T10", "实测辐照度", "Ppv/Pbat/Ibat/SOC/Vdc/Vc", "功率缺口由电池平滑", "图 18/19"],
    ], [0.65, 1.45, 1.35, 2.05, 1.00], font_size=7.0)
    heading(doc, "15.1 定量容差建议", 2)
    table(doc, ["指标", "初始模板阶段", "论文定量复现阶段", "计算方法"], [
        ["Vdc 稳态误差", "<±2%", "<±1%", "事件后稳定窗平均"],
        ["Vac RMS 误差", "<±5%", "<±2%", "一周波滑动 RMS"],
        ["P/Q 稳态误差", "<±10%", "<±5%", "一周波平均"],
        ["同步冲击电流", "低于额定 1.5 pu", "与论文波形同量级", "闭合前后峰值"],
        ["SOC 斜率", "符号正确", "与 Ibat/(3600Q) 误差<2%", "数值积分对比"],
        ["功率平衡残差", "<10% 额定", "<3% 额定", "Ppv−Pac−Pdc−Pbat−Pg−Ploss"],
    ], [1.35, 1.55, 1.65, 1.95], font_size=7.5)

    heading(doc, "16 综合排错树", 1)
    table(doc, ["症状", "先看什么", "高概率原因", "退回哪个阶段"], [
        ["Vdc 瞬间到几十/几百 kV", "Ipv、Ibat、隔离器、初始电流", "断开电感有非零初流；二极管仍连接；门极直通", "阶段 1/3/5"],
        ["PV MPPT 持续走向边界", "Vpv 与 Dpv 的两点方向", "占空比符号错；差分被纹波污染", "阶段 2/3"],
        ["电池控制越调越大", "Db→Ibat 方向；Ibat 传感器", "内环正反馈", "阶段 5"],
        ["离网 Vac 无法建立", "门极配对、Vref、Viαβ、LC 接线", "桥直通/相序错/候选电压错误", "阶段 7/8"],
        ["并网瞬间大冲击", "闭合前 Vac/Vg 相角和幅值", "同步参考或变压器相移未处理", "阶段 9"],
        ["并网 Vdc 与电池环互相振荡", "两个 Vdc 外环是否同时工作", "控制职责未切换", "阶段 6/10"],
        ["Qref 正确但电压不变", "物理电源是否真的跌落", "只改变了控制 profile；强电网容量过大", "阶段 10"],
        ["EMS 抖动", "Pac 是否瞬时值", "未做周期平均/滞环", "阶段 11"],
        ["模型慢", "步数、日志、非线性 PV", "直接跑 4 s 开关级", "阶段 12"],
    ], [1.55, 1.40, 2.10, 1.45], font_size=7.3, header_fill=PALE_RED)
    heading(doc, "16.1 排错的固定顺序", 2)
    for item in [
        "先断开所有调试隔离器，只验证 Cdc、测量和负载。",
        "每次只闭合一个源/变换器，固定门极或固定占空比，不先上闭环。",
        "先验证符号与能量方向，再整定 PI/MPC。",
        "出现高能量非物理量时立即退回并清零断开电感初值，不用饱和块掩盖。",
        "只有子系统开环通过后，才运行集成开环；集成开环通过后，才运行闭环和模式切换。",
    ]:
        number(doc, item)

    heading(doc, "附录 A：完整模块清单与关键参数", 1)
    table(doc, ["模型块", "库路径/类型", "模板参数", "用途/修改影响"], [
        ["powergui", "spspowerguiLib/powergui", "Discrete；2e-6；TBE", "SPS 网络；步长增大降低开关精度"],
        ["PV_Array", "spsPVArrayLib/PV Array", "SunPower 305；10s×656p", "PV 非线性源"],
        ["Cpv", "Parallel RLC Branch", "C=5 mF；IC=547 V", "PV 端滤波；越大 MPPT 越慢"],
        ["Lpv", "Series RLC Branch", "0.2 mH；2 mΩ；IC=0", "Boost 储能；断开时不可有初流"],
        ["PV_Boost", "spsBoostConverterLib/Boost Converter", "Switching devices", "真实升压开关"],
        ["PV_PWM", "PWM Generator (DC-DC)", "5 kHz；2 μs", "Dpv→门极"],
        ["Cdc", "Parallel RLC Branch", "6 mF；IC=1200 V", "论文母线电容"],
        ["DC_Load_0p6MW", "Variable Resistor", "2.4 Ω at 1200 V", "事件 3"],
        ["DC_Load_1MW", "Variable Resistor", "1.44 Ω at 1200 V", "事件 6/8"],
        ["Battery_Leg", "Two-Quadrant DC/DC Converter", "Switching devices", "双向半桥"],
        ["Lbat", "Series RLC Branch", "0.5 mH；2 mΩ", "电池电流动态"],
        ["Battery", "spsBatteryLib/Battery", "Li-ion 300 V 1300 Ah", "论文储能等效"],
        ["Interlinking_Converter", "spsUniversalBridgeLib/Universal Bridge", "3-arm IGBT/Diodes", "AC/DC 主桥"],
        ["Filter_Lf_Rf", "Three-Phase Series RLC Branch", "0.6 mH；0.019 Ω", "论文滤波电感/电阻"],
        ["Filter_Cf", "Three-Phase Parallel RLC Load", "Qc=ωCfVLL²", "等效 1338 μF"],
        ["AC_Load_Base", "Three-Phase Parallel RLC Load", "0.5 MW", "基础交流负载"],
        ["AC_Load_Step", "Three-Phase Parallel RLC Load", "0.5 MW", "事件 1"],
        ["Grid_Breaker", "Three-Phase Breaker", "open→1.6 s close", "并网"],
        ["Grid_Transformer", "Three-Phase Transformer Two Windings", "2.5 MVA；690/25 kV", "论文电网接口"],
        ["Utility_25kV", "Three-Phase Source", "25 kV；47 MVA；X/R=10", "工程等值电网"],
        ["IncCond_MPPT", "Level-2 MATLAB S-Function", "1 ms", "MPPT/Off-MPPT"],
        ["Battery_Controller", "Level-2 MATLAB S-Function", "100 μs", "双环/SOC 定向"],
        ["MPVC_MPPC_FCS", "Level-2 MATLAB S-Function", "50 μs", "有限控制集"],
        ["EMS_Fig9", "Level-2 MATLAB S-Function", "1 ms", "论文图 9"],
    ], [1.55, 2.05, 1.40, 1.50], font_size=6.9)

    heading(doc, "附录 B：关键接线总表", 1)
    table(doc, ["网络", "从", "到", "注意"], [
        ["PV 正端", "PV_Array +", "Ipv→Lpv→PV_Isolator→PV_Boost input", "Ipv 正向离开 PV"],
        ["PV 负端", "PV_Array −", "PV_Boost − / DC_Ground", "公共负端"],
        ["Boost 输出", "PV_Boost +", "PV_DC_Isolator→Cdc +", "隔离器默认断开"],
        ["母线正端", "Cdc +", "DC loads、Battery_DC_Isolator、Inverter_DC_Isolator", "公共节点"],
        ["电池", "Battery_Leg midpoint", "Battery_Isolator→Lbat→Ibat→Battery +", "Ibat>0 放电"],
        ["逆变器", "Universal Bridge A/B/C", "Filter L→Filter_VI→AC bus", "相序一致"],
        ["AC 分支", "AC bus", "Cf、Load_VI、Grid_Breaker", "Vc 测量节点"],
        ["负载", "Load_VI", "Base Load；Step Breaker→Step Load", "IL 为负载电流"],
        ["电网", "AC bus", "Grid Breaker→Grid_VI→Transformer→Source", "Grid_VI 位于电源侧"],
    ], [1.05, 1.60, 2.55, 1.30], font_size=7.4)

    heading(doc, "附录 C：门极状态表与符号检查", 1)
    table(doc, ["状态号", "Sa Sb Sc", "桥门极向量", "Viα / Vdc", "Viβ / Vdc"], [
        ["1", "000", "[0 1 0 1 0 1]", "0", "0"],
        ["2", "100", "[1 0 0 1 0 1]", "2/3", "0"],
        ["3", "110", "[1 0 1 0 0 1]", "1/3", "√3/3"],
        ["4", "010", "[0 1 1 0 0 1]", "−1/3", "√3/3"],
        ["5", "011", "[0 1 1 0 1 0]", "−2/3", "0"],
        ["6", "001", "[0 1 0 1 1 0]", "−1/3", "−√3/3"],
        ["7", "101", "[1 0 0 1 1 0]", "1/3", "−√3/3"],
        ["8", "111", "[1 0 1 0 1 0]", "0", "0"],
    ], [0.65, 0.90, 2.15, 1.40, 1.40], font_size=7.6)
    callout(doc, "门极顺序的实证检查",
            "R2025a 的 PWM Generator (2-Level) 输出中，互补对为 (1,2)、(3,4)、(5,6)。"
            "因此模板使用 [Sa,¬Sa,Sb,¬Sb,Sc,¬Sc]。换 MATLAB 版本或换桥模块时必须重新验证，"
            "不要仅凭器件编号记忆接线。", fill=PALE_YELLOW)

    heading(doc, "附录 D：已完成验证与已知限制", 1)
    table(doc, ["项目", "结果", "含义"], [
        ["模板保存态直接运行 10 ms", "Vdc=1199.93–1200.00 V；有限信号", "模型工作区、SPS 电路、求解器和日志可运行"],
        ["完整功率拓扑闭合 5 ms", "仿真完成；Vdc 最低约 488 V；Ibat 峰值约 933 A", "结构可编译，但未预充/整定，不能当论文结果"],
        ["MCP 结构检查", "本机 MATLAB 会话连接失败，改用实际编译仿真", "模型无编译错误；仍建议用户打开 Diagnostic Viewer 检查警告"],
        ["完整 4 s 论文场景", "未在本次交付中强行跑完", "需分段整定后执行，避免把启动问题与论文事件混在一起"],
    ], [1.85, 2.30, 2.35], font_size=7.6)
    heading(doc, "已知限制", 2)
    for item in [
        "论文没有公开全部器件与控制参数，因此当前 PI、Lpv/Cpv/Lbat、电网阻抗不是作者原值。",
        "模板的 GridVpu_profile 生成 Qref，但物理 25 kV Source 仍为固定幅值；图 17 复现必须换可编程电压源。",
        "三相变压器采用 Yg/Yg 工程简化，作者实际接线和相移未公开。",
        "Battery 库块使用通用锂电模型，完整 OCV/SOC 曲线和温度效应未由论文标定。",
        "EMS 对功率损耗 Ploss 暂取 0，且 loadShed 命令未自动接到全部物理负载。",
        "长时系统级研究建议另建平均模型；开关级模型用于控制器和功率电路的高频验证。",
    ]:
        bullet(doc, item)

    heading(doc, "附录 E：参考文献与可追溯来源", 1)
    paragraph(doc, "Hu, J., Xu, Y., Cheng, K. W., & Guerrero, J. M. (2018). "
                   "A model predictive control strategy of PV-Battery microgrid under variable power generations and load conditions. "
                   "Applied Energy, 221, 195–203. DOI: 10.1016/j.apenergy.2018.03.085.")
    paragraph(doc, "本手册中的论文参数、模式和事件均从上述论文的图 1–19、表 1–2 和式 (1)–(17) 提取；"
                   "工程补充值均在 pvess_init.m 中标注。")
    paragraph(doc, "本项目生成文件：" + str(MODEL) + "；脚本目录：" + str(SCRIPTS) + "。")

    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_manual()
