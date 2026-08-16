from pathlib import Path
from datetime import date
import sys

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\PV_MPPT")
DOCS = ROOT / "docs"
QA = DOCS / "_qa" / "full_paper_manual_switching_v2"
OUT = DOCS / "PV_Battery_MPC_Full_Paper_Reproduction_Manual_Switching_Topology_Detailed_v2.docx"
OUT_RAW = OUT
PROJECT = ROOT / "PV_Battery_MPC_Project"
SCRIPTS = PROJECT / "scripts"
ASSETS = PROJECT / "docs" / "report_assets"
PROTOTYPE = DOCS / "_qa" / "switching_prototype" / "prototype_overview.png"

BLUE = "2E74B5"
DARK_BLUE = "17365D"
MID_BLUE = "5B9BD5"
PALE_BLUE = "E8EEF5"
LIGHT_BLUE = "DDEBF7"
PALE_GREEN = "E2F0D9"
PALE_YELLOW = "FFF2CC"
PALE_RED = "FCE4D6"
LIGHT_GRAY = "F2F2F2"
GRAY = "666666"
WHITE = "FFFFFF"
BLACK = "222222"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, east_asia="Microsoft YaHei", ascii_font="Calibri"):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)


def set_keep(paragraph, keep_next=False, keep_lines=True):
    ppr = paragraph._p.get_or_add_pPr()
    if keep_next:
        el = OxmlElement("w:keepNext")
        el.set(qn("w:val"), "1")
        ppr.append(el)
    if keep_lines:
        el = OxmlElement("w:keepLines")
        el.set(qn("w:val"), "1")
        ppr.append(el)


def set_repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trpr.append(el)


def prevent_row_split(row):
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tcmar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="B7C9DC", size="4"):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.first_child_found_in("w:tblW")
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    total = sum(Inches(w).twips for w in widths)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    ind = tblpr.first_child_found_in("w:tblInd")
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tblpr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    layout = tblpr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = list(table._tbl.tblGrid)
    for idx, width in enumerate(widths):
        twips = Inches(width).twips
        if idx < len(grid):
            grid[idx].set(qn("w:w"), str(twips))
        for row in table.rows:
            cell = row.cells[idx]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.first_child_found_in("w:tcW")
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(twips))
            tcw.set(qn("w:type"), "dxa")


def add_field(paragraph, code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, end])


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True
        if name == "Heading 1":
            st.paragraph_format.page_break_before = True

    for name in ["List Bullet", "List Number", "List Bullet 2", "List Number 2"]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.paragraph_format.left_indent = Inches(0.375 if "2" not in name else 0.65)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.10)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(5)
    code.paragraph_format.line_spacing = 1.05

    # Editorial-cover header: short title on the left, stage family on the right.
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("PV–Battery MPC 论文复现")
    set_run_font(r)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = rgb(DARK_BLUE)
    r = hp.add_run("    开关级主电路 · 平均值预测模型 · 控制电路")
    set_run_font(r)
    r.font.size = Pt(8.2)
    r.font.color.rgb = rgb(GRAY)
    ppr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), MID_BLUE)
    pbdr.append(bottom)
    ppr.append(pbdr)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("第 ")
    set_run_font(r)
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(GRAY)
    add_field(fp, "PAGE")
    r = fp.add_run(" 页")
    set_run_font(r)
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(GRAY)


def paragraph(doc, text="", style=None, bold_prefix=None, align=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        set_run_font(a)
        a.bold = True
        b = p.add_run(text[len(bold_prefix):])
        set_run_font(b)
    else:
        r = p.add_run(text)
        set_run_font(r)
    if align is not None:
        p.alignment = align
    return p


def bullet(doc, text, level=0):
    return paragraph(doc, text, "List Bullet" if level == 0 else "List Bullet 2")


def number(doc, text, level=0):
    return paragraph(doc, text, "List Number" if level == 0 else "List Number 2")


def code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    r = p.add_run(text)
    set_run_font(r, ascii_font="Consolas")
    r.font.size = Pt(8.5)
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    ppr.append(shd)
    return p


def heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for r in p.runs:
        set_run_font(r)
    return p


def callout(doc, title, body, fill=PALE_YELLOW, border=MID_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 100, 140, 100, 140)
    set_cell_border(cell, border, "8")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r)
    r.bold = True
    r.font.color.rgb = rgb(DARK_BLUE)
    r = p.add_run("\n" + body)
    set_run_font(r)
    p.paragraph_format.space_after = Pt(0)
    prevent_row_split(table.rows[0])
    return table


def table(doc, headers, rows, widths, font_size=8.3, header_fill=PALE_BLUE):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    for idx, text in enumerate(headers):
        cell = t.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        set_run_font(r)
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = rgb(DARK_BLUE)
    set_repeat_header(t.rows[0])
    prevent_row_split(t.rows[0])
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ridx % 2:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            set_run_font(r)
            r.font.size = Pt(font_size)
        prevent_row_split(t.rows[-1])
    paragraph(doc, "", style=None)
    return t


def caption(doc, text):
    p = paragraph(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = rgb(GRAY)
    return p


def add_picture(doc, path, width=6.3, caption_text=None):
    if not Path(path).exists():
        callout(doc, "图像缺失", str(path), fill=PALE_RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    alt = caption_text if caption_text else Path(path).stem
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)
    if caption_text:
        caption(doc, caption_text)


def font(size=30, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_center(draw, rect, text, fnt, fill=BLACK):
    x1, y1, x2, y2 = rect
    box = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=5, align="center")
    w, h = box[2] - box[0], box[3] - box[1]
    draw.multiline_text(((x1 + x2 - w) / 2, (y1 + y2 - h) / 2), text, font=fnt,
                        fill="#" + fill, spacing=5, align="center")


def draw_box(draw, rect, title, subtitle="", fill="FFFFFF", outline=BLUE):
    draw.rounded_rectangle(rect, radius=16, fill="#" + fill, outline="#" + outline, width=4)
    text = title if not subtitle else title + "\n" + subtitle
    draw_center(draw, rect, text, font(28, True) if not subtitle else font(23, True))


def arrow(draw, start, end, color=BLUE, width=5):
    draw.line([start, end], fill="#" + color, width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    s = 14
    p1 = (x2 - s * ux + 0.55 * s * px, y2 - s * uy + 0.55 * s * py)
    p2 = (x2 - s * ux - 0.55 * s * px, y2 - s * uy - 0.55 * s * py)
    draw.polygon([end, p1, p2], fill="#" + color)


def create_diagrams():
    QA.mkdir(parents=True, exist_ok=True)

    # Three-layer architecture.
    path = QA / "three_layer_architecture.png"
    im = Image.new("RGB", (1800, 1050), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 35), "新版复现架构：物理对象与预测模型分层", font=font(42, True), fill="#" + DARK_BLUE)
    layers = [
        (90, 140, 1710, 390, PALE_GREEN, "A  开关级主电路（真实被控对象）",
         ["PV Array", "Boost + PWM", "DC Bus + Load", "双向半桥 + Battery"]),
        (90, 445, 1710, 680, LIGHT_BLUE, "B  测量与控制电路",
         ["传感器/滤波", "P&O / PMS", "ARIMA", "nlMPC / PI", "占空比限幅与门极"]),
        (90, 735, 1710, 970, PALE_YELLOW, "C  控制器内部平均值预测模型",
         ["PV 三维查表", "六状态方程", "RK4 预测", "代价函数/约束"]),
    ]
    for x1, y1, x2, y2, fillc, title, boxes in layers:
        d.rounded_rectangle((x1, y1, x2, y2), radius=20, fill="#" + fillc, outline="#" + MID_BLUE, width=4)
        d.text((x1 + 28, y1 + 20), title, font=font(29, True), fill="#" + DARK_BLUE)
        gap = (x2 - x1 - 80) / len(boxes)
        for i, label in enumerate(boxes):
            bx1 = x1 + 35 + i * gap
            bx2 = bx1 + gap - 25
            draw_box(d, (bx1, y1 + 85, bx2, y2 - 30), label, fill="FFFFFF", outline=BLUE)
            if i < len(boxes) - 1:
                arrow(d, (bx2 + 3, (y1 + y2) / 2 + 30), (bx2 + 20, (y1 + y2) / 2 + 30))
    arrow(d, (900, 680), (900, 735), color=DARK_BLUE, width=7)
    arrow(d, (850, 445), (850, 390), color=DARK_BLUE, width=7)
    im.save(path)

    # Sample-time ladder.
    path = QA / "sample_time_ladder.png"
    im = Image.new("RGB", (1800, 780), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 35), "三层时间尺度必须同时正确", font=font(42, True), fill="#" + DARK_BLUE)
    rows = [
        (150, "电力电子求解步长", "Ts_power = 1 µs", "解析 10 kHz 开关沿；改大后纹波和峰值失真", PALE_RED),
        (330, "PWM 周期", "Tsw = 100 µs", "Fsw = 10 kHz；每周期约 100 个求解点", PALE_YELLOW),
        (510, "控制器执行周期", "Ts_ctrl = 10 ms", "每次控制更新跨 100 个 PWM 周期", PALE_GREEN),
    ]
    for y, name, val, note, fillc in rows:
        draw_box(d, (90, y, 520, y + 115), name, val, fill=fillc, outline=BLUE)
        arrow(d, (535, y + 57), (750, y + 57), color=BLUE, width=7)
        d.text((790, y + 17), note, font=font(27), fill="#" + BLACK)
    d.text((100, 690), "Rate Transition / ZOH 只负责跨速率传递；它不能代替 PWM，也不能把 10 ms 当成开关周期。",
           font=font(27, True), fill="#" + DARK_BLUE)
    im.save(path)

    # Complementary gate logic.
    path = QA / "battery_gate_logic.png"
    im = Image.new("RGB", (1800, 850), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 35), "电池半桥门极：同一载波、互补、上升沿死区", font=font(40, True), fill="#" + DARK_BLUE)
    draw_box(d, (110, 155, 380, 275), "db", "Saturation [0.02,0.95]", fill=PALE_BLUE)
    draw_box(d, (470, 155, 790, 275), "PWM Generator", "10 kHz / Ts=1 µs", fill=PALE_GREEN)
    draw_box(d, (900, 95, 1130, 205), "NOT", "上管原始门极", fill=PALE_YELLOW)
    draw_box(d, (900, 255, 1130, 365), "直接支路", "下管原始门极", fill=PALE_YELLOW)
    draw_box(d, (1230, 80, 1570, 220), "Unit Delay + AND", "延迟上升沿 1 µs", fill=LIGHT_BLUE)
    draw_box(d, (1230, 245, 1570, 385), "Unit Delay + AND", "延迟上升沿 1 µs", fill=LIGHT_BLUE)
    arrow(d, (380, 215), (470, 215)); arrow(d, (790, 215), (900, 150)); arrow(d, (790, 215), (900, 310))
    arrow(d, (1130, 150), (1230, 150)); arrow(d, (1130, 310), (1230, 310))
    d.text((125, 485), "严禁：用两个独立 PWM 模块分别输入 db 与 1-db。它们同相起始，可能同时为 1，造成直通。",
           font=font(28, True), fill="#B22222")
    d.text((125, 560), "Mux 顺序：[上管 g_upper；下管 g_lower] → Data Type Conversion(double) → Two-Quadrant DC/DC Converter 的 g 口。",
           font=font(27), fill="#" + BLACK)
    d.text((125, 635), "平均值模型中 (1-db)Vdc 对应上管平均导通比例；下管平均导通比例为 db。",
           font=font(27), fill="#" + DARK_BLUE)
    im.save(path)

    # Stage roadmap.
    path = QA / "stage_roadmap.png"
    im = Image.new("RGB", (1800, 1100), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 35), "从主电路到闭环：必须按关卡推进", font=font(42, True), fill="#" + DARK_BLUE)
    stages = [
        ("0–2", "接口、参数、PV 源"), ("3–5", "PV Boost / Battery / 半桥"),
        ("6–8", "母线、测量、开环换能"), ("9–11", "预测模型、MPPT、PMS/ARIMA"),
        ("12–14", "nlMPC、PWM 集成、PI 对照"), ("15–17", "场景、敏感性、论文解释"),
    ]
    y = 150
    for i, (num, label) in enumerate(stages):
        x = 120 if i % 2 == 0 else 970
        if i % 2 == 0 and i > 0:
            y += 275
        draw_box(d, (x, y, x + 700, y + 170), f"阶段 {num}", label,
                 fill=PALE_GREEN if i < 3 else LIGHT_BLUE, outline=BLUE)
        if i < len(stages) - 1:
            if i % 2 == 0:
                arrow(d, (x + 700, y + 85), (970, y + 85), width=7)
            else:
                arrow(d, (x + 350, y + 170), (x + 350, y + 255), width=7)
    d.text((120, 970), "关卡规则：子系统开环通过 → 整机开环功率守恒 → 再接控制器；不能跳过。",
           font=font(30, True), fill="#" + DARK_BLUE)
    im.save(path)


def stage_header(doc, number_text, title, purpose, deliverable):
    heading(doc, f"阶段 {number_text}：{title}", 1, page_break=True)
    callout(doc, "本阶段意义", purpose, fill=LIGHT_BLUE)
    table(doc, ["开始前必须已有", "完成后应得到"], [["上一阶段检查点已通过；参数脚本可运行；所有信号单位和正方向已确认。", deliverable]], [3.25, 3.25], font_size=9)


def stage_checkpoint(doc, checks):
    heading(doc, "阶段检查点（全部满足才进入下一阶段）", 2)
    for item in checks:
        bullet(doc, "□ " + item)


def stage_faults(doc, rows):
    heading(doc, "本阶段常见故障：现象 → 原因 → 修正", 2)
    table(doc, ["现象", "最可能原因", "检查与修正"], rows, [1.55, 2.05, 2.90], font_size=8.1, header_fill=PALE_RED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    r = p.add_run("PV–Battery MPC 论文完整复现")
    set_run_font(r)
    r.bold = True
    r.font.size = Pt(27)
    r.font.color.rgb = rgb(DARK_BLUE)
    p = doc.add_paragraph()
    r = p.add_run("开关级主电路 + 平均值预测模型 + 控制电路\n详细搭建与排错手册 V2")
    set_run_font(r)
    r.bold = True
    r.font.size = Pt(19)
    r.font.color.rgb = rgb(BLUE)
    p.paragraph_format.space_after = Pt(18)

    callout(doc, "本版纠正的核心问题",
            "原模型的 PV、变换器、电池和母线主要由六状态平均值方程实现，只能作为控制算法验证对象。"
            "本版把论文图 2 的电气拓扑重新放回被控对象：PV Array、开关型 Boost、双向半桥、L/C、Battery、"
            "直流母线和可变电阻负载都用 Specialized Power Systems 的电气端口连接；原六状态模型只保留在 MPC 内部做预测。",
            fill=PALE_GREEN)
    meta = paragraph(doc, "适用环境：MATLAB / Simulink R2025a，Simscape Electrical，Model Predictive Control Toolbox。\n"
                     "工作目录：D:\\PV_MPPT；基准项目：D:\\PV_MPPT\\PV_Battery_MPC_Project。\n"
                     "编写日期：" + date.today().isoformat() + "。第一次从阶段 0 顺序操作；返工时按各阶段检查点定位。")
    meta.paragraph_format.space_after = Pt(0)
    for r in meta.runs:
        r.font.size = Pt(9.5)
        r.font.color.rgb = rgb(GRAY)
    doc.add_page_break()


def add_native_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run("目录")
    set_run_font(r)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = rgb(DARK_BLUE)
    p.paragraph_format.space_after = Pt(12)
    toc = doc.add_paragraph()
    add_field(toc, r'TOC \o "1-3" \h \z \u')
    doc.add_page_break()


def build_manual():
    create_diagrams()
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_native_toc(doc)

    heading(doc, "使用说明与总路线", 1)
    paragraph(doc, "这不是一份只讲公式的说明，而是一份可照着点击和接线的实施手册。每一阶段都回答六个问题：为什么做、用什么模块、模块怎么设、线怎么接、参数变化有什么影响、怎样判断做对了。")
    add_picture(doc, QA / "three_layer_architecture.png", 6.35, "图 1  新版三层架构：开关级主电路是真实对象，平均值方程只服务于预测")
    callout(doc, "必须坚持的分层原则",
            "主电路输出真实开关波形；控制器读取经过测量和滤波的状态；nlMPC 的状态转移仍使用平均值模型。"
            "不要把 PWM 门极接进平均值方程，也不要把占空比直接接到 IGBT 门极。",
            fill=PALE_YELLOW)
    add_picture(doc, QA / "stage_roadmap.png", 6.3, "图 2  全部阶段与关卡关系")

    heading(doc, "你最终应拥有的模型层次", 2)
    table(doc, ["模型/子系统", "类型", "用途", "是否含 PWM 纹波"], [
        ["PV_Battery_MPC_Switching.slx", "顶层闭环模型", "论文场景、MPC/PI 对比、最终结果", "是"],
        ["Switching_Power_Stage", "电气主电路 Subsystem", "PV、Boost、半桥、电池、母线和负载", "是"],
        ["Measurement_and_Filtering", "测量控制接口", "把电气测量变成控制器状态", "原始量有，滤波量弱"],
        ["Averaged_Predictor", "MATLAB/Simulink 数学模型", "nlMPC 预测、工作点、快速调参", "否"],
        ["P_and_O / PMS / ARIMA / MPC / PI", "控制电路", "参考值、模式、预测、占空比", "离散控制信号"],
    ], [1.70, 1.30, 2.55, 0.95])

    heading(doc, "论文参数、复现假设与本机验证边界", 2)
    table(doc, ["类别", "论文明确给出", "本项目采用但论文未给出", "处理原则"], [
        ["功率级", "Cpv=Cb=300 µF；Lpv=Lb=10 mH；rL=10 mΩ；Cdc=1500 µF；Vdc*=600 V；fsw=10 kHz", "器件 Ron、Vf、死区、ESR", "单独列为假设并做敏感性测试"],
        ["PV", "组件 213.15 W，Vmp=29 V，Imp=7.35 A，Voc=36.3 V，Isc=7.84 A，阵列约 9.5 kW", "9 串 × 5 并；温度系数；单二极管拟合参数", "沿用现有脚本并明确单位"],
        ["Battery", "300 V、20 Ah、SoC 20%–90%", "E0/K/A/B、Rbat；物理块完整放电曲线", "初版用现有 OCV 表/0.08 Ω 校准，不冒充论文原值"],
        ["预测与控制", "代价权重和主要约束", "ARIMA 阶次/系数、PI 完整实现、控制周期", "保留可替换字段，结果称为工程复现"],
    ], [1.05, 2.30, 2.15, 1.00], font_size=7.8)

    # Stage 0
    stage_header(doc, "0", "冻结接口、符号和模型边界",
                 "先冻结主电路与控制器之间的信号合同，防止做到后面才发现 db、iLb、Ib 的正方向与方程相反。",
                 "一个只有端口和占位子系统的顶层模型；所有输入/输出的名字、单位、采样时间和正方向固定。")
    heading(doc, "0.1 新建顶层模型，不覆盖旧模型", 2)
    number(doc, "在 D:\\PV_MPPT\\PV_Battery_MPC_Project\\models 中打开原 PV_Battery_MPC.slx。")
    number(doc, "执行 Save As，保存为 PV_Battery_MPC_Switching.slx。原文件继续作为平均值基准。")
    number(doc, "删除或断开原 PV_Battery_Averaged_Plant 与控制器之间的闭环线，但暂时保留该子系统，重命名为 Averaged_Predictor_Reference。")
    number(doc, "新建四个 Subsystem：Switching_Power_Stage、Measurement_and_Filtering、Control_System、Scenario_Profiles。")
    heading(doc, "0.2 冻结端口", 2)
    table(doc, ["分类", "信号", "单位/类型", "方向与定义", "采样时间"], [
        ["u 控制输入", "dpv", "pu, double", "PV 低侧开关平均导通比；增大通常使 Vpv 降低", "Ts_ctrl=10 ms"],
        ["u 控制输入", "db", "pu, double", "电池下管导通比；上管平均导通比为 1-db", "Ts_ctrl=10 ms"],
        ["w 外扰", "T, G, RLoad", "°C, W/m², Ω", "PV 温度、辐照度、等效负载电阻", "场景离散"],
        ["y 测量", "Vpv,iLpv,Vb,iLb,SoC,Vdc", "V,A,V,A,pu,V", "iLb 规定电池→开关节点/母线为正", "滤波后 Ts_ctrl"],
        ["z 真值", "Ipv,Ib,Ppv,Pbat,Pload,gates", "A/W/bool", "仅记录与排错，不直接给控制器", "Ts_power/raw"],
    ], [0.90, 1.25, 1.10, 2.30, 0.95], font_size=7.7)
    callout(doc, "论文公式勘误",
            "论文状态方程的 SoC 行写成 Ebat−x4，但按 x=[Vpv,iLpv,Vb,iLb,SoC,Vdc] 和前文式 (6)/(9)，应为 Ebat−x3。"
            "现有脚本采用 Vb，不要为了逐字抄论文而改错。",
            fill=PALE_RED)
    stage_checkpoint(doc, [
        "旧平均值模型仍可单独运行，新开关模型有独立文件名。",
        "dpv、db、iLb、Ib 的正方向已写在模型 Annotation 中。",
        "顶层只放四个主子系统，信号线不跨越多个功能域。",
    ])
    stage_faults(doc, [
        ["充电时 iLb 仍为正", "电流传感器方向与合同相反", "旋转 Current Measurement，使正箭头从电池指向半桥；控制器中不要靠随意乘 −1 掩盖拓扑。"],
        ["db 增大后的响应与预测相反", "上、下管定义颠倒", "确认上管平均导通比 1-db，下管平均导通比 db；检查 Gate Mux 顺序。"],
    ])

    # Stage 1
    stage_header(doc, "1", "参数脚本、求解器与三层时间尺度",
                 "开关级对象的最快事件是 10 kHz 开关沿，不再允许把原平均值模型的 10 µs 步长或 10 ms 控制周期混作 PWM 周期。",
                 "参数工作区中同时存在 Ts_power、Tsw、Ts_ctrl；powergui 和模型求解器可编译。")
    add_picture(doc, QA / "sample_time_ladder.png", 6.35, "图 3  电力电子步长、PWM 周期和控制周期的关系")
    heading(doc, "1.1 沿用 pvbatt_parameters.m，但增加字段而不改旧字段含义", 2)
    paragraph(doc, "打开 D:\\PV_MPPT\\PV_Battery_MPC_Project\\scripts\\pvbatt_parameters.m。在 p.converter.switchingFrequency = 10e3; 后加入以下内容。原 p.sim.sampleTime=10e-6 继续专供平均值模型。")
    code(doc, """% Switching physical plant (new fields; do not replace p.sim.sampleTime)
p.sim.powerStep = 1e-6;
p.sim.controllerSampleTime = p.controller.interval;
p.converter.switchingPeriod = 1/p.converter.switchingFrequency;
p.converter.deadTime = 1e-6;

% Power-device assumptions not published in the paper
p.device.Ron = 1e-3;
p.device.snubberResistance = 1e6;
p.device.snubberCapacitance = inf;
p.device.pvDiodeForwardVoltage = 0.8;
p.device.batteryDiodeForwardVoltage = 0;
p.batteryConverter.capacitorESR = 1e-3;

% Measurement filter; 500 Hz is below fsw but above 100 Hz controller rate
p.sensor.cutoffFrequency = 500;
p.sensor.alpha = exp(-2*pi*p.sensor.cutoffFrequency*p.sim.powerStep);""")
    heading(doc, "1.2 参数改变的影响", 2)
    table(doc, ["参数", "推荐起点", "增大后的主要影响", "减小后的主要影响/风险"], [
        ["Ts_power", "1 µs", "仿真快，但开关沿、峰值和纹波误差增大", "更准确但耗时与数据量显著增加"],
        ["fsw", "10 kHz", "同 L/C 下纹波减小，开关损耗实际会增大", "纹波增大，控制滤波压力变大"],
        ["deadTime", "1 µs", "直通更安全，但产生平均电压误差", "过小可能上下管重叠"],
        ["Ron", "1 mΩ", "损耗和压降增大，稳态占空比需补偿", "趋近理想器件，可能使数值条件变差"],
        ["Cb ESR", "1 mΩ 起步", "浪涌减小但损耗增加", "理想电压源与理想电容并联会报错或产生冲击"],
        ["fc_sensor", "500 Hz", "跟踪更快但残留开关纹波", "更平滑但相位滞后更大"],
    ], [1.15, 1.00, 2.15, 2.20], font_size=7.8)
    heading(doc, "1.3 powergui 与求解器精确设置", 2)
    table(doc, ["位置/模块", "字段", "设置", "作用"], [
        ["powergui", "Simulation type", "Discrete", "让 SPS 电气网络按离散步长求解"],
        ["powergui", "Sample time (s)", "p.sim.powerStep = 1e-6", "电力网络更新步长"],
        ["powergui", "Start simulation with initial electrical states from", "blocks", "使用 RLC 模块中设定的初始电压/电流"],
        ["Model Settings > Solver", "Type", "Fixed-step", "开关仿真时间网格固定"],
        ["Model Settings > Solver", "Solver", "ode3 (Bogacki-Shampine)", "R2025a Battery 含连续内部状态；FixedStepDiscrete 会报错"],
        ["Model Settings > Solver", "Fixed-step size", "1e-6", "与 powergui 一致"],
    ], [1.55, 1.65, 1.45, 1.85], font_size=7.7)
    callout(doc, "为什么不是 FixedStepDiscrete",
            "仅含纯离散 SPS 网络时可以使用离散求解器；但 R2025a 的 Battery 模块内部包含连续库仑计数/滤波状态。"
            "本机验证中 FixedStepDiscrete 会直接报“模型包含连续状态”，因此本手册采用固定步长 ode3。",
            fill=PALE_YELLOW)
    stage_checkpoint(doc, ["Fsw=10 kHz、Tsw=100 µs、Ts_power=1 µs、Ts_ctrl=10 ms 均可在工作区查看。", "模型 Update Diagram 无连续/离散求解器冲突。", "运行 1 ms 空模型不出现 powergui 缺失警告。"])
    stage_faults(doc, [
        ["FixedStepDiscrete 不能用于模型", "Battery 含连续状态", "改 Solver 为 ode3，Fixed-step=1e-6；不要删除 Battery 的动态状态。"],
        ["仿真极慢或文件巨大", "记录全部 1 µs 信号", "只保留必要 raw truth；控制量按 10 ms 记录；To Workspace 使用 decimation。"],
        ["PWM 波形只有少数点", "Ts_power 太大", "保证每个 100 µs 周期至少 50–100 点，推荐 1–2 µs。"],
    ])

    # Stage 2
    stage_header(doc, "2", "PV 静态查表与物理 PV Array 的职责分离",
                 "同一套 PV 参数需要服务两个不同对象：物理主电路用 PV Array；MPC 预测和工作点计算用三维查表。两者数值应一致，但不能互相替代。",
                 "PV 查表脚本可无参数或有参数运行；PV Array 的 I–V/P–V 曲线与查表在测试点一致。")
    heading(doc, "2.1 修正你之前遇到的“输入参数不足”", 2)
    paragraph(doc, "当前函数签名是 pvbatt_generate_pv_lookup(p)，所以直接在命令行输入函数名会缺少 p。推荐在函数开头加入默认参数，使两种调用都成立。")
    code(doc, """function [voltageBreakpoints, temperatureBreakpoints, ...
        irradianceBreakpoints, currentTable] = pvbatt_generate_pv_lookup(p)

if nargin < 1 || isempty(p)
    p = pvbatt_parameters();
end

% 原函数后续内容保持不变""")
    paragraph(doc, "修改后测试：")
    code(doc, """projectRoot = "D:\\PV_MPPT\\PV_Battery_MPC_Project";
addpath(fullfile(projectRoot,"scripts"),"-begin");
clear pvbatt_parameters pvbatt_generate_pv_lookup
[PV_V_BP,PV_T_BP,PV_G_BP,PV_I_TABLE] = pvbatt_generate_pv_lookup();
size(PV_I_TABLE)     % 期望 [191 7 8]
min(PV_I_TABLE,[],'all'), max(PV_I_TABLE,[],'all')""")
    heading(doc, "2.2 PV Array 模块", 2)
    paragraph(doc, "Library Browser 搜索 PV Array。程序化源路径：spsPVArrayLib/PV Array。把它放进 Switching_Power_Stage；输入端 1 接 G，输入端 2 接 T，右侧电气端为正、负极，m 口用于内部测量。")
    table(doc, ["字段", "设置", "作用", "改变时的影响"], [
        ["Parallel strings", "5", "总电流按并联串数放大", "增大时 Isc、Imp 和功率近似成比例增加"],
        ["Series-connected modules/string", "9", "总电压按串联数放大", "增大时 Voc、Vmp 增加，初始 dpv 必须重算"],
        ["Module", "User-defined", "使用论文组件额定点", "选择预置组件会改变整条 I–V 曲线"],
        ["Pm/Vmp/Imp/Voc/Isc", "213.15/29/7.35/36.3/7.84", "确定组件额定曲线", "任一项改变都应重新校验 MPP"],
        ["Cells/module", "60", "影响热电压和曲线斜率", "错误会使 Voc 附近曲线严重偏差"],
        ["βVoc", "−0.36 %/°C（假设）", "温度升高时 Voc 下降", "绝对值增大使高温 MPP 电压下降更快"],
        ["αIsc", "0.05 %/°C（沿用假设）", "温度对短路电流影响", "通常影响小于辐照度"],
        ["Robust discrete model", "Off", "保留外部温度输入", "On 时温度口可能被禁用，不适合论文温度场景"],
        ["Break algebraic loop", "On；Tc=1e-6；Tfilter=5e-5", "改善离散网络数值稳定性", "时间常数过大会扭曲快速 PV 响应"],
    ], [1.55, 1.55, 1.55, 1.85], font_size=7.5)
    heading(doc, "2.3 查表与物理块的一致性测试", 2)
    number(doc, "暂不连接变换器；用可变电阻从 1 Ω 扫到 1 kΩ，G=1000 W/m²、T=25 °C。")
    number(doc, "记录 Vpv、Ipv，计算 Ppv=Vpv×Ipv，绘制 I–V/P–V。")
    number(doc, "在查表中取 T=25、G=1000 的切片，叠加到同一图。")
    number(doc, "允许插值/块内部模型造成小差异，但 MPP 电压和功率应接近：查表网格约 Vmp=262 V、Pmax≈9.59 kW。")
    callout(doc, "单位检查",
            "PV Array 的 G 输入是 W/m²，不是 0–1 pu；T 输入是 °C，不是 K；物理模块输出已经是 9×5 阵列量。"
            "三维查表的第 1 维是阵列电压，不是单块组件电压。",
            fill=PALE_RED)
    stage_checkpoint(doc, ["PV_I_TABLE 尺寸为 191×7×8，电流范围非负。", "STC 物理块和查表的 MPP 功率均约 9.5 kW。", "辐照度从 1000 降至 500 时短路电流和最大功率近似减半。"])
    stage_faults(doc, [
        ["输入参数不足", "直接无参调用旧函数", "加入 nargin 默认参数，或先 p=pvbatt_parameters() 再传入。"],
        ["PV 功率大约只有 1/5", "Parallel strings 仍为默认 40 或设置错误", "设 Npar=5、Nser=9，并确认使用阵列量。"],
        ["T 变化没有作用", "Robust discrete model 开启后温度口禁用", "关闭该选项，T 接输入端 2。"],
        ["查表电流为 NaN/负数", "温度单位或 Newton 初值错误", "温度用 °C；保留指数上限 60 和电流限幅。"],
    ])

    # Stage 3
    stage_header(doc, "3", "搭建 PV 开关型 Boost 主电路",
                 "把平均式 L di/dt=Vpv−(1−dpv)Vdc 还原成两个真实开关状态，观察电感纹波、二极管导通和 PWM 对 Vpv 的作用。",
                 "PV Array、Cpv、Lpv/rLpv、Boost Converter、PWM、测量块组成独立可运行的 PV 支路。")
    add_picture(doc, PROTOTYPE, 6.5, "图 4  已在本机 R2025a 编译通过的开关级主电路布局参考")
    heading(doc, "3.1 所需模块与精确设置", 2)
    table(doc, ["模块（搜索名/源路径）", "数量", "关键设置", "用途"], [
        ["Parallel RLC Branch / spsParallelRLCBranchLib", "1", "Branch type=C；C=300e-6；初始电压先 0，后再设 Vpv0", "PV 端并联电容 Cpv"],
        ["Current Measurement / spsCurrentMeasurementLib", "1", "默认；正方向 PV→Boost", "测 Ipv 或 iLpv，取决于放置位置"],
        ["Series RLC Branch / spsSeriesRLCBranchLib", "1", "Branch type=RL；R=10e-3；L=10e-3", "Lpv 与铜耗 rLpv"],
        ["Boost Converter / spsBoostConverterLib", "1", "Model type=Switching devices；Ron=1e-3；Vf diode=0.8", "封装论文中的低侧开关和升压二极管"],
        ["PWM Generator (DC-DC) / spsPWMGeneratorDCDCLib", "1", "Fsw=10e3；Sample time=1e-6", "把 dpv 变成开关门极脉冲"],
        ["Constant", "1", "0", "接 Boost 的 BL（Block）端；0 表示不封锁"],
        ["Voltage Measurement", "1", "正端接 PV+，负端接母线−", "测 Vpv"],
    ], [2.25, 0.45, 2.35, 1.45], font_size=7.4)
    heading(doc, "3.2 接线顺序", 2)
    table(doc, ["序号", "从", "到", "接线意义"], [
        ["1", "PV Array +", "Cpv 上端、Vpv 测量 +、电感/电流测量入口", "形成 PV 正端节点"],
        ["2", "PV Array −", "Cpv 下端、Vpv 测量 −、DC−", "统一负母线"],
        ["3", "PV 正端", "Current Measurement → Lpv", "正电流定义 PV 向母线"],
        ["4", "Lpv 右端", "Boost Converter 的端口 1", "升压开关节点"],
        ["5", "Boost + / −", "DC+ / DC−", "输出接直流母线"],
        ["6", "dpv → Saturation", "PWM D 端", "占空比限制后调制"],
        ["7", "PWM P 端", "Boost g 端", "实际门极脉冲"],
        ["8", "Constant 0", "Boost BL 端", "正常运行不封锁"],
    ], [0.42, 1.65, 2.20, 2.23], font_size=7.7)
    heading(doc, "3.3 两个开关状态与平均式", 2)
    paragraph(doc, "开关 ON（持续 dpv·Tsw）：低侧开关闭合，二极管反偏，电感两端约为 Vpv−rL·i，电感储能，di/dt>0。")
    paragraph(doc, "开关 OFF（持续 (1−dpv)·Tsw）：低侧开关断开，二极管导通，电感向母线放能，电感电压约为 Vpv−rL·i−Vdc。")
    paragraph(doc, "一周期平均：L·di/dt=dpv(Vpv−ri)+(1−dpv)(Vpv−ri−Vdc)=Vpv−ri−(1−dpv)Vdc。稳态时它等于 0，并不代表每个瞬间电感电压为 0；ON/OFF 两段分别正、负，平均后才抵消。")
    callout(doc, "关于你之前提出的“不是相减为 0 吗”",
            "稳态电压平衡确实给出 Vpv≈(1−dpv)Vdc+rL·i，因此平均电感电压为 0；这正是稳态条件，不是方程重复。"
            "动态过程中 Vpv、Vdc、i 或 dpv 一变化，右端不再为 0，电流就上升或下降。开关级模型还能看到每个 100 µs 周期内的正负电压。",
            fill=PALE_GREEN)
    heading(doc, "3.4 开环验证", 2)
    number(doc, "先用受控/固定 600 V 母线替代后续电池和负载，仿真 5 ms；dpv 从 0.55 在 2 ms 阶跃到 0.60。")
    number(doc, "检查 g_pv 为 10 kHz；占空比 0.55/0.60；iLpv 为三角纹波。")
    number(doc, "比较 Vpv：dpv 增大后 PV 被拉向更低电压；若方向相反，检查开关端口和占空比定义。")
    number(doc, "比较理论纹波。比较工况 x0 下，ΔiLpv,pp≈Vpv/L·dpv·Tsw≈1.46 A。允许器件压降造成差异。")
    stage_checkpoint(doc, ["g_pv 频率为 10 kHz，且 duty 与 dpv 一致。", "iLpv 连续、非跳变，峰峰纹波约 1–2 A。", "ON 时电感电流上升、OFF 时下降；平均斜率在稳态接近 0。"])
    stage_faults(doc, [
        ["iLpv 瞬间发散", "L=10e-3 误写成 10e3 或 µH/mH 混淆；母线未预充", "确认单位 H；先用固定母线做子系统测试。"],
        ["dpv 增大时 Vpv 增大", "门极接错或把 Boost 端口 1 接到 DC+", "端口 1 只接电感后的开关节点；+/- 接母线。"],
        ["没有纹波", "仍在使用平均模型或 PWM 未接门极", "Boost Model type 必须 Switching devices；观察 g 端。"],
        ["二极管尖峰导致求解困难", "理想器件/无 snubber", "Ron=1e-3，Rs=1e6，Cs=inf；必要时减小步长。"],
    ])

    # Stage 4
    stage_header(doc, "4", "物理 Battery、SoC 与 Cb 支路",
                 "把电池从 OCV 查表代数源升级为带电气端口和内部 SoC 动态的 Battery 模块，同时保留平均模型的 OCV 表供预测器使用。",
                 "Battery 可开路运行，电压和 SoC 合理；Cb 不再触发理想电压源并联电容错误。")
    heading(doc, "4.1 Battery 模块设置", 2)
    paragraph(doc, "Library Browser 搜索 Battery；程序化源路径 spsBatteryLib/Battery。论文没有给出 E0、K、A、B，因此以下高级曲线参数是用于与现有 OCV 表对齐的工程起点，不是论文原始参数。")
    table(doc, ["字段", "起始设置", "作用与影响"], [
        ["Type", "Lithium-Ion", "选择论文电池类型；改变类型会改变极化/指数区方程"],
        ["Nominal voltage", "300 V", "电池平台标称电压；改变后半桥稳态占空比必须重算"],
        ["Rated capacity", "20 Ah", "SoC 积分基数；增大后同一电流下 SoC 变化更慢"],
        ["Initial SoC", "场景值×100，如 80", "该模块输入是百分数，不是 0.8"],
        ["Response time", "1 s（假设）", "极化电流滤波；增大后端电压动态更慢"],
        ["Preset model", "Off（对齐时）", "允许填写以下参数；初次学习可 On，但会自动给 R=0.15 Ω、FullV≈349 V，与现有 OCV 表不一致"],
        ["Maximum capacity", "20 Ah", "总可用电荷"],
        ["Cut-off / Full", "270 / 312 V", "来自现有 OCV 表端点；决定放电/满充边界"],
        ["Internal resistance", "0.08 Ω", "沿用当前复现假设；增大则压降/损耗更大"],
        ["Nominal-zone capacity", "18 Ah", "额定平台结束位置"],
        ["Exponential zone", "[303 V, 1 Ah]", "满充端快速电压下降区，需用数据校准"],
    ], [1.85, 1.80, 2.85], font_size=7.5)
    heading(doc, "4.2 Cb 为什么必须带 ESR", 2)
    paragraph(doc, "论文画的是电池端并联 Cb=300 µF。SPS Battery 内部含受控电压源；若直接并联理想 C 模块，R2025a 会报“voltage source cannot be connected in parallel with capacitive element”。")
    number(doc, "使用 Series RLC Branch，而不是理想 Parallel C。")
    number(doc, "Branch type=RC；R=1e-3 Ω（Cb_ESR）；C=300e-6 F。")
    number(doc, "该 RC 串联支路整体并联在电池正负端。初始电压符号由模块端口方向决定；t=0 用 Voltage Measurement 验证必须为 +Vb0。")
    callout(doc, "启动冲击警告",
            "Battery 在 80% SoC 的实际开路电压、Cb 初始电压和脚本算得 Vb0 必须接近。只差数伏也会通过毫欧级 ESR 形成很大瞬态。"
            "首次集成把 Cb_ESR 临时提高到 0.05–0.1 Ω或使用预充电支路；确认极性和稳态后再降回目标值。",
            fill=PALE_RED)
    heading(doc, "4.3 SoC 符号核对", 2)
    paragraph(doc, "规定 Ib>0 为电池放电，dSoC/dt=−Ib/(3600Q)。Battery 模块 m 端可输出内部测量；外部 Current Measurement 必须从电池正端指向半桥。用 5 A 放电 72 s，理论 SoC 下降约 5×72/(3600×20)=0.005，即 0.5 个百分点。")
    stage_checkpoint(doc, ["Battery 开路电压在 270–312 V 合理区间。", "5 A 放电时 SoC 单调下降，充电时上升。", "Cb 支路编译无理想源/电容并联错误，t=0 极性为正。"])
    stage_faults(doc, [
        ["电压源不能与电容并联", "Cb 使用理想 C", "改为 Series RLC Branch 的 RC 支路，加入 ESR。"],
        ["初始瞬间数千安", "Cb 极性反向或 Vb0 与 Battery OCV 不一致", "先看 t=0 Vb；反向则交换端口/改变 Initial Voltage 符号；临时增大 ESR。"],
        ["SoC 变化快 100 倍", "把 20 Ah 当 20 A·s，或初值 0.8 填进百分数口", "容量保持 Ah；Battery Initial SoC 填 80。"],
        ["Battery 满充电压约 349 V", "Preset model 自动缩放", "切换 Off，填写 270/312 V 和 0.08 Ω；标记为假设。"],
    ])

    # Stage 5
    stage_header(doc, "5", "搭建电池双向半桥及互补门极",
                 "还原论文中上、下两个开关，使能量可以双向流动，并保证 db 的定义与平均模型 (1−db)Vdc 完全一致。",
                 "Two-Quadrant DC/DC Converter、Lb/rLb、互补 PWM、死区和 double 门极向量可运行。")
    add_picture(doc, QA / "battery_gate_logic.png", 6.35, "图 5  电池半桥门极逻辑；一个载波产生两路互补脉冲")
    heading(doc, "5.1 功率模块与端口", 2)
    table(doc, ["模块", "设置", "端口/作用"], [
        ["Two-Quadrant DC/DC Converter", "Switching devices；Ron=1e-3；snubber Rs=1e6,Cs=inf", "+/− 接母线；端口 1（Out）接 Lb；g 接两元素门极向量；BL 接 0"],
        ["Series RLC Branch (Lb)", "RL；R=10e-3；L=10e-3", "连接半桥开关节点与电池正端"],
        ["Current Measurement", "正方向电池→半桥", "输出定义为 iLb；放电为正，充电为负"],
        ["PWM Generator (DC-DC)", "Fsw=10e3；Ts=1e-6", "输入 db，输出下管原始脉冲"],
        ["Logical NOT", "Operator=NOT", "同一脉冲取反得到上管原始脉冲"],
        ["Unit Delay ×2 + AND ×2", "Ts=1e-6；IC=0", "gate = raw AND z^-1(raw)，只延迟上升沿一个步长"],
        ["Mux + Data Type Conversion", "Mux Inputs=2；Output double", "向量顺序 [g_upper;g_lower]；解决 boolean 门极类型错误"],
    ], [2.15, 2.00, 2.35], font_size=7.5)
    heading(doc, "5.2 为什么不能使用两个独立 PWM", 2)
    paragraph(doc, "PWM(D) 通常在每周期同一相位开始为高。如果分别生成 PWM(db) 和 PWM(1−db)，两者在周期开始可能同时为高，并非互补，会让上下管直通。正确方式是一个 PWM 后取 NOT，再分别做上升沿死区。")
    heading(doc, "5.3 门极顺序和数据类型", 2)
    number(doc, "Mux 第 1 路接 g_upper（Sw1，上管），第 2 路接 g_lower（Sw2，下管）。")
    number(doc, "Mux 后放 Data Type Conversion；Output data type=double。逻辑块输出 boolean，而半桥内部 Saturation 要求数值型。")
    number(doc, "BL 接 Constant 0。若 BL>0.5，模块会封锁门极，适合故障停机。")
    heading(doc, "5.4 开环能量方向测试", 2)
    table(doc, ["测试", "条件", "期望"], [
        ["近零电感平均电压", "Vb≈300 V，Vdc≈600 V，db≈0.5", "iLb 平均变化慢；纹波约 1.5 A_pp"],
        ["提高 db", "db 0.50→0.55", "下管占空比增大，开关节点平均电压下降，电池侧电感电流趋向放电方向增加"],
        ["降低 db", "db 0.50→0.45", "上管占空比增大，电流趋向充电/负方向"],
        ["死区检查", "同时观察 g_upper、g_lower", "任意时刻不得同时为 1；每次换相有约 1 µs 两者都为 0"],
    ], [1.45, 2.35, 2.70], font_size=8.0)
    stage_checkpoint(doc, ["g_upper 与 g_lower 永不同时为 1，频率均为 10 kHz。", "Gate Mux 输出为 double[2×1]，不再出现 boolean 类型错误。", "db≈0.5 时 iLb 纹波约 1–2 A，方向与合同一致。"])
    stage_faults(doc, [
        ["半桥 Saturation 输入 boolean 错误", "Mux 输出继承 boolean", "Mux 后加 Data Type Conversion，设 double。"],
        ["母线瞬间过压/电流巨大", "上下管直通，或初始电压不一致", "检查同一 PWM+NOT+死区；先分离电池支路做开环。"],
        ["db 改变方向相反", "Mux 上下顺序反", "第 1 元素必须上管，第 2 元素下管；用内部图/门极状态确认。"],
        ["电池只能单向", "用普通 Boost 或二极管阻断充电", "使用 Two-Quadrant DC/DC Converter 或等效两只 IGBT/Diode 半桥。"],
    ])

    # Stage 6
    stage_header(doc, "6", "直流母线、可变电阻负载与功率守恒",
                 "建立论文式 Cdc–RLoad 公共母线，让 Vdc 成为功率不平衡的积分结果，而不是理想电压源强行固定。",
                 "Cdc、Variable Resistor、Vdc 测量和功率计算连接完成；整机开环满足能量平衡。")
    heading(doc, "6.1 模块与设置", 2)
    table(doc, ["模块", "设置", "作用与参数影响"], [
        ["Parallel RLC Branch (Cdc)", "C；1500e-6 F；初值先 600 V", "C 越大 Vdc 变化越慢、纹波越小，但启动能量更大"],
        ["Variable Resistor / spsVariableResistorLib", "Minimum resistance absolute value=1 Ω", "输入 RLoad 可随场景变化；过小会导致巨大负载电流"],
        ["Voltage Measurement", "+ 接 DC+，− 接 DC−", "输出 Vdc；反接会使控制器正反馈"],
        ["Product", "Vdc×Iload", "Pload；也可用 Vdc²/RLoad 交叉验证"],
        ["Ground", "接 DC−", "给 SPS 网络参考节点"],
    ], [2.35, 1.80, 2.35], font_size=7.8)
    heading(doc, "6.2 场景 RLoad 接法", 2)
    paragraph(doc, "Scenario_Profiles 已输出 w=[T,G,RLoad]。对 w Demux 后，RLoad 先经过 Saturation Lower limit=1，再接 Variable Resistor 的 R 控制口。负载 7.2 kW、Vdc=600 V 时 R=V²/P=50 Ω；14.4 kW 时 R=25 Ω。")
    heading(doc, "6.3 母线方程与功率解释", 2)
    paragraph(doc, "开关级瞬时 KCL：iCdc=iPV,bus+iBat,bus−iLoad。平均后得到 Cdc·dVdc/dt≈(1−dpv)iLpv+(1−db)iLb−Vdc/RLoad。")
    paragraph(doc, "乘以 Vdc 得到电容功率 d(½CdcVdc²)/dt≈Ppv,bus+Pbat,bus−Pload−损耗。Vdc 上升不是“控制器发散”的同义词：它首先说明输入功率在一段时间内大于负载与损耗；需要继续判断是占空比、初值还是控制逻辑造成。")
    heading(doc, "6.4 开环功率守恒检查", 2)
    number(doc, "记录 Vdc、Cdc 电流（可用母线支路电流求和）、Ppv、Pbat、Pload。")
    number(doc, "用有限差分计算 dEcdc/dt，其中 Ecdc=0.5*Cdc*Vdc^2。")
    number(doc, "检查 Pin−Pout−dE/dt 与开关/电阻损耗同量级；若差值巨大，通常是电流方向或功率符号错。")
    stage_checkpoint(doc, ["RLoad=50 Ω、Vdc≈600 V 时 Pload≈7.2 kW。", "RLoad 阶跃 50→25 Ω 后负载功率立即增加，Vdc 先下降。", "功率残差没有长期保持数千瓦的非物理偏差。"])
    stage_faults(doc, [
        ["RLoad=0 时仿真崩溃", "变量电阻无下限", "R 输入前 Saturation≥1 Ω；夜间负载不要用 0 Ω 表示断开。"],
        ["负载变大时 Vdc 上升", "Vdc 测量或负载接线/功率符号反", "检查 Variable Resistor 跨接 DC+/DC−，电流为从正母线到负母线。"],
        ["Cdc 初值显示 −600 V", "电容端口或 Initial Voltage 符号反", "交换端口或改初值符号，t=0 先验证再接控制器。"],
    ])

    # Stage 7
    stage_header(doc, "7", "测量、滤波、速率转换与状态总线",
                 "控制器需要平均状态而不是每个开关沿的尖峰。测量层把电气真值变成单位、方向、带宽和采样时间均明确的 y 信号。",
                 "raw 与 filtered 两套信号同时存在；x_state=[Vpv,iLpv,Vb,iLb,SoC,Vdc] 在 10 ms 控制速率更新。")
    heading(doc, "7.1 传感器放置", 2)
    table(doc, ["状态", "模块/位置", "正方向", "是否滤波"], [
        ["Vpv", "Voltage Measurement 跨 PV+/DC−", "PV+ 对 DC−", "是"],
        ["iLpv", "Current Measurement 串在 Lpv", "PV→母线", "是"],
        ["Vb", "Voltage Measurement 跨 Battery+/DC−", "Battery+ 对 DC−", "是"],
        ["iLb", "Current Measurement 串在 Lb", "Battery→半桥", "是"],
        ["SoC", "Battery m 口或独立库仑计", "0–1；若块输出 0–100 需 Gain=0.01", "只按控制周期采样"],
        ["Vdc", "Voltage Measurement 跨母线", "DC+ 对 DC−", "是"],
    ], [0.70, 2.45, 1.70, 1.65], font_size=7.9)
    heading(doc, "7.2 一阶离散滤波器精确搭法", 2)
    paragraph(doc, "每个快速模拟量用 Discrete Transfer Fcn：alpha=exp(−2πfc·Ts_power)。fc=500 Hz、Ts=1 µs 时 alpha≈0.996863。")
    table(doc, ["字段", "设置"], [["Numerator", "1-alpha"], ["Denominator", "[1 -alpha]"], ["Initial states", "对应工作点值，首次调试可 0"], ["Sample time", "Ts_power"]], [2.20, 4.30])
    paragraph(doc, "滤波器之后放 Rate Transition：输入快、输出慢，Output port sample time=Ts_ctrl；Ensure data integrity=On；Deterministic data transfer=On。然后用 Mux 按固定顺序组成 x_state。")
    heading(doc, "7.3 参数改变的影响", 2)
    table(doc, ["改动", "好处", "代价/症状"], [
        ["fc 提高", "更快反映负载阶跃", "更多 10 kHz 纹波进入 MPC，控制占空比抖动"],
        ["fc 降低", "波形更平滑", "相位滞后增大，Vdc 瞬态可能变差"],
        ["控制器直接接 raw", "看似响应快", "MPC 每次采到开关相位不同，预测残差大"],
        ["滤波初值设 0", "设置简单", "启动时产生虚假大误差，控制器立即饱和"],
    ], [1.65, 2.15, 2.70], font_size=8.0)
    stage_checkpoint(doc, ["raw iL 可看到约 10 kHz 纹波，filtered iL 平滑且平均值一致。", "x_state 顺序和平均预测模型完全一致。", "SoC 进入控制器前为 0–1，而不是 0–100。"])
    stage_faults(doc, [
        ["MPC 输出每次更新跳变很大", "直接采样开关纹波", "增加低通滤波和 Rate Transition；检查 fc。"],
        ["SoC 立刻触发上限", "Battery 输出百分数未缩放", "Battery SoC 若为 80，先乘 0.01。"],
        ["Update Diagram 报速率问题", "快慢离散线直接连接", "在测量层和 duty 门极层使用 Rate Transition。"],
    ])

    # Stage 8
    stage_header(doc, "8", "一致初值、预充电与整机开环换能",
                 "开关级电路储能元件具有极性和瞬时能量；平均模型的 x0 只能提供目标工作点，必须经过极性、Battery OCV 和预充电一致性检查。",
                 "整机在固定 u0 下可稳定运行若干毫秒，没有上千安浪涌、母线翻极或开关直通。")
    heading(doc, "8.1 沿用 pvbatt_operating_point.m 计算目标", 2)
    code(doc, """projectRoot = "D:\\PV_MPPT\\PV_Battery_MPC_Project";
addpath(fullfile(projectRoot,"scripts"),"-begin");
pvbatt_initialize;
[x0,u0,w0,op] = pvbatt_operating_point(PVBATT_P,0.8,35,500,7200);
format long g
x0, u0, w0""")
    table(doc, ["量", "比较场景目标"], [
        ["x0", "[248.117469 V; 18.443148 A; 299.298439 V; 8.769514 A; 0.8; 600 V]"],
        ["u0", "[dpv;db]=[0.586778271;0.501315427]"],
        ["w0", "[35 °C;500 W/m²;50 Ω]"],
        ["功率", "Ppv≈4576.07 W；电池需约 2623.93 W"],
    ], [1.20, 5.30], font_size=8.4)
    heading(doc, "8.2 物理块初值映射", 2)
    table(doc, ["x0 元素", "物理模块字段", "设置前验证"], [
        ["Vpv", "Cpv Initial capacitor voltage", "端口正方向与 Vpv Measurement 一致"],
        ["iLpv", "Lpv Initial inductor current", "正方向 PV→母线"],
        ["Vb", "Cb Initial capacitor voltage", "必须接近 Battery 在初始 SoC 的开路电压"],
        ["iLb", "Lb Initial inductor current", "正方向 Battery→半桥；模块旋转后符号可能变化"],
        ["SoC", "Battery Initial SoC (%)", "0.8 要填 80"],
        ["Vdc", "Cdc Initial capacitor voltage", "t=0 Vdc 测量必须 +600"],
    ], [1.10, 2.85, 2.55], font_size=7.8)
    heading(doc, "8.3 推荐的三步启动法", 2)
    number(doc, "子系统启动：先以固定 600 V DC 源分别测试 PV Boost 和电池半桥，确认门极、电流方向和占空比。")
    number(doc, "整机预充：在电池支路或母线加入临时 Breaker/预充电电阻，使 Cdc、Cb 和 Battery OCV 先接近；控制门极保持封锁。")
    number(doc, "软启：解除 BL 后，用 Ramp 在 2–5 ms 内把 duty 从安全值过渡到 u0；确认无大冲击后再接 MPC。最终论文结果可从稳定点保存/加载初始状态。")
    callout(doc, "不要机械照抄 x0",
            "若 Battery 的 80% SoC 开路电压不是 299.3 V，把 Cb 强行初始化为 299.3 V 会制造大电压差。"
            "正确做法是先校准 Battery 曲线，或让预充过程把电容自然充到一致电压，再保存 operating point。",
            fill=PALE_RED)
    heading(doc, "8.4 开环验收数字", 2)
    table(doc, ["观察量", "正常特征", "异常阈值/处理"], [
        ["g_upper & g_lower", "不重叠，1 µs 死区", "同时为 1：立即停止，修正门极"],
        ["iLpv", "平均约 18 A，纹波约 1.46 A_pp", ">100 A：检查初值、极性和单位"],
        ["iLb", "平均约 8.8 A，纹波约 1.50 A_pp", "方向反或持续斜坡：检查 db/门极顺序"],
        ["Vdc", "短时围绕 600 V，无单调快速上升", "数 ms 超过 660 V：先断控制，查功率/初值"],
        ["功率残差", "接近器件损耗和电容能量变化", "长期几千瓦不平衡：查符号/接线"],
    ], [1.20, 2.85, 2.45], font_size=7.8)
    stage_checkpoint(doc, ["固定 u0 运行 5–20 ms 无大浪涌。", "电感纹波与理论量级一致，平均值靠近 x0。", "Vdc、Vpv、Vb 的 t=0 极性均为正，Battery OCV 与 Cb 一致。"])
    stage_faults(doc, [
        ["Vdc 数 ms 升到 1000 V", "电池/电容初值不一致、门极直通或占空比方向错", "先封锁两变换器，逐支路恢复；核对 Cb OCV、死区、Mux 顺序。"],
        ["PV 电流数百安", "PV 端电容极性反或开关节点接错", "只保留 PV 子系统；检查 t=0 Vpv 和 PV Array +/-。"],
        ["固定 duty 仍缓慢漂移", "器件压降使平均模型 u0 不再精确", "用小范围 duty 扫描/trim 找物理模型稳态，保留两套 operating point。"],
    ])

    # Stage 9
    stage_header(doc, "9", "保留并校准六状态平均值预测模型",
                 "nlMPC 不应在每次优化中仿真 IGBT；它使用论文六状态平均模型预测未来。该模型是控制器内部镜像，不再冒充真实主电路。",
                 "pvbatt_state_transition.m 继续运行，预测一步的平均状态能与开关模型的周期平均结果对齐。")
    add_picture(doc, ASSETS / "averaged_plant.png", 6.2, "图 6  原平均值对象现在作为预测器参考，而非最终主电路")
    heading(doc, "9.1 不应修改的脚本", 2)
    table(doc, ["脚本", "是否保留", "原因"], [
        ["pvbatt_generate_pv_lookup.m", "保留并加默认参数", "预测器需要快速求 Ipv(V,T,G)"],
        ["pvbatt_state_derivatives.m", "保留", "实现论文平均微分方程"],
        ["pvbatt_state_transition.m", "保留", "RK4 子步积分用于 nlMPC 的离散状态函数"],
        ["pvbatt_operating_point.m", "保留", "给控制器初值和 duty 起点；物理模型需二次 trim"],
        ["PV_Battery_Averaged_Plant", "保留但改名/隔离", "快速回归和定位控制器问题"],
    ], [2.25, 1.15, 3.10], font_size=8.0)
    heading(doc, "9.2 平均模型与开关模型的对齐方法", 2)
    number(doc, "在相同 T/G/RLoad、相同 duty、相同稳定初值下分别运行。")
    number(doc, "对开关模型信号按一个或多个 Tsw 做平均，再以 Ts_ctrl=10 ms 取样。")
    number(doc, "比较一控制周期后的 Vpv、iLpv、Vb、iLb、Vdc；差异来源包括 Ron、Vf、dead time、Battery 曲线和滤波。")
    number(doc, "如果稳态有固定偏差，优先把器件压降纳入预测模型或在物理模型上重新求 u0，而不是随意增大 MPC 权重。")
    heading(doc, "9.3 pvbatt_initialize.m 应增加的工作区变量", 2)
    paragraph(doc, "在 variables 单元数组中加入以下条目，避免每个模块写魔法数字。")
    code(doc, """'PWR_TS', p.sim.powerStep;
'CTRL_TS', p.sim.controllerSampleTime;
'PWM_FSW', p.converter.switchingFrequency;
'PWM_TSW', p.converter.switchingPeriod;
'PWM_DEADTIME', p.converter.deadTime;
'SW_RON', p.device.Ron;
'CB_ESR', p.batteryConverter.capacitorESR;
'SENSOR_ALPHA', p.sensor.alpha;""")
    stage_checkpoint(doc, ["平均模型仍能运行原 comparison 场景。", "相同 duty 下，开关模型周期平均趋势与平均模型一致。", "模型差异有明确来源表，不用调整权重掩盖硬件模型偏差。"])
    stage_faults(doc, [
        ["把 PWM 接进 state_transition", "混淆对象与预测模型", "预测函数只接 duty 平均值；PWM 只进入物理主电路。"],
        ["预测与物理稳态固定偏差", "Ron/Vf/deadtime/Battery 参数未建模", "先校准或加等效压降，再调控制器。"],
        ["预测运行很慢", "在优化里调用 PV Array/开关仿真", "恢复三维查表+RK4 平均模型。"],
    ])

    # Stage 10
    stage_header(doc, "10", "P&O MPPT 参考电压发生器",
                 "P&O 不直接产生开关门极，而是根据滤波后的 Vpv、Ipv 更新 Vpv_ref；MPC/PI 再把参考电压转换为 duty。",
                 "P&O 每 10 ms 更新一次，Vpv_ref 在限幅范围内向 MPP 收敛，开关纹波不会触发误判。")
    add_picture(doc, ASSETS / "po_mppt.png", 6.2, "图 7  现有 P&O 标准模块结构可继续沿用")
    heading(doc, "10.1 模块链与参数", 2)
    table(doc, ["模块", "设置", "作用/改变影响"], [
        ["Zero-Order Hold ×2", "Sample time=0.01", "按控制周期采 Vpv/Ipv；减小周期更快但更受纹波影响"],
        ["Product", "V×I", "计算 Ppv"],
        ["Unit Delay ×3", "IC: P0,V0,Vref0；Ts=0.01", "保存上次功率、电压和参考值"],
        ["Sum ΔP/ΔV", "+−", "得到变化量"],
        ["Product ΔPΔV + Relational >=0", "阈值 0", "同号继续、异号反向"],
        ["Abs + Compare", "deadband=1 W", "功率变化太小时保持，避免噪声游走"],
        ["Step gains", "+0.25 / −0.25 V", "步长大收敛快但振荡大；小则慢"],
        ["Saturation", "180–315 V", "限制在阵列安全电压区间"],
    ], [2.00, 1.80, 2.70], font_size=7.8)
    heading(doc, "10.2 实际周期怎样改", 2)
    paragraph(doc, "若要把 MPPT 周期从 10 ms 改为 5 ms，必须同时修改 p.controller.interval、p.mppt.sampleTime、P&O 内所有 ZOH/Unit Delay、MPC SampleTime、ARIMA Unit Delay 和控制速率 Rate Transition。只改一个模块会造成多速率不一致。")
    table(doc, ["周期变化", "主要好处", "必须重新检查"], [
        ["10 ms→5 ms", "更快跟踪辐照度变化", "P&O 步长、MPC rate constraint、计算耗时、滤波带宽"],
        ["10 ms→20 ms", "计算量下降", "负载阶跃时 Vdc 峰值可能增大，P&O 变慢"],
    ], [1.20, 2.25, 3.05])
    stage_checkpoint(doc, ["固定 G/T 下 Vpv_ref 在 MPP 附近小幅摆动。", "P&O 输入来自 filtered Vpv/Ipv，不是 raw 开关波形。", "所有 P&O 记忆块采样时间一致。"])
    stage_faults(doc, [
        ["Vref 单向跑到上限", "V/I 极性错、ΔV 符号错或初值为 0", "检查 P=V×I 为正；Unit Delay IC 使用 operating point。"],
        ["Vref 高频抖动", "输入未滤波或 deadband 太小", "使用 filtered 信号；提高 deadband/减小步长。"],
        ["改周期后 MPC 报速率", "只改 P&O", "按本节列出的所有控制速率同步修改。"],
    ])

    # Stage 11
    stage_header(doc, "11", "功率管理模式与 ARIMA 外扰预测",
                 "PMS 决定何时 MPPT、何时充放电、何时限发或切负载；ARIMA 提供下一控制周期的 T/G/RLoad 预测给 nlMPC。",
                 "模式 I–V 在边界条件下无重叠；w_hat 顺序固定为 [T,G,RLoad]。")
    add_picture(doc, ASSETS / "power_management.png", 6.1, "图 8  现有功率管理逻辑可继续使用，但输入改为开关模型的滤波状态")
    heading(doc, "11.1 模式定义", 2)
    table(doc, ["模式", "条件", "目标"], [
        ["I", "Ppv_avail≥Pload，SoC<0.9，白天", "PV MPPT，多余功率充电"],
        ["II", "Ppv_avail<Pload，SoC>0.2，白天", "PV MPPT，电池补功率"],
        ["III", "Ppv_avail≥Pload，SoC≥0.9，白天", "电池禁止继续充，PV 限发"],
        ["IV", "夜间且 SoC>0.2", "电池供电"],
        ["V", "无有效供能条件", "负载切除/安全停机"],
    ], [0.60, 3.55, 2.35], font_size=8.0)
    paragraph(doc, "现有逻辑用 Available PV ≈ arrayRatedPower/1000×G 判断余缺。它忽略温度与曲线细节，正式研究可改为查表最大功率；改动后必须重新测试模式边界。Mode III 使用 latch 防止在 SoC=0.9 附近退出/进入抖动。")
    heading(doc, "11.2 ARIMA(1,1,0) 标准块结构", 2)
    add_picture(doc, ASSETS / "arima_forecast.png", 6.1, "图 9  w_hat(k+1)=w(k)+phi[w(k)−w(k−1)]")
    table(doc, ["模块", "设置", "作用"], [
        ["ZOH", "0.01 s", "按控制周期读取实际外扰"],
        ["Unit Delay", "IC=w0(1/2/3)，Ts=0.01", "保存上次外扰"],
        ["Sum", "+−", "一阶差分"],
        ["Gain", "phi_T/phi_G/phi_R", "控制趋势外推；当前 phi=0 等于持久性预测"],
        ["Sum", "++", "当前值+趋势"],
        ["Mux", "[T,G,R]", "顺序必须与 nlMPC MD 一致"],
    ], [1.65, 2.10, 2.75], font_size=8.0)
    stage_checkpoint(doc, ["任一时刻只有一个模式为真，模式码稳定。", "Mode III 进入后在仍过剩且白天时保持。", "phi=0 时 w_hat(k+1)=w(k)，且 Mux 顺序为 T/G/R。"])
    stage_faults(doc, [
        ["模式 I/II 高频切换", "功率比较无滞环且测量含纹波", "用 filtered 功率；在阈值两侧加入功率滞环。"],
        ["负载预测进了 G 端", "Mux 顺序错", "全项目固定 [T,G,RLoad]，逐端口标注单位。"],
        ["预测比实测更差", "phi 未由历史数据辨识", "先用 phi=0 基线；有数据后再估计，不凭感觉填系数。"],
    ])

    # Stage 12
    stage_header(doc, "12", "nlMPC 对象、参考、权重与约束",
                 "控制器以平均模型预测 Vdc、Vpv、iLb，输出 dpv/db；开关级对象只反馈周期平均状态。",
                 "nlmpc 对象通过 validateFcns；正常模式和限发模式的在线权重切换正确。")
    add_picture(doc, ASSETS / "mpc_controller.png", 6.2, "图 10  现有 nlMPC 结构可复用，plant 输入换成 filtered switching states")
    heading(doc, "12.1 现有对象定义", 2)
    table(doc, ["项目", "当前设置", "意义/改变影响"], [
        ["nlmpc(6,3)", "MV=[1,2]；MD=[3,4,5]", "6 状态、3 输出；两个 duty、三个外扰"],
        ["Ts", "0.01 s", "控制更新周期；减小会显著增加优化次数"],
        ["Prediction/Control horizon", "1 / 1", "计算快但前瞻短；增大可改善预见性但代价大"],
        ["StateFcn", "pvbatt_state_transition", "10 个 RK4 子步；预测器平均模型"],
        ["OutputFcn", "[Vdc,Vpv,iLb]", "输出顺序决定参考和权重顺序"],
        ["MV limits", "0–1；applied 0.02–0.95", "理论约束与实际开关安全限幅"],
        ["MV rate", "±0.05/update", "限制每 10 ms duty 跳变；改周期需按秒率重算"],
        ["Weights normal", "[0.75,0.15,0]", "优先 Vdc，其次 Vpv"],
        ["Weights Mode III", "[0.75,0,0.15]", "限发时关注 Vdc 和 iLb≈0"],
        ["SoC guard", "0.1995–0.9005", "避免数值误差频繁撞约束"],
    ], [2.10, 1.90, 2.50], font_size=7.4)
    heading(doc, "12.2 Simulink Nonlinear MPC Controller 端口", 2)
    table(doc, ["端口", "接入", "形状/顺序"], [
        ["x", "filtered x_state", "6×1 [Vpv,iLpv,Vb,iLb,SoC,Vdc]"],
        ["ref", "Constant/Mux", "1×3 [600,Vpv_ref,0]"],
        ["last_mv", "Unit Delay", "2×1，IC=u0，Ts=0.01"],
        ["md", "ARIMA forecast", "1×3 [T,G,RLoad]"],
        ["y.wt", "Mode weight switch", "1×3 normal/curtail"],
        ["mv", "后续 duty 安全层", "2×1 [dpv,db]"],
        ["status", "记录/诊断", "优化状态，不要丢弃"],
    ], [1.10, 2.70, 2.70], font_size=8.0)
    heading(doc, "12.3 Mode III 约束加强层", 2)
    paragraph(doc, "现有模型在 Mode III 对上一时刻 duty 加小修正：dpv 增量与 600−Vdc 相关，db 增量与 −iLb 相关，然后限幅 0.02–0.95。它用于抑制 horizon=1 的局部解抖动；若扩大预测域，可重新评估是否保留。")
    stage_checkpoint(doc, ["validateFcns 无维度/状态函数错误。", "状态、输出、ref、MD 顺序逐项一致。", "MPC mv 在 0–1，最终 applied duty 在 0.02–0.95。"])
    stage_faults(doc, [
        ["MPC 一开始饱和", "滤波初值为 0、x 顺序错或物理初值未稳定", "先开环达到稳态；给滤波器/last_mv 正确 IC。"],
        ["Vdc 好但 Vpv 发散", "Output/ref/weight 顺序错", "确认 y=[Vdc,Vpv,iLb]，ref=[600,Vpv_ref,0]。"],
        ["优化超时", "开关模型被放入 StateFcn 或 horizon 太大", "StateFcn 只使用平均模型；先保持 horizon=1。"],
    ])

    # Stage 13
    stage_header(doc, "13", "把 duty 接到实际 PWM 与开关级对象",
                 "这是数学控制输出变成门极脉冲的边界。正确的限幅、速率转换、PWM 和故障封锁决定闭环是否安全。",
                 "MPC/PI selector 输出经安全层后驱动 PV PWM 和电池半桥，控制速率与功率速率清晰分离。")
    heading(doc, "13.1 duty 安全链", 2)
    table(doc, ["顺序", "模块", "设置", "作用"], [
        ["1", "Manual Switch / ControllerSelect", "MPC 或 PI", "选择控制器，不并联两个输出"],
        ["2", "Saturation ×2", "0.02–0.95", "避免 0/1 极限造成无换相和数值问题"],
        ["3", "Rate Limiter 或 MPC MV Rate", "±0.05/10 ms", "限制命令突变"],
        ["4", "Rate Transition", "slow→fast；out Ts=1e-6", "把 10 ms duty 稳定送到 PWM"],
        ["5", "PWM Generator", "10 kHz/1 µs", "dpv 产生 PV 门极；db 产生下管原始门极"],
        ["6", "NOT+deadtime+Mux+double", "按阶段 5", "产生电池上下管"],
        ["7", "BL fault logic", "正常 0；故障 1", "过压/过流/SoC 极限时封锁"],
    ], [0.55, 1.90, 1.65, 2.40], font_size=7.8)
    heading(doc, "13.2 故障封锁建议", 2)
    paragraph(doc, "用 Compare To Constant 与 OR 生成 fault：Vdc>660 V、|iLpv|/|iLb|超过工程上限、SoC 越界或优化器严重失败。fault=1 接两个 Converter 的 BL。首次调试阈值可保守，但不要把正常纹波误判为故障。")
    heading(doc, "13.3 闭环接入顺序", 2)
    number(doc, "只闭合 PV 电压环，电池 duty 固定，检查 Vpv 跟踪。")
    number(doc, "只闭合电池 Vdc 环，PV duty 固定，检查 Vdc 跟踪。")
    number(doc, "两环均闭合但禁用模式切换，运行恒定 G/T/R。")
    number(doc, "启用 P&O，再启用 PMS，最后启用 ARIMA 和论文场景阶跃。")
    stage_checkpoint(doc, ["duty 每 10 ms 更新，gate 每 100 µs 切换。", "控制器 selector 任何时刻只选一组 duty。", "BL 正常为 0，故障测试可在一个功率步长内封锁门极。"])
    stage_faults(doc, [
        ["duty 波形直接是 10 kHz 方波", "把 gate 当 duty 回馈", "MPC last_mv 和日志用限幅前/后的平均 duty，不用门极。"],
        ["控制器更新时出现窄脉冲", "slow/fast 直接相连", "PWM 前加 Rate Transition，保持 duty 整个控制周期稳定。"],
        ["BL 一直封锁", "故障逻辑极性反或阈值单位错", "正常状态 BL=0；逐项显示 fault 条件。"],
    ])

    # Stage 14
    stage_header(doc, "14", "PI 基线控制器与抗饱和",
                 "PI 不是最终算法，而是判断主电路/测量/符号是否正确的低复杂度基线。若 PI 也发散，通常先查对象而不是怪 MPC。",
                 "PV 电压 PI、母线 PI、限幅和抗饱和可单独闭环；与 MPC 使用同一主电路、PWM 和测量。")
    add_picture(doc, ASSETS / "pi_controller.png", 6.1, "图 11  现有 PI 基线可复用，但必须加入开关模型需要的限幅和抗饱和")
    heading(doc, "14.1 两个误差的符号", 2)
    table(doc, ["环", "误差", "当前增益起点", "为什么这样定义"], [
        ["PV", "e_pv=Vpv−Vpv_ref", "Kp=0.002, Ki=0.20", "dpv 增大通常使 Vpv 降低；Vpv 过高时应增大 dpv"],
        ["DC", "e_dc=600−Vdc", "Kp=0.001, Ki=0.05", "当前实现基于平均模型；接物理半桥后必须用小阶跃验证 db 方向"],
    ], [0.80, 2.25, 1.65, 1.80], font_size=7.8)
    heading(doc, "14.2 离散积分和抗饱和", 2)
    paragraph(doc, "积分器用 Unit Delay：I[k]=I[k−1]+Ki·Ts_ctrl·e[k]。输出 duty=bias+Kp·e+I。Saturation 后将 (duty_sat−duty_unsat) 乘抗饱和系数反馈到积分状态，避免长时间饱和后恢复缓慢。")
    heading(doc, "14.3 调试顺序", 2)
    number(doc, "Ki=0，只调 Kp 到响应方向正确且无持续振荡。")
    number(doc, "逐步增加 Ki 消除稳态误差；每次只改一个环。")
    number(doc, "先恒定环境，再做小负载阶跃，最后做论文大阶跃。")
    number(doc, "开关级对象的器件压降会改变 bias，优先重新求物理 u0，不要靠积分器长期补大偏差。")
    stage_checkpoint(doc, ["两个 PI 单环小阶跃方向正确。", "duty 触及限幅时积分状态不继续无界增长。", "MPC/PI 对比使用完全相同的主电路、滤波和场景。"])
    stage_faults(doc, [
        ["PI 和 MPC 都发散", "对象方向/初值/测量错误", "退回阶段 8 开环；不要继续调权重。"],
        ["PI 长时间卡限幅", "无抗饱和或 bias 不对", "加入 back-calculation；用物理稳态重新求 bias。"],
        ["PV 环正反馈", "误差写成 Vref−Vpv 但 Kp 正", "按 e=Vpv−Vref，或同时反转增益符号。"],
    ])

    # Stage 15
    stage_header(doc, "15", "脚本化运行、论文场景与数据记录",
                 "长时间开关仿真必须可重复、可切换控制器、可限制日志；不能依赖手工点击后无法复现实验。",
                 "run_switching_scenario.m 使用 SimulationInput 运行新模型，输出平均与纹波指标。")
    heading(doc, "15.1 推荐新增脚本，而不是改坏原 run_paper_scenario.m", 2)
    paragraph(doc, "在 scripts 新建 run_switching_scenario.m。原 run_paper_scenario.m 继续跑平均模型，便于回归。")
    code(doc, """function result = run_switching_scenario(scenarioName,controllerName)
arguments
    scenarioName (1,1) string = "comparison"
    controllerName (1,1) string {mustBeMember(controllerName,["MPC","PI"])} = "MPC"
end

projectRoot = fileparts(fileparts(mfilename('fullpath')));
model = "PV_Battery_MPC_Switching";
addpath(fullfile(projectRoot,"scripts"),fullfile(projectRoot,"models"));

controllerSelect = double(controllerName=="MPC");
in = Simulink.SimulationInput(model);
in = in.setVariable("ScenarioName",scenarioName);
in = in.setVariable("ControllerSelect",controllerSelect);
in = in.setPreSimFcn(@(~) pvbatt_initialize());
in = in.setModelParameter("StopTime","ScenarioStopTime");
result = sim(in);
end""")
    heading(doc, "15.2 论文场景复现顺序", 2)
    table(doc, ["场景", "初始/变化", "主要验证"], [
        ["comparison", "0–1 s: T35,G500,Pload7.2kW；1 s G→1000,T→25；2 s Pload→14.4kW", "MPC/PI 的 Vdc 偏差与模式变化"],
        ["mode1", "PV 过剩、SoC 未满", "MPPT+充电"],
        ["mode2", "PV 不足、SoC>min", "电池放电"],
        ["mode3", "PV 过剩、SoC≈max", "PV 限发、iLb≈0"],
        ["mode4", "夜间、SoC>min", "电池供电"],
        ["mode5", "夜间或能源不足、SoC≤min", "负载切除/保护"],
    ], [1.15, 3.35, 2.00], font_size=7.9)
    heading(doc, "15.3 两类指标必须同时给出", 2)
    table(doc, ["平均控制性能", "开关/器件性能"], [
        ["Vdc 平均误差、VRI、超调、恢复时间", "Vdc PWM 纹波、峰值过压"],
        ["Vpv 对 Vref 跟踪、Ppv 利用率", "iLpv/iLb 峰峰纹波与峰值电流"],
        ["模式正确性、SoC 约束", "门极重叠、死区、开关频率"],
        ["MPC 与 PI 对比", "仿真耗时、日志体积、数值鲁棒性"],
    ], [3.25, 3.25])
    callout(doc, "长场景运行策略",
            "先用平均模型跑完整 3 s 验证控制逻辑；开关模型先跑 20–100 ms 局部窗口。"
            "确认后再跑完整场景，并对 raw 信号设置 decimation。否则 1 µs 步长下 3 s 会产生约 300 万个基本步。",
            fill=PALE_YELLOW)
    stage_checkpoint(doc, ["原平均模型和新开关模型有两个独立 run 脚本。", "每次运行都通过 SimulationInput 设置场景与控制器。", "结果同时包含平均性能和开关纹波指标。"])
    stage_faults(doc, [
        ["ScenarioStopTime 未定义", "PreSimFcn 未初始化工作区", "在 SimulationInput 中设置 PreSimFcn 调用 pvbatt_initialize。"],
        ["3 s 开关仿真内存不足", "全部 raw 信号逐步保存", "先短窗；设置 decimation；只记录关键器件波形。"],
        ["MPC/PI 对比不公平", "使用不同滤波/初值/限幅", "两者只在 ControllerSelect 处分支，其余路径共享。"],
    ])

    # Stage 16
    stage_header(doc, "16", "分层验证与参数敏感性",
                 "没有实测数据时，准确性来自可追溯参数、物理守恒、平均/开关一致性和敏感性，而不是只看一张“波形像论文”的图。",
                 "子系统、整机开环、闭环三层测试全部通过；关键假设有 ±变化结果。")
    heading(doc, "16.1 验证矩阵", 2)
    table(doc, ["层级", "测试", "输入", "验收"], [
        ["PV 源", "I–V/P–V", "G/T 扫描", "MPP 与额定点一致"],
        ["PV Boost", "固定母线 duty 阶跃", "0.55→0.60", "Vpv 方向正确；纹波量级正确"],
        ["Battery", "恒流充放电", "±5 A", "SoC 斜率与 Q 一致"],
        ["半桥", "互补门极与 duty 扫描", "db 0.45/0.5/0.55", "无直通；iLb 方向正确"],
        ["母线", "负载阶跃", "R 50→25 Ω", "Vdc 先下降；能量守恒"],
        ["整机开环", "固定 u0", "comparison 初值", "无大浪涌，周期平均接近 x0"],
        ["整机闭环", "论文场景", "MPC/PI", "稳定、模式正确、约束满足"],
    ], [1.05, 2.05, 1.45, 1.95], font_size=7.7)
    heading(doc, "16.2 必做敏感性", 2)
    table(doc, ["参数", "范围", "观察"], [
        ["Ts_power", "0.5/1/2 µs", "平均结果、纹波峰值、仿真时间"],
        ["Ron / Vf", "0.5×/1×/2×", "稳态 duty 偏置、效率、Vdc"],
        ["Cb ESR", "1 mΩ–0.1 Ω", "启动峰值、阻尼、损耗"],
        ["Battery Rbat", "0.04/0.08/0.12 Ω", "母线支撑、电池压降、功率限制"],
        ["传感器 fc", "200/500/1000 Hz", "纹波泄漏、相位延迟、控制抖动"],
        ["dead time", "0.5/1/2 µs", "直通裕度、平均电压误差"],
        ["PV 温度系数", "资料范围", "高温 MPP 与模式判断"],
    ], [1.55, 1.65, 3.30], font_size=8.0)
    heading(doc, "16.3 数值鲁棒性", 2)
    paragraph(doc, "同一短场景用 Ts_power=0.5/1/2 µs 重跑。控制层平均量若差异超过约 1–2%，说明步长未收敛；器件尖峰允许更敏感，但必须说明带宽和测量定义。")
    stage_checkpoint(doc, ["每个子系统都有独立开环测试结果。", "关键假设参数至少做低/中/高三点。", "步长减半后控制层关键指标变化很小。"])
    stage_faults(doc, [
        ["波形像论文但功率不守恒", "符号或记录错误", "优先修守恒，不以视觉相似替代验证。"],
        ["只在 1 µs 稳定", "模型数值依赖特定步长", "做 0.5/2 µs；检查理想器件、初值和滤波。"],
        ["参数变化后立即发散", "控制器鲁棒性不足或假设范围不物理", "先确认范围，再分别检查对象和控制器。"],
    ])

    # Stage 17
    stage_header(doc, "17", "发散排查决策树与最终解释",
                 "把“MPPT 发散”拆成物理启动、门极/拓扑、测量、平均预测偏差和优化器五类问题，按证据逐层缩小。",
                 "能够明确说出发散发生在哪一层、哪一信号最先异常、怎样复现实验和怎样修正。")
    heading(doc, "17.1 最短排查路径", 2)
    number(doc, "封锁所有门极：确认 Battery/Cpv/Cb/Cdc 初值和极性，没有自发巨大电流。")
    number(doc, "固定 duty：确认主电路开环稳定，电感纹波/功率方向正确。")
    number(doc, "用 PI 单环：若仍发散，问题在对象、符号、滤波或 PI 方向。")
    number(doc, "切回平均对象：若 MPC 在平均对象稳定、开关对象不稳定，查模型失配/采样/纹波。")
    number(doc, "最后才检查 MPC 权重、horizon、rate constraint 和优化状态。")
    heading(doc, "17.2 第一异常信号定位表", 2)
    table(doc, ["最先异常", "优先怀疑", "立即检查"], [
        ["g_upper 与 g_lower 重叠", "PWM/互补/死区", "同一 PWM、NOT、Unit Delay+AND、Mux 顺序"],
        ["iL 在第一个周期暴增", "电压极性/初值/电感单位", "t=0 电压；L=10e-3 H；BL/门极"],
        ["Vdc 慢慢单调升", "输入功率长期大于负载", "Ppv+Pbat−Pload−dE/dt；模式/占空比"],
        ["Vpv_ref 自己跑边界", "P&O 输入/符号/噪声", "filtered V/I、P 正号、ΔPΔV、deadband"],
        ["duty 首次更新就饱和", "x/ref/md 顺序或初始状态", "六状态顺序、ref=[600,Vref,0]、w=[T,G,R]"],
        ["status 显示优化失败", "约束不可行/模型失配", "SoC/状态边界、rate limit、物理初值"],
    ], [1.70, 2.10, 2.70], font_size=7.8)
    heading(doc, "17.3 本复现能声明什么，不能声明什么", 2)
    table(doc, ["可以声明", "不能直接声明"], [
        ["论文拓扑、主要方程、模式、MPC 结构与额定参数已实现", "数值结果与作者代码一一相同"],
        ["开关级模型能观察 PWM、电感纹波、器件压降和启动浪涌", "论文未公开的电池 E0/K/A/B、ARIMA 系数已被准确恢复"],
        ["平均模型与开关模型之间做了周期平均对齐", "所有硬件寄生、热效应和传感器均与真实装置一致"],
        ["假设字段可替换，并可做敏感性测试", "未校准参数产生的某个峰值就是论文真实峰值"],
    ], [3.25, 3.25], font_size=8.1)
    callout(doc, "最终判断",
            "把开关主电路补回来后，平均值模型并没有失去价值：它承担 MPC 快速预测；开关模型承担物理验证。"
            "两者分工清楚，才能既复现论文控制思想，又看见实际 PWM、电感纹波和器件/初值问题。",
            fill=PALE_GREEN)
    stage_checkpoint(doc, ["发散时能找到第一个异常信号，而非只看最后 Vdc。", "平均对象、开关对象、PI、MPC 四种组合均可切换。", "论文未公开参数全部标注为假设，不混入 published 参数。"])

    # Appendices
    heading(doc, "附录 A：R2025a 模块速查表", 1, page_break=True)
    table(doc, ["界面搜索名", "程序化源路径", "本项目用途"], [
        ["powergui", "spspowerguiLib/powergui", "SPS 电气网络求解"],
        ["PV Array", "spsPVArrayLib/PV Array", "物理 PV 阵列"],
        ["Battery", "spsBatteryLib/Battery", "物理锂电池与 SoC"],
        ["Boost Converter", "spsBoostConverterLib/Boost Converter", "PV 开关升压器"],
        ["Two-Quadrant DC/DC Converter", "spsTwoQuadrantDCDCConverterLib/...", "电池双向半桥"],
        ["Series RLC Branch", "spsSeriesRLCBranchLib/Series RLC Branch", "L/rL、带 ESR 的 Cb"],
        ["Parallel RLC Branch", "spsParallelRLCBranchLib/Parallel RLC Branch", "Cpv、Cdc"],
        ["Variable Resistor", "spsVariableResistorLib/Variable Resistor", "RLoad(t)"],
        ["Voltage Measurement", "spsVoltageMeasurementLib/Voltage Measurement", "Vpv/Vb/Vdc"],
        ["Current Measurement", "spsCurrentMeasurementLib/Current Measurement", "iLpv/iLb/支路电流"],
        ["PWM Generator (DC-DC)", "spsPWMGeneratorDCDCLib/...", "10 kHz 门极"],
        ["Ground", "spsGroundLib/Ground", "负母线参考"],
    ], [2.30, 2.65, 1.55], font_size=7.7)
    callout(doc, "找不到模块时",
            "先在 Library Browser 直接搜索界面名称；本手册的程序化源路径已在本机 R2025a 查询。不同版本若库层级变化，搜索名称通常仍有效。",
            fill=PALE_YELLOW)

    heading(doc, "附录 B：原脚本的保留、修改与新增清单", 1, page_break=True)
    table(doc, ["文件", "动作", "具体变化"], [
        ["pvbatt_parameters.m", "修改（追加）", "新增 Ts_power、Tsw、deadtime、器件、ESR、滤波参数；不删除 p.sim.sampleTime"],
        ["pvbatt_generate_pv_lookup.m", "修改 3 行", "增加 nargin 默认 p，修复直接调用报错"],
        ["pvbatt_initialize.m", "修改（追加）", "向 base workspace 发布 PWR_TS/PWM_FSW/CB_ESR 等变量"],
        ["pvbatt_operating_point.m", "保留", "继续生成预测模型 x0/u0；物理模型需要二次一致性检查/trim"],
        ["pvbatt_state_derivatives.m", "保留", "仍是 nlMPC 平均预测核心"],
        ["pvbatt_state_transition.m", "保留", "仍以 RK4 子步预测，不接 PWM"],
        ["run_paper_scenario.m", "保留", "继续运行平均模型基准"],
        ["run_switching_scenario.m", "新增", "用 SimulationInput 运行开关级模型"],
        ["PV_Battery_MPC_Switching.slx", "新增", "主电路、控制电路和预测模型分层集成"],
    ], [2.35, 1.20, 2.95], font_size=7.8)

    heading(doc, "附录 C：逐阶段完成清单", 1, page_break=True)
    checklist = [
        "阶段 0：接口、单位、状态顺序和正方向冻结",
        "阶段 1：三层时间尺度、powergui、ode3 设置完成",
        "阶段 2：查表可运行，物理 PV Array 与查表 MPP 对齐",
        "阶段 3：PV Boost 开关波形和纹波验证",
        "阶段 4：Battery/SoC/Cb ESR 与初值验证",
        "阶段 5：半桥互补门极、死区、double 向量验证",
        "阶段 6：Cdc/Variable Resistor/功率守恒验证",
        "阶段 7：raw/filtered/Rate Transition/状态总线验证",
        "阶段 8：预充电、一致初值、整机开环通过",
        "阶段 9：平均预测模型与开关周期平均对齐",
        "阶段 10：P&O 参考电压收敛",
        "阶段 11：PMS 模式与 ARIMA 顺序正确",
        "阶段 12：nlMPC 对象和端口通过验证",
        "阶段 13：duty 安全链、PWM、BL 故障层完成",
        "阶段 14：PI 单环/双环稳定并有抗饱和",
        "阶段 15：脚本化场景和日志完成",
        "阶段 16：三层验证和敏感性完成",
        "阶段 17：发散定位和论文边界说明完成",
    ]
    for item in checklist:
        bullet(doc, "□ " + item)

    heading(doc, "附录 D：第一次实际操作的最短命令", 1, page_break=True)
    code(doc, """cd('D:\\PV_MPPT\\PV_Battery_MPC_Project')
openProject(pwd)
addpath(fullfile(pwd,'scripts'),fullfile(pwd,'models'))

% 1) 初始化并确认平均模型仍能工作
START_HERE
baseline = run_paper_scenario("comparison","MPC");

% 2) 生成/检查 PV 查表
p = pvbatt_parameters();
[Vbp,Tbp,Gbp,Itable] = pvbatt_generate_pv_lookup(p);
size(Itable)

% 3) 求比较场景目标工作点
pvbatt_initialize;
[x0,u0,w0,op] = pvbatt_operating_point(PVBATT_P,0.8,35,500,7200);
disp(x0); disp(u0); disp(w0);

% 4) 开关模型建好后先短场景
sw = run_switching_scenario("comparison","PI");""")
    callout(doc, "建议的学习顺序",
            "不要从完整 MPC 闭环开始。先做 PV Boost 开环，再做电池半桥开环，再做母线功率守恒；"
            "PI 稳定后才接 MPC。任何一步失败都退回最近通过的检查点。",
            fill=PALE_GREEN)

    QA.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_RAW)
    print(OUT_RAW)


if __name__ == "__main__":
    build_manual()
