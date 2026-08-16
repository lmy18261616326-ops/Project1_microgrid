from pathlib import Path
from datetime import date
import sys

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\PV_MPPT")
DOCS = ROOT / "docs"
QA = DOCS / "_qa" / "full_paper_manual"
OUT = DOCS / "PV_Battery_MPC_Full_Paper_Reproduction_Manual.docx"

sys.path.insert(0, str(DOCS))
import build_phase2_manual as base


BLUE = base.BLUE
DARK_BLUE = base.DARK_BLUE
MID_BLUE = base.MID_BLUE
PALE_BLUE = base.PALE_BLUE
LIGHT_BLUE = base.LIGHT_BLUE
PALE_GREEN = base.PALE_GREEN
PALE_YELLOW = base.PALE_YELLOW
PALE_RED = base.PALE_RED
GRAY = base.GRAY
LIGHT_GRAY = base.LIGHT_GRAY
WHITE = base.WHITE
BLACK = base.BLACK

ARCH_DIAGRAM = QA / "full_architecture.png"
ROADMAP_DIAGRAM = QA / "reproduction_roadmap.png"
MODE_DIAGRAM = QA / "mode_matrix.png"
MPC_DIAGRAM = QA / "mpc_receding_horizon.png"


def configure_document(doc):
    base.configure_document(doc)
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("PV—BATTERY MPC  |  整篇文献复现")
    base.set_run_font(r)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r2 = p.add_run("    原理 · 搭建 · 验证 · 排错")
    base.set_run_font(r2)
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string(GRAY)
    p_pr = p._p.get_or_add_pPr()
    old = p_pr.find(qn("w:pBdr"))
    if old is not None:
        p_pr.remove(old)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), MID_BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    doc.core_properties.title = "PV—电池系统预测控制整篇文献复现说明书"
    doc.core_properties.subject = "论文意义、控制原理、Simulink逐步搭建、验证与故障排查"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = (
        "Layout preset: compact_reference_guide; "
        "header pattern: editorial_cover; source summarized and paraphrased."
    )


def pil_color(value):
    return base.pil_color(value)


def draw_box(draw, rect, title, subtitle="", fill=WHITE, outline=BLUE):
    base.draw_box(draw, rect, title, subtitle, fill, outline)


def draw_arrow(draw, start, end, color=BLUE, width=5):
    base.draw_arrow(draw, start, end, color, width)


def create_architecture_diagram():
    QA.mkdir(parents=True, exist_ok=True)
    im = Image.new("RGB", (2400, 1250), "white")
    d = ImageDraw.Draw(im)
    d.text((80, 45), "独立 PV—电池直流微网：能量流与控制信息流", font=base.font(42, True), fill=pil_color(DARK_BLUE))
    d.text((80, 105), "实线蓝箭头为控制/测量，粗绿色箭头为能量流；对象采用平均值模型", font=base.font(23), fill=pil_color(GRAY))

    draw_box(d, (100, 260, 520, 470), "PV 阵列", "T、G → Ipv(vpv,T,G)", LIGHT_BLUE)
    draw_box(d, (720, 260, 1140, 470), "PV Boost", "dpv；Cpv、Lpv、rLpv", PALE_GREEN)
    draw_box(d, (1320, 260, 1730, 470), "600 V 直流母线", "Cdc 与 RLoad", PALE_YELLOW)
    draw_box(d, (1900, 220, 2300, 430), "负载", "PLoad = vdc²/RLoad", LIGHT_GRAY)
    draw_box(d, (1320, 620, 1730, 830), "双向电池变换器", "db；Cb、Lb、rLb", PALE_GREEN)
    draw_box(d, (1900, 620, 2300, 830), "锂电池", "OCV(SoC)、Rbat、20 Ah", LIGHT_BLUE)

    draw_arrow(d, (520, 365), (720, 365), PALE_GREEN, 9)
    draw_arrow(d, (1140, 365), (1320, 365), PALE_GREEN, 9)
    draw_arrow(d, (1730, 325), (1900, 325), PALE_GREEN, 9)
    draw_arrow(d, (1900, 725), (1730, 725), PALE_GREEN, 9)
    draw_arrow(d, (1525, 620), (1525, 470), PALE_GREEN, 9)

    draw_box(d, (300, 930, 790, 1120), "P&O + 功率管理", "Vpv_ref、α、模式号", PALE_RED)
    draw_box(d, (980, 930, 1480, 1120), "ARIMA 一步预测", "T、G、RLoad 的预测值", PALE_BLUE)
    draw_box(d, (1670, 930, 2210, 1120), "非线性 MPC", "x、ŵ、参考 → [dpv,db]", PALE_YELLOW)
    draw_arrow(d, (790, 1025), (1670, 1025))
    draw_arrow(d, (1480, 1025), (1670, 1025))
    draw_arrow(d, (1940, 930), (1010, 470), BLUE, 5)
    draw_arrow(d, (2010, 930), (1525, 830), BLUE, 5)
    d.line([(1525, 470), (1525, 880), (520, 880), (520, 930)], fill=pil_color(GRAY), width=4)
    draw_arrow(d, (520, 880), (520, 930), GRAY, 4)
    d.text((610, 842), "x=[vpv,iLpv,vb,iLb,SoC,vdc]；aux=[Ipv,Ib,Ebat,Ppv,Pbat,Pload]", font=base.font(22), fill=pil_color(GRAY))
    im.save(ARCH_DIAGRAM, quality=95)


def create_roadmap_diagram():
    QA.mkdir(parents=True, exist_ok=True)
    im = Image.new("RGB", (2400, 1380), "white")
    d = ImageDraw.Draw(im)
    d.text((80, 45), "推荐复现顺序：每一阶段都先独立验收，再接入下一阶段", font=base.font(40, True), fill=pil_color(DARK_BLUE))
    stages = [
        ("1", "参数与求解器", "统一变量、10 μs 对象步长"),
        ("2", "PV 静态模型", "I–V/P–V 曲线与查表"),
        ("3", "PV Boost", "vpv、iLpv 两个动态状态"),
        ("4", "电池与双向变换器", "vb、iLb、SoC 三个状态"),
        ("5", "直流母线", "vdc、负载、六状态闭合"),
        ("6", "一致初值", "求工作点，消除启动冲击"),
        ("7", "P&O", "每 10 ms 生成 Vpv_ref"),
        ("8", "模式逻辑", "充电/放电/限功率/夜间"),
        ("9", "ARIMA", "一步扰动预测"),
        ("10", "非线性 MPC", "参考、权重、约束、占空比"),
        ("11", "PI 基线", "同工况对比"),
        ("12", "论文场景", "四模式、VRI、300 s 数据"),
    ]
    left = 120
    top = 180
    w = 600
    h = 155
    xgap = 165
    ygap = 85
    colors = [LIGHT_BLUE, PALE_GREEN, PALE_YELLOW, PALE_RED]
    positions = []
    for i, (num, title, sub) in enumerate(stages):
        row = i // 3
        col = i % 3 if row % 2 == 0 else 2 - (i % 3)
        x1 = left + col * (w + xgap)
        y1 = top + row * (h + ygap)
        positions.append((x1, y1, x1 + w, y1 + h))
        draw_box(d, positions[-1], f"阶段 {num}　{title}", sub, colors[row % len(colors)])
    for i in range(len(positions) - 1):
        a = positions[i]
        b = positions[i + 1]
        if abs(a[1] - b[1]) < 5:
            start = (a[2], (a[1] + a[3]) // 2) if b[0] > a[0] else (a[0], (a[1] + a[3]) // 2)
            end = (b[0], (b[1] + b[3]) // 2) if b[0] > a[0] else (b[2], (b[1] + b[3]) // 2)
        else:
            start = ((a[0] + a[2]) // 2, a[3])
            end = ((b[0] + b[2]) // 2, b[1])
        draw_arrow(d, start, end, BLUE, 5)
    d.rounded_rectangle((120, 1195, 2280, 1315), radius=15, fill="#F7F9FC", outline=pil_color(MID_BLUE), width=3)
    d.text((160, 1228), "不要先搭 MPC 再找对象问题：若静态模型、符号或初值错误，优化器只会更快地把占空比推到边界。", font=base.font(25, True), fill=pil_color(DARK_BLUE))
    im.save(ROADMAP_DIAGRAM, quality=95)


def create_mode_diagram():
    QA.mkdir(parents=True, exist_ok=True)
    im = Image.new("RGB", (2300, 1280), "white")
    d = ImageDraw.Draw(im)
    d.text((80, 45), "功率管理模式矩阵", font=base.font(42, True), fill=pil_color(DARK_BLUE))
    d.text((80, 105), "先判断白天/夜间，再比较 Ppv 与 Pload，最后检查 SoC 边界", font=base.font(24), fill=pil_color(GRAY))
    cards = [
        ((120, 220, 1080, 620), "模式 I：PV 有富余，电池可充电", "条件：Ppv ≥ Pload，SoC < 0.90，G > Gmin\nPV：MPPT　电池：吸收多余功率\nMPC 权重：跟踪 600 V + Vpv_ref", PALE_GREEN),
        ((1220, 220, 2180, 620), "模式 II：PV 不足，电池放电", "条件：Ppv < Pload，SoC > 0.20，G > Gmin\nPV：MPPT　电池：补足功率缺口\nMPC 权重：跟踪 600 V + Vpv_ref", LIGHT_BLUE),
        ((120, 730, 1080, 1130), "模式 III：电池满，PV 限功率", "条件：Ppv ≥ Pload，SoC ≥ 0.90，G > Gmin\n电池：Ib → 0　PV：离开 MPP 调节母线\nα=1；权重切换为 600 V + 电池电流", PALE_RED),
        ((1220, 730, 2180, 1130), "模式 IV：无光，电池独供", "条件：G ≤ Gmin，SoC > 0.20\nPV：不可用　电池：维持 600 V\nSoC 到下限后应进入负载切除（论文未展开）", PALE_YELLOW),
    ]
    for rect, title, body, fill in cards:
        d.rounded_rectangle(rect, radius=18, fill=pil_color(fill), outline=pil_color(BLUE), width=4)
        d.text((rect[0] + 35, rect[1] + 35), title, font=base.font(29, True), fill=pil_color(DARK_BLUE))
        y = rect[1] + 105
        for line in body.split("\n"):
            d.text((rect[0] + 35, y), line, font=base.font(23), fill=pil_color(BLACK))
            y += 58
    im.save(MODE_DIAGRAM, quality=95)


def create_mpc_diagram():
    QA.mkdir(parents=True, exist_ok=True)
    im = Image.new("RGB", (2400, 1160), "white")
    d = ImageDraw.Draw(im)
    d.text((80, 45), "非线性 MPC 的一次控制循环（每 10 ms 重复）", font=base.font(42, True), fill=pil_color(DARK_BLUE))
    boxes = [
        ((110, 300, 480, 520), "① 采样", "x(k)、w(k)"),
        ((650, 300, 1040, 520), "② 预测", "ARIMA 得到 ŵ(k+1)"),
        ((1210, 300, 1620, 520), "③ 求解", "动态模型 + 代价 + 约束"),
        ((1790, 300, 2260, 520), "④ 只施加第一步", "u*(k)=[dpv,db]"),
    ]
    fills = [LIGHT_BLUE, PALE_BLUE, PALE_YELLOW, PALE_GREEN]
    for (rect, title, sub), fill in zip(boxes, fills):
        draw_box(d, rect, title, sub, fill)
    for i in range(3):
        draw_arrow(d, (boxes[i][0][2], 410), (boxes[i + 1][0][0], 410), BLUE, 6)
    d.line([(2025, 520), (2025, 760), (300, 760), (300, 520)], fill=pil_color(GRAY), width=5)
    draw_arrow(d, (300, 760), (300, 520), GRAY, 5)
    d.text((700, 710), "对象演化到 k+1，反馈新状态，然后重新优化", font=base.font(25, True), fill=pil_color(GRAY))
    d.rounded_rectangle((130, 850, 2270, 1050), radius=15, fill="#F7F9FC", outline=pil_color(MID_BLUE), width=3)
    d.text((170, 880), "当前工程：预测时域 Np=1、控制时域 Nc=1；对象控制周期 10 ms；预测模型内部 RK4 分 10 个子步。", font=base.font(24), fill=pil_color(DARK_BLUE))
    d.text((170, 935), "意义：实现快速、计算量较小；代价是前瞻性有限，ARIMA 只影响下一步，无法体现长时域能量调度。", font=base.font(24), fill=pil_color(DARK_BLUE))
    im.save(MPC_DIAGRAM, quality=95)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(45)
    r = p.add_run("SIMULINK 文献复现总手册")
    base.set_run_font(r)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MID_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("负载与环境不确定性下")
    base.set_run_font(r)
    r.bold = True
    r.font.size = Pt(27)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("光伏—电池系统的预测控制")
    base.set_run_font(r)
    r.bold = True
    r.font.size = Pt(27)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("内容意义、核心原理、逐步搭建、结果解释与问题排查")
    base.set_run_font(r)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(4.35)
    rows = [
        ("主文献", "Batiyah 等，Energies 2022, 15, 4100"),
        ("复现对象", "9.5 kW PV、20 Ah 电池、600 V 独立直流母线"),
        ("核心方法", "P&O + 四模式功率管理 + ARIMA + 非线性 MPC"),
        ("对应工程", r"D:\PV_MPPT\PV_Battery_MPC_Project"),
        ("文档版本", f"V1.0　{date(2026, 7, 31).isoformat()}"),
    ]
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        base.prevent_row_split(row)
        c0, c1 = row.cells
        base.set_cell_shading(c0, PALE_BLUE)
        base.set_cell_shading(c1, "F8FAFC")
        for cell in (c0, c1):
            base.set_cell_margins(cell, top=105, start=150, bottom=105, end=150)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r0 = c0.paragraphs[0].add_run(label)
        base.set_run_font(r0)
        r0.bold = True
        r1 = c1.paragraphs[0].add_run(value)
        base.set_run_font(r1)

    doc.add_paragraph()
    base.add_callout(
        doc,
        "本手册的定位",
        "它不是对论文逐句翻译，而是把论文的研究问题、公式、控制思想和实验安排转换成可执行的 Simulink 搭建步骤。论文未公开的参数会明确标记为“复现假设”，避免把工程近似误认为论文原值。",
        fill=PALE_YELLOW,
        border="BF9000",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("建议先通读第 1～4 章，再从第 5 章按阶段动手。")
    base.set_run_font(r)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()


def add_stage_intro(doc, purpose, meaning, deliverable):
    base.add_table(
        doc,
        ["本阶段做什么", "为什么要做", "完成后得到什么"],
        [(purpose, meaning, deliverable)],
        widths=[2.10, 2.25, 2.15],
        font_size=8.8,
        header_fill=PALE_BLUE,
    )


def add_acceptance(doc, items):
    base.add_heading(doc, "", "完成判据", level=3)
    for item in items:
        base.add_bullet(doc, item)


def add_problems(doc, rows):
    base.add_heading(doc, "", "常见问题与处理", level=3)
    base.add_table(
        doc,
        ["现象", "原因", "处理"],
        rows,
        widths=[1.65, 2.20, 2.65],
        font_size=7.9,
        header_fill=PALE_RED,
    )


def add_equation(doc, text):
    p = doc.add_paragraph(style="Code Block")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    base.set_run_font(r, east_asia="Microsoft YaHei", ascii_font="Consolas")
    r.font.size = Pt(9.2)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)


def add_numbered_list(doc, items):
    """Add a real Word numbered list that restarts at 1 for each group."""
    numbering = doc.part.numbering_part.element
    base_num_id = str(doc.styles["List Number"].element.pPr.numPr.numId.val)
    abstract_num_id = None
    existing_ids = []
    for num in numbering.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        if num_id is not None:
            existing_ids.append(int(num_id))
        if num_id == base_num_id:
            abstract = num.find(qn("w:abstractNumId"))
            abstract_num_id = abstract.get(qn("w:val"))
    if abstract_num_id is None:
        raise RuntimeError("Unable to resolve the List Number abstract numbering definition.")

    new_num_id = max(existing_ids, default=0) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    new_num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    new_num.append(override)
    numbering.append(new_num)

    for item in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = num_pr.find(qn("w:ilvl"))
        if ilvl is None:
            ilvl = OxmlElement("w:ilvl")
            num_pr.append(ilvl)
        ilvl.set(qn("w:val"), "0")
        num_id = num_pr.find(qn("w:numId"))
        if num_id is None:
            num_id = OxmlElement("w:numId")
            num_pr.append(num_id)
        num_id.set(qn("w:val"), str(new_num_id))
        r = p.add_run(item)
        base.set_run_font(r)


def build_manual():
    create_architecture_diagram()
    create_roadmap_diagram()
    create_mode_diagram()
    create_mpc_diagram()

    doc = Document()
    configure_document(doc)
    add_cover(doc)

    base.add_heading(doc, "0", "怎么使用这份说明书")
    base.add_text(doc, "这份手册同时服务于三种需求：理解论文、从空白模型手工复现、排查仿真发散。建议按下面的顺序使用。")
    add_numbered_list(doc, [
        "先读第 1～4 章，弄清论文解决的矛盾、信号定义、控制分工和工程边界。",
        "从第 5 章开始按阶段搭建。每完成一阶段，先执行该节的完成判据，不要一次把所有子系统接完。",
        "模型能运行后，用第 17 章的论文工况验证，再用第 19 章从信号链顺序排查异常。",
        "若你的目标只是修改 MPPT 周期，可直接看第 11 章；但改变周期后必须同步检查离散状态和 MPC 控制周期。",
    ])
    base.add_callout(
        doc,
        "最重要的复现原则",
        "先保证对象模型守恒、符号正确、初值一致，再调控制器。占空比打满、母线电压飞走、MPC 不收敛，很多时候并不是优化器的问题，而是对象方程、单位或初始工作点错误。",
        fill=PALE_GREEN,
        border="70AD47",
    )

    base.add_heading(doc, "0.1", "章节导航", level=2)
    base.add_table(
        doc,
        ["阅读目的", "重点章节"],
        [
            ("理解论文讲了什么、有什么意义", "第 1～4 章"),
            ("从空白 Simulink 模型开始搭建", "第 5～16 章"),
            ("理解论文图 5～14", "第 17 章"),
            ("确认当前工程与论文差异", "第 18 章"),
            ("排查发散、不收敛和模式误切换", "第 19 章"),
            ("查参数、信号顺序和文件作用", "附录 A～C"),
        ],
        widths=[3.15, 3.35],
        font_size=9,
    )

    base.add_heading(doc, "1", "整篇论文在解决什么问题")
    base.add_text(doc, "研究对象是一个没有大电网支撑的独立直流微网：光伏、锂电池和直流负载共用 600 V 直流母线。光照、温度和负载都在变化，而光伏本身没有惯性，任何功率不平衡都会立即体现在母线电容的充放电上，表现为 vdc 波动。")
    base.add_table(
        doc,
        ["矛盾", "如果不处理", "论文的处理方式"],
        [
            ("希望 PV 始终工作在最大功率点", "电池已满时仍有多余功率，母线电压上升", "SoC 达上限后退出 MPPT，主动限功率"),
            ("负载和光照突然变化", "母线电压产生较大过冲/跌落", "电池快速补偿，MPC提前利用扰动预测"),
            ("对象是非线性、多输入多输出", "多个 PI 独立调节时耦合与约束难处理", "将两个占空比放入统一非线性 MPC"),
            ("电池不能过充或过放", "寿命下降，甚至失去供电能力", "把 SoC 上下限纳入模式逻辑与约束"),
        ],
        widths=[2.05, 2.15, 2.30],
        font_size=8.4,
    )
    base.add_heading(doc, "1.1", "论文的核心贡献与意义", level=2)
    for title, text in [
        ("把功率管理和变换器动态放在同一模型中", "不是只做慢速能量调度，也不是只做单个变换器电流环，而是把 PV、电池、母线和占空比耦合起来。"),
        ("把“最大功率”变成有条件的目标", "在电池可充电时最大化 PV 利用率；电池满且功率过剩时，以系统安全为先，牺牲一部分光伏功率。"),
        ("模式切换不是另一个控制器", "通过 α 改变代价函数：正常模式跟踪 Vpv_ref，限功率模式改为压低电池电流，同时始终调节 vdc。"),
        ("引入未来扰动", "ARIMA 的意义是让优化器不只根据当前误差动作，而是利用未来温度、辐照度和负载预测计算下一步占空比。"),
    ]:
        base.add_callout(doc, title, text, fill=LIGHT_BLUE)
    base.add_heading(doc, "1.2", "论文没有完全解决的内容", level=2)
    for item in [
        "负载切除只在流程图中出现，没有给出详细控制和仿真。",
        "ARIMA 的阶次、训练数据、系数和预测误差没有公开，无法数值一比一复现。",
        "电池 E0、K、A、B、Rbat 以及 PV 阵列串并联结构没有完整给出。",
        "实验以仿真为主，没有给出实时处理器、求解时间和硬件闭环验证。",
        "预测时域设为 1，前瞻能力有限，更接近带预测扰动的一步非线性优化控制。",
    ]:
        base.add_bullet(doc, item)

    base.add_heading(doc, "2", "系统结构、能量流和信号定义")
    doc.add_picture(str(ARCH_DIAGRAM), width=Inches(6.55))
    base.add_caption(doc, "图 1　依据论文与当前工程重绘的系统与控制信息流")
    base.add_heading(doc, "2.1", "两个占空比各控制什么", level=2)
    base.add_table(
        doc,
        ["控制量", "物理对象", "主要作用", "方向直觉"],
        [
            ("dpv", "PV 单向 Boost", "改变 PV 端工作电压与输出功率", "理想稳态近似 vpv≈(1−dpv)vdc；dpv 增大通常使 vpv 降低"),
            ("db", "电池双向 Boost", "改变电池向母线注入/吸收的功率", "正负功率由 iLb 符号决定；必须统一充放电正方向"),
        ],
        widths=[0.75, 1.35, 2.05, 2.35],
        font_size=8.6,
    )
    base.add_heading(doc, "2.2", "向量顺序必须固定", level=2)
    base.add_table(
        doc,
        ["向量", "顺序", "单位/含义"],
        [
            ("x_state（6）", "[vpv, iLpv, vb, iLb, SoC, vdc]", "V, A, V, A, 0～1, V"),
            ("u_duty（2）", "[dpv, db]", "两个变换器占空比"),
            ("w_actual（3）", "[T, G, RLoad]", "°C, W/m², Ω"),
            ("aux（6）", "[Ipv, Ib, Ebat, Ppv, Pbat, Pload]", "A, A, V, W, W, W"),
            ("mode_bus（2）", "[α, modeID]", "α=1 仅限功率；模式号 1～5"),
        ],
        widths=[1.35, 2.55, 2.60],
        font_size=8.8,
    )
    base.add_callout(
        doc,
        "功率正负号",
        "当前工程中 Pbat>0 表示电池向母线放电，Pbat<0 表示电池吸收功率充电；SoC 导数与电池端电流 Ib 的正方向必须保持一致。画图前先确认符号，否则会把“充电”误判成“放电”。",
        fill=PALE_YELLOW,
    )

    base.add_heading(doc, "3", "数学模型的原理与物理意义")
    base.add_heading(doc, "3.1", "PV 单二极管模型", level=2)
    base.add_text(doc, "PV 电流不是一个简单常数，而是 vpv、温度 T 和辐照度 G 的非线性函数。辐照度主要改变短路电流和最大功率，温度主要改变开路电压。论文用单二极管模型并建议 Newton–Raphson 求隐式电流；当前工程先离线生成三维电流表，仿真时查表，以提高稳定性和速度。")
    add_equation(doc, "Ipv = Iph − Io{exp[(vpv + Ipv·Rs)/(a·Ns·Vt)] − 1} − (vpv + Ipv·Rs)/Rsh")
    base.add_text(doc, "工程意义：这个方程决定 P–V 曲线的形状，也决定 P&O 能否看到正确的最大功率点。若静态 PV 曲线不对，后续任何 MPPT 或 MPC 调参都没有意义。")

    base.add_heading(doc, "3.2", "PV Boost 平均值动态", level=2)
    add_equation(doc, "dvpv/dt = (Ipv − iLpv)/Cpv")
    add_equation(doc, "diLpv/dt = [−rLpv·iLpv + vpv − (1−dpv)·vdc]/Lpv")
    base.add_text(doc, "第一式是 PV 端电容的电流平衡；第二式是电感电压平衡。采用平均值模型意味着不显式产生 10 kHz PWM 纹波，只保留一个开关周期的平均效应。")

    base.add_heading(doc, "3.3", "电池、SoC 与双向变换器", level=2)
    add_equation(doc, "Ib = (Ebat(SoC) − vb)/Rbat；dSoC/dt = −Ib/(3600·Q)")
    add_equation(doc, "dvb/dt = (Ib − iLb)/Cb")
    add_equation(doc, "diLb/dt = [−rLb·iLb + vb − (1−db)·vdc]/Lb")
    base.add_text(doc, "Ebat 是开路电压，vb 是电池端电容电压，iLb 是变换器电感电流。电池容量 Q 用 Ah 给出，所以积分时必须乘 3600 转成 A·s。")
    base.add_callout(
        doc,
        "论文公式中的一个重要疑点",
        "论文式 (12) 的 SoC 行写成了 Ebat−x4，但按状态定义 x4=iLbat，单位不相容；应使用电池端电压 x3，或直接使用 Ib。当前工程采用 dSoC/dt=−Ib/(3600Q)。自己搭建时不要机械照抄该处。",
        fill=PALE_RED,
        border="C00000",
    )

    base.add_heading(doc, "3.4", "直流母线与功率平衡", level=2)
    add_equation(doc, "dvdc/dt = [(1−dpv)iLpv + (1−db)iLb − vdc/RLoad]/Cdc")
    base.add_text(doc, "母线电压本质上是功率不平衡的积分结果：进入母线的电流大于负载电流，Cdc 被充电，vdc 上升；反之下降。因此母线电压是判断功率管理是否成功的核心信号。")

    base.add_heading(doc, "4", "从论文到工程的完整复现路线")
    doc.add_picture(str(ROADMAP_DIAGRAM), width=Inches(6.55))
    base.add_caption(doc, "图 2　推荐的十二阶段复现顺序")
    base.add_table(
        doc,
        ["阶段", "保存点建议", "不要进入下一阶段的条件"],
        [
            ("1～2 参数/PV", "PV_Static_Check.slx 或独立脚本", "I–V/P–V 曲线无最大值、出现 NaN"),
            ("3～5 六状态对象", "PV_Battery_Plant.slx", "恒定占空比下状态发散或功率不守恒"),
            ("6 工作点", "记录 x0、u0、w0", "t=0 出现巨大冲击，初始母线非 600 V"),
            ("7～9 离散监督层", "MPPT_PMS_Forecast.slx", "采样周期不一致、模式互相重叠"),
            ("10～11 控制器", "MPC_PI_Comparison.slx", "MPC 不可行、PI 单独都无法稳定"),
            ("12 场景验证", "最终 PV_Battery_MPC.slx", "没有逐模式验收就直接跑 300 s"),
        ],
        widths=[1.25, 2.30, 2.95],
        font_size=8.4,
    )
    base.add_callout(
        doc,
        "当前可直接参考的工程",
        r"模型：D:\PV_MPPT\PV_Battery_MPC_Project\models\PV_Battery_MPC.slx；脚本：D:\PV_MPPT\PV_Battery_MPC_Project\scripts。建议自己搭建时新建模型，不要覆盖已验证工程。",
        fill=LIGHT_BLUE,
    )

    doc.add_page_break()
    base.add_heading(doc, "5", "阶段 1：建立工程、参数和求解器")
    add_stage_intro(doc, "创建工程目录、初始化脚本和空白模型。", "所有子系统引用同一组变量，避免单位和数值散落在模块里。", "一个能加载参数、能编译的空白顶层模型。")
    base.add_heading(doc, "5.1", "文件结构", level=2)
    base.add_table(
        doc,
        ["文件/目录", "作用"],
        [
            ("models/PV_Battery_MPC.slx", "最终顶层模型"),
            ("scripts/pvbatt_parameters.m", "集中保存论文参数与复现假设"),
            ("scripts/pvbatt_initialize.m", "生成查表、工况、工作点和 nlmpc 对象"),
            ("scripts/pvbatt_*model*.m", "PV、电池和六状态预测函数"),
            ("scripts/pvbatt_build_scenario.m", "生成比较、四模式和 300 s 场景"),
            ("scripts/run_paper_scenario.m", "使用 SimulationInput 启动仿真"),
        ],
        widths=[2.55, 3.95],
        font_size=8.8,
    )
    base.add_heading(doc, "5.2", "模型配置参数", level=2)
    base.add_table(
        doc,
        ["配置项", "设置", "用处"],
        [
            ("Solver type", "Fixed-step", "保证对象步长、控制周期和日志对齐"),
            ("Solver", "ode4 (Runge-Kutta)", "连续平均值对象的固定步长积分"),
            ("Fixed-step size", "1e-5 s", "对应论文样本时间 10 μs"),
            ("Start time", "0", "统一场景时间轴"),
            ("Stop time", "ScenarioStopTime", "由场景脚本切换 1/3/16/300 s"),
            ("Return workspace outputs", "on", "通过 SimulationOutput 读取日志"),
        ],
        widths=[1.55, 1.70, 3.25],
        font_size=8.7,
    )
    base.add_heading(doc, "5.3", "操作步骤", level=2)
    add_numbered_list(doc, [
        "新建空白模型，命名为 PV_Battery_MPC_manual.slx；把模型和 scripts 目录加入 MATLAB 路径。",
        "复制或自行建立 pvbatt_parameters.m。所有值使用 SI 单位：F、H、Ω、V、A、s；SoC 用 0～1。",
        "在仿真前执行 pvbatt_initialize。不要把大量参数写在模型回调里，调试时应能从命令行单独运行初始化。",
        "先设置 fixed-step ode4 和 1e-5 s，再开始搭对象；后续不要让某个连续模块自行继承成不一致步长。",
    ])
    add_acceptance(doc, ["初始化后工作区存在 PVBATT_P、x0、u0、w0、scenario_w。", "模型更新图时没有“未定义变量”错误。", "所有参数单位已统一为 SI。"])
    add_problems(doc, [
        ("变量未定义", "脚本路径没加入，或初始化未执行", "先运行 pvbatt_initialize，再更新模型"),
        ("仿真极慢", "误用了开关级器件或可变步长", "本阶段使用平均值模型与 fixed-step ode4"),
        ("结果相差 1000 倍", "mH/μF/kW 未换算", "参数脚本中统一乘 1e−3、1e−6、1e3"),
    ])

    base.add_heading(doc, "6", "阶段 2：建立 PV 静态模型和三维查表")
    add_stage_intro(doc, "实现 Ipv(vpv,T,G)。", "为对象和预测器提供连续、快速、可重复的 PV 电流。", "三维查表及经过验证的 I–V/P–V 曲线。")
    base.add_heading(doc, "6.1", "推荐做法", level=2)
    base.add_text(doc, "先用脚本按照单二极管方程在电压、温度和辐照度网格上求 Ipv，再把结果保存为 PV_I_TABLE。当前工程使用 9 串 × 5 并的阵列假设，额定功率 9×5×213.15=9.59175 kW。论文只给出约 9.5 kW，没有公开阵列排布。")
    base.add_table(
        doc,
        ["模块", "关键设置", "连接"],
        [
            ("n-D Lookup Table", "Dimensions=3；Table=PV_I_TABLE", "输入 1=vpv，输入 2=T，输入 3=G"),
            ("Breakpoints 1", "PV_V_BP，严格递增", "覆盖 0 到阵列开路电压附近"),
            ("Breakpoints 2", "PV_T_BP，严格递增", "覆盖所有场景温度"),
            ("Breakpoints 3", "PV_G_BP，严格递增", "包含 0 与 1000 W/m²"),
            ("插值/外推", "线性插值；边界夹紧", "避免超出表格时产生异常电流"),
        ],
        widths=[1.55, 2.45, 2.50],
        font_size=8.5,
    )
    base.add_heading(doc, "6.2", "检查曲线", level=2)
    add_numbered_list(doc, [
        "在 T=25°C、G=1000 W/m² 下扫 vpv，绘制 I–V 与 P–V。",
        "阵列最大功率应接近 9.59 kW，最大功率电压约 9×29=261 V。",
        "G 降低时电流和最大功率明显降低；温度升高时开路电压通常降低。",
        "电压超过开路电压附近时，电流应接近 0，不能成为巨大负值。",
    ])
    add_acceptance(doc, ["P–V 曲线只有一个主峰。", "STC 最大功率和额定值误差在可接受范围内。", "查表在整个场景范围内不产生 NaN/Inf。"])
    add_problems(doc, [
        ("查表维度不匹配", "表格维度顺序与三个断点顺序不同", "固定为 V×T×G，并检查 size(PV_I_TABLE)"),
        ("MPP 电压错误", "模块电压与阵列电压混用", "串联数只放大电压，并联数只放大电流"),
        ("Newton 求解不收敛", "初值差或指数溢出", "离线生成表并限制指数；仿真中不迭代求根"),
    ])

    doc.add_page_break()
    base.add_heading(doc, "7", "阶段 3：搭建 PV Boost 平均值子系统")
    add_stage_intro(doc, "用 Integrator、Sum、Gain 和 Product 实现 vpv、iLpv 动态。", "把占空比对 PV 工作点的影响变成可控对象。", "输入 [dpv,w]，输出 PV 两个状态和 Ppv。")
    base.add_heading(doc, "7.1", "端口", level=2)
    base.add_table(
        doc,
        ["端口", "宽度", "内容"],
        [("u_duty", "2", "[dpv,db]，本分支取第 1 个"), ("w_actual", "3", "[T,G,RLoad]，本分支取 T、G"), ("内部输出", "—", "vpv、iLpv、Ipv、Ppv")],
        widths=[1.65, 0.80, 4.05],
        font_size=8.8,
    )
    base.add_heading(doc, "7.2", "模块与设置", level=2)
    base.add_table(
        doc,
        ["模块名", "类型/设置", "表达式或连接"],
        [
            ("Duty_Split", "Demux，2 输出", "取 dpv"),
            ("One_Minus_Dpv", "Sum，符号 +−；Constant=1", "1−dpv"),
            ("PV_Current_Lookup", "3-D Lookup Table", "Ipv(vpv,T,G)"),
            ("PV_Capacitor_KCL", "Sum，+−", "Ipv−iLpv"),
            ("Inv_Cpv", "Gain=1/Cpv", "得到 dvpv/dt"),
            ("PV_Voltage_State", "Integrator，IC=x0(1)", "输出 vpv"),
            ("PV_Inductor_R_Drop", "Gain=rLpv", "rLpv·iLpv"),
            ("PV_Bus_Voltage_Term", "Product", "(1−dpv)·vdc"),
            ("PV_Inductor_KVL", "Sum，+−−", "vpv−rLpv·iLpv−(1−dpv)vdc"),
            ("Inv_Lpv", "Gain=1/Lpv", "得到 diLpv/dt"),
            ("PV_Inductor_Current_State", "Integrator，IC=x0(2)", "输出 iLpv"),
            ("PV_Source_Power", "Product", "Ppv=vpv·Ipv"),
        ],
        widths=[2.00, 1.90, 2.60],
        font_size=7.8,
    )
    add_acceptance(doc, ["恒定 T、G、vdc、dpv 下 vpv 和 iLpv 收敛。", "Ppv 与 vpv×Ipv 完全一致。", "增大 dpv 时 vpv 的变化方向符合模型直觉。"])
    add_problems(doc, [
        ("vpv 立即变负", "KVL 中 (1−d)vdc 符号或 Sum 顺序错误", "逐项观察三路电压，确认 +−−"),
        ("iLpv 高频发散", "Lpv/Cpv 单位错或步长过大", "确认 10 mH、300 μF、1e−5 s"),
        ("Ppv 大于额定数倍", "阵列并联数重复放大", "查表输出应已经是阵列总电流"),
    ])

    base.add_heading(doc, "8", "阶段 4：搭建电池、SoC 和双向变换器")
    add_stage_intro(doc, "实现 Ebat(SoC)、Ib、vb、iLb 和 SoC 积分。", "电池是维持功率平衡和母线电压的执行器。", "能正确表现充电、放电和 SoC 变化方向的电池分支。")
    base.add_heading(doc, "8.1", "当前工程的电池近似", level=2)
    base.add_text(doc, "论文给出了较完整的充放电电压公式，但没有提供 E0、K、A、B、Rbat。当前工程因此使用可替换的 OCV–SoC 一维查表和 Rbat=0.08 Ω。这个近似适合控制结构复现，不适合声称精确复现电池端电压。")
    base.add_table(
        doc,
        ["模块名", "设置", "表达式"],
        [
            ("Battery_OCV_Lookup", "1-D Lookup；BP=BAT_SOC_BP；Table=BAT_OCV_TABLE", "Ebat(SoC)"),
            ("Battery_Voltage_Drop", "Sum，+−", "Ebat−vb"),
            ("Battery_Resistance", "Gain=1/Rbat", "Ib=(Ebat−vb)/Rbat"),
            ("Battery_Capacitor_KCL", "Sum，+−", "Ib−iLb"),
            ("Inv_Cb", "Gain=1/Cb", "dvb/dt"),
            ("Battery_Voltage_State", "Integrator，IC=x0(3)", "vb"),
            ("One_Minus_Db", "Sum，+−", "1−db"),
            ("Battery_Inductor_KVL", "Sum，+−−", "vb−rLb·iLb−(1−db)vdc"),
            ("Inv_Lb", "Gain=1/Lb", "diLb/dt"),
            ("Battery_Inductor_Current_State", "Integrator，IC=x0(4)", "iLb"),
            ("SoC_Coulomb_Counting", "Gain=−1/(3600Q)", "dSoC/dt"),
            ("Battery_SoC_State", "Integrator；IC=x0(5)；限制 0～1", "SoC"),
        ],
        widths=[2.05, 2.35, 2.10],
        font_size=7.6,
    )
    base.add_heading(doc, "8.2", "必须先做的符号试验", level=2)
    add_numbered_list(doc, [
        "令电池向母线输出正功率，检查 Ib>0、Pbat>0、SoC 下降。",
        "令电池吸收母线多余功率，检查 Ib<0、Pbat<0、SoC 上升。",
        "若上述任一方向相反，先修正符号，不要用负增益在绘图阶段补救。",
    ])
    add_acceptance(doc, ["充电与放电时 SoC 方向正确。", "SoC 始终限制在 0～1。", "电池端电压处于合理范围且没有瞬间跳成极大值。"])
    add_problems(doc, [
        ("充电时 SoC 下降", "Ib 正方向或 SoC 增益符号错", "统一“放电 Ib 为正”，SoC 增益使用负号"),
        ("vb 快速振荡", "Cb 很小且初值不一致", "用工作点 x0(3)；对象步长保持 10 μs"),
        ("电池无法满足负载", "工作点二次方程判别式小于零", "降低初始负载或更换电池参数/额定能力"),
    ])

    doc.add_page_break()
    base.add_heading(doc, "9", "阶段 5：搭建直流母线并组成六状态对象")
    add_stage_intro(doc, "把 PV、电池和负载电流汇入 Cdc。", "母线是所有能量耦合的节点，也是主控制目标。", "完整六状态平均值对象 PV_Battery_Averaged_Plant。")
    base.add_table(
        doc,
        ["模块名", "设置", "表达式"],
        [
            ("PV_Bus_Current", "Product", "(1−dpv)iLpv"),
            ("Battery_Bus_Current", "Product", "(1−db)iLb"),
            ("Load_Current", "Product 设为除法", "vdc/RLoad"),
            ("DC_Bus_KCL", "Sum，++−", "Ipv_bus+Ib_bus−Iload"),
            ("Inv_Cdc", "Gain=1/Cdc", "dvdc/dt"),
            ("DC_Bus_Voltage_State", "Integrator，IC=x0(6)", "vdc"),
            ("Load_Power", "Product", "vdc²/RLoad"),
            ("Battery_Bus_Power", "Product", "(1−db)vdc·iLb"),
            ("State_Vector", "Mux，6 输入", "[vpv,iLpv,vb,iLb,SoC,vdc]"),
            ("Aux_Vector", "Mux，6 输入", "[Ipv,Ib,Ebat,Ppv,Pbat,Pload]"),
        ],
        widths=[2.00, 2.05, 2.45],
        font_size=7.9,
    )
    base.add_heading(doc, "9.1", "功率守恒检查", level=2)
    add_equation(doc, "误差 ≈ Ppv + Pbat − Pload − d(½Cdc·vdc²)/dt − 变换器电阻损耗")
    base.add_text(doc, "稳态时母线电容储能变化接近 0，Ppv+Pbat 应接近 Pload 加损耗。瞬态时不能只用三种功率直接相等，因为 Cdc 正在吸收或释放能量。")
    add_acceptance(doc, ["六个状态的顺序与脚本完全一致。", "恒定工况下 vdc 能在工作点附近保持有限。", "稳态功率差可由电阻损耗解释。"])
    doc.add_page_break()
    add_problems(doc, [
        ("vdc 单调上升", "进入母线的两路电流都取正但电池实际在充电", "检查 iLb 与 Pbat 正方向、(1−db)乘积"),
        ("vdc 单调下降", "负载用 P/v 而场景传入的是 R，或反之", "当前 w(3) 必须是 RLoad；Iload=vdc/RLoad"),
        ("母线初始冲击很大", "x0/u0 不满足稳态方程", "进入阶段 6 求一致工作点"),
    ])

    doc.add_page_break()
    base.add_heading(doc, "10", "阶段 6：计算一致初始工作点")
    add_stage_intro(doc, "根据初始 T、G、负载和 SoC 求 x0、u0。", "避免积分器从互相矛盾的状态起步，引发假瞬态甚至数值发散。", "一组满足静态方程的六状态和两个占空比。")
    base.add_heading(doc, "10.1", "计算顺序", level=2)
    add_numbered_list(doc, [
        "固定 vdc=600 V、初始 SoC、温度、辐照度和负载功率。",
        "扫描 PV P–V 曲线找到 MPP，得到 vpv、Ipv、Ppv，并令 iLpv≈Ipv。",
        "由负载功率减去 PV 功率，得到初始电池所需功率。",
        "结合 OCV、Rbat 和电感电阻求 iLb、vb。",
        "由稳态电感方程求 dpv 和 db；最后检查两者在允许范围内。",
    ])
    base.add_heading(doc, "10.2", "当前 comparison 场景的参考初值", level=2)
    base.add_table(
        doc,
        ["变量", "当前工程值", "说明"],
        [
            ("x0", "[248.117469, 18.443148, 299.298439, 8.769514, 0.8, 600]", "T=35°C、G=500 W/m²、负载 7.2 kW"),
            ("u0", "[0.586778271, 0.501315427]", "PV 与电池占空比"),
            ("w0", "[35, 500, 50]", "RLoad=600²/7200=50 Ω"),
            ("Ppv0", "4576.067 W", "温度和半辐照度下的阵列功率"),
            ("电池需补功率", "2623.933 W", "负载减去 PV 功率"),
        ],
        widths=[1.20, 3.15, 2.15],
        font_size=8.4,
    )
    add_acceptance(doc, ["所有 x0/u0/w0 为有限实数。", "0<dpv,db<1，且建议留有 0.02～0.95 的实际裕度。", "用 u0 启动时 vdc 不出现无物理意义的大尖峰。"])
    doc.add_page_break()
    add_problems(doc, [
        ("判别式≤0", "初始负载超过假设电池能力", "降低负载、增大电池电压/能力或修改等效参数"),
        ("占空比越界", "目标工作点无法由当前拓扑实现", "检查 600 V 母线与源侧电压变比"),
        ("启动仍有尖峰", "查表插值点与工作点计算方法不同", "对象和工作点必须调用同一 PV/OCV 函数"),
    ])

    base.add_heading(doc, "11", "阶段 7：搭建 P&O 最大功率点参考发生器")
    add_stage_intro(doc, "用 vpv、Ipv 每 10 ms 更新 Vpv_ref。", "让 MPC 知道正常模式下 PV 应跟踪哪个工作电压。", "有边界、有死区、无代数环的离散 P&O 子系统。")
    base.add_text(doc, "论文中的 P&O 输出是 PV 参考电压，不是直接输出占空比。MPC 再通过 dpv 让 vpv 跟踪这个参考。因此本项目的 P&O 与前一份“直接扰动占空比”的第二阶段模型用途不同。")
    base.add_table(
        doc,
        ["模块组", "设置", "作用"],
        [
            ("Vpv_Sample / Ipv_Sample", "Zero-Order Hold，Ts=0.01 s", "把连续对象信号送入离散算法"),
            ("PV_Power", "Product", "P=V×I"),
            ("Previous_Power", "Unit Delay；Ts=0.01；IC=MPPT_P_init", "保存 P(k−1)"),
            ("Previous_Voltage", "Unit Delay；Ts=0.01；IC=MPPT_V_init", "保存 V(k−1)"),
            ("Delta_Power / Delta_Voltage", "Sum，+−", "ΔP、ΔV"),
            ("PO_Direction_Indicator", "Product", "ΔP·ΔV"),
            ("Move_Right", "Relational >=0", "决定参考电压增/减"),
            ("Power_Deadband", "1 W", "变化过小时保持"),
            ("Positive/Negative Step", "±0.25 V", "每次扰动幅值"),
            ("Previous_Voltage_Reference", "Unit Delay；Ts=0.01；IC=MPPT_V_init", "保存 Vref(k−1)"),
            ("Voltage_Reference_Limits", "Saturate：180～315 V", "防止参考跑出可用范围"),
        ],
        widths=[2.15, 2.00, 2.35],
        font_size=7.8,
    )
    base.add_heading(doc, "11.1", "实际周期怎么改", level=2)
    base.add_text(doc, "当前工程的实际 MPPT 周期由 p.mppt.sampleTime=p.controller.interval=0.01 s 决定，并体现在所有 ZOH 和 Unit Delay 的 SampleTime=0.01。若要改为 2 ms，应同时改控制器间隔、P&O 的 ZOH/Delay、ARIMA 的 ZOH/Delay、PI 离散积分增量以及 nlmpcobj.Ts；不能只改一个模块。")
    base.add_table(
        doc,
        ["要改的位置", "从 10 ms 改到 2 ms 时"],
        [
            ("p.controller.interval", "0.01 → 0.002"),
            ("p.mppt.sampleTime", "保持等于 controller.interval"),
            ("P&O 所有 ZOH/Unit Delay", "0.01 → 0.002，初值不变"),
            ("ARIMA 所有 ZOH/Unit Delay", "0.01 → 0.002"),
            ("PI 积分增益离散化", "Ki×0.01 → Ki×0.002"),
            ("nlmpcobj.Ts", "自动跟随 controller.interval"),
            ("预测器子步", "重新选择 predictionSubsteps，保证内部积分步长足够小"),
        ],
        widths=[3.10, 3.40],
        font_size=8.7,
    )
    base.add_callout(
        doc,
        "周期、步长与抖动",
        "周期更短不一定更好：测量噪声尚未衰减时频繁扰动，会使 ΔP 符号反复变化。先保持 10 ms / 0.25 V 跑通；再单独扫描周期与电压步长，并比较跟踪时间和稳态功率纹波。",
        fill=PALE_YELLOW,
    )
    add_acceptance(doc, ["Vpv_ref 始终位于 180～315 V。", "恒定环境下 Vpv_ref 最终在 MPP 附近小幅往返。", "离散环内没有直接馈通代数环。"])
    add_problems(doc, [
        ("参考电压单向跑到边界", "ΔP/ΔV 符号或 Switch 真/假输入接反", "手工制造一次升压/降压并检查方向"),
        ("完全不更新", "死区过大或 ZOH/Delay 周期未执行", "先把死区设小并显示采样命中"),
        ("MPP 附近抖动大", "周期太短或步长太大", "增加周期、减小 0.25 V 步长或采用自适应步长"),
    ])

    base.add_heading(doc, "12", "阶段 8：搭建四模式功率管理逻辑")
    add_stage_intro(doc, "根据 Ppv、Pload、SoC 和 G 产生 α 与 modeID。", "决定当前应优先 MPPT、电池充放电、限功率还是夜间供电。", "互斥、可解释、能覆盖边界条件的模式逻辑。")
    doc.add_picture(str(MODE_DIAGRAM), width=Inches(6.45))
    base.add_caption(doc, "图 3　四种工作模式及其控制目标")
    base.add_heading(doc, "12.1", "模块设置", level=2)
    base.add_table(
        doc,
        ["判断", "模块/阈值", "输出"],
        [
            ("PV_Surplus", "Relational：Ppv ≥ Pload", "功率有富余"),
            ("SoC_High", "Relational：SoC ≥ 0.90", "达到充电上限"),
            ("SoC_Above_Min", "Relational：SoC > 0.20", "允许放电"),
            ("Night_or_Cloud", "Relational：G ≤ 1e−3", "PV 不可用"),
            ("Daylight", "NOT Night_or_Cloud", "白天有效"),
            ("Mode I", "Surplus AND NOT High AND Daylight", "modeID=1"),
            ("Mode II", "NOT Surplus AND AboveMin AND Daylight", "modeID=2"),
            ("Mode III", "Surplus AND High AND Daylight", "modeID=3，α=1"),
            ("Mode IV", "Night AND AboveMin", "modeID=4"),
            ("Load shedding", "NOT(Mode I OR II OR III OR IV)", "modeID=5"),
        ],
        widths=[1.60, 3.25, 1.65],
        font_size=7.9,
    )
    base.add_heading(doc, "12.2", "建议增加但当前模型未加的工程措施", level=2)
    for item in [
        "SoC 上下限加入滞环，例如进入限功率用 0.90，退出用 0.895，避免边界抖动。",
        "Ppv−Pload 比较加入功率死区，避免两者接近时 Mode I/II 频繁切换。",
        "模式最小保持时间或权重渐变，减少硬切换引起的占空比突变。",
        "把负载切除实现为真实负载分级逻辑，而不是只输出 modeID=5。",
    ]:
        base.add_bullet(doc, item)
    add_acceptance(doc, ["任意输入组合最多只有一个模式为真。", "Mode III 时 α=1，其余模式 α=0。", "SoC≤0.20 且功率不足时进入 modeID=5，而不是继续放电。"])
    add_problems(doc, [
        ("模式号为 0", "逻辑存在未覆盖组合", "增加 Any_Valid_Mode 和负载切除兜底"),
        ("模式号跳动", "阈值无滞环、功率噪声大", "加入迟滞/死区/保持时间"),
        ("夜间仍进入 Mode II", "Daylight 条件漏接", "Mode I～III 均与 Daylight 相与"),
    ])

    base.add_heading(doc, "13", "阶段 9：搭建 ARIMA 一步预测")
    add_stage_intro(doc, "对 T、G、RLoad 生成下一控制步预测。", "MPC 的状态转移需要未来扰动，而不是只知道当前值。", "与控制周期同步的 w_forecast=[T̂,Ĝ,R̂]。")
    base.add_text(doc, "完整 ARIMA(p,d,q) 包含差分、自回归和移动平均误差项。论文只说明采用 ARIMA，没有公开阶次和系数。当前工程用一个透明的一阶差分预测分支：")
    add_equation(doc, "ŵ(k+1) = w(k) + φ·[w(k) − w(k−1)]")
    base.add_table(
        doc,
        ["模块组", "设置"],
        [
            ("Disturbance_Split", "Demux：T、G、RLoad"),
            ("三个 Sample", "Zero-Order Hold，Ts=0.01 s"),
            ("三个 Previous", "Unit Delay，Ts=0.01；IC=w0 对应分量"),
            ("三个 Delta", "当前值−上一拍"),
            ("三个 AR Gain", "ARIMA_Phi_T/G/R"),
            ("三个 Forecast Sum", "当前值+φ·差分"),
            ("Forecast_Vector", "Mux：[T̂,Ĝ,R̂]"),
        ],
        widths=[2.65, 3.85],
        font_size=8.8,
    )
    base.add_callout(
        doc,
        "当前 φ=0 的意义",
        "φ=0 时预测值等于当前测量，即 ARIMA(0,1,0) 的持久性基线。它能让接口和控制流程完整，但不代表论文真实预测器已经复现。若要研究预测价值，必须用历史数据识别阶次和系数，并单独报告预测误差。",
        fill=PALE_RED,
        border="C00000",
    )
    add_acceptance(doc, ["阶跃前后预测值有限且单位不变。", "φ=0 时 w_forecast 与当前采样值一致。", "改变 φ 后先检查预测误差，再接入 MPC。"])
    add_problems(doc, [
        ("预测超出物理范围", "φ 大、阶跃大、无边界", "对 G≥0、RLoad≥最小值、温度范围加限幅"),
        ("预测没有作用", "Np=1 且 φ=0", "这属于当前基线；增大时域前先识别预测器"),
        ("启动第一拍异常", "Previous 初值不是 w0", "Unit Delay 初值逐项使用 w0(1:3)"),
    ])

    doc.add_page_break()
    base.add_heading(doc, "14", "阶段 10：建立非线性 MPC 预测模型")
    add_stage_intro(doc, "建立 nlmpc 对象、离散状态函数、输出函数和约束。", "把六状态非线性对象用于下一步预测与优化。", "可通过 validateFcns 的 nlmpcobj_pvbatt。")
    base.add_heading(doc, "14.1", "对象配置", level=2)
    base.add_table(
        doc,
        ["项目", "当前设置", "意义"],
        [
            ("状态数", "6", "vpv、iLpv、vb、iLb、SoC、vdc"),
            ("输出数", "3", "[vdc,vpv,iLb]"),
            ("MV", "2；索引 [1,2]", "dpv、db"),
            ("MD", "3；索引 [3,4,5]", "T、G、RLoad"),
            ("Ts", "0.01 s", "控制间隔"),
            ("Prediction horizon", "1", "只预测下一步"),
            ("Control horizon", "1", "只优化当前一步"),
            ("MV 范围", "0～1；实际施加 0.02～0.95", "理论与工程裕度"),
            ("MV rate", "每步 ±0.05", "限制占空比突变"),
            ("SoC 状态约束", "0.20～0.90", "防过充/过放"),
            ("Solver", "SQP；MaxIterations=20", "非线性约束优化"),
            ("次优解", "UseSuboptimalSolution=true", "迭代未完全收敛时维持控制连续性"),
        ],
        widths=[1.75, 2.05, 2.70],
        font_size=8.1,
    )
    base.add_heading(doc, "14.2", "预测状态函数", level=2)
    base.add_text(doc, "论文式 (13) 说用前向 Euler 离散。当前工程为了数值稳健，控制周期 10 ms 内用 RK4 分 10 个子步积分；同时把电池端电容在预测器中近似为准稳态 vb=OCV−Rbat·iLb，并令预测模型的 dvb/dt=0。真实对象仍保留六个连续状态。")
    base.add_callout(
        doc,
        "这是结构复现，不是离散算法一比一复现",
        "RK4、十子步和电池电容准稳态是当前工程的数值处理。它们提高 Np=1 预测器的条件性，但会使当前结果与论文的前向 Euler 实现产生差异。若做学术复现实验，应把两种离散方法都测试并报告。",
        fill=PALE_YELLOW,
    )
    base.add_heading(doc, "14.3", "输出函数与模式相关权重", level=2)
    base.add_table(
        doc,
        ["模式", "参考向量", "输出权重"],
        [
            ("正常：I/II/IV", "[600, Vpv_ref, 0]", "[0.75, 0.15, 0]"),
            ("限功率：III", "[600, Vpv_ref, 0]", "[0.75, 0, 0.15]"),
        ],
        widths=[1.75, 2.35, 2.40],
        font_size=8.8,
    )
    base.add_text(doc, "始终调节 vdc；正常模式调节 vpv 跟踪 P&O 参考；Mode III 不再追求 MPP，而是让电池电流接近 0，由 PV 限功率承担母线调节。占空比变化率权重为 [0.10,0.10]。")
    add_acceptance(doc, ["validateFcns 无状态/输出维度错误。", "x、MV、MD 的顺序与 Simulink 端口一致。", "SoC 和占空比约束在对象的同一单位体系中。"])
    add_problems(doc, [
        ("validateFcns 失败", "StateFcn 输入维度或输出列向量不对", "combinedInput 固定为 [dpv,db,T,G,R]，返回 6×1"),
        ("优化器不可行", "工作点靠近约束、速率限制太紧", "先放宽 RateMin/Max，检查 x0 与 SoC 边界"),
        ("求解很慢", "状态尺度差异大、预测函数过重", "查表代替求根；设置 ScaleFactor；减少不必要迭代"),
        ("出现零权重警告", "3 个 OV 但只有 2 个 MV，且某模式关闭一个目标", "当前是有意的模式权重；调试时确认未误关 vdc 权重"),
    ])

    doc.add_page_break()
    base.add_heading(doc, "15", "阶段 11：在 Simulink 中连接 Nonlinear MPC Controller")
    add_stage_intro(doc, "把状态、参考、上一拍控制、预测扰动和在线权重接入 MPC 模块。", "将脚本中的优化器对象真正用于闭环。", "每 10 ms 输出 [dpv,db] 的 MPC 子系统。")
    doc.add_picture(str(MPC_DIAGRAM), width=Inches(6.55))
    base.add_caption(doc, "图 4　滚动时域控制的一次循环")
    base.add_table(
        doc,
        ["模块/端口", "设置与连接"],
        [
            ("State_Sample", "Zero-Order Hold，0.01 s；x_state → MPC x"),
            ("Output_Reference_Vector", "Mux：[600,Vpv_ref,0]；再 Reshape 成行向量 → ref"),
            ("Previous_Duty", "Unit Delay，0.01 s，IC=u0；MPC 输出反馈 → last_mv"),
            ("Forecast_Row", "w_forecast Reshape 成行向量 → md"),
            ("Mode_Split", "从 mode_bus 取 α"),
            ("Mode_Dependent_Weights", "Switch：α≥0.5 选 [0.75,0,0.15]，否则 [0.75,0.15,0]"),
            ("Weight_Row", "在线 OV 权重 Reshape → weight 端口"),
            ("Nonlinear MPC Controller", "nlmpc object=nlmpcobj_pvbatt"),
            ("Enable_MPC", "ControllerSelect=1 时执行；避免 PI 与 MPC 同时计算"),
            ("MPC status", "当前接 Terminator；调试时建议接日志"),
        ],
        widths=[2.25, 4.25],
        font_size=8.1,
    )
    base.add_heading(doc, "15.1", "顶层控制选择与安全限幅", level=2)
    base.add_table(
        doc,
        ["模块", "设置"],
        [
            ("Controller_Select", "Constant=ControllerSelect；1=MPC，0=PI"),
            ("Controller_Selector", "Switch threshold=0.5；u1=MPC，u3=PI"),
            ("Enable_PI_When_Not_MPC", "Logical NOT"),
            ("Duty_Limits", "Saturation：0.02～0.95"),
            ("对象输入", "Duty_Limits 输出同时接对象、Scope 和 sim_u 日志"),
        ],
        widths=[2.50, 4.00],
        font_size=8.7,
    )
    add_acceptance(doc, ["MPC 开启时 PI 子系统不执行，PI 开启时 MPC 不执行。", "实际占空比全程位于 0.02～0.95。", "控制更新点与 10 ms 周期对齐。"])
    add_problems(doc, [
        ("MPC 输出维度错", "ref/md/weight 方向为列向量", "按当前模块用 Reshape 转成行向量"),
        ("占空比卡边界", "对象符号错、权重失衡或初值差", "先看未限幅 MPC 输出和状态预测，不要先放大边界"),
        ("仿真每步都求解", "MPC 输入未离散采样", "State_Sample 与所有离散状态设 0.01 s"),
    ])

    base.add_heading(doc, "16", "阶段 12：建立 PI 基线、日志和运行脚本")
    add_stage_intro(doc, "搭建透明的双回路 PI，统一记录状态、功率、占空比和模式。", "给 MPC 提供同对象、同扰动的比较基线。", "能用一个参数切换 MPC/PI 并计算 VRI。")
    base.add_heading(doc, "16.1", "当前 PI 子系统", level=2)
    base.add_table(
        doc,
        ["回路", "误差", "参数与离散实现", "输出"],
        [
            ("PV 电压 PI", "vpv−Vpv_ref", "Kp=0.002；Ki=0.20；积分增量 Ki×0.01；IC=0", "校正量 + u0(1)"),
            ("母线电压 PI", "600−vdc", "Kp=0.001；Ki=0.05；积分增量 Ki×0.01；IC=0", "校正量 + u0(2)"),
        ],
        widths=[1.25, 1.40, 2.70, 1.15],
        font_size=8.0,
    )
    base.add_callout(
        doc,
        "PI 对比的复现边界",
        "论文公布的是基于 LQR 设计的一组状态反馈增益，但没有给出完整实现。当前工程使用可解释的双 PI 近似，因此只能用于“当前工程内 MPC 与 PI 的相对比较”，不能宣称精确复现论文 PI 曲线。",
        fill=PALE_RED,
        border="C00000",
    )
    base.add_heading(doc, "16.2", "日志设置", level=2)
    base.add_table(
        doc,
        ["To Workspace", "变量", "格式", "Decimation"],
        [
            ("State_Log", "sim_x", "Structure With Time", "100"),
            ("Aux_Log", "sim_aux", "Structure With Time", "100"),
            ("Duty_Log", "sim_u", "Structure With Time", "100"),
            ("Mode_Log", "sim_mode", "Structure With Time", "100"),
            ("Disturbance_Log", "sim_w", "Structure With Time", "100"),
        ],
        widths=[1.65, 1.15, 2.40, 1.30],
        font_size=8.7,
    )
    base.add_text(doc, "对象步长 10 μs，Decimation=100，因此连续日志约每 1 ms 记录一次；这比 10 ms 控制周期更细，足以观察瞬态。")
    base.add_heading(doc, "16.3", "运行与指标", level=2)
    base.add_code(doc, 'result = run_paper_scenario("comparison","MPC");\nplot_paper_scenario(result);\nfprintf("mean VRI = %.6f %%\\n", result.meanVRI);')
    add_equation(doc, "VRI(%) = |vdc − 600| / 600 × 100")
    add_acceptance(doc, ["MPC/PI 可通过 ControllerSelect 互斥切换。", "sim_x、sim_aux、sim_u、sim_mode、sim_w 都有时间轴。", "平均 VRI 与最大 VRI 从同一时间范围计算。"])
    add_problems(doc, [
        ("日志长度不同", "不同采样率或 Decimation", "绘图时用各自 time，不要按数组下标硬对齐"),
        ("PI 积分越积越大", "输出限幅后无抗饱和", "增加 anti-windup，或在比较中记录积分状态"),
        ("切换控制器有冲击", "PI/MPC 内部状态未做无扰切换", "比较试验从 t=0 固定选择，不在运行中切换"),
    ])

    base.add_heading(doc, "17", "论文每组仿真做了什么、说明了什么")
    base.add_heading(doc, "17.1", "论文统一参数", level=2)
    base.add_table(
        doc,
        ["参数", "论文值"],
        [
            ("PV 额定功率", "约 9.5 kW"),
            ("电池", "300 V，20 Ah，SoC 20%～90%"),
            ("直流母线", "600 V，Cdc=1500 μF"),
            ("PV/电池侧电容", "各 300 μF"),
            ("PV/电池侧电感", "各 10 mH，电阻各 10 mΩ"),
            ("开关频率", "10 kHz；但控制分析使用平均值动态"),
            ("对象样本时间", "10 μs"),
            ("控制间隔", "10 ms"),
            ("预测时域", "1"),
            ("权重 F/P/R/Q", "0.75 / 0.15 / 0.10 / 0.15"),
        ],
        widths=[3.20, 3.30],
        font_size=8.8,
    )
    base.add_heading(doc, "17.2", "Mode I：由放电转充电", level=2)
    base.add_text(doc, "论文在约 4 s 改变环境和负载，使系统从“PV 略不足、由电池补偿”转为“PV 有富余、电池吸收”。意义是验证同一正常模式目标下，电池功率能跨过零点，而 PV 始终保持 MPPT，vdc 保持 600 V。")
    base.add_table(
        doc,
        ["观察量", "正确趋势", "失败表现"],
        [
            ("Ppv", "跟随环境变化并保持接近 MPP", "长期偏离可用最大功率"),
            ("Pbat", "由正变负", "符号不变或功率不平衡"),
            ("SoC", "先微降后上升", "充电时仍下降"),
            ("vdc", "扰动后快速回 600 V", "持续漂移或过冲过大"),
        ],
        widths=[1.25, 2.75, 2.50],
        font_size=8.5,
    )
    base.add_heading(doc, "17.3", "Mode II：由充电转放电", level=2)
    base.add_text(doc, "约 4 s 后负载超过 PV 可用功率且 SoC 高于下限。PV 继续 MPPT，电池补足缺口。意义是验证欠功率时的供电可靠性和母线调节能力。")
    base.add_heading(doc, "17.4", "Mode III：电池满后的 PV 限功率", level=2)
    base.add_text(doc, "SoC 在约 4.2 s 达 90%，电池不能继续吸收功率；α 从 0 切到 1，目标由跟踪 Vpv_ref 改为让电池电流接近 0，PV 离开 MPP 并主动降功率。约 12.2 s 负载变大后重新进入 MPPT。")
    base.add_callout(
        doc,
        "Mode III 是论文最有价值的部分",
        "只做 MPPT 会在电池满时制造能量无处可去的问题。Mode III 说明“最大化发电”不是任何时候都正确，独立系统必须把能量平衡和储能边界置于局部 MPPT 目标之上。",
        fill=PALE_GREEN,
        border="70AD47",
    )
    base.add_heading(doc, "17.5", "Mode IV：PV 失去功率", level=2)
    base.add_text(doc, "约 4.2 s 将 PV 功率降为 0，电池独立供负载，SoC 持续下降。意义是验证夜间/遮云情况下的孤岛供电。SoC 到 20% 后应切除非关键负载，但论文没有进一步实现。")
    base.add_heading(doc, "17.6", "MPC 与 PI 对比", level=2)
    base.add_text(doc, "论文在 t=1 s 同时把 G 从 500 提高到 1000 W/m²、T 从 35°C 降到 25°C，在 t=2 s 把负载从 7.2 kW 提到 14.4 kW。两者都恢复 600 V，但论文报告 MPC 瞬态更平滑，平均母线误差指标优于 PI。")
    base.add_table(
        doc,
        ["项目", "论文报告", "当前工程实测"],
        [
            ("MPC 平均指标", "0.001356（论文原文称 average error）", "mean VRI≈0.121544%"),
            ("PI 平均指标", "0.011267", "mean VRI≈0.110473%"),
            ("MPC 最大 VRI", "未单独给出", "≈3.945106%"),
            ("PI 最大 VRI", "未单独给出", "≈5.989661%"),
            ("MPC 终值", "约 600 V", "600.001 V"),
            ("PI 终值", "约 600 V", "600.000 V"),
        ],
        widths=[1.55, 2.30, 2.65],
        font_size=8.2,
    )
    base.add_callout(
        doc,
        "如何解释当前对比",
        "当前 MPC 的峰值优于 PI，但平均 VRI 略高；这与论文结论并不完全一致。主要原因包括电池参数缺失、PI 实现不同、ARIMA 系数为基线、预测离散法不同和代价函数实现差异。正确做法是公开这些差异，而不是继续调图直到“看起来一样”。",
        fill=PALE_RED,
        border="C00000",
    )
    base.add_heading(doc, "17.7", "300 s 真实日曲线", level=2)
    base.add_text(doc, "论文把 24 小时的负载、温度和辐照度压缩到 300 s，验证从清晨、白天到傍晚的连续运行。白天 PV 富余时电池充电，早晚 PV 不足时电池放电，vdc 维持 600 V。")
    base.add_text(doc, "当前工程的 realworld 场景是根据论文形状重构的合成曲线，不是论文所用 NREL 原始数据。若要严谨复现，应下载相同日期和站点数据，完成单位、时区、缺失值和缩放处理后替换 scenario_w。")

    base.add_heading(doc, "18", "论文原值、工程假设和实现差异")
    base.add_text(doc, "下表是判断复现可信度的关键。A 类可直接对应论文；B 类是工程等效实现；C 类是由于论文信息不足而采用的假设。")
    base.add_table(
        doc,
        ["等级", "项目", "论文/工程对应", "使用时怎么表述"],
        [
            ("A", "C、L、r、Cdc、600 V、10 μs、10 ms、Np=1", "按论文原值", "可称参数复现"),
            ("A", "PV 模块四个额定点、300 V/20 Ah 电池", "按论文原值", "可称额定参数复现"),
            ("B", "PV 隐式方程 → 三维查表", "数学关系等效，数值实现不同", "称工程等效实现"),
            ("B", "前向 Euler → RK4 十子步", "离散方法不同", "称稳健复现，不称同算法"),
            ("B", "论文范数代价 → nlmpc 输出跟踪权重", "目标对应但形式可能为平方代价", "称目标映射"),
            ("C", "9 串 5 并阵列结构", "论文未公开", "明确为额定功率匹配假设"),
            ("C", "PV 二极管参数", "由公开点拟合", "明确为拟合值"),
            ("C", "电池 OCV 表、Rbat=0.08 Ω", "论文未公开", "明确为可替换假设"),
            ("C", "ARIMA φ=0", "论文未公开阶次/系数", "称持久性基线"),
            ("C", "PI Kp/Ki", "论文只公开另一组反馈增益", "称透明对比基线"),
            ("C", "realworld 曲线", "当前为合成重构", "不可称原始数据复现"),
        ],
        widths=[0.55, 1.85, 2.45, 1.65],
        font_size=7.4,
    )
    base.add_heading(doc, "18.1", "如果目标是论文级严谨复现，下一步要补什么", level=2)
    for item in [
        "获取或辨识论文使用的完整 PV 阵列布局和单二极管参数。",
        "获取电池 E0、K、A、B、Rbat、OCV 曲线及充放电效率。",
        "恢复论文 ARIMA 阶次、训练窗口、系数和预测误差评价。",
        "按论文原始代价函数直接编写自定义 fmincon，并与 nlmpc 映射实现对照。",
        "按相同场景和相同 VRI 定义复算论文数值，区分比例、百分数和平均方式。",
        "用相同 NREL 数据替换合成 realworld 场景。",
    ]:
        base.add_bullet(doc, item)

    base.add_heading(doc, "19", "发散、不收敛和结果不对的系统排查")
    base.add_heading(doc, "19.1", "推荐排查顺序", level=2)
    add_numbered_list(doc, [
        "先看是否出现 NaN/Inf，以及第一次出现的时间和信号。",
        "再看 u_duty 是否撞到 0.02/0.95；若撞边界，查看未限幅控制输出。",
        "断开控制器，用一致 u0 和恒定 w0 运行对象，确认六状态本身稳定。",
        "检查 Ppv、Pbat、Pload 与母线电容储能是否满足功率平衡。",
        "检查 P&O 的 ΔP、ΔV、Vpv_ref 和采样命中，确认参考没有单向跑边。",
        "检查 modeID 和 α，确认没有在阈值附近高速切换。",
        "最后看 MPC status、迭代次数、约束和预测状态；不要一开始就调权重。",
    ])
    base.add_heading(doc, "19.2", "症状—原因—处理总表", level=2)
    base.add_table(
        doc,
        ["症状", "优先检查", "典型原因", "处理"],
        [
            ("t=0 立即大尖峰", "x0/u0/w0", "初值不满足稳态方程", "重新求工作点；所有 Integrator/Delay 使用一致初值"),
            ("vpv 或 iLpv 发散", "PV KVL/KCL", "Sum 符号、L/C 单位或查表异常", "断开 MPC，逐项显示导数"),
            ("vb 高频振荡", "电池电容分支", "Cb 快动态与 10 ms 预测不协调", "保持对象 10 μs；预测器可用准稳态近似"),
            ("vdc 持续上升", "功率符号", "充电功率被当作注入", "统一 Pbat 正方向并做功率守恒"),
            ("vdc 持续下降", "负载输入", "把功率当电阻或把电阻当功率", "w(3)=RLoad；Pload=vdc²/RLoad"),
            ("P&O 参考撞边界", "ΔP·ΔV 逻辑", "方向反、延时缺失、噪声大", "恢复 10 ms/0.25 V，逐拍检查"),
            ("Mode I/II 抖动", "Ppv−Pload 比较", "无死区", "加入功率滞环和最小保持时间"),
            ("Mode III 尖峰", "α 权重硬切换", "目标瞬时改变", "权重渐变、MV rate 限制、SoC 滞环"),
            ("MPC 不可行", "状态/MV 约束", "初值越界、速率限制过紧", "先放宽 Rate，检查 SoC 和上一拍 MV"),
            ("MPC 很慢", "预测函数和求解器", "每步求根、尺度差、300 s 全程优化", "使用查表、尺度化、短场景逐步验证"),
            ("零 OV 权重警告", "在线权重", "某模式故意关闭一个目标", "确认 vdc 权重始终非零；调试时记录权重"),
            ("PI 比 MPC 更好", "比较公平性", "PI 实现与论文不同、权重未调、指标范围不同", "固定对象/场景/采样/指标，再分别整定"),
            ("300 s 跑不完", "步数与控制求解", "300 s/10 μs=3000 万对象步", "先用缩短场景；确认后再加速模式或离线数据"),
        ],
        widths=[1.35, 1.40, 2.05, 1.70],
        font_size=6.9,
        header_fill=PALE_RED,
    )
    base.add_heading(doc, "19.3", "最小化问题的断开法", level=2)
    base.add_table(
        doc,
        ["试验", "临时替换", "说明"],
        [
            ("对象稳定性", "u_duty=u0，w=w0", "失败说明对象/初值问题，与 MPPT/MPC 无关"),
            ("PV 支路", "固定 db 和负载，仅改变 dpv", "观察 vpv/iLpv/Ppv 方向"),
            ("电池支路", "固定 dpv，给小负载阶跃", "观察 Pbat、SoC、vdc"),
            ("P&O", "对象稳定后只启用 Vpv_ref，不接 MPC", "观察参考是否有界"),
            ("MPC", "固定 Vpv_ref、固定模式、固定扰动", "先验证单一 Mode I"),
            ("模式切换", "MPC 稳定后再跨 SoC/P 功率阈值", "把控制问题与切换问题分离"),
        ],
        widths=[1.35, 2.60, 2.55],
        font_size=8.3,
    )

    base.add_heading(doc, "20", "独立复现完成清单")
    checklist = [
        ("待确认", "参数脚本区分论文原值、拟合值和复现假设"),
        ("待确认", "求解器 fixed-step ode4，步长 1e−5 s"),
        ("待确认", "PV 查表维度为 V×T×G，STC 最大功率约 9.59 kW"),
        ("待确认", "x_state 顺序为 [vpv,iLpv,vb,iLb,SoC,vdc]"),
        ("待确认", "aux 顺序为 [Ipv,Ib,Ebat,Ppv,Pbat,Pload]"),
        ("待确认", "充电时 Pbat<0 且 SoC 上升；放电时相反"),
        ("待确认", "恒定 u0/w0 下六状态对象有限且功率平衡"),
        ("待确认", "P&O 周期 10 ms，步长 0.25 V，参考限幅 180～315 V"),
        ("待确认", "Mode III 时 α=1，其余模式 α=0"),
        ("待确认", "ARIMA 基线明确为 φ=0，不冒充论文系数"),
        ("待确认", "nlmpc：6 状态、3 输出、2 MV、3 MD、Ts=0.01、Np=Nc=1"),
        ("待确认", "MV 范围 0～1、Rate±0.05，实际 Duty 限幅 0.02～0.95"),
        ("待确认", "MPC/PI 互斥启用，日志变量完整"),
        ("待确认", "Mode I～IV 分别单独通过后再跑 comparison/realworld"),
        ("待确认", "报告中明确当前数值与论文不一致的原因"),
    ]
    base.add_table(doc, ["完成", "检查项"], checklist, widths=[0.75, 5.75], font_size=8.5, header_fill=PALE_GREEN)
    base.add_callout(
        doc,
        "结构检查结果",
        "当前 PV_Battery_MPC.slx 已检查：无未连接端口、无悬空连线。comparison 场景的 MPC 和 PI 均能回到约 600 V；这证明工程可运行，但不等于所有论文数值已经一比一复现。",
        fill=LIGHT_BLUE,
    )

    base.add_heading(doc, "附录 A", "论文参数与当前工程参数速查")
    base.add_table(
        doc,
        ["类别", "参数", "当前值", "来源"],
        [
            ("PV 变换器", "Cpv / Lpv / rLpv", "300 μF / 10 mH / 10 mΩ", "论文"),
            ("电池变换器", "Cb / Lb / rLb", "300 μF / 10 mH / 10 mΩ", "论文"),
            ("母线", "Cdc / Vdc_ref", "1500 μF / 600 V", "论文"),
            ("PV 模块", "Pmp/Vmp/Imp/Voc/Isc", "213.15 W / 29 V / 7.35 A / 36.3 V / 7.84 A", "论文"),
            ("PV 阵列", "Ns/Np", "9 / 5", "复现假设"),
            ("电池", "额定电压/容量/SoC", "300 V / 20 Ah / 0.20～0.90", "论文"),
            ("电池", "Rbat/OCV 表", "0.08 Ω / 270～312 V", "复现假设"),
            ("控制", "对象步长/控制周期", "10 μs / 10 ms", "论文"),
            ("控制", "Np/Nc", "1 / 1", "论文/工程"),
            ("控制", "输出权重正常", "[0.75,0.15,0]", "论文目标映射"),
            ("控制", "输出权重限功率", "[0.75,0,0.15]", "论文目标映射"),
            ("控制", "MV rate/实际限幅", "±0.05 / 0.02～0.95", "工程措施"),
            ("P&O", "周期/步长/死区/范围", "10 ms / 0.25 V / 1 W / 180～315 V", "工程设置"),
            ("ARIMA", "φT/φG/φR", "0 / 0 / 0", "复现基线"),
        ],
        widths=[1.20, 2.05, 2.20, 1.05],
        font_size=7.2,
    )

    base.add_heading(doc, "附录 B", "顶层连线速查")
    base.add_table(
        doc,
        ["源", "目标"],
        [
            ("Scenario_Profiles.w_actual", "ARIMA、Power_Management_Logic、Plant、sim_w"),
            ("Plant.x_state", "MPC、PI、P&O、PMS、Scope、sim_x"),
            ("Plant.aux", "P&O、PMS、Scope、sim_aux"),
            ("ARIMA.w_forecast", "MPC.md"),
            ("P&O.Vpv_ref", "MPC.ref、PI PV 电压回路"),
            ("PMS.mode_bus", "MPC 在线权重、sim_mode"),
            ("MPC.u_mpc / PI.u_pi", "Controller_Selector"),
            ("Controller_Selector", "Duty_Limits"),
            ("Duty_Limits", "Plant.u_duty、Scope、sim_u"),
        ],
        widths=[2.70, 3.80],
        font_size=8.6,
    )

    base.add_heading(doc, "附录 C", "现有脚本作用速查")
    base.add_table(
        doc,
        ["脚本", "作用"],
        [
            ("pvbatt_parameters.m", "所有论文值、拟合值和工程假设"),
            ("pvbatt_generate_pv_lookup.m", "生成 PV 三维电流查表"),
            ("pvbatt_pv_current.m", "读取 PV 查表"),
            ("pvbatt_battery_ocv.m", "读取 OCV–SoC 表"),
            ("pvbatt_operating_point.m", "计算一致 x0/u0/w0"),
            ("pvbatt_state_derivatives.m", "六状态连续平均值方程"),
            ("pvbatt_state_transition.m", "MPC 离散预测状态函数"),
            ("pvbatt_mpc_output.m", "输出 [vdc,vpv,iLb]"),
            ("pvbatt_create_nlmpc.m", "创建并验证 nlmpc 对象"),
            ("pvbatt_build_scenario.m", "生成 comparison、mode1～4、realworld"),
            ("pvbatt_initialize.m", "统一初始化并写入工作区"),
            ("run_paper_scenario.m", "使用 SimulationInput 仿真并计算 VRI"),
            ("plot_paper_scenario.m", "绘制功率、母线电压和 SoC"),
        ],
        widths=[2.65, 3.85],
        font_size=8.3,
    )

    base.add_heading(doc, "参考文献", "主文献与说明")
    base.add_text(doc, "Batiyah, S.; Sharma, R.; Abdelwahed, S.; Alhosaini, W.; Aldosari, O. Predictive Control of PV/Battery System under Load and Environmental Uncertainty. Energies 2022, 15, 4100. DOI: 10.3390/en15114100.")
    base.add_text(doc, "本手册对主文献内容进行了归纳、解释和工程化重构，没有逐句复制论文。模型参数和实现差异以当前 D:\\PV_MPPT 工作区为准。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Created {OUT}")


if __name__ == "__main__":
    build_manual()
