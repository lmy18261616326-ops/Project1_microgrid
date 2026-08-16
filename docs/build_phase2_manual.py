from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\PV_MPPT")
DOCS = ROOT / "docs"
QA = DOCS / "_qa"
OUT = DOCS / "PV_MPPT_Phase2_Simulink_Build_Manual.docx"
ACTUAL_DIAGRAM = QA / "PV_MPPT_v2_actual_diagram.png"
ACTUAL_CROP = QA / "PV_MPPT_v2_mppt_chain_crop.png"
FLOW_DIAGRAM = QA / "phase2_po_control_flow.png"

BLUE = "2E74B5"
DARK_BLUE = "17365D"
MID_BLUE = "5B9BD5"
PALE_BLUE = "E8EEF5"
LIGHT_BLUE = "DDEBF7"
PALE_GREEN = "E2F0D9"
PALE_YELLOW = "FFF2CC"
PALE_RED = "FCE4D6"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"
BLACK = "222222"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                edge.set(qn(f"w:{key}"), str(edge_data[key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_fixed_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_twips = sum(Inches(width).twips for width in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid_cols = list(table._tbl.tblGrid)
    for idx, width in enumerate(widths):
        twips = Inches(width).twips
        if idx < len(grid_cols):
            grid_cols[idx].set(qn("w:w"), str(twips))
        table.columns[idx].width = Inches(width)
        for row in table.rows:
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(twips))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        keep_next_el = OxmlElement("w:keepNext")
        keep_next_el.set(qn("w:val"), "1")
        p_pr.append(keep_next_el)
    if keep_lines:
        keep_lines_el = OxmlElement("w:keepLines")
        keep_lines_el.set(qn("w:val"), "1")
        p_pr.append(keep_lines_el)


def set_run_font(run, east_asia="Microsoft YaHei", ascii_font="Calibri"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.20

    for style_name, size, color, before, after in [
        ("Title", 28, DARK_BLUE, 0, 12),
        ("Subtitle", 13, GRAY, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name in {"Title", "Heading 1", "Heading 2", "Heading 3"}
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.paragraph_format.left_indent = Inches(0.22)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(5)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.20

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("PV_MPPT  |  Phase 2")
    set_run_font(run)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    run2 = p.add_run("    Simulink 搭建说明书")
    set_run_font(run2)
    run2.font.size = Pt(8.5)
    run2.font.color.rgb = RGBColor.from_string(GRAY)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), MID_BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("第 ")
    set_run_font(fr)
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string(GRAY)
    add_field(fp, "PAGE")
    fr2 = fp.add_run(" 页")
    set_run_font(fr2)
    fr2.font.size = Pt(8)
    fr2.font.color.rgb = RGBColor.from_string(GRAY)


def add_text(doc, text="", bold_prefix=None, style=None, align=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    if align is not None:
        p.alignment = align
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text, level=0):
    p = doc.add_paragraph(style="List Number" if level == 0 else "List Number 2")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    return p


def add_callout(doc, title, body, fill=PALE_YELLOW, border=MID_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.65)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=110, start=150, bottom=110, end=150)
    set_cell_border(
        cell,
        start={"val": "single", "sz": "18", "color": border},
        top={"val": "single", "sz": "4", "color": fill},
        bottom={"val": "single", "sz": "4", "color": fill},
        end={"val": "single", "sz": "4", "color": fill},
    )
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run_font(r2)
    p2.paragraph_format.space_after = Pt(0)
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, widths=None, font_size=8.5, header_fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        set_run_font(r)
        r.bold = True
        r.font.size = Pt(font_size)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r)
            r.font.size = Pt(font_size)
    if widths:
        set_fixed_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, number, title, level=1):
    p = doc.add_heading(f"{number}  {title}", level=level)
    set_paragraph_keep(p, keep_next=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_run_font(r)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    return p


def font(size=28, bold=False):
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    if not font_path.exists():
        font_path = Path(r"C:\Windows\Fonts\simhei.ttf")
    return ImageFont.truetype(str(font_path), size=size, index=0)


def pil_color(value):
    if isinstance(value, str) and len(value) == 6 and not value.startswith("#"):
        return f"#{value}"
    return value


def draw_centered(draw, xy, text, fnt, fill=BLACK):
    x, y = xy
    box = draw.multiline_textbbox((0, 0), text, font=fnt, align="center", spacing=4)
    w = box[2] - box[0]
    h = box[3] - box[1]
    draw.multiline_text((x - w / 2, y - h / 2), text, font=fnt, fill=pil_color(fill), align="center", spacing=4)


def draw_box(draw, rect, title, subtitle="", fill=WHITE, outline=BLUE):
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle(rect, radius=16, fill=pil_color(fill), outline=pil_color(outline), width=4)
    draw_centered(draw, ((x1 + x2) / 2, (y1 + y2) / 2 - (16 if subtitle else 0)), title, font(27, True), DARK_BLUE)
    if subtitle:
        draw_centered(draw, ((x1 + x2) / 2, (y1 + y2) / 2 + 30), subtitle, font(20), GRAY)


def draw_arrow(draw, start, end, color=BLUE, width=5):
    draw.line([start, end], fill=pil_color(color), width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 17
    p1 = (x2, y2)
    p2 = (x2 - ux * size + px * size * 0.55, y2 - uy * size + py * size * 0.55)
    p3 = (x2 - ux * size - px * size * 0.55, y2 - uy * size - py * size * 0.55)
    draw.polygon([p1, p2, p3], fill=pil_color(color))


def create_flow_diagram():
    QA.mkdir(parents=True, exist_ok=True)
    canvas = Image.new("RGB", (2400, 1180), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((80, 45), "第二阶段 P&O MPPT 信号流与离散状态", font=font(40, True), fill=pil_color(DARK_BLUE))
    draw.text((80, 102), "蓝色箭头：前向信号　灰色箭头：一拍反馈　所有控制模块编译周期均为 2 ms", font=font(23), fill=pil_color(GRAY))

    boxes = {
        "vpv": (70, 220, 350, 355),
        "ipv": (70, 425, 350, 560),
        "vdelay": (470, 220, 800, 355),
        "idelay": (470, 425, 800, 560),
        "power": (930, 315, 1220, 460),
        "delta": (1340, 315, 1640, 460),
        "judge": (1770, 315, 2170, 460),
        "direction": (1770, 610, 2170, 770),
        "step": (1340, 620, 1640, 760),
        "candidate": (930, 620, 1220, 760),
        "saturation": (470, 620, 800, 760),
        "plant": (70, 620, 350, 760),
    }
    draw_box(draw, boxes["vpv"], "Vpv", "x(1)", LIGHT_BLUE)
    draw_box(draw, boxes["ipv"], "Ipv", "aux(1)", LIGHT_BLUE)
    draw_box(draw, boxes["vdelay"], "Unit Delay", "Ts=2 ms, IC=261 V", PALE_YELLOW)
    draw_box(draw, boxes["idelay"], "Unit Delay", "Ts=2 ms, IC=36.75 A", PALE_YELLOW)
    draw_box(draw, boxes["power"], "P = V × I", "PV_Power", PALE_GREEN)
    draw_box(draw, boxes["delta"], "ΔP = P − Pz⁻¹", "Delta_Power", PALE_GREEN)
    draw_box(draw, boxes["judge"], "ΔP < −0.1 W ?", "Power_Decreased", PALE_RED)
    draw_box(draw, boxes["direction"], "方向记忆 / 反向", "s[k] = ±s[k−1]", PALE_YELLOW)
    draw_box(draw, boxes["step"], "ΔD = s × 0.0005", "Duty_Perturbation", PALE_GREEN)
    draw_box(draw, boxes["candidate"], "D候选 = Dz⁻¹ + ΔD", "Duty_Candidate", PALE_GREEN)
    draw_box(draw, boxes["saturation"], "占空比限幅", "0.05…0.90", PALE_RED)
    draw_box(draw, boxes["plant"], "平均值对象模型", "Dpv → Control_Vector(1)", LIGHT_BLUE)

    draw_arrow(draw, (350, 287), (470, 287))
    draw_arrow(draw, (350, 492), (470, 492))
    draw_arrow(draw, (800, 287), (930, 350))
    draw_arrow(draw, (800, 492), (930, 425))
    draw_arrow(draw, (1220, 387), (1340, 387))
    draw_arrow(draw, (1640, 387), (1770, 387))
    draw_arrow(draw, (1970, 460), (1970, 610))
    draw_arrow(draw, (1770, 690), (1640, 690))
    draw_arrow(draw, (1340, 690), (1220, 690))
    draw_arrow(draw, (930, 690), (800, 690))
    draw_arrow(draw, (470, 690), (350, 690))

    feedback = "#7F7F7F"
    draw.line([(635, 760), (635, 965), (1075, 965), (1075, 760)], fill=feedback, width=4)
    draw_arrow(draw, (1075, 965), (1075, 760), color=feedback, width=4)
    draw.text((700, 920), "Previous_Duty：保存限幅后的 D[k−1]", font=font(21), fill=pil_color(GRAY))
    draw.line([(1075, 315), (1075, 175), (1490, 175), (1490, 315)], fill=feedback, width=4)
    draw_arrow(draw, (1490, 175), (1490, 315), color=feedback, width=4)
    draw.text((1130, 130), "Previous_Power：保存 P[k−1]", font=font(21), fill=pil_color(GRAY))
    draw.line([(1970, 770), (1970, 930), (1490, 930), (1490, 760)], fill=feedback, width=4)
    draw_arrow(draw, (1490, 930), (1490, 760), color=feedback, width=4)
    draw.text((1570, 875), "Direction_State：保存 s[k−1]", font=font(21), fill=pil_color(GRAY))

    draw.rounded_rectangle((70, 1010, 2325, 1120), radius=12, fill="#F7F9FC", outline=pil_color(MID_BLUE), width=3)
    draw.text((100, 1038), "关键：Unit Delay 使控制器只在采样点更新，并切断 D → 对象 → V/I → D 的直接馈通。若改成直接馈通的 Zero-Order Hold，模型可能形成代数环。", font=font(22), fill=pil_color(DARK_BLUE))
    canvas.save(FLOW_DIAGRAM, quality=95)


def create_actual_crop():
    if not ACTUAL_DIAGRAM.exists():
        return
    im = Image.open(ACTUAL_DIAGRAM).convert("RGB")
    crop = im.crop((1450, 40, 3300, 980))
    crop.save(ACTUAL_CROP, quality=95)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(58)
    r = p.add_run("SIMULINK 实践手册")
    set_run_font(r)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MID_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("PV_MPPT 第二阶段")
    set_run_font(r)
    r.bold = True
    r.font.size = Pt(29)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("离散 P&O MPPT 控制器搭建说明书")
    set_run_font(r)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("从 PV_MPPT_v1 手工复现到 PV_MPPT_v2")
    set_run_font(r)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(4.05)
    cover_rows = [
        ("对象模型", "连续平均值 PV—电池系统（6 个状态）"),
        ("控制算法", "离散扰动观察法（P&O），直接调节光伏 Boost 占空比"),
        ("推荐设置", "MPPT 周期 2 ms；占空比步长 0.0005"),
        ("文档版本", f"V1.0　{date(2026, 7, 31).isoformat()}"),
    ]
    for i, (label, value) in enumerate(cover_rows):
        row = table.rows[i]
        prevent_row_split(row)
        c0, c1 = row.cells
        set_cell_shading(c0, PALE_BLUE)
        set_cell_shading(c1, "F8FAFC")
        for c in [c0, c1]:
            set_cell_margins(c, top=110, start=150, bottom=110, end=150)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr0 = p0.add_run(label)
        set_run_font(rr0)
        rr0.bold = True
        rr0.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        rr1 = c1.paragraphs[0].add_run(value)
        set_run_font(rr1)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_callout(
        doc,
        "这份说明书解决什么",
        "解释第二阶段做了什么、为什么要这样做，并给出可以从第一阶段模型逐块复现的 Simulink 操作步骤、参数、连线、验证标准和常见错误处理方法。",
        fill=LIGHT_BLUE,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(25)
    r = p.add_run(r"配套模型：D:\PV_MPPT\model\PV_MPPT_v2.slx")
    set_run_font(r)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()


def build_manual():
    create_flow_diagram()
    create_actual_crop()

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "PV_MPPT 第二阶段离散 P&O MPPT 控制器搭建说明书"
    doc.core_properties.subject = "Simulink 模块级搭建与验证"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "PV, MPPT, P&O, Simulink, Boost, 占空比"
    add_cover(doc)

    add_heading(doc, "0", "阅读路线")
    add_text(doc, "如果你只想马上动手，先看第 5～11 章；如果你希望先理解为什么这么搭建，先看第 1～4 章。")
    add_table(
        doc,
        ["目标", "建议阅读"],
        [
            ("弄清第二阶段做了什么", "第 1、2 章"),
            ("理解控制公式和数据流", "第 3、4 章"),
            ("从第一阶段亲手搭建", "第 5～11 章"),
            ("检查自己的模型是否正确", "第 12、13 章"),
            ("定位发散或不跟踪问题", "第 14 章"),
        ],
        widths=[2.65, 3.85],
        font_size=9.3,
    )
    add_callout(
        doc,
        "最重要的时间设置",
        "MPPT 实际更新周期是 MPPT_Ts = 0.002 s（2 ms）；模型求解器固定步长仍是 1e-5 s（10 μs）。不要把整个模型的固定步长改成 2 ms。",
        fill=PALE_YELLOW,
        border="BF9000",
    )

    add_heading(doc, "1", "第二阶段做了什么")
    add_text(
        doc,
        "第一阶段只有被控对象和固定占空比：模型能运行，但光伏变换器没有根据功率变化主动寻找最大功率点。第二阶段把固定的光伏占空比替换为一个离散 P&O 控制器，使占空比形成闭环。",
    )
    add_table(
        doc,
        ["项目", "第一阶段", "第二阶段"],
        [
            ("光伏占空比", "固定常数", "由 P&O 每 2 ms 更新"),
            ("反馈量", "无 MPPT 反馈", "Vpv、Ipv → Ppv"),
            ("离散状态", "无", "前一拍功率、方向和占空比"),
            ("安全限制", "依赖固定值", "Dpv 限制为 0.05～0.90"),
            ("新增日志", "phase1_x、phase1_aux", "增加 phase2_power、phase2_duty"),
            ("电池侧控制", "固定占空比", "仍为固定占空比，本阶段不改"),
        ],
        widths=[1.55, 2.35, 2.60],
        font_size=8.8,
    )
    add_text(doc, "第二阶段新增的功能链可以概括为：")
    for item in [
        "从 6 维状态向量中取出 Vpv，从 10 维辅助量中取出 Ipv。",
        "以 2 ms 周期离散化两路测量，并引入一拍延迟。",
        "计算 Ppv = Vpv × Ipv，再与前一拍功率比较。",
        "功率明显下降时反转扰动方向，否则保持方向。",
        "用方向乘以占空比步长，累加到上一拍占空比。",
        "对占空比限幅，并送入对象模型的光伏控制输入。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2", "第二阶段的意义")
    add_table(
        doc,
        ["意义", "具体体现"],
        [
            ("从开环变为闭环", "占空比不再靠人工猜测，而是根据实际功率反馈自动修正。"),
            ("验证离散控制结构", "为后续代码生成、控制器上板和多速率设计建立基础。"),
            ("建立抗发散边界", "Unit Delay 切断代数环，Saturation 防止占空比越界。"),
            ("形成可量化基线", "可以用稳态功率、占空比范围和负载阶跃恢复情况评价控制器。"),
            ("明确后续阶段边界", "本阶段只完成 MPPT；母线稳压仍需电池侧闭环。"),
        ],
        widths=[1.55, 5.00],
        font_size=9,
    )
    add_callout(
        doc,
        "它不等于完整的能量管理系统",
        "第二阶段解决的是“光伏侧尽量取到最大功率”。它没有让电池变换器调节直流母线，因此负载从 50 Ω 变为 25 Ω 后，Vdc 出现瞬态并最终约为 597 V 是合理现象，不代表 MPPT 本身失败。",
        fill=PALE_RED,
        border="C65911",
    )

    add_heading(doc, "3", "控制原理与离散公式")
    add_text(doc, "本模型采用“直接扰动占空比”的 P&O 形式，不显式计算 ΔV。方向状态 s 记录占空比上一次移动的方向。只要功率没有显著下降，就沿原方向继续走；如果功率下降，就反向。")
    add_code(
        doc,
        "V[k] = Vpv_meas[k−1]\n"
        "I[k] = Ipv_meas[k−1]\n"
        "P[k] = V[k] × I[k]\n"
        "ΔP[k] = P[k] − P[k−1]",
    )
    add_code(
        doc,
        "若 ΔP[k] < −MPPT_dP_threshold：s[k] = −s[k−1]\n"
        "否则：                           s[k] =  s[k−1]\n\n"
        "Dcandidate[k] = D[k−1] + s[k] × MPPT_D_step\n"
        "D[k] = sat(Dcandidate[k], MPPT_D_min, MPPT_D_max)",
    )
    add_table(
        doc,
        ["符号", "含义", "在模型中的实现"],
        [
            ("P[k]", "当前离散光伏功率", "PV_Power"),
            ("P[k−1]", "前一拍功率", "Previous_Power"),
            ("s[k]", "当前扰动方向，±1", "Direction_Update 输出"),
            ("s[k−1]", "前一拍扰动方向", "Direction_State"),
            ("D[k−1]", "前一拍限幅后占空比", "Previous_Duty"),
            ("sat(·)", "上下限裁剪", "Duty_Limits"),
        ],
        widths=[1.15, 2.45, 2.90],
        font_size=9,
    )
    add_callout(
        doc,
        "为什么比较的是 ΔP < −0.1 W",
        "如果只判断 ΔP < 0，求解误差和很小的功率抖动也会触发反向。加入 0.1 W 的下降阈值可形成数值死区；阈值模块必须填 −MPPT_dP_threshold，而 Relational Operator 必须选择“<”。",
        fill=PALE_YELLOW,
        border="BF9000",
    )

    add_heading(doc, "4", "第二阶段模型架构")
    if FLOW_DIAGRAM.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FLOW_DIAGRAM), width=Inches(6.65))
        add_caption(doc, "图 1　第二阶段 P&O MPPT 的前向链与三条离散反馈")
    add_text(doc, "从实现角度看，控制器包含三类模块：")
    add_bullet(doc, "接口模块：Demux、Terminator，用于从向量中提取 Vpv/Ipv。")
    add_bullet(doc, "算法模块：Unit Delay、Product、Sum、Relational Operator、Switch、Gain。")
    add_bullet(doc, "约束与观察模块：Saturation、To Workspace、Control_Vector Mux。")
    add_callout(
        doc,
        "代数环防护",
        "对象输出 Vpv/Ipv 依赖占空比 D；如果测量链具有直接馈通，D 又在同一时刻由 Vpv/Ipv 计算，就可能形成 D → 对象 → 测量 → D 的代数环。两只 Unit Delay 让控制器使用上一拍测量，从结构上切断该环。",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "5", "开始前的准备")
    add_number(doc, r"在 Simulink 中打开 D:\PV_MPPT\model\PV_MPPT_v1.slx。")
    add_number(doc, "选择“文件 → 另存为”，保存为 PV_MPPT_v2.slx。不要直接覆盖第一阶段模型。")
    add_number(doc, "确认第一阶段模型可以执行 Ctrl+D（更新模型）且没有报错。")
    add_number(doc, "在模型空白区域规划控制链位置：建议把测量提取放在对象右侧，功率判断放中间，占空比更新靠近 Control_Vector。")
    add_number(doc, "打开“建模 → 模型资源管理器（Model Explorer）”，准备在 Model Workspace 中建立参数。")
    add_callout(
        doc,
        "建议的搭建策略",
        "一次只完成一个小闭环：先提取信号，再做功率计算，再做方向判断，最后接占空比反馈。每完成一组就按 Ctrl+D。这样一旦出错，范围很小。",
        fill=PALE_GREEN,
        border="70AD47",
    )
    add_heading(doc, "5.1", "需要保留的第一阶段模块", level=2)
    add_table(
        doc,
        ["模块", "保持设置", "用途"],
        [
            ("Plant_Derivatives", "不修改内部算法", "连续平均值对象的状态导数与辅助量"),
            ("Plant_State_Integrator", "IC 保持 6 维初始状态", "积分得到 x"),
            ("Disturbance_Vector", "3 输入", "温度、辐照度、负载电阻"),
            ("Battery_Duty", "0.49920850480429", "电池侧占空比，本阶段仍开环"),
            ("Control_Vector", "2 输入", "u(1)=Dpv，u(2)=Dbat"),
        ],
        widths=[2.10, 2.05, 2.35],
        font_size=8.7,
    )

    add_heading(doc, "6", "建立 Model Workspace 参数")
    add_text(doc, "在 Model Explorer 左侧选择 PV_MPPT_v2 → Model Workspace。逐项添加下表变量；名称必须完全一致，包括大小写。")
    add_table(
        doc,
        ["变量名", "值", "单位", "模块用途"],
        [
            ("MPPT_Ts", "0.002", "s", "所有 MPPT Unit Delay 的 Sample time"),
            ("MPPT_D_init", "0.5656125", "1", "Previous_Duty 初值"),
            ("MPPT_D_step", "0.0005", "1", "Duty_Perturbation 增益"),
            ("MPPT_D_min", "0.05", "1", "Duty_Limits 下限"),
            ("MPPT_D_max", "0.90", "1", "Duty_Limits 上限"),
            ("MPPT_dP_threshold", "0.1", "W", "功率下降判定死区"),
            ("MPPT_P_init", "9591.75", "W", "Previous_Power 初值"),
        ],
        widths=[2.05, 1.25, 0.70, 2.50],
        font_size=8.6,
    )
    add_text(doc, "这些值的含义：")
    add_bullet(doc, "2 ms 是控制器的实际更新周期；相当于 500 Hz 更新率。")
    add_bullet(doc, "0.0005 表示每次只改变 0.05% 的占空比，兼顾收敛速度和稳态抖动。")
    add_bullet(doc, "D_init 接近理想 Boost 关系 1 − Vpv/Vdc ≈ 1 − 261/600 = 0.565。")
    add_bullet(doc, "P_init 使用额定最大功率 261 × 36.75 = 9591.75 W，减小第一拍的非物理跳变。")

    add_heading(doc, "7", "搭建测量提取与离散采样")
    add_heading(doc, "7.1", "放置模块", level=2)
    add_table(
        doc,
        ["模块名", "库路径", "参数设置", "作用"],
        [
            ("State_Vpv_Split", "Simulink / Signal Routing / Demux", "Outputs = [1 5]", "将 6 维 x 拆成 Vpv=x(1) 和其余 5 维"),
            ("Unused_State_Terminator", "Simulink / Sinks / Terminator", "默认", "终止未使用的 5 维状态"),
            ("Aux_Ipv_Split", "Simulink / Signal Routing / Demux", "Outputs = [1 9]", "将 10 维 aux 拆成 Ipv=aux(1) 和其余 9 维"),
            ("Unused_Aux_Terminator", "Simulink / Sinks / Terminator", "默认", "终止未使用的 9 维辅助量"),
            ("Vpv_Sample_Hold", "Simulink / Discrete / Unit Delay", "IC=261；Sample time=MPPT_Ts", "把连续 Vpv 转为 2 ms 离散值并延迟一拍"),
            ("Ipv_Sample_Hold", "Simulink / Discrete / Unit Delay", "IC=36.75；Sample time=MPPT_Ts", "把连续 Ipv 转为 2 ms 离散值并延迟一拍"),
        ],
        widths=[1.62, 1.68, 1.55, 1.65],
        font_size=7.8,
    )
    add_heading(doc, "7.2", "连线", level=2)
    for item in [
        "从 Plant_State_Integrator 输出 x 分支一根线，连接 State_Vpv_Split 输入。",
        "State_Vpv_Split 第 1 输出连接 Vpv_Sample_Hold；第 2 输出连接 Unused_State_Terminator。",
        "从 Plant_Derivatives 的 aux 输出分支一根线，连接 Aux_Ipv_Split 输入。",
        "Aux_Ipv_Split 第 1 输出连接 Ipv_Sample_Hold；第 2 输出连接 Unused_Aux_Terminator。",
    ]:
        add_number(doc, item)
    add_callout(
        doc,
        "模块名称里的 Sample_Hold 只是功能命名",
        "实际块类型必须是 Unit Delay。不要因为名称而换成 Zero-Order Hold；后者具有直接馈通，在此闭环中可能重新引入代数环。",
        fill=PALE_RED,
        border="C65911",
    )

    add_heading(doc, "8", "搭建功率计算与功率下降判断")
    add_heading(doc, "8.1", "功率计算", level=2)
    add_table(
        doc,
        ["模块名", "库路径", "参数设置", "作用"],
        [
            ("PV_Power", "Simulink / Math Operations / Product", "Inputs=2；Multiplication=Element-wise (.*)", "计算 P=V×I"),
            ("Previous_Power", "Simulink / Discrete / Unit Delay", "IC=MPPT_P_init；Ts=MPPT_Ts", "保存 P[k−1]"),
            ("Delta_Power", "Simulink / Math Operations / Sum", "List of signs = +−", "计算 P[k]−P[k−1]"),
        ],
        widths=[1.55, 1.95, 1.55, 1.45],
        font_size=8.0,
    )
    for item in [
        "Vpv_Sample_Hold 输出连接 PV_Power 输入 1。",
        "Ipv_Sample_Hold 输出连接 PV_Power 输入 2。",
        "PV_Power 输出分支连接 Previous_Power 输入和 Delta_Power 的“+”端。",
        "Previous_Power 输出连接 Delta_Power 的“−”端。",
    ]:
        add_number(doc, item)
    add_heading(doc, "8.2", "功率下降判定", level=2)
    add_table(
        doc,
        ["模块名", "库路径", "参数设置", "作用"],
        [
            ("Negative_Power_Threshold", "Simulink / Sources / Constant", "Value=−MPPT_dP_threshold；Sample time=inf", "生成 −0.1 W 判据"),
            ("Power_Decreased", "Simulink / Logic and Bit Operations / Relational Operator", "Operator=<", "判断 ΔP 是否低于 −0.1 W"),
        ],
        widths=[1.70, 2.25, 1.35, 1.20],
        font_size=7.9,
    )
    add_number(doc, "Delta_Power 输出连接 Power_Decreased 输入 1。")
    add_number(doc, "Negative_Power_Threshold 输出连接 Power_Decreased 输入 2。")
    add_text(doc, "Power_Decreased 输出为布尔量：真表示本次沿原方向扰动后功率显著变差，需要反向。")

    add_heading(doc, "9", "搭建扰动方向状态机")
    add_table(
        doc,
        ["模块名", "库路径", "参数设置", "作用"],
        [
            ("Direction_State", "Simulink / Discrete / Unit Delay", "IC=1；Ts=MPPT_Ts", "保存上一拍方向 s[k−1]"),
            ("Reverse_Direction", "Simulink / Math Operations / Gain", "Gain=−1", "把方向 +1/−1 取反"),
            ("Direction_Update", "Simulink / Signal Routing / Switch", "Criteria=u2 >= Threshold；Threshold=0.5", "功率下降时选反向，否则维持原方向"),
        ],
        widths=[1.58, 1.88, 1.67, 1.37],
        font_size=7.9,
    )
    add_text(doc, "Switch 三个输入端的连线顺序不能错：")
    add_table(
        doc,
        ["Switch 端口", "连接信号", "含义"],
        [
            ("u1（上）", "Reverse_Direction 输出", "条件为真时采用反向方向"),
            ("u2（中）", "Power_Decreased 输出", "控制条件，布尔真≈1"),
            ("u3（下）", "Direction_State 输出", "条件为假时保持上一方向"),
        ],
        widths=[1.30, 2.65, 2.55],
        font_size=9,
    )
    add_text(doc, "随后把 Direction_Update 输出同时分支到：")
    add_bullet(doc, "Direction_State 输入——在下一拍保存本次方向。")
    add_bullet(doc, "Duty_Perturbation 输入——把方向转换成带符号的占空比增量。")
    add_callout(
        doc,
        "判断逻辑快速自检",
        "Power_Decreased=1 时，Switch 必须输出 −s[k−1]；Power_Decreased=0 时，必须输出 s[k−1]。若结果相反，检查 u1/u3 是否接反。",
        fill=PALE_YELLOW,
        border="BF9000",
    )

    add_heading(doc, "10", "搭建占空比更新与限幅")
    add_table(
        doc,
        ["模块名", "库路径", "参数设置", "作用"],
        [
            ("Duty_Perturbation", "Simulink / Math Operations / Gain", "Gain=MPPT_D_step", "得到 ±0.0005 的增量"),
            ("Previous_Duty", "Simulink / Discrete / Unit Delay", "IC=MPPT_D_init；Ts=MPPT_Ts", "保存上一拍限幅后占空比"),
            ("Duty_Candidate", "Simulink / Math Operations / Sum", "List of signs = ++", "D[k−1]+ΔD"),
            ("Duty_Limits", "Simulink / Discontinuities / Saturation", "Lower=MPPT_D_min；Upper=MPPT_D_max", "限制占空比到 0.05～0.90"),
        ],
        widths=[1.55, 1.90, 1.55, 1.50],
        font_size=7.9,
    )
    for item in [
        "Direction_Update 输出连接 Duty_Perturbation。",
        "Previous_Duty 输出连接 Duty_Candidate 输入 1。",
        "Duty_Perturbation 输出连接 Duty_Candidate 输入 2。",
        "Duty_Candidate 输出连接 Duty_Limits。",
        "Duty_Limits 输出反馈到 Previous_Duty 输入。",
    ]:
        add_number(doc, item)
    add_callout(
        doc,
        "反馈点必须在限幅之后",
        "Previous_Duty 要保存 Duty_Limits 的输出，而不是 Duty_Candidate。否则一旦候选值越界，内部记忆会继续跑到范围之外，控制器可能长期贴住限幅甚至恢复缓慢。",
        fill=PALE_RED,
        border="C65911",
    )

    add_heading(doc, "11", "接入对象模型并增加日志")
    add_heading(doc, "11.1", "替换固定光伏占空比", level=2)
    add_number(doc, "找到原先连接 Control_Vector 输入 1 的固定光伏占空比常数。")
    add_number(doc, "删除该常数到 Control_Vector 输入 1 的连线；若该常数只为此用途，可删除该常数块。")
    add_number(doc, "把 Duty_Limits 输出连接 Control_Vector 输入 1。")
    add_number(doc, "保持 Battery_Duty=0.49920850480429 连接 Control_Vector 输入 2。")
    add_number(doc, "确认 Control_Vector 的 Inputs=2，输出连接 Plant_Derivatives 的控制输入 u。")
    add_heading(doc, "11.2", "日志模块", level=2)
    add_table(
        doc,
        ["模块名", "库路径", "设置", "连接"],
        [
            ("Power_Log", "Simulink / Sinks / To Workspace", "Variable=phase2_power；Save format=Timeseries；Decimation=1；Max data points=inf", "PV_Power 输出"),
            ("Duty_Log", "Simulink / Sinks / To Workspace", "Variable=phase2_duty；Save format=Timeseries；Decimation=1；Max data points=inf", "Duty_Limits 输出"),
        ],
        widths=[1.25, 1.70, 2.65, 0.90],
        font_size=7.8,
    )
    add_text(doc, "第一阶段的 State_Log（phase1_x）和 Aux_Log（phase1_aux）保留不动，这样可以同时检查状态、辅助量、功率和占空比。")
    if ACTUAL_CROP.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ACTUAL_CROP), width=Inches(6.65))
        add_caption(doc, "图 2　现有 PV_MPPT_v2 中的 MPPT 控制链局部截图（用于核对模块相对关系）")

    add_heading(doc, "12", "完整连线核对表")
    wiring_rows = [
        ("Plant_State_Integrator.y", "State_Vpv_Split.u1", "提取 Vpv"),
        ("State_Vpv_Split.y1", "Vpv_Sample_Hold.u1", "标量 Vpv"),
        ("State_Vpv_Split.y2", "Unused_State_Terminator.u1", "终止剩余 5 维"),
        ("Plant_Derivatives.aux", "Aux_Ipv_Split.u1", "提取 Ipv"),
        ("Aux_Ipv_Split.y1", "Ipv_Sample_Hold.u1", "标量 Ipv"),
        ("Aux_Ipv_Split.y2", "Unused_Aux_Terminator.u1", "终止剩余 9 维"),
        ("Vpv_Sample_Hold.y1", "PV_Power.u1", "功率乘法电压端"),
        ("Ipv_Sample_Hold.y1", "PV_Power.u2", "功率乘法电流端"),
        ("PV_Power.y1", "Previous_Power.u1", "保存当前功率"),
        ("PV_Power.y1", "Delta_Power.u1 (+)", "当前功率"),
        ("Previous_Power.y1", "Delta_Power.u2 (−)", "前一拍功率"),
        ("Delta_Power.y1", "Power_Decreased.u1", "被比较量"),
        ("Negative_Power_Threshold.y1", "Power_Decreased.u2", "−0.1 W 阈值"),
        ("Direction_State.y1", "Reverse_Direction.u1", "待反向方向"),
        ("Reverse_Direction.y1", "Direction_Update.u1", "Switch 真值输入"),
        ("Power_Decreased.y1", "Direction_Update.u2", "Switch 控制输入"),
        ("Direction_State.y1", "Direction_Update.u3", "Switch 假值输入"),
        ("Direction_Update.y1", "Direction_State.u1", "方向状态反馈"),
        ("Direction_Update.y1", "Duty_Perturbation.u1", "带符号步进方向"),
        ("Previous_Duty.y1", "Duty_Candidate.u1", "上一拍占空比"),
        ("Duty_Perturbation.y1", "Duty_Candidate.u2", "占空比增量"),
        ("Duty_Candidate.y1", "Duty_Limits.u1", "候选占空比"),
        ("Duty_Limits.y1", "Previous_Duty.u1", "限幅后状态反馈"),
        ("Duty_Limits.y1", "Control_Vector.u1", "光伏占空比输出"),
        ("Battery_Duty.y1", "Control_Vector.u2", "固定电池占空比"),
        ("Control_Vector.y1", "Plant_Derivatives.u", "[Dpv, Dbat]"),
        ("PV_Power.y1", "Power_Log.u1", "功率日志"),
        ("Duty_Limits.y1", "Duty_Log.u1", "占空比日志"),
    ]
    add_table(
        doc,
        ["源端", "目标端", "用途"],
        wiring_rows,
        widths=[2.20, 2.35, 1.95],
        font_size=7.45,
    )

    add_heading(doc, "13", "求解器、采样周期与仿真设置")
    add_table(
        doc,
        ["设置位置", "项目", "值", "用途"],
        [
            ("模型设置 → Solver", "Type", "Fixed-step", "保证离散控制更新时间确定"),
            ("模型设置 → Solver", "Solver", "ode4 (Runge-Kutta)", "积分连续平均值对象"),
            ("模型设置 → Solver", "Fixed-step size", "1e-5 s", "连续对象数值积分步长"),
            ("模型工具栏", "Stop time", "8 s", "覆盖 4 s 负载阶跃及恢复"),
            ("MPPT Unit Delay", "Sample time", "MPPT_Ts=0.002 s", "控制器实际更新周期"),
            ("Load_Resistance_Step", "Time/Before/After", "4 s / 50 Ω / 25 Ω", "验证负载突变下的稳定性"),
        ],
        widths=[1.70, 1.55, 1.35, 1.90],
        font_size=8.3,
    )
    add_text(doc, "时间尺度关系：每个 MPPT 周期包含 0.002/0.00001 = 200 个对象积分步。控制器在两次更新之间保持占空比不变，而连续对象仍以 10 μs 步长演化。")
    add_callout(
        doc,
        "如何修改 MPPT 实际周期",
        "只修改 Model Workspace 中的 MPPT_Ts，并保持所有 MPPT Unit Delay 的 Sample time 都引用 MPPT_Ts。不要通过改 Solver 的 Fixed-step size 来修改 MPPT 周期。修改后按 Ctrl+D，并显示 Sample Time Colors 检查整条控制链是否为同一离散速率。",
        fill=LIGHT_BLUE,
    )
    add_heading(doc, "13.1", "为什么最终选择 2 ms / 0.0005", level=2)
    add_table(
        doc,
        ["MPPT 周期", "步长", "稳态平均功率", "功率标准差", "占空比跨度", "结论"],
        [
            ("1 ms", "0.0010", "9562.35 W", "46.08 W", "0.037", "扰动偏激进"),
            ("1 ms", "0.0005", "9557.54 W", "44.27 W", "0.023", "仍有较大波动"),
            ("2 ms", "0.0010", "9602.70 W", "1.81 W", "0.007", "可用"),
            ("2 ms", "0.0005", "9603.78 W", "0.70 W", "0.005", "选定"),
            ("5 ms", "0.0010", "9526.35 W", "44.39 W", "—", "跟踪偏慢"),
            ("5 ms", "0.0005", "9274.65 W", "31.28 W", "—", "步进不足"),
        ],
        widths=[0.85, 0.72, 1.26, 1.05, 0.92, 1.20],
        font_size=7.6,
    )
    add_text(doc, "这些结果仅适用于当前平均值对象、初始点和负载阶跃；换成开关模型、不同 PV 阵列或不同采样滤波后，应重新扫参。")

    add_heading(doc, "14", "编译、运行和验收")
    add_heading(doc, "14.1", "编译检查", level=2)
    for item in [
        "按 Ctrl+D 更新模型；不得出现 algebraic loop、端口维度、采样时间或变量未定义错误。",
        "打开“调试/显示 → Sample Time → Colors”（不同版本菜单名略有差异），确认 Vpv_Sample_Hold 之后的控制链均为 0.002 s。",
        "检查 State_Vpv_Split 的两路宽度为 1 和 5；Aux_Ipv_Split 的两路宽度为 1 和 9。",
        "确认所有未使用的 Demux 输出已接 Terminator。",
    ]:
        add_number(doc, item)
    add_heading(doc, "14.2", "运行检查", level=2)
    for item in [
        "仿真 Stop time 设为 8 s，点击 Run。",
        "在 MATLAB Workspace 中确认生成 phase1_x、phase1_aux、phase2_power、phase2_duty。",
        "检查 phase2_duty 全程位于 0.05～0.90，且没有 NaN/Inf。",
        "检查 phase2_power 在负载阶跃后重新回到约 9.60 kW，而不是持续下降。",
        "查看 x 的第 6 个分量 Vdc：阶跃后允许有瞬态和小偏差，因为电池侧仍开环。",
    ]:
        add_number(doc, item)
    add_heading(doc, "14.3", "当前模型的参考结果", level=2)
    add_table(
        doc,
        ["时间区间", "Vpv 均值", "Ipv 均值", "Ppv 均值", "Vdc 均值", "Dpv 均值"],
        [
            ("3.5～3.9 s（阶跃前）", "264.875 V", "36.258 A", "9603.75 W", "600.003 V", "0.55914"),
            ("4.0～4.5 s（阶跃恢复）", "265.407 V", "36.127 A", "9585.79 W", "596.251 V", "0.55547"),
            ("7.5～8.0 s（最终）", "264.855 V", "36.261 A", "9603.78 W", "597.179 V", "0.55710"),
        ],
        widths=[1.70, 1.02, 1.02, 1.12, 1.05, 0.95],
        font_size=8,
    )
    add_text(doc, "全程占空比参考范围为 0.5471125～0.5666125；所有状态、辅助量、功率和占空比均应为有限值。")
    add_callout(
        doc,
        "验收结论怎么判断",
        "MPPT 是否成功，主要看光伏功率是否稳定回到最大功率附近、占空比是否有界以及信号是否有限；不要用“Vdc 是否精确等于 600 V”作为本阶段唯一判据。",
        fill=PALE_GREEN,
        border="70AD47",
    )

    add_heading(doc, "15", "常见问题与发散排查")
    troubleshooting = [
        ("报 algebraic loop", "测量链使用了直接馈通模块，或 Duty 回路缺少离散状态。", "Vpv/Ipv 入口改用 Unit Delay；确认 Previous_Duty 和 Direction_State 都存在。"),
        ("占空比迅速打到 0.05 或 0.90", "Delta_Power 符号、Switch u1/u3 或增益符号接反。", "核对“当前 P − 前一拍 P”；真时取 −s，假时取 s。"),
        ("占空比完全不动", "Direction_Update 未反馈，Gain 为 0，或控制输出仍接旧常数。", "检查 Direction_Update→Direction_State、MPPT_D_step 和 Control_Vector.u1。"),
        ("最大功率点附近抖动很大", "MPPT_Ts 太短或 MPPT_D_step 太大。", "先恢复 2 ms / 0.0005，再一次只改一个参数。"),
        ("跟踪很慢或停在非最大点", "周期太长、步长太小或死区太大。", "减小 MPPT_Ts、增大 D_step 或减小 dP_threshold，并重新扫参。"),
        ("第一拍出现大跳变", "Unit Delay 初值与对象初值不匹配。", "使用 V=261 V、I=36.75 A、P=9591.75 W、D=0.5656125。"),
        ("负载阶跃后 Vdc 不回 600 V", "电池占空比仍为固定常数。", "这是阶段边界；下一阶段增加电池侧母线电压闭环。"),
        ("日志采样数不对", "To Workspace 接在连续信号上或控制链采样时间未传播。", "日志接 PV_Power 和 Duty_Limits；编译后确认两者为 0.002 s。"),
    ]
    add_table(
        doc,
        ["现象", "最可能原因", "处理方法"],
        troubleshooting,
        widths=[1.65, 2.25, 2.60],
        font_size=7.8,
    )
    add_heading(doc, "15.1", "推荐排查顺序", level=2)
    for item in [
        "先断言占空比是否有限且在 0.05～0.90；若否，先查限幅与反馈。",
        "再看 P、ΔP 和 Power_Decreased；确认功率下降时布尔量变为 1。",
        "再看 Direction_State；它应只在 +1 与 −1 之间切换。",
        "再看 Duty_Candidate 与 Duty_Limits；确认每拍变化量约为 0.0005。",
        "最后再看对象状态和 Vdc；避免把电池侧开环问题误认为 MPPT 发散。",
    ]:
        add_number(doc, item)

    add_heading(doc, "16", "独立复现完成清单")
    checklist = [
        ("待确认", "已从 PV_MPPT_v1 另存为 PV_MPPT_v2，第一阶段文件未被覆盖"),
        ("待确认", "Model Workspace 中 7 个 MPPT 参数名称和值正确"),
        ("待确认", "x 拆分为 [1 5]，aux 拆分为 [1 9]，剩余输出均有 Terminator"),
        ("待确认", "Vpv/Ipv 使用 Unit Delay，周期 2 ms，初值分别为 261 和 36.75"),
        ("待确认", "P=V×I，ΔP=当前 P−前一拍 P"),
        ("待确认", "功率下降判据为 ΔP<−0.1 W"),
        ("待确认", "Switch 真时反向、假时保持，Direction_State 初值为 +1"),
        ("待确认", "占空比步长为 0.0005，Previous_Duty 初值为 0.5656125"),
        ("待确认", "Duty_Limits 为 0.05～0.90，Previous_Duty 反馈取限幅后信号"),
        ("待确认", "Duty_Limits 接 Control_Vector.u1，Battery_Duty 接 u2"),
        ("待确认", "求解器为 fixed-step ode4，步长 1e-5 s，Stop time=8 s"),
        ("待确认", "Ctrl+D 无结构错误，控制链编译采样周期为 0.002 s"),
        ("待确认", "8 s 仿真所有日志有限，功率最终约 9.60 kW"),
    ]
    add_table(doc, ["完成", "检查项"], checklist, widths=[0.75, 5.75], font_size=9.2, header_fill=PALE_GREEN)

    add_heading(doc, "17", "下一阶段建议")
    add_text(doc, "在本阶段稳定后，下一步建议增加电池侧直流母线电压闭环：以 Vdc−Vdc_ref 为误差，通过离散 PI 或更高阶控制器调节 Battery_Duty，并加入功率/电流/SOC 限制。")
    add_table(
        doc,
        ["后续功能", "为什么需要", "与第二阶段的关系"],
        [
            ("电池侧 Vdc 闭环", "负载变化时恢复直流母线电压", "与 MPPT 并行，作用于 Control_Vector.u2"),
            ("辐照度变化测试", "验证经典 P&O 的误判与恢复", "检验当前方向逻辑的边界"),
            ("自适应步长", "兼顾快速跟踪和低稳态纹波", "替换固定 MPPT_D_step"),
            ("开关级模型验证", "确认 PWM、纹波和采样同步影响", "需要重新选择采样周期与滤波"),
        ],
        widths=[1.55, 2.35, 2.60],
        font_size=8.8,
    )
    add_callout(
        doc,
        "本说明书对应的已验证模型",
        r"D:\PV_MPPT\model\PV_MPPT_v2.slx。模型结构检查结果为 healthy：无未连接端口、无悬空连线、无 Stateflow 编辑期错误。",
        fill=LIGHT_BLUE,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Created {OUT}")


if __name__ == "__main__":
    build_manual()
